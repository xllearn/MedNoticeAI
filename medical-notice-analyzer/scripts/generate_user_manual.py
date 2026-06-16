from __future__ import annotations

from datetime import date
from pathlib import Path
from xml.sax.saxutils import escape
from zipfile import BadZipFile, ZipFile

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from docx.table import _Cell
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
ASSET_DIR = ROOT / "manual-assets"
OUTPUT = ROOT / "医械公告智析报告系统 V1.0 操作手册.docx"

TITLE = "医械公告智析报告系统 V1.0 操作手册"
SYSTEM_NAME = "医械公告智析报告系统 V1.0"


def cjk_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        Path(r"C:\Windows\Fonts\msyhbd.ttc") if bold else Path(r"C:\Windows\Fonts\msyh.ttc"),
        Path(r"C:\Windows\Fonts\simhei.ttf"),
        Path(r"C:\Windows\Fonts\simsun.ttc"),
    ]
    for path in candidates:
        if path.exists():
            return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


def text_width(text: str, font: ImageFont.ImageFont) -> int:
    box = font.getbbox(text)
    return box[2] - box[0]


def wrap_text(text: str, font: ImageFont.ImageFont, max_width: int) -> list[str]:
    lines: list[str] = []
    current = ""
    for ch in text:
        test = current + ch
        if text_width(test, font) <= max_width or not current:
            current = test
        else:
            lines.append(current)
            current = ch
    if current:
        lines.append(current)
    return lines


def draw_center_text(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    text: str,
    font: ImageFont.ImageFont,
    fill: str = "#111827",
    line_gap: int = 6,
) -> None:
    x1, y1, x2, y2 = box
    lines = wrap_text(text, font, x2 - x1 - 32)
    heights = [font.getbbox(line)[3] - font.getbbox(line)[1] for line in lines]
    total_h = sum(heights) + max(0, len(lines) - 1) * line_gap
    y = y1 + (y2 - y1 - total_h) // 2
    for line, h in zip(lines, heights):
        x = x1 + (x2 - x1 - text_width(line, font)) // 2
        draw.text((x, y), line, font=font, fill=fill)
        y += h + line_gap


def rounded_box(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    text: str,
    fill: str,
    outline: str = "#2563eb",
    text_fill: str = "#111827",
    radius: int = 14,
    font: ImageFont.ImageFont | None = None,
) -> None:
    font = font or cjk_font(30)
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=3)
    draw_center_text(draw, box, text, font, fill=text_fill)


def arrow(
    draw: ImageDraw.ImageDraw,
    start: tuple[int, int],
    end: tuple[int, int],
    fill: str = "#334155",
    width: int = 4,
) -> None:
    draw.line([start, end], fill=fill, width=width)
    sx, sy = start
    ex, ey = end
    if abs(ex - sx) >= abs(ey - sy):
        direction = 1 if ex >= sx else -1
        points = [(ex, ey), (ex - direction * 16, ey - 10), (ex - direction * 16, ey + 10)]
    else:
        direction = 1 if ey >= sy else -1
        points = [(ex, ey), (ex - 10, ey - direction * 16), (ex + 10, ey - direction * 16)]
    draw.polygon(points, fill=fill)


def image_canvas(width: int = 1600, height: int = 900) -> tuple[Image.Image, ImageDraw.ImageDraw]:
    img = Image.new("RGB", (width, height), "#ffffff")
    draw = ImageDraw.Draw(img)
    return img, draw


def save_diagram(filename: str, painter) -> Path:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    path = ASSET_DIR / filename
    img, draw = image_canvas()
    painter(img, draw)
    img.save(path, quality=95)
    return path


def draw_title(draw: ImageDraw.ImageDraw, text: str) -> None:
    font = cjk_font(44, bold=True)
    draw.text((80, 55), text, font=font, fill="#0f172a")
    draw.line([(80, 122), (1520, 122)], fill="#94a3b8", width=3)


def paint_architecture(_: Image.Image, draw: ImageDraw.ImageDraw) -> None:
    draw_title(draw, "系统总体架构图")
    f = cjk_font(30)
    boxes = {
        "user": (95, 250, 315, 390),
        "dify": (430, 230, 690, 410),
        "service": (820, 220, 1120, 420),
        "source": (280, 610, 570, 770),
        "llm": (700, 610, 980, 770),
        "word": (1120, 610, 1460, 770),
    }
    rounded_box(draw, boxes["user"], "业务人员\n输入公告 URL", "#eff6ff", "#2563eb", font=f)
    rounded_box(draw, boxes["dify"], "Dify 应用\n流程编排与模型调用", "#f0fdf4", "#16a34a", font=f)
    rounded_box(draw, boxes["service"], "本地辅助服务\nFastAPI / Docker", "#fff7ed", "#ea580c", font=f)
    rounded_box(draw, boxes["source"], "公告网页与附件\nHTML / Word / Excel / PDF / CSV", "#f8fafc", "#64748b", font=f)
    rounded_box(draw, boxes["llm"], "LLM 报告生成\nReportIR + Markdown", "#faf5ff", "#9333ea", font=f)
    rounded_box(draw, boxes["word"], "输出结果\nMarkdown 展示 + Word 下载", "#fef2f2", "#dc2626", font=f)
    arrow(draw, (315, 320), (430, 320))
    arrow(draw, (690, 320), (820, 320))
    arrow(draw, (970, 420), (970, 610))
    arrow(draw, (570, 690), (820, 360))
    arrow(draw, (1120, 690), (1120, 420))
    arrow(draw, (980, 690), (1120, 690))
    draw.text((92, 830), "说明：Dify Docker 容器访问本地服务时使用 http://host.docker.internal:8099。", font=cjk_font(26), fill="#475569")


def paint_business_flow(_: Image.Image, draw: ImageDraw.ImageDraw) -> None:
    draw_title(draw, "公告分析与报告导出流程图")
    steps = [
        "输入公告 URL",
        "抓取网页正文",
        "发现并解析附件",
        "生成证据包",
        "生成 ReportIR 与 Markdown",
        "报告质检",
        "自动修复与二次质检",
        "导出 Word 并返回下载链接",
    ]
    x0, y, w, h, gap = 90, 220, 300, 100, 42
    for i, step in enumerate(steps):
        row = i // 4
        col = i % 4
        x = x0 + col * (w + gap)
        yy = y + row * 250
        rounded_box(draw, (x, yy, x + w, yy + h), f"{i + 1}. {step}", "#eef2ff", "#4f46e5", font=cjk_font(28))
        if col < 3:
            arrow(draw, (x + w, yy + h // 2), (x + w + gap, yy + h // 2), "#475569")
        elif row == 0:
            arrow(draw, (x + w // 2, yy + h), (x + w // 2, yy + 250), "#475569")
    arrow(draw, (x0 + 3 * (w + gap), y + 250 + h // 2), (x0 + 2 * (w + gap) + w, y + 250 + h // 2), "#475569")
    arrow(draw, (x0 + 2 * (w + gap), y + 250 + h // 2), (x0 + 1 * (w + gap) + w, y + 250 + h // 2), "#475569")
    arrow(draw, (x0 + 1 * (w + gap), y + 250 + h // 2), (x0 + w, y + 250 + h // 2), "#475569")
    rounded_box(draw, (1135, 725, 1505, 820), "质检阻断时：返回问题摘要，不生成 Word", "#fff1f2", "#e11d48", font=cjk_font(25))


def paint_modules(_: Image.Image, draw: ImageDraw.ImageDraw) -> None:
    draw_title(draw, "系统功能模块图")
    center = (600, 330, 1000, 500)
    rounded_box(draw, center, SYSTEM_NAME, "#eff6ff", "#2563eb", font=cjk_font(34, bold=True))
    modules = [
        ((105, 220, 420, 350), "公告采集模块\n网页正文、标题、日期"),
        ((105, 560, 420, 690), "附件解析模块\nWord / Excel / PDF / CSV"),
        ((650, 620, 950, 760), "证据包整理模块\n正文与表格摘要"),
        ((1180, 220, 1495, 350), "报告生成模块\nReportIR / Markdown"),
        ((1180, 560, 1495, 690), "质检修复模块\nQA / 修复 / 二次 QA"),
        ((650, 130, 950, 250), "Word 导出模块\n正式报告下载"),
    ]
    for box, label in modules:
        rounded_box(draw, box, label, "#f8fafc", "#64748b", font=cjk_font(27))
        arrow(draw, ((box[0] + box[2]) // 2, (box[1] + box[3]) // 2), ((center[0] + center[2]) // 2, (center[1] + center[3]) // 2), "#94a3b8", 3)


def paint_input_ui(_: Image.Image, draw: ImageDraw.ImageDraw) -> None:
    draw_title(draw, "公告链接输入界面示意图")
    draw.rounded_rectangle((120, 165, 1480, 805), 20, fill="#f8fafc", outline="#cbd5e1", width=3)
    draw.rectangle((120, 165, 1480, 235), fill="#0f172a")
    draw.text((155, 184), "Dify - 医械公告智析报告系统", font=cjk_font(28, bold=True), fill="#ffffff")
    draw.rounded_rectangle((160, 285, 440, 735), 12, fill="#ffffff", outline="#e2e8f0", width=2)
    for i, item in enumerate(["应用工作台", "公告分析", "历史记录", "运行日志"]):
        y = 330 + i * 80
        fill = "#dbeafe" if item == "公告分析" else "#ffffff"
        draw.rounded_rectangle((190, y, 410, y + 48), 8, fill=fill, outline="#dbeafe")
        draw.text((215, y + 10), item, font=cjk_font(24), fill="#1e293b")
    draw.text((505, 305), "输入公告链接", font=cjk_font(34, bold=True), fill="#0f172a")
    draw.text((505, 365), "公告 URL", font=cjk_font(26), fill="#334155")
    draw.rounded_rectangle((505, 405, 1355, 475), 10, fill="#ffffff", outline="#94a3b8", width=2)
    draw.text((530, 425), "https://example.gov.cn/notice/detail.html", font=cjk_font(25), fill="#475569")
    draw.rounded_rectangle((505, 525, 720, 590), 10, fill="#2563eb", outline="#1d4ed8")
    draw.text((560, 541), "开始分析", font=cjk_font(27, bold=True), fill="#ffffff")
    draw.text((505, 645), "提示：如在 Docker 内调用本地服务，应配置 host.docker.internal:8099。", font=cjk_font(23), fill="#64748b")


def paint_result_ui(_: Image.Image, draw: ImageDraw.ImageDraw) -> None:
    draw_title(draw, "报告结果与下载界面示意图")
    draw.rounded_rectangle((120, 165, 1480, 805), 20, fill="#ffffff", outline="#cbd5e1", width=3)
    draw.rectangle((120, 165, 1480, 235), fill="#0f172a")
    draw.text((155, 184), "Dify - 运行结果", font=cjk_font(28, bold=True), fill="#ffffff")
    draw.text((170, 285), "分析报告预览", font=cjk_font(32, bold=True), fill="#0f172a")
    for i, line in enumerate(["一、项目基本情况", "二、采购品种范围与产品分类", "三、企业报价与中选规则", "四、企业关注事项"]):
        draw.text((190, 350 + i * 58), line, font=cjk_font(26), fill="#334155")
    draw.rounded_rectangle((980, 300, 1380, 555), 16, fill="#f8fafc", outline="#cbd5e1", width=2)
    draw.text((1025, 335), "质检状态", font=cjk_font(28, bold=True), fill="#0f172a")
    draw.rounded_rectangle((1025, 390, 1260, 450), 12, fill="#dcfce7", outline="#16a34a")
    draw.text((1075, 405), "已通过", font=cjk_font(27, bold=True), fill="#166534")
    draw.text((1025, 485), "已生成 Word 文件", font=cjk_font(25), fill="#475569")
    draw.rounded_rectangle((980, 610, 1380, 685), 12, fill="#2563eb", outline="#1d4ed8")
    draw.text((1075, 632), "下载 Word 报告", font=cjk_font(28, bold=True), fill="#ffffff")
    draw.text((170, 735), "质检阻断时，页面显示问题摘要和修复建议，不显示下载按钮。", font=cjk_font(24), fill="#64748b")


def paint_app_home_ui(_: Image.Image, draw: ImageDraw.ImageDraw) -> None:
    draw_title(draw, "Dify 应用入口界面示意图")
    draw.rounded_rectangle((95, 150, 1505, 815), 20, fill="#f8fafc", outline="#cbd5e1", width=3)
    draw.rectangle((95, 150, 1505, 225), fill="#111827")
    draw.text((135, 172), "Dify 应用工作台", font=cjk_font(28, bold=True), fill="#ffffff")
    draw.rounded_rectangle((130, 270, 430, 735), 12, fill="#ffffff", outline="#e2e8f0", width=2)
    for i, item in enumerate(["全部应用", "Workflow", "Chatflow", "工具", "设置"]):
        y = 315 + i * 68
        fill = "#dbeafe" if item == "Workflow" else "#ffffff"
        draw.rounded_rectangle((160, y, 400, y + 44), 8, fill=fill, outline="#dbeafe")
        draw.text((190, y + 8), item, font=cjk_font(23), fill="#1e293b")
    draw.text((500, 295), "选择应用", font=cjk_font(34, bold=True), fill="#0f172a")
    cards = [
        ((500, 360, 880, 520), "医械公告智析报告系统", "公告链接分析、报告生成、质检导出"),
        ((930, 360, 1310, 520), "历史报告修订助手", "历史稿参考与多轮修订蓝本"),
        ((500, 570, 880, 730), "通用文本分析", "文本摘要与结构化抽取"),
    ]
    for box, title, desc in cards:
        draw.rounded_rectangle(box, 12, fill="#ffffff", outline="#cbd5e1", width=2)
        draw.text((box[0] + 28, box[1] + 30), title, font=cjk_font(27, bold=True), fill="#0f172a")
        draw.text((box[0] + 28, box[1] + 85), desc, font=cjk_font(22), fill="#64748b")
        if "医械" in title:
            draw.rounded_rectangle((box[2] - 130, box[3] - 58, box[2] - 30, box[3] - 22), 8, fill="#2563eb")
            draw.text((box[2] - 106, box[3] - 52), "进入", font=cjk_font(21, bold=True), fill="#ffffff")


def paint_workflow_running_ui(_: Image.Image, draw: ImageDraw.ImageDraw) -> None:
    draw_title(draw, "工作流运行状态界面示意图")
    draw.rounded_rectangle((95, 150, 1505, 815), 20, fill="#ffffff", outline="#cbd5e1", width=3)
    draw.rectangle((95, 150, 1505, 225), fill="#111827")
    draw.text((135, 172), "Workflow 运行监控", font=cjk_font(28, bold=True), fill="#ffffff")
    nodes = [
        ((150, 330, 370, 430), "抓取正文\n与附件", "#dcfce7"),
        ((470, 330, 690, 430), "生成采购\n分析报告", "#dcfce7"),
        ((790, 330, 1010, 430), "报告质检", "#fef9c3"),
        ((1110, 330, 1330, 430), "导出 Word", "#f8fafc"),
        ((470, 570, 690, 670), "质检修复", "#f8fafc"),
        ((790, 570, 1010, 670), "二次质检", "#f8fafc"),
    ]
    for box, label, fill in nodes:
        rounded_box(draw, box, label, fill, "#64748b", font=cjk_font(25))
    arrow(draw, (370, 380), (470, 380))
    arrow(draw, (690, 380), (790, 380))
    arrow(draw, (1010, 380), (1110, 380))
    arrow(draw, (900, 430), (590, 570))
    arrow(draw, (690, 620), (790, 620))
    arrow(draw, (1010, 620), (1220, 430))
    draw.rounded_rectangle((150, 720, 1330, 780), 10, fill="#f8fafc", outline="#e2e8f0")
    draw.text((175, 735), "状态说明：绿色为已完成，黄色为执行中，灰色为等待执行。", font=cjk_font(23), fill="#475569")


def paint_evidence_ui(_: Image.Image, draw: ImageDraw.ImageDraw) -> None:
    draw_title(draw, "公告解析结果界面示意图")
    draw.rounded_rectangle((110, 155, 1490, 810), 18, fill="#ffffff", outline="#cbd5e1", width=3)
    draw.rectangle((110, 155, 1490, 225), fill="#0f172a")
    draw.text((150, 176), "解析结果 / Evidence Package", font=cjk_font(27, bold=True), fill="#ffffff")
    draw.text((160, 280), "公告标题", font=cjk_font(25, bold=True), fill="#0f172a")
    draw.rounded_rectangle((160, 320, 1390, 375), 8, fill="#f8fafc", outline="#e2e8f0")
    draw.text((185, 335), "关于开展医用耗材阳光挂网价格联动工作的通知", font=cjk_font(23), fill="#475569")
    draw.text((160, 425), "网页正文摘要", font=cjk_font(25, bold=True), fill="#0f172a")
    for i in range(4):
        draw.rounded_rectangle((160, 470 + i * 42, 780, 495 + i * 42), 5, fill="#e2e8f0")
    draw.text((870, 425), "附件解析结果", font=cjk_font(25, bold=True), fill="#0f172a")
    rows = [("附件 1.docx", "已解析"), ("申报表.xlsx", "表格摘要"), ("采购文件.pdf", "已解析")]
    for i, (name, status) in enumerate(rows):
        y = 470 + i * 68
        draw.rounded_rectangle((870, y, 1390, y + 46), 8, fill="#f8fafc", outline="#e2e8f0")
        draw.text((895, y + 9), name, font=cjk_font(22), fill="#334155")
        draw.text((1250, y + 9), status, font=cjk_font(22, bold=True), fill="#166534")
    draw.rounded_rectangle((160, 705, 1390, 765), 10, fill="#fff7ed", outline="#fed7aa")
    draw.text((185, 722), "规则完整性提示：发现 企业报价要求、价格联动、非中选产品管理 等关键规则。", font=cjk_font(23), fill="#9a3412")


def paint_qa_block_ui(_: Image.Image, draw: ImageDraw.ImageDraw) -> None:
    draw_title(draw, "质检阻断结果界面示意图")
    draw.rounded_rectangle((115, 155, 1485, 810), 20, fill="#ffffff", outline="#cbd5e1", width=3)
    draw.rectangle((115, 155, 1485, 225), fill="#111827")
    draw.text((150, 176), "报告质检结果", font=cjk_font(28, bold=True), fill="#ffffff")
    draw.rounded_rectangle((165, 290, 540, 380), 14, fill="#fee2e2", outline="#ef4444", width=2)
    draw.text((205, 316), "导出已阻断", font=cjk_font(32, bold=True), fill="#991b1b")
    draw.text((165, 440), "问题列表", font=cjk_font(28, bold=True), fill="#0f172a")
    issues = [
        "报告存在未在本次证据中出现的价格信息",
        "出现“历史分析”独立标题，不符合历史稿使用规则",
        "质检 JSON 指示仍需修复，不能直接导出 Word",
    ]
    for i, issue in enumerate(issues):
        y = 500 + i * 68
        draw.rounded_rectangle((165, y, 1320, y + 48), 8, fill="#fff1f2", outline="#fecdd3")
        draw.text((195, y + 10), f"{i + 1}. {issue}", font=cjk_font(23), fill="#7f1d1d")
    draw.rounded_rectangle((970, 300, 1320, 380), 12, fill="#f8fafc", outline="#cbd5e1")
    draw.text((1010, 323), "下载按钮隐藏", font=cjk_font(27, bold=True), fill="#475569")


def paint_health_ui(_: Image.Image, draw: ImageDraw.ImageDraw) -> None:
    draw_title(draw, "服务健康检查界面示意图")
    draw.rounded_rectangle((130, 185, 1470, 790), 18, fill="#0f172a", outline="#334155", width=3)
    draw.rectangle((130, 185, 1470, 250), fill="#111827")
    draw.text((170, 205), "PowerShell / 本地服务检查", font=cjk_font(27, bold=True), fill="#ffffff")
    lines = [
        "PS> Invoke-WebRequest http://127.0.0.1:8099/health",
        "",
        "StatusCode        : 200",
        "StatusDescription : OK",
        "Content           : {\"status\":\"ok\"}",
        "",
        "说明：返回 200 且 status 为 ok，表示本地辅助服务可用。",
    ]
    y = 315
    for line in lines:
        draw.text((180, y), line, font=cjk_font(27), fill="#d1fae5" if "200" in line or "ok" in line else "#e5e7eb")
        y += 58


def paint_qa_flow(_: Image.Image, draw: ImageDraw.ImageDraw) -> None:
    draw_title(draw, "质检与导出控制流程图")
    boxes = [
        ((130, 245, 430, 365), "接收 ReportIR\n和 Markdown"),
        ((540, 245, 840, 365), "质检模型输出\nQA JSON"),
        ((950, 245, 1250, 365), "本地解析与\n历史泄漏检查"),
        ((540, 545, 840, 665), "一次自动修复"),
        ((950, 545, 1250, 665), "二次质检"),
        ((1320, 395, 1520, 525), "通过后\n导出 Word"),
    ]
    for box, text in boxes:
        rounded_box(draw, box, text, "#f8fafc", "#475569", font=cjk_font(27))
    arrow(draw, (430, 305), (540, 305))
    arrow(draw, (840, 305), (950, 305))
    arrow(draw, (1100, 365), (690, 545))
    arrow(draw, (840, 605), (950, 605))
    arrow(draw, (1250, 605), (1420, 525))
    arrow(draw, (1250, 305), (1420, 395))
    rounded_box(draw, (130, 610, 430, 735), "严重问题\n阻断导出", "#fff1f2", "#e11d48", font=cjk_font(27))
    arrow(draw, (950, 330), (430, 665), "#e11d48")
    draw.text((130, 805), "阻断条件示例：无证据事实、历史内容误写、本地无法解析 QA JSON、二次质检仍存在严重问题。", font=cjk_font(25), fill="#475569")


def drawio_file(filename: str, nodes: list[tuple[str, str, int, int, int, int]], edges: list[tuple[str, str]]) -> Path:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    path = ASSET_DIR / filename
    cells = [
        '<mxCell id="0"/>',
        '<mxCell id="1" parent="0"/>',
    ]
    for node_id, label, x, y, w, h in nodes:
        cells.append(
            f'<mxCell id="{escape(node_id)}" value="{escape(label)}" '
            'style="rounded=1;whiteSpace=wrap;html=1;fillColor=#dae8fc;strokeColor=#6c8ebf;fontFamily=Microsoft YaHei;fontSize=14;" '
            'vertex="1" parent="1">'
            f'<mxGeometry x="{x}" y="{y}" width="{w}" height="{h}" as="geometry"/>'
            '</mxCell>'
        )
    for index, (source, target) in enumerate(edges, start=1):
        cells.append(
            f'<mxCell id="edge-{index}" value="" '
            'style="endArrow=block;html=1;rounded=0;strokeWidth=2;" '
            f'edge="1" parent="1" source="{escape(source)}" target="{escape(target)}">'
            '<mxGeometry relative="1" as="geometry"/>'
            '</mxCell>'
        )
    xml = (
        '<mxfile host="app.diagrams.net" modified="2026-06-03T00:00:00.000Z" agent="Codex" version="24.7.17">'
        '<diagram id="manual-diagram" name="Page-1">'
        '<mxGraphModel dx="1200" dy="800" grid="1" gridSize="10" guides="1" tooltips="1" connect="1" arrows="1" '
        'fold="1" page="1" pageScale="1" pageWidth="1169" pageHeight="827" math="0" shadow="0">'
        f'<root>{"".join(cells)}</root>'
        '</mxGraphModel>'
        '</diagram>'
        '</mxfile>'
    )
    path.write_text(xml, encoding="utf-8")
    return path


def make_drawio_sources() -> dict[str, Path]:
    return {
        "architecture": drawio_file(
            "01-system-architecture.drawio",
            [
                ("user", "业务人员\\n输入公告 URL", 40, 120, 160, 70),
                ("dify", "Dify 应用\\n流程编排与模型调用", 270, 110, 190, 90),
                ("service", "本地辅助服务\\nFastAPI / Docker", 540, 110, 210, 90),
                ("source", "公告网页与附件\\nHTML / Word / Excel / PDF / CSV", 210, 330, 260, 90),
                ("llm", "LLM 报告生成\\nReportIR + Markdown", 520, 330, 220, 90),
                ("word", "输出结果\\nMarkdown 展示 + Word 下载", 810, 330, 240, 90),
            ],
            [("user", "dify"), ("dify", "service"), ("source", "service"), ("service", "llm"), ("llm", "word")],
        ),
        "business_flow": drawio_file(
            "02-business-flow.drawio",
            [
                ("s1", "输入公告 URL", 40, 120, 150, 60),
                ("s2", "抓取网页正文", 240, 120, 150, 60),
                ("s3", "发现并解析附件", 440, 120, 170, 60),
                ("s4", "生成证据包", 660, 120, 150, 60),
                ("s5", "生成 ReportIR 与 Markdown", 40, 300, 210, 60),
                ("s6", "报告质检", 300, 300, 140, 60),
                ("s7", "自动修复与二次质检", 500, 300, 210, 60),
                ("s8", "导出 Word 并返回下载链接", 760, 300, 220, 60),
            ],
            [("s1", "s2"), ("s2", "s3"), ("s3", "s4"), ("s4", "s5"), ("s5", "s6"), ("s6", "s7"), ("s7", "s8")],
        ),
        "modules": drawio_file(
            "03-function-modules.drawio",
            [
                ("center", SYSTEM_NAME, 420, 230, 240, 80),
                ("m1", "公告采集模块", 70, 80, 170, 60),
                ("m2", "附件解析模块", 70, 420, 170, 60),
                ("m3", "证据包整理模块", 420, 480, 190, 60),
                ("m4", "报告生成模块", 820, 80, 170, 60),
                ("m5", "质检修复模块", 820, 420, 170, 60),
                ("m6", "Word 导出模块", 430, 40, 170, 60),
            ],
            [("center", "m1"), ("center", "m2"), ("center", "m3"), ("center", "m4"), ("center", "m5"), ("center", "m6")],
        ),
        "qa_flow": drawio_file(
            "06-qa-flow.drawio",
            [
                ("q1", "接收 ReportIR 和 Markdown", 40, 130, 190, 70),
                ("q2", "质检模型输出 QA JSON", 290, 130, 190, 70),
                ("q3", "本地解析与历史泄漏检查", 540, 130, 220, 70),
                ("q4", "一次自动修复", 290, 330, 180, 70),
                ("q5", "二次质检", 540, 330, 180, 70),
                ("q6", "通过后导出 Word", 820, 220, 180, 70),
                ("q7", "严重问题阻断导出", 40, 330, 190, 70),
            ],
            [("q1", "q2"), ("q2", "q3"), ("q3", "q4"), ("q4", "q5"), ("q5", "q6"), ("q3", "q6"), ("q3", "q7")],
        ),
    }


def make_images() -> dict[str, Path]:
    return {
        "architecture": save_diagram("01-system-architecture.png", paint_architecture),
        "business_flow": save_diagram("02-business-flow.png", paint_business_flow),
        "modules": save_diagram("03-function-modules.png", paint_modules),
        "app_home_ui": save_diagram("04-app-home-ui.png", paint_app_home_ui),
        "input_ui": save_diagram("04-input-ui.png", paint_input_ui),
        "workflow_running_ui": save_diagram("05-workflow-running-ui.png", paint_workflow_running_ui),
        "evidence_ui": save_diagram("06-evidence-ui.png", paint_evidence_ui),
        "result_ui": save_diagram("05-result-ui.png", paint_result_ui),
        "qa_block_ui": save_diagram("07-qa-block-ui.png", paint_qa_block_ui),
        "health_ui": save_diagram("08-health-ui.png", paint_health_ui),
        "qa_flow": save_diagram("06-qa-flow.png", paint_qa_flow),
    }


def set_cell_shading(cell: _Cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell: _Cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if bold else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(10.5)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, "D9EAF7")
        set_cell_text(cell, header, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    doc.add_paragraph()


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)

    for name, size in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)]:
        style = styles[name]
        style.font.name = "黑体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.bold = True


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def add_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = header.add_run(SYSTEM_NAME)
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(9)
    footer = section.footer.paragraphs[0]
    footer.add_run("第 ")
    add_page_number(footer)
    footer.add_run(" 页")


def add_paragraph(doc: Document, text: str, first_line: bool = True) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.35
    if first_line:
        p.paragraph_format.first_line_indent = Pt(21)
    run = p.add_run(text)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(10.5)


def add_bullets(doc: Document, items: list[str]) -> None:
    for index, item in enumerate(items, start=1):
        p = doc.add_paragraph(style=None)
        p.paragraph_format.left_indent = Pt(18)
        p.paragraph_format.first_line_indent = Pt(-18)
        run = p.add_run(f"（{index}）{item}")
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(10.5)


def add_code(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(18)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(79, 79, 79)


def add_picture(doc: Document, path: Path, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(6.3))
    add_caption(doc, caption)


def add_picture_with_source(doc: Document, path: Path, caption: str, source: Path | None = None) -> None:
    add_picture(doc, path, caption)
    if source is not None:
        add_caption(doc, f"可编辑源文件：manual-assets/{source.name}")


def soft_page_break(doc: Document) -> None:
    doc.add_page_break()


def add_cover(doc: Document) -> None:
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(SYSTEM_NAME)
    r.font.name = "黑体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    r.font.size = Pt(24)
    r.font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("操作手册")
    r.font.name = "黑体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    r.font.size = Pt(26)
    r.font.bold = True

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("版本：V1.0")
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(14)

    for _ in range(9):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    today = date.today().strftime("%Y年%m月%d日")
    r = p.add_run(f"编制日期：{today}")
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(12)
    doc.add_page_break()


def add_document_info(doc: Document) -> None:
    doc.add_heading("文档信息", level=1)
    add_table(
        doc,
        ["项目", "内容"],
        [
            ["文档名称", TITLE],
            ["系统名称", SYSTEM_NAME],
            ["文档版本", "V1.0"],
            ["适用对象", "系统管理员、Dify 应用维护人员、业务操作人员、软著材料审查人员"],
            ["主要用途", "说明系统安装、启动、公告分析、报告生成、质检导出、异常处理等操作流程。"],
        ],
    )


def add_toc(doc: Document) -> None:
    doc.add_heading("目录", level=1)
    entries = [
        "1 引言",
        "2 系统概述",
        "3 运行环境与启动",
        "4 系统输入与输出",
        "5 工作流程说明",
        "6 功能操作说明",
        "7 后台接口说明",
        "8 输出文件与数据结构说明",
        "9 异常处理与常见问题",
        "10 附录",
        "11 附图与可编辑源文件",
    ]
    for entry in entries:
        add_paragraph(doc, entry, first_line=False)
    doc.add_page_break()


def build_manual() -> None:
    images = make_images()
    drawio_sources = make_drawio_sources()
    doc = Document()
    configure_doc(doc)
    add_header_footer(doc)
    add_cover(doc)
    add_document_info(doc)
    add_toc(doc)

    doc.add_heading("1 引言", level=1)
    doc.add_heading("1.1 编写目的", level=2)
    add_paragraph(
        doc,
        "本操作手册用于说明医械公告智析报告系统 V1.0 的功能、运行环境、输入输出、主要业务流程和具体操作方法。"
        "本系统面向医药器械采购、集采、挂网、价格联动等公告分析场景，帮助用户从公告链接中自动提取正文和附件信息，"
        "生成规范化项目分析报告，并通过质检机制控制报告导出质量。",
    )
    doc.add_heading("1.2 适用范围", level=2)
    add_bullets(
        doc,
        [
            "适用于需要分析政府或公共平台医药器械公告的业务人员。",
            "适用于负责部署本地辅助服务和维护 Dify 工作流的技术人员。",
            "适用于软件著作权登记材料中对系统使用流程、功能结构和操作步骤的说明。",
        ],
    )
    doc.add_heading("1.3 术语说明", level=2)
    add_table(
        doc,
        ["术语", "说明"],
        [
            ["公告 URL", "医药器械采购、集采、挂网、价格联动等公告页面链接。"],
            ["Dify", "用于搭建 LLM 应用和工作流的低代码平台，本系统通过 Dify 编排报告生成流程。"],
            ["本地辅助服务", "运行在本地或服务器 Docker 容器中的 FastAPI 服务，负责网页抓取、附件解析、质检解析和 Word 导出。"],
            ["证据包", "由网页正文、附件文本、表格摘要、来源信息组成的结构化输入，用于支持 LLM 生成报告。"],
            ["ReportIR", "报告中间表示结构，是 Word 导出的规范化数据格式。"],
            ["QA", "报告质量检查流程，用于发现无证据事实、历史内容误用、格式错误等问题。"],
        ],
    )
    soft_page_break(doc)

    doc.add_heading("2 系统概述", level=1)
    doc.add_heading("2.1 系统简介", level=2)
    add_paragraph(
        doc,
        "医械公告智析报告系统 V1.0 是一套配合 Dify 使用的公告分析与报告生成辅助系统。用户输入公告链接后，"
        "系统自动抓取网页正文，识别并解析公告附件，将正文与附件整理为 LLM 可使用的证据包；Dify 根据提示词生成结构化 ReportIR "
        "和 Markdown 报告；本地服务根据 ReportIR 导出正式 Word 文件。系统还提供质检、自动修复、二次质检和阻断导出机制，"
        "用于降低报告出现无依据内容或历史材料误用的风险。",
    )
    add_picture_with_source(doc, images["architecture"], "图 2-1 系统总体架构图", drawio_sources["architecture"])

    doc.add_heading("2.2 功能模块", level=2)
    add_paragraph(
        doc,
        "系统由公告采集、附件解析、证据包整理、报告生成、质检修复、Word 导出等模块组成。各模块通过 Dify 工作流和本地接口协同工作，"
        "既能支持网页公告正文提取，也能处理 Word、Excel、PDF、CSV 等常见附件格式。",
    )
    add_picture_with_source(doc, images["modules"], "图 2-2 系统功能模块图", drawio_sources["modules"])
    add_table(
        doc,
        ["模块", "主要功能"],
        [
            ["公告采集模块", "校验公告链接，抓取网页内容，提取标题、发布时间候选和地区候选。"],
            ["附件解析模块", "发现附件链接，下载并解析 Word、Excel、PDF、CSV 文件，对大表进行摘要化处理。"],
            ["证据包整理模块", "将网页正文和附件信息合并为可供 LLM 使用的证据包。"],
            ["报告生成模块", "通过 Dify LLM 节点生成 ReportIR 和 Markdown 正式报告。"],
            ["质检修复模块", "解析 QA JSON，检查严重问题，必要时触发一次自动修复和二次质检。"],
            ["Word 导出模块", "根据 ReportIR 生成正式 Word 报告，并返回中文文件名下载链接。"],
        ],
    )
    soft_page_break(doc)

    doc.add_heading("3 运行环境与启动", level=1)
    doc.add_heading("3.1 运行环境", level=2)
    add_table(
        doc,
        ["类别", "要求"],
        [
            ["操作系统", "Windows 环境或支持 Docker 的服务器环境。"],
            ["容器环境", "Docker / Docker Compose。"],
            ["本地服务端口", "默认使用 8099 端口。"],
            ["Dify 访问地址", "http://localhost/apps。"],
            ["Dify 容器访问本地服务", "http://host.docker.internal:8099。"],
            ["可选服务", "如配置 FIRECRAWL_API_KEY，可优先使用 Firecrawl 抽取网页 Markdown。"],
        ],
    )
    doc.add_heading("3.2 启动本地辅助服务", level=2)
    add_paragraph(doc, "进入项目根目录后，执行以下命令启动系统服务：")
    add_code(doc, r"cd C:\Users\admin\Documents\htmldataconclusion\medical-notice-analyzer")
    add_code(doc, "docker compose up -d --build")
    add_paragraph(doc, "服务启动后，可通过健康检查确认本地辅助服务是否可用：")
    add_code(doc, "Invoke-WebRequest http://127.0.0.1:8099/health")
    add_paragraph(doc, "如返回健康状态，说明本地辅助服务已启动并监听 8099 端口。")
    add_picture(doc, images["health_ui"], "图 3-1 服务健康检查界面示意图")

    doc.add_heading("3.3 Dify 应用访问", level=2)
    add_paragraph(doc, "浏览器打开 Dify 应用列表页面：")
    add_code(doc, "http://localhost/apps")
    add_paragraph(doc, "当前项目文档记录的 Dify 工作流页面为：")
    add_code(doc, "http://localhost/app/a09ac7aa-3f49-4fc2-aea4-55471a3c6802/workflow")
    add_paragraph(
        doc,
        "Dify 工作流运行在 Docker 容器中时，不能使用 127.0.0.1 调用宿主机服务，应使用 http://host.docker.internal:8099 作为本地辅助服务地址。",
    )
    doc.add_heading("3.4 运行前检查", level=2)
    add_table(
        doc,
        ["检查项", "检查方法", "预期结果"],
        [
            ["本地服务", "访问 /health 或执行 PowerShell 健康检查命令。", "返回 status 为 ok。"],
            ["Dify 页面", "浏览器打开 http://localhost/apps。", "可进入应用列表。"],
            ["Workflow 配置", "检查 HTTP 请求节点地址。", "容器内地址为 host.docker.internal:8099。"],
            ["报告目录", "查看 reports 目录或 Docker volume。", "质检通过后可看到生成的 .docx 文件。"],
        ],
    )
    soft_page_break(doc)

    doc.add_heading("4 系统输入与输出", level=1)
    doc.add_heading("4.1 输入内容", level=2)
    add_table(
        doc,
        ["输入项", "是否必填", "说明"],
        [
            ["公告 URL", "必填", "需要分析的医药器械公告网页链接。"],
            ["历史分析报告", "可选", "仅在 Chatflow 扩展模式中使用，用于风格参考和历史承接，不作为本次公告事实来源。"],
            ["用户修订要求", "可选", "在多轮 Chatflow 模式下，用于对当前报告进行局部修订。"],
            ["QA JSON", "系统生成", "质检模型输出的结构化检查结果，由本地服务解析。"],
        ],
    )
    doc.add_heading("4.2 输出内容", level=2)
    add_table(
        doc,
        ["输出项", "说明"],
        [
            ["证据包", "包含公告正文、附件解析结果、表格摘要、来源链接和规则完整性提示。"],
            ["Markdown 报告", "用于在 Dify 页面直接展示的正式报告内容。"],
            ["ReportIR", "用于 Word 导出的结构化报告数据。"],
            ["QA 摘要", "展示质检结果、问题列表和修复建议。"],
            ["Word 报告", "质检通过后生成的 .docx 文件。"],
            ["下载链接", "通过 /download/{filename} 获取 Word 文件。"],
        ],
    )
    doc.add_heading("4.3 输入输出对应关系", level=2)
    add_paragraph(
        doc,
        "在正常业务流程中，业务人员只需要输入公告 URL。系统内部会自动完成网页访问、附件下载、证据包构造、模型生成、质检和导出。"
        "历史报告和用户修订要求属于扩展 Chatflow 场景，当前 Workflow 可作为单轮公告分析和导出使用。",
    )
    add_table(
        doc,
        ["业务动作", "系统内部处理", "用户可见结果"],
        [
            ["提交公告链接", "调用 /analyze 抓取正文和附件。", "Dify 显示运行状态。"],
            ["等待分析完成", "生成证据包并交给 LLM 生成报告。", "页面出现 Markdown 报告草稿。"],
            ["执行质检", "解析 QA JSON，检查历史泄漏和严重问题。", "显示通过、需修复或阻断状态。"],
            ["下载报告", "调用 /report/export_checked 生成 Word。", "获得 .docx 下载链接。"],
        ],
    )
    soft_page_break(doc)

    doc.add_heading("5 工作流程说明", level=1)
    add_picture_with_source(doc, images["business_flow"], "图 5-1 公告分析与报告导出流程图", drawio_sources["business_flow"])
    add_paragraph(
        doc,
        "系统从公告 URL 开始工作。首先由本地辅助服务抓取网页正文并发现附件，然后解析 Word、Excel、PDF、CSV 等附件内容，"
        "形成证据包。Dify LLM 节点基于证据包生成 ReportIR 和 Markdown 报告。随后系统执行质检，如发现可修复问题则进行一次自动修复并再次质检；"
        "若仍存在严重问题，则阻断 Word 导出；若通过质检，则生成正式 Word 报告并返回下载链接。",
    )
    add_picture_with_source(doc, images["qa_flow"], "图 5-2 质检与导出控制流程图", drawio_sources["qa_flow"])
    doc.add_heading("5.1 正常导出流程", level=2)
    add_bullets(
        doc,
        [
            "用户提交公告 URL，Dify 调用本地 /analyze 接口。",
            "本地服务返回证据包，Dify 根据证据包生成 ReportIR 和 Markdown。",
            "质检节点输出 pass，系统进入导出阶段。",
            "本地 /report/export_checked 接口生成 Word 并返回下载链接。",
        ],
    )
    doc.add_heading("5.2 自动修复流程", level=2)
    add_bullets(
        doc,
        [
            "质检节点输出 needs_fix 时，Dify 将问题列表和修复指令交给修复节点。",
            "修复节点只允许依据原始证据和当前报告进行局部修订，不得新增无来源事实。",
            "修复完成后进入二次质检，二次质检通过后才允许导出 Word。",
        ],
    )
    doc.add_heading("5.3 阻断导出流程", level=2)
    add_bullets(
        doc,
        [
            "质检结果为 block，或 QA JSON 无法被本地服务解析时，系统阻断 Word 导出。",
            "如果历史报告中的价格、周期、日期等内容被误写成本次公告事实，系统会标记历史泄漏。",
            "阻断时页面展示 qa_summary 和问题列表，用户应根据提示重新生成或修订报告。",
        ],
    )
    soft_page_break(doc)

    doc.add_heading("6 功能操作说明", level=1)
    doc.add_heading("6.1 打开系统应用", level=2)
    add_bullets(
        doc,
        [
            "确认 Docker 服务已启动，本地辅助服务健康检查正常。",
            "在浏览器中打开 Dify 应用页面。",
            "进入医械公告智析报告系统对应的 Workflow 或 Chatflow 应用。",
        ],
    )
    add_picture(doc, images["app_home_ui"], "图 6-1 Dify 应用入口界面示意图")
    doc.add_heading("6.2 输入公告链接", level=2)
    add_paragraph(doc, "在系统运行页面中，将需要分析的公告网页地址填写到公告 URL 输入框，然后点击开始分析按钮。")
    add_picture(doc, images["input_ui"], "图 6-2 公告链接输入界面示意图")
    add_table(
        doc,
        ["操作项", "说明"],
        [
            ["公告 URL 输入框", "填写公告页面完整链接，应包含 http:// 或 https://。"],
            ["开始分析按钮", "触发 Dify 工作流，调用本地 /analyze 接口抓取正文和附件。"],
            ["运行状态", "Dify 页面显示节点执行状态，便于查看是否出现调用失败或超时。"],
        ],
    )
    add_picture(doc, images["workflow_running_ui"], "图 6-3 工作流运行状态界面示意图")
    soft_page_break(doc)

    doc.add_heading("6.3 分析公告正文与附件", level=2)
    add_paragraph(
        doc,
        "用户提交公告链接后，系统自动执行网页抓取和附件解析。普通用户无需手工下载附件。若公告包含大表格，系统会提取关键行列并生成摘要，"
        "以避免模型输入过长。对于部分特殊公告站点，系统采用站点结构适配方式获取正文，不针对单个 URL 写死正文内容。",
    )
    add_bullets(
        doc,
        [
            "网页正文：提取公告主体内容、标题、发布时间候选、地区候选。",
            "附件文档：解析 Word、Excel、PDF、CSV 文件。",
            "表格数据：对大表进行摘要化展示，保留业务分析所需关键信息。",
            "异常提示：附件下载或解析失败时，证据包中记录 warning 信息，便于排查。",
        ],
    )
    add_picture(doc, images["evidence_ui"], "图 6-4 公告解析结果界面示意图")

    doc.add_heading("6.4 生成分析报告", level=2)
    add_paragraph(
        doc,
        "Dify LLM 节点根据证据包和提示词生成结构化 ReportIR 与 Markdown 报告。报告内容应围绕公告原文展开，"
        "重点呈现采购品种范围、产品分类、报价规则、中选规则、协议采购量、采购执行、价格联动、非中选产品管理、企业关注事项等信息。"
    )
    add_paragraph(
        doc,
        "系统不要求所有报告套用固定五段式结构，若公告原文缺少某类内容，报告不应强行编造。ReportIR 是 Word 导出的优先格式，"
        "Markdown 主要用于 Dify 页面展示。",
    )

    doc.add_heading("6.5 执行质检与自动修复", level=2)
    add_paragraph(
        doc,
        "报告生成后，系统执行质检流程。质检结果分为 pass、needs_fix 和 block。pass 表示通过，可继续导出；needs_fix 表示存在可修复问题，"
        "系统会根据修复指令自动修订一次；block 表示存在严重问题，应停止导出并向用户展示问题摘要。",
    )
    add_table(
        doc,
        ["质检状态", "系统处理方式"],
        [
            ["pass", "进入 Word 导出流程。"],
            ["needs_fix", "执行一次自动修复，修复后进行二次质检。"],
            ["block", "阻断 Word 导出，返回 QA 摘要和问题列表。"],
            ["JSON 非法", "本地服务无法解析质检结果时，按阻断处理，不生成 Word。"],
        ],
    )
    add_picture(doc, images["qa_block_ui"], "图 6-5 质检阻断结果界面示意图")

    doc.add_heading("6.6 查看报告并下载 Word", level=2)
    add_paragraph(
        doc,
        "质检通过后，系统生成 Word 文件，并在 Dify 结果区域展示下载链接。用户点击下载链接即可获取正式 .docx 报告。"
        "Word 文件名优先使用 ReportIR 中的 suggested_filename，其次使用标题、文档名或系统默认名称，并自动清理 Windows 非法字符。",
    )
    add_picture(doc, images["result_ui"], "图 6-6 报告结果与下载界面示意图")
    soft_page_break(doc)

    doc.add_heading("6.7 多轮修订与历史报告使用", level=2)
    add_paragraph(
        doc,
        "当前项目文档记录的 Dify 应用仍为 Workflow 模式，已支持质检、修复、二次质检和阻断导出；但 Workflow 模式不适合自然支持同一会话中的多轮修改和历史 Word 上传。"
        "如需完整实现“公告链接 + 可选历史 Word + 多轮反馈修改”，建议新建 Dify Chatflow 应用。",
    )
    add_bullets(
        doc,
        [
            "历史 Word 只能用于历史对照、风格参考、项目延续性观察和企业关注点补充。",
            "历史 Word 不得作为本次公告事实来源。",
            "不得新增“历史对照”“历史分析”等独立标题。",
            "用户后续修订要求应基于原始证据、当前 ReportIR 和当前 Markdown 局部修改，不得脱离原文新增事实。",
        ],
    )
    doc.add_heading("6.8 典型使用场景", level=2)
    add_table(
        doc,
        ["场景", "操作步骤", "预期结果"],
        [
            ["采购文件分析", "输入集采或接续采购公告链接，等待解析附件和采购文件。", "生成包含采购范围、报价规则、中选规则、执行要求的分析报告。"],
            ["挂网流程调整分析", "输入挂网操作流程调整通知链接。", "生成流程变化、企业影响和申报注意事项说明。"],
            ["价格联动通知分析", "输入价格治理或价格联动公告链接。", "生成价格规则、企业申报要求和风险提示。"],
            ["质检阻断复核", "查看 QA 摘要，确认无证据事实或历史泄漏位置。", "根据问题重新生成或修订后再次导出。"],
        ],
    )
    add_paragraph(
        doc,
        "上述场景中，系统均以本次公告网页和附件作为事实来源。若使用历史报告作为参考，历史内容只能形成简短承接或风格参考，"
        "不得替代本次公告证据，也不得把历史项目中的价格、周期、范围、日期等内容写成本次公告事实。",
    )
    soft_page_break(doc)

    doc.add_heading("7 后台接口说明", level=1)
    add_table(
        doc,
        ["接口", "方法", "用途"],
        [
            ["/health", "GET", "健康检查，确认服务是否启动。"],
            ["/analyze", "POST", "接收公告 URL，抓取正文并解析附件，返回证据包。"],
            ["/report/export", "POST", "根据 ReportIR 或 final_report 导出 Word。"],
            ["/report/qa", "POST", "解析质检模型输出，返回结构化 QA 和阻断状态。"],
            ["/report/export_checked", "POST", "在质检通过后导出 Word；质检阻断时不生成文件。"],
            ["/download/{filename}", "GET", "下载已生成的 Word 文件。"],
        ],
    )
    doc.add_heading("7.1 /analyze 请求示例", level=2)
    add_code(doc, 'POST http://127.0.0.1:8099/analyze')
    add_code(doc, '{"url":"https://example.gov.cn/notice/detail.html"}')
    add_table(
        doc,
        ["字段", "类型", "说明"],
        [
            ["url", "string", "公告页面链接，必须是 http 或 https 地址。"],
            ["max_attachments", "integer", "可选，限制附件解析数量，避免超长公告一次解析过多文件。"],
            ["use_cache", "boolean", "可选，是否优先读取 site-cache 中的缓存页面和附件。"],
        ],
    )
    doc.add_heading("7.2 /analyze 返回内容", level=2)
    add_table(
        doc,
        ["字段", "说明"],
        [
            ["title", "公告标题或页面标题。"],
            ["url", "规范化后的最终公告链接。"],
            ["content", "面向 LLM 的合并证据文本。"],
            ["page_text", "网页正文抽取结果。"],
            ["attachments", "附件解析结果列表。"],
            ["warnings", "下载、解析、站点适配等过程中的提示信息。"],
            ["publish_date_candidates", "系统识别到的发布时间候选。"],
            ["region_candidates", "系统识别到的地区候选。"],
        ],
    )
    doc.add_heading("7.3 /report/export_checked 行为", level=2)
    add_paragraph(
        doc,
        "当 QA 结果通过时，该接口生成 Word 并返回 filename 和 download_url；当 QA 结果阻断时，该接口返回 success:false、blocked:true 和 qa_summary，不生成 Word。",
    )
    add_table(
        doc,
        ["返回状态", "含义", "用户处理方式"],
        [
            ["success=true, blocked=false", "报告已通过质检并导出 Word。", "点击下载链接获取文件。"],
            ["success=false, blocked=true", "质检阻断，不生成 Word。", "查看 qa_summary，修订后重新运行。"],
            ["422 或解析错误", "请求结构或 QA JSON 不符合要求。", "检查 Dify 节点输出格式，保证为合法 JSON。"],
        ],
    )
    soft_page_break(doc)

    doc.add_heading("8 输出文件与数据结构说明", level=1)
    doc.add_heading("8.1 Word 报告文件", level=2)
    add_paragraph(
        doc,
        "系统导出的 Word 报告采用正式报告样式，包含标题、引言段落、正文分节、表格、企业提示和固定声明。"
        "导出逻辑会清理模型思考内容、Dify 变量、Markdown 代码块和 JSON 外壳，只保留正式报告内容。"
    )
    add_bullets(
        doc,
        [
            "中文文件名通过 UTF-8 Content-Disposition 支持浏览器下载。",
            "notice_type 仅作为内部分类和命名参考，不在 Word 正文中机械展示。",
            "表格按主题渲染，表头加底纹，表格边框清晰。",
            "固定声明始终作为报告最后部分输出。",
        ],
    )
    doc.add_heading("8.2 ReportIR 结构", level=2)
    add_code(
        doc,
        '{ "title": "", "suggested_filename": "", "notice_type": "", "publish_date": "", "source_agency": "", "document_name": "", "lead_paragraphs": [], "sections": [], "enterprise_tips": [], "disclaimer": "" }',
    )
    add_paragraph(
        doc,
        "ReportIR 中 sections 可包含多个章节，每个章节可包含 paragraphs、tables 和 highlights。tables 包含 title、headers、rows、notes。"
        "系统优先使用 ReportIR 导出 Word，只有在 ReportIR 无法解析时才尝试使用 final_report 区域作为回退。",
    )
    doc.add_heading("8.3 QA JSON 结构", level=2)
    add_code(
        doc,
        '{ "status": "pass | needs_fix | block", "issues": [], "unsupported_claims": [], "history_leakage": [], "missing_rules": [], "language_issues": [], "fix_instructions": [], "summary": "" }',
    )
    add_paragraph(
        doc,
        "QA JSON 由 Dify 质检模型节点输出，本地 /report/qa 接口负责解析。若 JSON 非法，系统不应继续导出 Word。"
        "issues、unsupported_claims、history_leakage 等字段用于呈现问题明细，fix_instructions 用于指导一次自动修复。",
    )
    doc.add_heading("8.4 文件目录", level=2)
    add_table(
        doc,
        ["目录或文件", "说明"],
        [
            ["reports", "保存系统生成的 Word 报告。"],
            ["site-cache", "保存公告页面和附件缓存，用于重复分析或调试。"],
            ["prompts", "保存 Dify 使用的系统提示词、用户提示词、历史提示词、修订提示词和质检提示词。"],
            ["dify-workflow-medical-notice-report.yml", "Workflow 配置文件。"],
            ["dify-chatflow-medical-notice-report.yml", "Chatflow 蓝本配置文件。"],
            ["manual-assets", "操作手册配套图片和 drawio 可编辑源文件。"],
        ],
    )
    soft_page_break(doc)

    doc.add_heading("9 异常处理与常见问题", level=1)
    add_table(
        doc,
        ["问题", "可能原因", "处理方法"],
        [
            ["健康检查失败", "Docker 服务未启动或端口被占用。", "执行 docker compose up -d --build，检查 8099 端口占用情况。"],
            ["Dify 无法调用本地服务", "容器内使用了 127.0.0.1。", "将接口地址改为 http://host.docker.internal:8099。"],
            ["公告正文为空", "页面为动态加载或站点结构特殊。", "检查 URL 是否有效，必要时增加站点结构适配。"],
            ["附件解析失败", "附件格式异常、链接失效或下载超时。", "查看证据包 warnings，确认附件链接和文件格式。"],
            ["Word 未生成", "质检阻断、QA JSON 非法或 ReportIR 校验失败。", "查看 qa_summary 和问题列表，修复报告或调整质检输出格式。"],
            ["中文文件名异常", "浏览器或代理未正确处理 UTF-8 文件名。", "使用支持 UTF-8 下载文件名的浏览器，或从 reports 目录直接获取文件。"],
        ],
    )
    doc.add_heading("9.1 推荐排查顺序", level=2)
    add_bullets(
        doc,
        [
            "先检查本地 /health 是否返回 ok，确认服务处于可用状态。",
            "再检查 Dify HTTP 节点地址是否使用 host.docker.internal:8099。",
            "若网页解析失败，查看 /analyze 返回的 warnings 和 page_text。",
            "若附件解析失败，确认附件链接是否可访问、文件格式是否为系统支持类型。",
            "若 Word 未生成，优先查看 QA 摘要、ReportIR 结构和 final_report 区域。",
        ],
    )
    doc.add_heading("9.2 管理员维护建议", level=2)
    add_table(
        doc,
        ["维护项", "建议频率", "说明"],
        [
            ["服务健康检查", "每次使用前", "确认本地服务和 Dify 页面均可访问。"],
            ["报告目录清理", "按需", "清理 reports 中过期或重复生成的测试文件。"],
            ["站点适配维护", "遇到新站点解析失败时", "根据网页结构补充站点级适配逻辑，不针对单个 URL 写死正文。"],
            ["提示词同步", "修改 Dify 节点后", "将 prompts 目录中的提示词同步到线上 Dify 节点。"],
            ["回归测试", "代码或提示词调整后", "执行 tests/test_report_export.py 覆盖核心导出和质检行为。"],
        ],
    )
    soft_page_break(doc)

    doc.add_heading("10 附录", level=1)
    doc.add_heading("10.1 规则完整性关注项", level=2)
    add_paragraph(
        doc,
        "当公告原文包含下列规则时，系统提示词会要求模型尽量完整保留，不应压缩为几句概括：采购品种范围、产品分类、最高有效申报价、参考价、"
        "企业报价要求、有效报价、拟中选产品确定、中选产品确定、协议采购量、首年协议采购量、采购执行、价格联动、非中选产品管理、新获批产品管理、"
        "名词解释、信用评价、失信约束、取消中选资格、暂不予挂网等。",
    )
    doc.add_heading("10.2 测试与维护", level=2)
    add_paragraph(doc, "项目当前测试文件为 tests/test_report_export.py，主要覆盖 Word 导出、质检解析、历史泄漏检测、特殊站点解析等行为。")
    add_code(doc, "docker exec medical-notice-analyzer python -m unittest discover -s tests -v")
    add_paragraph(doc, "项目文档记录最近一次测试结果为 Ran 26 tests，OK。维护人员修改后端逻辑或提示词后，应重新执行回归测试。")
    doc.add_heading("10.3 安全与合规说明", level=2)
    add_bullets(
        doc,
        [
            "系统生成报告基于互联网公开资料和用户输入的公告链接，不替代人工审核和正式法律意见。",
            "报告内容应以本次公告正文和附件为事实依据，不应引入无来源推断。",
            "历史报告仅用于风格和经验参考，不得作为本次事实来源。",
            "导出的 Word 报告末尾固定加入声明，提示信息准确性和完整性需由读者自行判断。",
        ],
    )
    soft_page_break(doc)

    doc.add_heading("11 附图与可编辑源文件", level=1)
    add_paragraph(
        doc,
        "本手册中的网页界面图为自绘示意图，用于表达系统操作页面和运行状态。流程图、架构图、模块图和质检控制图同时生成 PNG 图片和 drawio 可编辑源文件；"
        "如需修改图形内容，可使用 diagrams.net 或 draw.io 打开对应 .drawio 文件后编辑。",
    )
    add_table(
        doc,
        ["图名", "Word 内嵌图片", "可编辑源文件"],
        [
            ["系统总体架构图", "manual-assets/01-system-architecture.png", "manual-assets/01-system-architecture.drawio"],
            ["公告分析与报告导出流程图", "manual-assets/02-business-flow.png", "manual-assets/02-business-flow.drawio"],
            ["系统功能模块图", "manual-assets/03-function-modules.png", "manual-assets/03-function-modules.drawio"],
            ["质检与导出控制流程图", "manual-assets/06-qa-flow.png", "manual-assets/06-qa-flow.drawio"],
            ["Dify 应用入口界面", "manual-assets/04-app-home-ui.png", "无，界面示意图为 PNG 自绘图"],
            ["公告链接输入界面", "manual-assets/04-input-ui.png", "无，界面示意图为 PNG 自绘图"],
            ["工作流运行状态界面", "manual-assets/05-workflow-running-ui.png", "无，界面示意图为 PNG 自绘图"],
            ["公告解析结果界面", "manual-assets/06-evidence-ui.png", "无，界面示意图为 PNG 自绘图"],
            ["质检阻断结果界面", "manual-assets/07-qa-block-ui.png", "无，界面示意图为 PNG 自绘图"],
            ["服务健康检查界面", "manual-assets/08-health-ui.png", "无，界面示意图为 PNG 自绘图"],
        ],
    )

    doc.core_properties.title = TITLE
    doc.core_properties.subject = "软件著作权操作手册"
    doc.core_properties.keywords = "医械公告, Dify, 报告生成, Word导出, 操作手册"
    doc.save(OUTPUT)


def verify_docx(path: Path) -> None:
    try:
        with ZipFile(path) as archive:
            names = set(archive.namelist())
            if "word/document.xml" not in names:
                raise RuntimeError("word/document.xml missing")
            media = [name for name in names if name.startswith("word/media/")]
            if len(media) < 11:
                raise RuntimeError(f"expected at least 11 embedded images, got {len(media)}")
    except BadZipFile as exc:
        raise RuntimeError(f"invalid docx zip: {exc}") from exc
    drawio_files = list(ASSET_DIR.glob("*.drawio"))
    if len(drawio_files) < 4:
        raise RuntimeError(f"expected at least 4 drawio source files, got {len(drawio_files)}")


if __name__ == "__main__":
    build_manual()
    verify_docx(OUTPUT)
    print(OUTPUT)
