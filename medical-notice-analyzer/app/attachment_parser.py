from __future__ import annotations

import csv
import io
import os
import re
import shutil
import subprocess
import tempfile
import zipfile
from html.parser import HTMLParser
from pathlib import Path
from typing import Any

import xlrd
from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader


KEY_COLUMN_PATTERNS = [
    "企业名称",
    "申报企业",
    "生产企业",
    "配送企业",
    "产品名称",
    "通用名",
    "耗材名称",
    "项目名称",
    "项目编码",
    "服务项目",
    "注册证号",
    "备案号",
    "医保编码",
    "医保耗材代码",
    "医保代码",
    "医保支付",
    "支付类别",
    "支付标准",
    "规格型号",
    "申报价",
    "报价",
    "价格",
    "中选价",
    "挂网价",
    "收费",
    "采购量",
    "报量",
    "需求量",
    "约定采购量",
    "地区",
    "分组",
    "分类",
    "中选状态",
    # Historical mojibake test fixtures kept for compatibility.
    "浼佷笟鍚嶇О",
    "鐢虫姤浼佷笟",
    "鐢熶骇浼佷笟",
    "浜у搧鍚嶇О",
    "娉ㄥ唽璇佸彿",
    "鍖讳繚缂栫爜",
    "瑙勬牸鍨嬪彿",
    "浠锋牸",
    "涓€変环",
    "閲囪喘閲",
]

IMPORTANT_SECTION_KEYWORDS = [
    "申报",
    "价格",
    "采购",
    "企业",
    "产品",
    "执行",
    "医疗机构",
    "医疗服务价格项目",
    "项目内涵",
    "计价单位",
    "医保支付",
    "支付类别",
    "呼吸系统",
    "神经系统",
    "泌尿系统",
    "妇科",
    "心血管",
    "收费",
]

SECTION_TOPIC_KEYWORDS = [
    ("政策背景", ["背景", "依据", "规范整合", "制定"]),
    ("项目范围", ["范围", "十二类", "呼吸系统", "神经系统", "泌尿系统", "项目目录", "项目名称"]),
    ("价格和支付规则", ["价格", "收费", "医保支付", "支付类别", "支付标准", "计价单位"]),
    ("执行时间和实施安排", ["执行", "实施", "时间", "日期", "起", "过渡"]),
    ("医疗机构要求", ["医疗机构", "公立医疗机构", "定点医疗机构", "执行要求"]),
    ("企业和产品要求", ["企业", "产品", "注册证", "医保编码", "申报"]),
]


def _contains_any(headers: list[str], keywords: list[str]) -> bool:
    text = " ".join(str(header or "") for header in headers)
    return any(keyword in text for keyword in keywords)


def _business_table_flags(headers: list[str]) -> dict[str, bool]:
    return {
        "contains_enterprise": _contains_any(headers, ["企业", "申报企业", "生产企业", "配送企业", "浼佷笟"]),
        "contains_product": _contains_any(headers, ["产品", "耗材", "品名", "通用名", "医疗服务价格项目", "项目名称", "项目编码", "浜у搧"]),
        "contains_registration_cert": _contains_any(headers, ["注册证", "注册证号", "备案号", "娉ㄥ唽璇"]),
        "contains_medical_insurance_code": _contains_any(headers, ["医保编码", "医保耗材代码", "医保代码", "医保支付", "支付类别", "鍖讳繚"]),
        "contains_specification": _contains_any(headers, ["规格", "型号", "规格型号", "瑙勬牸"]),
        "contains_price": _contains_any(headers, ["价格", "报价", "申报价", "中选价", "挂网价", "收费", "支付标准", "浠锋牸"]),
        "contains_selected_status": _contains_any(headers, ["中选", "拟中选", "入围", "挂网状态", "状态", "涓€"]),
        "contains_purchase_volume": _contains_any(headers, ["采购量", "报量", "需求量", "约定采购量", "閲囪喘閲"]),
    }


FIELD_COLUMN_KEYWORDS = {
    "enterprise": ["enterprise", "company", "manufacturer", "supplier", "distributor", "企业", "申报企业", "生产企业", "配送企业", "浼佷笟"],
    "product": ["product", "item", "goods", "耗材", "产品", "通用名", "项目名称", "浜у搧", "椤圭洰"],
    "registration_cert": ["registration", "certificate", "cert", "注册证", "备案号", "娉ㄥ唽璇", "澶囨"],
    "medical_insurance_code": ["medical_insurance", "insurance_code", "医保编码", "医保耗材代码", "医保代码", "鍖讳繚"],
    "price": ["price", "selected_price", "bid_price", "申报价", "报价", "中选价", "挂网价", "价格", "浠锋牸", "鎶ヤ环", "涓€変环"],
    "purchase_volume": ["purchase_volume", "volume", "quantity", "采购量", "报量", "需求量", "约定采购量", "閲囪喘閲", "鎶ラ噺"],
    "selected_status": ["selected_status", "status", "selected", "winner", "中选", "拟中选", "入围", "状态", "涓€", "鐘舵€"],
    "region": ["region", "province", "city", "area", "地区", "省", "市", "地市", "鍦板尯"],
    "group": ["group", "category", "分组", "分类", "组别", "鍒嗙粍", "鍒嗙被"],
}


def _column_type_map(headers: list[str]) -> dict[str, list[str]]:
    result: dict[str, list[str]] = {field: [] for field in FIELD_COLUMN_KEYWORDS}
    for header in headers:
        normalized = str(header or "")
        normalized_lower = normalized.lower()
        for field, keywords in FIELD_COLUMN_KEYWORDS.items():
            if any(keyword.lower() in normalized_lower for keyword in keywords):
                result[field].append(normalized)
    return result


def _numeric_value(value: Any) -> float | None:
    text = _clean_text(value)
    if not text:
        return None
    match = re.search(r"-?\d+(?:\.\d+)?", text.replace(",", ""))
    if not match:
        return None
    try:
        return float(match.group(0))
    except ValueError:
        return None


def _field_stats(headers: list[str], data_rows: list[list[str]], column_map: dict[str, list[str]]) -> tuple[dict[str, Any], list[str]]:
    warnings: list[str] = []
    header_index = {header: index for index, header in enumerate(headers)}
    stats: dict[str, Any] = {}
    for field, columns in column_map.items():
        if not columns:
            continue
        non_empty = 0
        unique_values: set[str] = set()
        sample_values: list[str] = []
        numeric_values: list[float] = []
        for row in data_rows:
            for column in columns:
                index = header_index.get(column)
                if index is None or index >= len(row):
                    continue
                value = _clean_text(row[index])
                if not value:
                    continue
                non_empty += 1
                unique_values.add(value)
                if len(sample_values) < 5 and value not in sample_values:
                    sample_values.append(value)
                if field in {"price", "purchase_volume"}:
                    number = _numeric_value(value)
                    if number is not None:
                        numeric_values.append(number)
        item: dict[str, Any] = {
            "columns": columns,
            "non_empty_count": non_empty,
            "unique_count": len(unique_values),
            "sample_values": sample_values,
        }
        if field in {"price", "purchase_volume"}:
            if numeric_values:
                item["min"] = min(numeric_values)
                item["max"] = max(numeric_values)
                item["sample_numeric_values"] = numeric_values[:5]
            elif non_empty:
                warnings.append(f"{field} columns could not be parsed as numeric values")
        stats[field] = item
    return stats, warnings


def _table_evidence_score(flags: dict[str, bool], field_stats: dict[str, Any], *, rows: int, table_heavy: bool) -> int:
    score = min(35, sum(8 for value in flags.values() if value))
    score += min(35, len(field_stats) * 6)
    if table_heavy:
        score += 20
    if rows >= 5000:
        score += 10
    return max(0, score)


def _float_env(name: str, default: float) -> float:
    try:
        return float((os.getenv(name) or "").strip() or default)
    except ValueError:
        return default


def _int_env(name: str, default: int) -> int:
    try:
        return int(float((os.getenv(name) or "").strip() or default))
    except ValueError:
        return default


def _clean_text(text: Any) -> str:
    return re.sub(r"\s+", " ", str(text or "")).strip()


def _section_summaries_from_text(text: str, *, per_topic_limit: int = 260) -> list[str]:
    clean = _clean_text(text)
    if not clean:
        return []
    sentences = [item.strip() for item in re.split(r"[。；;\n]", clean) if len(item.strip()) >= 12]
    summaries: list[str] = []
    for title, keywords in SECTION_TOPIC_KEYWORDS:
        matches = [sentence for sentence in sentences if any(keyword in sentence for keyword in keywords)]
        if not matches:
            continue
        merged = "；".join(matches[:3])
        summaries.append(f"{title}：{merged[:per_topic_limit]}")
    return summaries


def _summary_from_text(text: str, limit: int = 1200) -> str:
    clean = _clean_text(text)
    if len(clean) <= limit:
        return clean
    structured = _section_summaries_from_text(clean)
    if structured:
        summary = "；".join(structured)
        if len(summary) <= limit:
            return summary
        return summary[:limit].rstrip() + "..."
    return clean[:limit].rstrip() + "..."


def _key_columns(headers: list[str]) -> list[str]:
    result = []
    for header in headers:
        if any(pattern in str(header) for pattern in KEY_COLUMN_PATTERNS):
            result.append(str(header))
    return result


def _row_values(row: Any) -> list[str]:
    return [_clean_text(value) for value in list(row or [])]


def _header_score(row: list[str]) -> tuple[int, int, int]:
    non_empty = sum(1 for value in row if value)
    key_hits = len(_key_columns(row))
    business_hits = sum(1 for value in _business_table_flags(row).values() if value)
    return key_hits * 10 + business_hits * 4 + min(non_empty, 12), key_hits, non_empty


def _select_header_row(rows: list[list[str]]) -> tuple[list[str], int]:
    if not rows:
        return [], 0
    best_index = 0
    best_score = (-1, -1, -1)
    for index, row in enumerate(rows[:25]):
        normalized = _row_values(row)
        score = _header_score(normalized)
        if score > best_score:
            best_index = index
            best_score = score
    return _row_values(rows[best_index]), best_index


def _table_summary(
    *,
    sheet_name: str,
    rows: int,
    columns: int,
    headers: list[str],
    sample_count: int,
    default_business_value: str,
    data_rows: list[list[str]] | None = None,
) -> dict[str, Any]:
    key_columns = _key_columns(headers)
    column_map = _column_type_map(headers)
    field_stats, stats_warnings = _field_stats(headers, data_rows or [], column_map)
    flags = _business_table_flags(headers)
    table_heavy = rows >= max(1, _int_env("TABLE_HEAVY_ROW_THRESHOLD", 5000))
    evidence_value_score = _table_evidence_score(flags, field_stats, rows=rows, table_heavy=table_heavy)
    return {
        "sheet_name": sheet_name,
        "rows": rows,
        "columns_count": columns,
        "headers": headers,
        "key_columns": key_columns,
        "enterprise_columns": column_map["enterprise"],
        "product_columns": column_map["product"],
        "registration_cert_columns": column_map["registration_cert"],
        "medical_insurance_code_columns": column_map["medical_insurance_code"],
        "price_columns": column_map["price"],
        "purchase_volume_columns": column_map["purchase_volume"],
        "selected_status_columns": column_map["selected_status"],
        "region_columns": column_map["region"],
        "group_columns": column_map["group"],
        "sample_rows_count": sample_count,
        "summary": f"该表主要包含 {rows} 行、{columns} 列，关键字段包括：{', '.join(key_columns or headers[:6])}。",
        "business_value": default_business_value,
        "field_stats": field_stats,
        "field_stats_warnings": stats_warnings,
        "table_heavy": table_heavy,
        "evidence_value_score": evidence_value_score,
        "data_completeness_hint": f"field_stats scanned {len(data_rows or [])} data rows; rows keeps the worksheet total row count.",
        "recommended_report_usage": "Use the field structure, key columns, statistics, and sample values to analyze product scope, enterprise scope, price, purchase-volume, or selected-result fields; do not list full rows or invent missing row-level values.",
        **flags,
    }


def _extract_key_facts(text: str) -> list[dict[str, str]]:
    clean = _clean_text(text)
    facts: list[dict[str, str]] = []
    for pattern, label in [
        (r"\d{4}年\d{1,2}月\d{1,2}日|\d{4}-\d{1,2}-\d{1,2}", "时间"),
        (r"(?:申报|报价|采购|执行|递交|实施|规范整合).{0,40}(?:时间|截止|周期|要求|日期|项目)", "要求"),
        (r"(?:价格|报价|申报价|中选价|挂网价|收费标准|医保支付|支付类别).{0,60}", "价格/支付规则"),
        (r"(?:企业|医疗机构|产品|注册证|医保编码|医疗服务价格项目|项目内涵|计价单位).{0,60}", "主体/产品/项目"),
    ]:
        for match in re.finditer(pattern, clean):
            value = match.group(0).strip()
            if value and all(item["value"] != value for item in facts):
                facts.append({"name": label, "value": value})
            if len(facts) >= 12:
                return facts
    return facts


def _important_sections_from_text(text: str) -> list[str]:
    sections = _section_summaries_from_text(text, per_topic_limit=500)
    for line in re.split(r"[。；;\n]", _clean_text(text)):
        clean = line.strip()
        if len(clean) < 12:
            continue
        if any(clean in section for section in sections):
            continue
        if any(word in clean for word in IMPORTANT_SECTION_KEYWORDS):
            sections.append(clean[:500])
        if len(sections) >= 12:
            break
    return sections


class _HTMLTextParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.parts: list[str] = []

    def handle_data(self, data: str) -> None:
        clean = _clean_text(data)
        if clean:
            self.parts.append(clean)


def _parse_docx(content: bytes) -> tuple[str, list[dict[str, Any]]]:
    doc = Document(io.BytesIO(content))
    parts: list[str] = [paragraph.text for paragraph in doc.paragraphs if _clean_text(paragraph.text)]
    table_summaries: list[dict[str, Any]] = []
    for index, table in enumerate(doc.tables, start=1):
        rows = [[_clean_text(cell.text) for cell in row.cells] for row in table.rows]
        if not rows:
            continue
        headers, header_index = _select_header_row(rows)
        table_summaries.append(
            _table_summary(
                sheet_name=f"表格{index}",
                rows=len(rows),
                columns=len(headers),
                headers=headers,
                sample_count=max(0, min(len(rows) - header_index - 1, 20)),
                default_business_value="可用于补充附件中的结构化规则或清单信息。",
            )
        )
        for row in rows[:10]:
            parts.append(" | ".join(row))
    return "\n".join(parts), table_summaries


def _libreoffice_binary() -> str:
    configured = (os.getenv("LIBREOFFICE_BIN") or "").strip()
    if configured:
        return configured
    return shutil.which("libreoffice") or shutil.which("soffice") or ""


def _convert_doc_to_docx(content: bytes, filename: str) -> tuple[bytes, list[str]]:
    binary = _libreoffice_binary()
    if not binary:
        raise RuntimeError("LibreOffice is not installed or LIBREOFFICE_BIN is not configured")
    safe_stem = re.sub(r"[^A-Za-z0-9._-]+", "_", Path(filename or "legacy.doc").stem).strip("._") or "legacy"
    with tempfile.TemporaryDirectory(prefix="medical_notice_doc_") as tmpdir:
        tmp_path = Path(tmpdir)
        source = tmp_path / f"{safe_stem}.doc"
        source.write_bytes(content)
        command = [binary, "--headless", "--convert-to", "docx", "--outdir", str(tmp_path), str(source)]
        result = subprocess.run(command, capture_output=True, text=True, timeout=90, check=False)
        if result.returncode != 0:
            message = (result.stderr or result.stdout or "LibreOffice failed to convert .doc").strip()
            raise RuntimeError(message[:500])
        converted = tmp_path / f"{safe_stem}.docx"
        if not converted.exists():
            candidates = sorted(tmp_path.glob("*.docx"))
            if candidates:
                converted = candidates[0]
        if not converted.exists():
            raise RuntimeError("LibreOffice did not create a .docx file")
        return converted.read_bytes(), ["DOC 文件已通过 LibreOffice 临时转换为 DOCX 后解析，未长期保存原文件。"]


def _parse_xlsx(content: bytes) -> list[dict[str, Any]]:
    wb = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    summaries: list[dict[str, Any]] = []
    max_scan_rows = max(1, _int_env("TABLE_STATS_MAX_SCAN_ROWS", 5000))
    for ws in wb.worksheets:
        preview_rows = [_row_values(row) for row in ws.iter_rows(values_only=True, max_row=25)]
        headers, header_index = _select_header_row(preview_rows)
        rows = ws.max_row or len(preview_rows)
        columns = ws.max_column or len(headers)
        data_rows: list[list[str]] = []
        if headers:
            for row in ws.iter_rows(values_only=True, min_row=header_index + 2, max_row=min(rows, header_index + 1 + max_scan_rows)):
                data_rows.append(_row_values(row))
        summaries.append(
            _table_summary(
                sheet_name=ws.title,
                rows=rows,
                columns=columns,
                headers=headers,
                sample_count=min(max(rows - header_index - 1, 0), 20),
                data_rows=data_rows,
                default_business_value="可用于分析产品范围、企业申报口径、价格字段或采购清单结构。",
            )
        )
    return summaries


def _parse_xls(content: bytes) -> list[dict[str, Any]]:
    book = xlrd.open_workbook(file_contents=content)
    summaries: list[dict[str, Any]] = []
    max_scan_rows = max(1, _int_env("TABLE_STATS_MAX_SCAN_ROWS", 5000))
    for sheet in book.sheets():
        preview_rows = [
            [str(sheet.cell_value(row, col)).strip() for col in range(sheet.ncols)]
            for row in range(min(sheet.nrows, 25))
        ]
        headers, header_index = _select_header_row(preview_rows)
        data_rows = [
            [str(sheet.cell_value(row, col)).strip() for col in range(sheet.ncols)]
            for row in range(header_index + 1, min(sheet.nrows, header_index + 1 + max_scan_rows))
        ]
        summaries.append(
            _table_summary(
                sheet_name=sheet.name,
                rows=sheet.nrows,
                columns=sheet.ncols,
                headers=headers,
                sample_count=min(max(sheet.nrows - header_index - 1, 0), 20),
                data_rows=data_rows,
                default_business_value="可用于分析产品范围、企业申报口径、价格字段或采购清单结构。",
            )
        )
    return summaries


def _parse_csv(content: bytes) -> list[dict[str, Any]]:
    text = content.decode("utf-8-sig", errors="replace")
    reader = csv.reader(io.StringIO(text))
    preview_rows: list[list[str]] = []
    scanned_rows: list[list[str]] = []
    row_count = 0
    max_scan_rows = max(1, _int_env("TABLE_STATS_MAX_SCAN_ROWS", 5000))
    for row in reader:
        normalized = _row_values(row)
        if row_count < 25:
            preview_rows.append(normalized)
        if len(scanned_rows) < max_scan_rows:
            scanned_rows.append(normalized)
        row_count += 1
    headers, header_index = _select_header_row(preview_rows)
    data_rows = scanned_rows[header_index + 1 : header_index + 1 + max_scan_rows]
    return [
        _table_summary(
            sheet_name="CSV",
            rows=row_count,
            columns=len(headers),
            headers=headers,
            sample_count=min(max(row_count - header_index - 1, 0), 20),
            data_rows=data_rows,
            default_business_value="可用于分析清单类结构化信息。",
        )
    ]


def _parse_pdf(content: bytes) -> str:
    reader = PdfReader(io.BytesIO(content))
    max_pages = max(1, _int_env("ATTACHMENT_PDF_MAX_PAGES", 120))
    max_chars = max(5000, _int_env("ATTACHMENT_PDF_MAX_EXTRACT_CHARS", 60000))
    texts = []
    for page in reader.pages[:max_pages]:
        texts.append(page.extract_text() or "")
        if sum(len(item) for item in texts) >= max_chars:
            break
    return "\n".join(texts)


def _env_bool(name: str, default: bool) -> bool:
    raw = (os.getenv(name) or "").strip().lower()
    if not raw:
        return default
    return raw in {"1", "true", "yes", "on"}


def _ocr_pdf_text(content: bytes) -> tuple[str, list[str]]:
    warnings: list[str] = []
    if not _env_bool("ENABLE_PDF_OCR", True):
        return "", ["PDF OCR 未启用"]
    pdftoppm = shutil.which("pdftoppm")
    tesseract = shutil.which("tesseract")
    if not pdftoppm or not tesseract:
        missing = []
        if not pdftoppm:
            missing.append("pdftoppm")
        if not tesseract:
            missing.append("tesseract")
        return "", [f"PDF OCR 依赖缺失: {', '.join(missing)}"]

    max_pages = max(1, _int_env("ATTACHMENT_PDF_OCR_MAX_PAGES", 8))
    dpi = max(120, _int_env("ATTACHMENT_PDF_OCR_DPI", 160))
    lang = (os.getenv("ATTACHMENT_PDF_OCR_LANG") or "chi_sim+eng").strip() or "chi_sim+eng"
    texts: list[str] = []
    with tempfile.TemporaryDirectory(prefix="medical_notice_pdf_ocr_") as tmpdir:
        tmp_path = Path(tmpdir)
        pdf_path = tmp_path / "source.pdf"
        pdf_path.write_bytes(content)
        prefix = tmp_path / "page"
        convert_cmd = [pdftoppm, "-f", "1", "-l", str(max_pages), "-r", str(dpi), "-png", str(pdf_path), str(prefix)]
        converted = subprocess.run(convert_cmd, capture_output=True, text=True, timeout=180, check=False)
        if converted.returncode != 0:
            message = (converted.stderr or converted.stdout or "pdftoppm failed").strip()
            return "", [f"PDF OCR 页面转换失败: {message[:200]}"]
        images = sorted(tmp_path.glob("page-*.png"))
        if not images:
            return "", ["PDF OCR 页面转换未生成图片"]
        for image in images[:max_pages]:
            ocr_cmd = [tesseract, str(image), "stdout", "-l", lang, "--psm", "6"]
            result = subprocess.run(ocr_cmd, capture_output=True, text=True, timeout=90, check=False)
            if result.returncode != 0:
                message = (result.stderr or result.stdout or "tesseract failed").strip()
                warnings.append(f"PDF OCR 页面失败: {image.name}: {message[:160]}")
                continue
            page_text = _clean_text(result.stdout)
            if page_text:
                texts.append(page_text)
    if texts:
        warnings.append(f"PDF 文本抽取为空，已对前 {min(max_pages, len(texts))} 页执行 OCR。")
    return "\n".join(texts), warnings


def _parse_zip(content: bytes, filename: str) -> tuple[list[dict[str, Any]], list[str]]:
    warnings: list[str] = []
    child_summaries: list[dict[str, Any]] = []
    with zipfile.ZipFile(io.BytesIO(content)) as archive:
        total = 0
        for info in archive.infolist()[:20]:
            if info.is_dir():
                continue
            if ".." in info.filename.replace("\\", "/").split("/"):
                warnings.append(f"ZIP 内文件存在路径穿越风险，已跳过: {info.filename}")
                continue
            total += info.file_size
            if total > int(_float_env("ATTACHMENT_MAX_PARSE_MB", 30) * 1024 * 1024):
                warnings.append("ZIP 内文件总大小超过解析限制，已截断处理。")
                break
            child = parse_attachment_bytes(archive.read(info), info.filename, os.path.splitext(info.filename)[1], info.file_size)
            child_summaries.append(
                {
                    "filename": info.filename,
                    "summary": child.get("summary", ""),
                    "parse_statuses": child.get("parse_statuses", []),
                    "table_summaries": child.get("table_summaries", []),
                }
            )
    return child_summaries, warnings


def _failure_result(summary: str, warning: str) -> dict[str, Any]:
    return {
        "parse_statuses": ["parse_failed"],
        "text_length": 0,
        "summary": summary,
        "key_facts": [],
        "important_sections": [],
        "table_summaries": [],
        "warnings": [warning],
    }


def _too_large_metadata_result(size: int) -> dict[str, Any]:
    return {
        "parse_statuses": ["too_large_summary_only"],
        "text_length": 0,
        "summary": f"附件大小 {size} 字节，超过解析限制，仅保留元数据和结构信息。",
        "key_facts": [],
        "important_sections": [],
        "table_summaries": [],
        "warnings": ["附件超过最大解析大小限制"],
    }


def parse_attachment_bytes(content: bytes, filename: str, fileext: str, filesize: int | str | None = None) -> dict[str, Any]:
    ext = (fileext or os.path.splitext(filename)[1] or "").lower()
    size = int(filesize or len(content) or 0)
    max_parse_bytes = int(_float_env("ATTACHMENT_MAX_PARSE_MB", 30) * 1024 * 1024)
    warnings: list[str] = []
    parse_statuses: list[str] = []
    text = ""
    table_summaries: list[dict[str, Any]] = []

    if size > max_parse_bytes:
        if ext != ".pdf":
            return _too_large_metadata_result(size)
        warnings.append("PDF 附件超过常规解析大小限制，已按前若干页抽取文本并生成摘要。")
        parse_statuses.append("too_large_summary_only")
        try:
            text = _parse_pdf(content)
        except Exception as exc:  # noqa: BLE001
            return _failure_result("PDF 附件超过解析限制，且限页文本抽取失败。", f"PDF 限页解析失败: {exc.__class__.__name__}")
        if not _clean_text(text):
            ocr_text, ocr_warnings = _ocr_pdf_text(content)
            warnings.extend(ocr_warnings)
            text = ocr_text
        if text:
            parse_statuses.append("parsed_text")
    else:
        try:
            if ext == ".pdf":
                text = _parse_pdf(content)
                parse_statuses.append("parsed_text")
            elif ext == ".docx":
                text, table_summaries = _parse_docx(content)
                parse_statuses.append("parsed_text")
                if table_summaries:
                    parse_statuses.append("parsed_table_summary")
            elif ext == ".doc":
                try:
                    converted, conversion_warnings = _convert_doc_to_docx(content, filename)
                except Exception as exc:  # noqa: BLE001
                    return _failure_result(
                        "DOC 格式需要 LibreOffice 转换后解析，当前转换失败。",
                        f"DOC 格式解析失败: {exc.__class__.__name__}",
                    )
                warnings.extend(conversion_warnings)
                text, table_summaries = _parse_docx(converted)
                parse_statuses.extend(["temp_file_parsed", "parsed_text"])
                if table_summaries:
                    parse_statuses.append("parsed_table_summary")
            elif ext in {".xlsx", ".xlsm"}:
                table_summaries = _parse_xlsx(content)
                parse_statuses.append("parsed_table_summary")
            elif ext == ".xls":
                table_summaries = _parse_xls(content)
                parse_statuses.append("parsed_table_summary")
            elif ext == ".csv":
                table_summaries = _parse_csv(content)
                text = content.decode("utf-8-sig", errors="replace")[:5000]
                parse_statuses.extend(["parsed_text", "parsed_table_summary"])
            elif ext in {".txt", ".text"}:
                text = content.decode("utf-8", errors="replace")
                parse_statuses.append("parsed_text")
            elif ext in {".html", ".htm"}:
                parser = _HTMLTextParser()
                parser.feed(content.decode("utf-8", errors="replace"))
                text = "\n".join(parser.parts)
                parse_statuses.append("parsed_text")
            elif ext == ".zip":
                children, zip_warnings = _parse_zip(content, filename)
                warnings.extend(zip_warnings)
                table_summaries = [summary for child in children for summary in child.get("table_summaries", [])]
                text = "\n".join(child.get("summary", "") for child in children)
                parse_statuses.append("parsed_summary")
                if table_summaries:
                    parse_statuses.append("parsed_table_summary")
            else:
                return {
                    "parse_statuses": ["unsupported"],
                    "text_length": 0,
                    "summary": f"{ext or '未知'} 格式暂不支持解析。",
                    "key_facts": [],
                    "important_sections": [],
                    "table_summaries": [],
                    "warnings": ["附件格式暂不支持"],
                }
        except Exception as exc:  # noqa: BLE001
            return _failure_result("附件解析失败。", f"附件解析失败: {exc.__class__.__name__}")

    if table_summaries and not text:
        summary = "；".join(item.get("summary", "") for item in table_summaries[:3])
    else:
        summary = _summary_from_text(text)
    if not summary and "too_large_summary_only" in parse_statuses:
        return {
            "parse_statuses": [*parse_statuses, "parse_failed"],
            "text_length": 0,
            "summary": "PDF 附件超过常规解析大小限制，文本抽取和 OCR 未获得可用内容。",
            "key_facts": [],
            "important_sections": [],
            "table_summaries": table_summaries,
            "warnings": warnings or ["PDF 附件未获得可用文本"],
        }
    if summary and "parsed_summary" not in parse_statuses:
        parse_statuses.append("parsed_summary")

    return {
        "parse_statuses": parse_statuses or ["metadata_only"],
        "text_length": len(_clean_text(text)),
        "summary": summary,
        "key_facts": _extract_key_facts(text),
        "important_sections": _important_sections_from_text(text),
        "table_summaries": table_summaries,
        "warnings": warnings,
    }
