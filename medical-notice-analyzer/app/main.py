from __future__ import annotations

import asyncio
import copy
import csv
import hashlib
import html as html_lib
import io
import json
import logging
import os
import re
import subprocess
import tempfile
import threading
import time
import uuid
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any
from urllib.parse import quote, urljoin, urlparse
from urllib.parse import parse_qs

import httpx
import xlrd
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_COLOR_INDEX, WD_LINE_SPACING
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Inches, Pt, RGBColor
from fastapi import FastAPI, HTTPException, Query, Request
from fastapi.responses import FileResponse, JSONResponse
from openpyxl import load_workbook
from PIL import Image, ImageDraw, ImageFont
from pydantic import BaseModel, Field, HttpUrl
from pypdf import PdfReader

from app.attachment_cache import cleanup_cache, load_cached_result, store_cached_result
from app.attachment_fetcher import fetch_attachment_bytes
from app.attachment_parser import parse_attachment_bytes
from app.diagnostics import build_pack_diagnostics, build_run_diagnostics, build_run_progress

try:
    import pdfplumber  # type: ignore[import-not-found]
except Exception:  # noqa: BLE001
    pdfplumber = None


USER_AGENT = (
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
    "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/125.0 Safari/537.36"
)
ATTACHMENT_EXTENSIONS = {
    ".pdf",
    ".doc",
    ".docx",
    ".xls",
    ".xlsx",
    ".xlsm",
    ".csv",
    ".txt",
    ".text",
    ".html",
    ".htm",
    ".zip",
}
MAX_ATTACHMENT_BYTES = 60 * 1024 * 1024
DEFAULT_MAX_COMBINED_CHARS = 120_000
CHATFLOW_SAFE_MAX_COMBINED_CHARS = 60_000
DEFAULT_MAX_ATTACHMENTS = 25
DEFAULT_REPORT_TITLE = "医药器械采购项目分析报告"
DEFAULT_DISCLAIMER = (
    "本文基于互联网公开资料进行整理，目的在于传递分享信息，仅供读者参考之用。"
    "本网站不保证信息的准确性、有效性、及时性和完整性。"
    "本公司及其雇员一概毋须以任何方式就任何信息传递或传送的失误、不准确或错误，"
    "对用户或任何其他人士负任何直接或间接责任。"
    "在法律允许的范围内，本公司在此声明，"
    "不承担用户或任何人士就使用或未能使用本网站所提供的信息或任何链接所引致的任何直接、间接、附带、从属、特殊、惩罚性或惩戒性的损害赔偿。"
)
REPORT_WATERMARK_TEXT = "易联器械"
RULE_COMPLETENESS_KEYWORDS = [
    "采购品种范围",
    "产品分类",
    "最高有效申报价",
    "参考价",
    "企业报价要求",
    "有效报价",
    "拟中选产品确定",
    "中选产品确定",
    "协议采购量",
    "首年协议采购量",
    "采购执行",
    "价格联动",
    "非中选产品管理",
    "新获批产品管理",
    "名词解释",
    "带量最低价",
    "非带量最低价",
    "非中选产品",
    "申报材料",
    "配送要求",
    "信用评价",
    "失信约束",
    "取消中选资格",
    "暂不予挂网",
]
MODELIZED_LANGUAGE_REPLACEMENTS = {
    "绝对红线": "报价上限",
    "全国最低价刚性护城河": "价格联动要求",
    "刚性护城河": "价格联动要求",
    "不给价格虚高留任何死角": "企业需关注后续价格调整义务",
    "中选即触雷": "不接受价格联动的，可能影响对应产品中选资格",
    "全面压缩": "按文件要求调整",
    "强势倒逼": "按文件要求约束",
    "重大冲击": "影响",
    "必然": "可能",
    "显然": "文件显示",
    "唯一选择": "可选路径之一",
    "全国底价": "全国较低价格",
    "全国最低价": "全国较低价格",
    "实时联动": "价格联动",
}


class AnalyzeRequest(BaseModel):
    url: HttpUrl
    firecrawl_api_key: str | None = Field(default=None)
    max_attachments: int = Field(default=DEFAULT_MAX_ATTACHMENTS, ge=0, le=80)
    max_combined_chars: int = Field(default=DEFAULT_MAX_COMBINED_CHARS, ge=20_000, le=600_000)


class AnalyzeV2Request(AnalyzeRequest):
    evidence_mode: bool = True
    uploaded_files: list[Any] = Field(default_factory=list)


class AnalyzeResponse(BaseModel):
    source_url: str
    final_url: str
    title: str
    publish_date_candidates: list[str]
    region_candidates: list[str]
    page_text: str
    attachments: list[dict[str, Any]]
    combined_text: str
    evidence_for_llm: str
    warnings: list[str]


class AnalyzeV2Response(BaseModel):
    success: bool
    run_id: str = ""
    evidence_id: str = ""
    content_hash: str = ""
    source_url: str = ""
    fetch_status: dict[str, Any] = Field(default_factory=dict)
    notice_summary: dict[str, Any] = Field(default_factory=dict)
    evidence_pack: dict[str, Any] = Field(default_factory=dict)
    llm_input_text: str = ""
    raw_storage: dict[str, Any] = Field(default_factory=dict)
    error_type: str = ""
    message: str = ""


class ReportTable(BaseModel):
    title: str = ""
    headers: list[str] = Field(default_factory=list)
    rows: list[list[str]] = Field(default_factory=list)
    notes: list[str] = Field(default_factory=list)


class ReportSection(BaseModel):
    heading: str = ""
    paragraphs: list[str] = Field(default_factory=list)
    tables: list[ReportTable] = Field(default_factory=list)
    highlights: list[str] = Field(default_factory=list)


class ReportIR(BaseModel):
    title: str = ""
    suggested_filename: str = ""
    notice_type: str = ""
    publish_date: str = ""
    source_agency: str = ""
    document_name: str = ""
    lead_paragraphs: list[str] = Field(default_factory=list)
    sections: list[ReportSection] = Field(default_factory=list)
    enterprise_tips: list[str] = Field(default_factory=list)
    disclaimer: str = ""


class ExportReportRequest(BaseModel):
    title: str = Field(default=DEFAULT_REPORT_TITLE)
    markdown: str = ""
    report_ir: ReportIR | None = None
    strict_quality: bool = Field(default=True)


class ExportReportResponse(BaseModel):
    success: bool = True
    filename: str
    download_url: str


class RenderReportRequest(ExportReportRequest):
    pass


class RenderReportResponse(BaseModel):
    success: bool
    report_ir: ReportIR | None = None
    report_markdown: str = ""
    error: str = ""
    user_message: str = ""
    quality_warnings: list[str] = Field(default_factory=list)


class CheckedExportReportRequest(ExportReportRequest):
    qa_output: str = ""
    qa_status: str = ""
    qa_result: dict[str, Any] | None = None
    report_text: str = ""
    history_text: str = ""
    evidence_text: str = ""


class CheckedExportReportResponse(BaseModel):
    success: bool
    filename: str = ""
    download_url: str = ""
    blocked: bool = False
    qa_summary: str = ""
    report_markdown: str = ""


class ReportQAIssue(BaseModel):
    severity: str = "minor"
    category: str = ""
    report_text: str = ""
    source_quote: str = ""
    fix_instruction: str = ""


class ReportQA(BaseModel):
    status: str = "pass"
    issues: list[ReportQAIssue] = Field(default_factory=list)
    unsupported_claims: list[ReportQAIssue] = Field(default_factory=list)
    history_leakage: list[ReportQAIssue] = Field(default_factory=list)
    missing_rules: list[str] = Field(default_factory=list)
    language_issues: list[ReportQAIssue] = Field(default_factory=list)
    fix_instructions: list[str] = Field(default_factory=list)
    summary: str = ""


class ReportQAParseRequest(BaseModel):
    qa_output: str = ""
    report_text: str = ""
    report_ir: ReportIR | None = None
    history_text: str = ""
    evidence_text: str = ""


class ReportQAParseResponse(BaseModel):
    success: bool = True
    blocked: bool
    needs_fix: bool = False
    status: str = "pass"
    issues: list[ReportQAIssue] = Field(default_factory=list)
    unsupported_claims: list[ReportQAIssue] = Field(default_factory=list)
    history_leakage: list[ReportQAIssue] = Field(default_factory=list)
    missing_rules: list[str] = Field(default_factory=list)
    language_issues: list[ReportQAIssue] = Field(default_factory=list)
    fix_instructions: list[str] = Field(default_factory=list)
    summary: str = ""
    qa: ReportQA
    qa_summary: str


class RecordListItem(BaseModel):
    menu_code: str = ""
    articleid: str = ""
    title: str = ""
    audittime: Any = ""
    menu_name: str = ""
    areaname: str = ""
    source: str = ""
    publicorg: str = ""
    projectphase: str = ""
    projecttype: str = ""
    category: str = ""
    sourceurl: str = ""
    summary: str = ""
    attachment_count: int = 0
    list_attach_filename: str = ""


class RecordListResponse(BaseModel):
    items: list[RecordListItem]
    page: int
    page_size: int
    total: int
    total_pages: int


class RecordAttachment(BaseModel):
    articleattid: str = ""
    filename: str = ""
    filepath: str = ""
    fileext: str = ""
    filesize: Any = ""
    uploadtime: Any = ""
    sortnum: Any = ""


class RecordDetailResponse(BaseModel):
    menu_code: str = ""
    articleid: str = ""
    title: str = ""
    audittime: Any = ""
    updatetime: Any = ""
    menu_name: str = ""
    areaname: str = ""
    source: str = ""
    sourceurl: str = ""
    publicorg: str = ""
    projectphase: str = ""
    projecttype: str = ""
    dl_project_type: str = ""
    category: str = ""
    referencenumber: str = ""
    policytype: str = ""
    belongproject: str = ""
    projectabbreviation: str = ""
    summary: str = ""
    content: str = ""
    content_text: str = ""
    attachments: list[RecordAttachment] = Field(default_factory=list)


class MaterialRef(BaseModel):
    menu_code: str = Field(min_length=1)
    articleid: str = Field(min_length=1)


class SelectionPreviewRequest(BaseModel):
    primary_materials: list[MaterialRef] = Field(default_factory=list)
    auxiliary_materials: list[MaterialRef] = Field(default_factory=list)
    enable_attachment_download: bool | None = None
    force_refresh_attachments: bool = False
    attachment_cookie: str = ""
    attachment_headers: dict[str, str] = Field(default_factory=dict)


class SelectionPreviewResponse(BaseModel):
    success: bool
    primary_materials: list[RecordListItem]
    auxiliary_materials: list[RecordListItem]


class AnalysisPrepareResponse(BaseModel):
    success: bool
    pack_id: str
    primary_count: int
    auxiliary_count: int
    attachment_count: int
    warnings: list[str] = Field(default_factory=list)
    evidence_pack: dict[str, Any] = Field(default_factory=dict)


class AnalysisRunRequest(BaseModel):
    pack_id: str = Field(min_length=1)


class AnalysisRunResponse(BaseModel):
    success: bool
    run_id: str
    pack_id: str
    status: str
    workflow_run_id: str = ""
    report_title: str = ""
    quality_passed: bool | None = None
    version: int = 1
    warnings: list[str] = Field(default_factory=list)


class AnalysisRunReportResponse(BaseModel):
    success: bool
    run_id: str
    pack_id: str
    report_title: str = ""
    report_markdown: str = ""
    quality_check: dict[str, Any] = Field(default_factory=dict)
    quality_gate: dict[str, Any] = Field(default_factory=dict)
    version: int = 1
    warnings: list[str] = Field(default_factory=list)
    remaining_issues: list[Any] = Field(default_factory=list)


class AnalysisRunReviseRequest(BaseModel):
    feedback: str = Field(min_length=1, max_length=4000)
    mode: str = "user_feedback"
    analysis_highlight: bool | None = None


class AnalysisRunReviseResponse(BaseModel):
    success: bool
    revision_id: str
    run_id: str
    pack_id: str
    version: int
    report_title: str = ""
    report_markdown: str = ""
    quality_check: dict[str, Any] = Field(default_factory=dict)
    warnings: list[str] = Field(default_factory=list)


@dataclass
class DownloadedFile:
    url: str
    filename: str
    content_type: str
    content: bytes


logging.basicConfig(
    level=os.getenv("LOG_LEVEL", "INFO").upper(),
    format="%(asctime)s %(levelname)s %(name)s %(message)s",
)
logger = logging.getLogger("medical_notice_analyzer")

app = FastAPI(title="Medical Notice Analyzer", version="0.1.0")
REPORT_DIR = Path(os.getenv("REPORT_DIR", "/tmp/medical-notice-reports"))
SITE_CACHE_DIR = Path(os.getenv("SITE_CACHE_DIR", "/app/site-cache"))
DEFAULT_PUBLIC_BASE_URL = "http://192.168.34.88:8099"


@app.on_event("startup")
def _cleanup_attachment_parse_cache_on_start() -> None:
    if not _env_bool("ATTACHMENT_PARSE_CACHE_CLEANUP_ON_START", True):
        return
    try:
        result = cleanup_cache()
        if result.get("deleted_files"):
            logger.info(
                "attachment_parse_cache_cleanup deleted_files=%s freed_bytes=%s",
                result.get("deleted_files"),
                result.get("freed_bytes"),
            )
    except Exception as exc:  # noqa: BLE001
        logger.warning("attachment_parse_cache_cleanup_failed error_type=%s", exc.__class__.__name__)

ARTICLE_LIST_FIELDS = [
    "menu_code",
    "articleid",
    "title",
    "audittime",
    "menu_name",
    "areaname",
    "source",
    "publicorg",
    "projectphase",
    "projecttype",
    "category",
    "sourceurl",
    "summary",
    "list_attach_filename",
]
ARTICLE_DETAIL_FIELDS = [
    "menu_code",
    "articleid",
    "title",
    "audittime",
    "updatetime",
    "menu_name",
    "areaname",
    "source",
    "sourceurl",
    "publicorg",
    "projectphase",
    "projecttype",
    "dl_project_type",
    "category",
    "referencenumber",
    "policytype",
    "belongproject",
    "projectabbreviation",
    "summary",
    "content",
]
ATTACHMENT_FIELDS = [
    "articleattid",
    "filename",
    "filepath",
    "fileext",
    "filesize",
    "uploadtime",
    "sortnum",
    "fileerrortype",
]


@app.middleware("http")
async def log_requests(request: Request, call_next):
    started = time.perf_counter()
    try:
        response = await call_next(request)
    except Exception:
        elapsed_ms = int((time.perf_counter() - started) * 1000)
        logger.exception(
            "request_failed method=%s path=%s elapsed_ms=%s",
            request.method,
            request.url.path,
            elapsed_ms,
        )
        raise
    elapsed_ms = int((time.perf_counter() - started) * 1000)
    logger.info(
        "request_completed method=%s path=%s status_code=%s elapsed_ms=%s",
        request.method,
        request.url.path,
        response.status_code,
        elapsed_ms,
    )
    return response


@app.get("/health")
def health() -> dict[str, Any]:
    public_base_url = _configured_public_base_url()
    return {
        "status": "ok",
        "service": "medical-notice-analyzer",
        "version": app.version,
        "public_base_url_configured": bool((os.getenv("PUBLIC_BASE_URL") or "").strip()),
        "public_base_url": public_base_url,
        "report_dir_configured": bool((os.getenv("REPORT_DIR") or "").strip()),
        "report_dir": str(REPORT_DIR),
        "site_cache_dir_configured": bool((os.getenv("SITE_CACHE_DIR") or "").strip()),
        "max_attachment_bytes": MAX_ATTACHMENT_BYTES,
    }


def _db_config() -> dict[str, Any]:
    return {
        "host": os.getenv("DB_HOST", "192.168.36.36"),
        "port": int(os.getenv("DB_PORT", "3306")),
        "database": os.getenv("DB_NAME", "wangfanqi_test"),
        "user": os.getenv("DB_USER", ""),
        "password": os.getenv("DB_PASSWORD", ""),
        "charset": os.getenv("DB_CHARSET", "utf8mb4"),
    }


def _env_bool(name: str, default: bool = False) -> bool:
    value = (os.getenv(name) or "").strip().lower()
    if not value:
        return default
    return value in {"1", "true", "yes", "on"}


def _db_connect():
    cfg = _db_config()
    if not cfg["database"] or not cfg["user"]:
        raise HTTPException(status_code=500, detail="Database configuration is incomplete")
    try:
        import pymysql
        from pymysql.cursors import DictCursor
    except Exception as exc:  # noqa: BLE001
        raise HTTPException(status_code=500, detail="Database driver pymysql is not installed") from exc
    try:
        return pymysql.connect(
            host=cfg["host"],
            port=cfg["port"],
            user=cfg["user"],
            password=cfg["password"],
            database=cfg["database"],
            charset=cfg["charset"],
            cursorclass=DictCursor,
            read_timeout=20,
            write_timeout=20,
            connect_timeout=10,
        )
    except Exception as exc:  # noqa: BLE001
        logger.warning("database_connection_failed host=%s port=%s database=%s error_type=%s", cfg["host"], cfg["port"], cfg["database"], exc.__class__.__name__)
        raise HTTPException(status_code=503, detail="Database connection failed") from exc


def _db_fetch_all(sql: str, params: list[object]) -> list[dict[str, Any]]:
    with _db_connect() as conn:
        with conn.cursor() as cursor:
            cursor.execute(sql, params)
            return list(cursor.fetchall())


def _db_fetch_one(sql: str, params: list[object]) -> dict[str, Any] | None:
    with _db_connect() as conn:
        with conn.cursor() as cursor:
            cursor.execute(sql, params)
            row = cursor.fetchone()
            return dict(row) if row else None


def _article_text_from_html(content: str) -> str:
    if not content:
        return ""
    soup = BeautifulSoup(content, "html.parser")
    for tag in soup(["script", "style", "noscript", "svg"]):
        tag.decompose()
    return re.sub(r"\s+", " ", _normalize_text(soup.get_text(" "))).strip()


def _normalize_db_row(row: dict[str, Any]) -> dict[str, Any]:
    result = dict(row)
    for key, value in list(result.items()):
        if isinstance(value, datetime):
            result[key] = value.isoformat(sep=" ")
        elif value is None:
            result[key] = ""
    if "attachment_count" in result:
        result["attachment_count"] = int(result.get("attachment_count") or 0)
    return result


def _record_key(material: MaterialRef) -> tuple[str, str]:
    return material.menu_code, material.articleid


def _record_list_item(row: dict[str, Any]) -> RecordListItem:
    return RecordListItem(**_normalize_db_row(row))


def _build_records_where(
    *,
    keyword: str = "",
    menu_code: str = "",
    areaname: str = "",
    projectphase: str = "",
    projecttype: str = "",
    start_date: str = "",
    end_date: str = "",
) -> tuple[str, list[object]]:
    clauses = ["a.status = %s"]
    params: list[object] = [0]
    if keyword:
        clauses.append("a.title LIKE %s")
        params.append(f"%{keyword}%")
    if menu_code:
        clauses.append("(a.menu_code = %s OR a.menu_name LIKE %s)")
        params.extend([menu_code, f"%{menu_code}%"])
    if areaname:
        clauses.append("a.areaname = %s")
        params.append(areaname)
    if projectphase:
        clauses.append("a.projectphase = %s")
        params.append(projectphase)
    if projecttype:
        clauses.append("a.projecttype = %s")
        params.append(projecttype)
    if start_date:
        clauses.append("a.audittime >= %s")
        params.append(start_date)
    if end_date:
        clauses.append("a.audittime <= %s")
        params.append(f"{end_date} 23:59:59" if len(end_date) == 10 else end_date)
    return " AND ".join(clauses), params


def _records_by_keys(keys: list[tuple[str, str]]) -> dict[tuple[str, str], RecordListItem]:
    if not keys:
        return {}
    placeholders = ", ".join(["(%s, %s)"] * len(keys))
    params: list[object] = []
    for menu_code, articleid in keys:
        params.extend([menu_code, articleid])
    selected_fields = ", ".join(f"a.{field}" for field in ARTICLE_LIST_FIELDS)
    group_fields = ", ".join(f"a.{field}" for field in ARTICLE_LIST_FIELDS)
    sql = f"""
        SELECT {selected_fields}, COUNT(att.articleattid) AS attachment_count
        FROM sample_article_wide a
        LEFT JOIN sample_article_attach att
          ON a.menu_code = att.menu_code AND a.articleid = att.articleid
        WHERE a.status = %s AND (a.menu_code, a.articleid) IN ({placeholders})
        GROUP BY {group_fields}
    """
    rows = _db_fetch_all(sql, [0, *params])
    return {(str(row.get("menu_code") or ""), str(row.get("articleid") or "")): _record_list_item(row) for row in rows}


def _selection_error(status_code: int, code: str, message: str) -> JSONResponse:
    return JSONResponse(
        status_code=status_code,
        content={
            "success": False,
            "error": {
                "code": code,
                "message": message,
            },
        },
    )


def _analysis_error(status_code: int, code: str, message: str, detail: str = "") -> JSONResponse:
    error: dict[str, Any] = {
        "code": code,
        "message": message,
    }
    if detail:
        error["detail"] = _truncate(str(detail), 500)
    return JSONResponse(status_code=status_code, content={"success": False, "error": error})


def _validate_material_selection(req: SelectionPreviewRequest) -> tuple[list[tuple[str, str]], list[tuple[str, str]], JSONResponse | None]:
    if not (1 <= len(req.primary_materials) <= 3):
        return [], [], _selection_error(422, "INVALID_SELECTION", "主分析材料必须选择 1-3 条")
    if len(req.auxiliary_materials) > 10:
        return [], [], _selection_error(422, "INVALID_SELECTION", "辅助分析材料最多选择 10 条")

    primary_keys = [_record_key(item) for item in req.primary_materials]
    auxiliary_keys = [_record_key(item) for item in req.auxiliary_materials]
    overlap = set(primary_keys) & set(auxiliary_keys)
    if overlap:
        return [], [], _selection_error(422, "INVALID_SELECTION", "同一条文章不能同时作为主分析材料和辅助分析材料")
    return primary_keys, auxiliary_keys, None


def _database_evidence_pack_dir() -> Path:
    configured = (os.getenv("EVIDENCE_PACK_DIR") or "").strip()
    if configured:
        return Path(configured)
    return _evidence_cache_dir() / "database_packs"


def _safe_pack_id(pack_id: str) -> str:
    if not re.fullmatch(r"pack_[A-Za-z0-9_-]{8,80}", pack_id or ""):
        raise HTTPException(status_code=404, detail="evidence pack not found")
    return pack_id


def _database_pack_path(pack_id: str) -> Path:
    return _database_evidence_pack_dir() / f"{_safe_pack_id(pack_id)}.json"


def _write_database_evidence_pack(pack: dict[str, Any]) -> None:
    pack_id = str(pack.get("pack_id") or "")
    path = _database_pack_path(pack_id)
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(pack, ensure_ascii=False, indent=2), encoding="utf-8")
    logger.info("database_evidence_pack_saved pack_id=%s path=%s", pack_id, path)


def _read_database_evidence_pack(pack_id: str) -> dict[str, Any]:
    path = _database_pack_path(pack_id)
    if not path.exists():
        raise HTTPException(status_code=404, detail="evidence pack not found")
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except Exception as exc:  # noqa: BLE001
        logger.warning("database_evidence_pack_read_failed pack_id=%s error_type=%s", pack_id, exc.__class__.__name__)
        raise HTTPException(status_code=500, detail="evidence pack read failed") from exc


def _limit_text(value: Any, limit: int) -> str:
    text = _normalize_text(str(value or ""))
    if limit <= 0 or len(text) <= limit:
        return text
    return f"{text[:limit]}……"


_DIFY_USABLE_ATTACHMENT_STATUSES = {
    "stream_parsed",
    "temp_file_parsed",
    "parsed_text",
    "parsed_summary",
    "parsed_table_summary",
    "too_large_summary_only",
}


def _material_content_text_length(material: dict[str, Any]) -> int:
    explicit = material.get("content_text_length")
    if isinstance(explicit, int) and explicit >= 0:
        return explicit
    return len(str(material.get("content_text") or "").strip())


def _attachment_is_usable_for_dify(attachment: dict[str, Any]) -> bool:
    parse_status = str(attachment.get("parse_status") or "").strip()
    parse_statuses = attachment.get("parse_statuses") if isinstance(attachment.get("parse_statuses"), list) else []
    return parse_status in _DIFY_USABLE_ATTACHMENT_STATUSES or any(item in _DIFY_USABLE_ATTACHMENT_STATUSES for item in parse_statuses)


def _material_attachment_evidence_chars(material: dict[str, Any]) -> int:
    total = 0
    for attachment in material.get("attachments") or []:
        if not isinstance(attachment, dict) or not _attachment_is_usable_for_dify(attachment):
            continue
        total += len(str(attachment.get("summary") or "").strip())
        total += len(json.dumps(attachment.get("key_facts") or [], ensure_ascii=False))
        total += len(json.dumps(attachment.get("important_sections") or [], ensure_ascii=False))
        for table in attachment.get("table_summaries") or []:
            if isinstance(table, dict):
                total += len(
                    json.dumps(
                        {
                            "summary": table.get("summary"),
                            "business_value": table.get("business_value"),
                            "headers": table.get("headers"),
                            "key_columns": table.get("key_columns"),
                        },
                        ensure_ascii=False,
                    )
                )
    return total


def _is_attachment_led_primary_material(material: dict[str, Any]) -> bool:
    if _material_content_text_length(material) >= 800:
        return False
    usable_core_attachment = any(
        isinstance(attachment, dict)
        and bool(attachment.get("core_attachment"))
        and _attachment_is_usable_for_dify(attachment)
        for attachment in material.get("attachments") or []
    )
    return usable_core_attachment and _material_attachment_evidence_chars(material) >= 800


def _env_int(name: str, default: int) -> int:
    try:
        return int(str(os.getenv(name) or "").strip() or default)
    except ValueError:
        return default


def _env_float(name: str, default: float) -> float:
    try:
        return float(str(os.getenv(name) or "").strip() or default)
    except ValueError:
        return default


def _dify_input_thresholds() -> dict[str, int]:
    hard_limit = max(10_000, _env_int("DIFY_VARIABLE_HARD_MAX_CHARS", 80_000))
    full_input = min(hard_limit, max(1000, _env_int("DIFY_FULL_INPUT_MAX_CHARS", 65_000)))
    safe_compact = min(hard_limit, max(full_input, _env_int("DIFY_SAFE_COMPACT_MAX_CHARS", 75_000)))
    return {
        "hard_limit": hard_limit,
        "full_input": full_input,
        "safe_compact": safe_compact,
        "light_compact": safe_compact,
    }


_EVIDENCE_VALUE_KEYWORDS = [
    "申报",
    "采购",
    "价格",
    "报价",
    "申报价",
    "中选价",
    "挂网价",
    "中选",
    "拟中选",
    "产品",
    "企业",
    "注册证",
    "医保编码",
    "医保耗材代码",
    "采购量",
    "报量",
    "需求量",
    "协议采购量",
    "执行",
    "周期",
    "配送",
    "信用评价",
    "价格联动",
    "selected",
    "enterprise",
    "product",
    "price",
    "purchase",
    "volume",
]


def _table_summary_has_business_flags(table: dict[str, Any]) -> bool:
    return any(
        bool(table.get(key))
        for key in [
            "contains_enterprise",
            "contains_product",
            "contains_registration_cert",
            "contains_medical_insurance_code",
            "contains_price",
            "contains_purchase_volume",
            "contains_selected_status",
        ]
    )


def _table_summary_is_heavy(table: dict[str, Any]) -> bool:
    if bool(table.get("table_heavy")):
        return True
    rows = int(table.get("rows") or 0)
    return rows >= max(1, _env_int("TABLE_HEAVY_ROW_THRESHOLD", 5000))


def _pack_has_table_heavy_primary(primary_source: list[dict[str, Any]]) -> bool:
    for material in primary_source:
        for attachment in material.get("attachments") or []:
            if not isinstance(attachment, dict):
                continue
            for table in attachment.get("table_summaries") or []:
                if isinstance(table, dict) and _table_summary_is_heavy(table):
                    return True
    return False


def _evidence_keyword_hits(value: Any) -> int:
    text = json.dumps(value, ensure_ascii=False) if isinstance(value, (dict, list)) else str(value or "")
    return sum(1 for keyword in _EVIDENCE_VALUE_KEYWORDS if keyword.lower() in text.lower())


def _evidence_value_score_attachment(attachment: dict[str, Any], *, role: str) -> int:
    score = 40 if role == "primary" else 12
    if attachment.get("core_attachment"):
        score += 35
    business_type = str(attachment.get("business_type") or attachment.get("filename") or "")
    if any(keyword in business_type for keyword in ["中选", "产品", "价格", "采购", "申报", "企业", "配送", "报量", "selected", "price"]):
        score += 20
    score += min(25, _evidence_keyword_hits(attachment) * 2)
    parse_status = str(attachment.get("parse_status") or "")
    if parse_status in {"parsed_text", "parsed_summary", "parsed_table_summary", "temp_file_parsed"}:
        score += 15
    elif parse_status in {"metadata_only", "download_failed", "parse_failed", "unsupported"}:
        score -= 8
    table_score = 0
    for table in attachment.get("table_summaries") or []:
        if not isinstance(table, dict):
            continue
        if _table_summary_has_business_flags(table):
            table_score += 12
        if _table_summary_is_heavy(table):
            table_score += 10
        table_score += int(table.get("evidence_value_score") or 0) // 5
    return max(0, score + min(40, table_score))


def _evidence_value_score_material(material: dict[str, Any], *, role: str) -> int:
    score = 60 if role == "primary" else 15
    score += min(30, _evidence_keyword_hits({k: material.get(k) for k in ["title", "summary", "content_text", "content_summary"]}) * 2)
    if role == "auxiliary":
        try:
            score += int(float(material.get("relevance_score") or 0) * 20)
        except (TypeError, ValueError):
            pass
    attachment_scores = [
        _evidence_value_score_attachment(item, role=role)
        for item in material.get("attachments") or []
        if isinstance(item, dict)
    ]
    if attachment_scores:
        score += min(70, max(attachment_scores) + sum(attachment_scores[:3]) // 5)
    return max(0, score)


def _dify_relevant_input_chars(pack: dict[str, Any], primary_source: list[dict[str, Any]], auxiliary_source: list[dict[str, Any]]) -> int:
    primary_payload = []
    for material in primary_source:
        primary_payload.append(
            {
                "metadata": {key: material.get(key) for key in ["menu_code", "articleid", "title", "audittime", "areaname", "publicorg", "projecttype", "category"]},
                "summary": material.get("summary"),
                "content_text": material.get("content_text"),
                "content_summary": material.get("content_summary"),
                "key_facts": material.get("key_facts"),
                "important_passages": material.get("important_passages"),
                "policy_rules": material.get("policy_rules"),
                "price_rules": material.get("price_rules"),
                "time_requirements": material.get("time_requirements"),
                "product_scope": material.get("product_scope"),
                "enterprise_requirements": material.get("enterprise_requirements"),
                "execution_requirements": material.get("execution_requirements"),
                "attachments": [
                    {
                        "filename": attachment.get("filename"),
                        "core_attachment": attachment.get("core_attachment"),
                        "business_type": attachment.get("business_type"),
                        "parse_status": attachment.get("parse_status"),
                        "download_status": attachment.get("download_status"),
                        "summary": attachment.get("summary"),
                        "key_facts": attachment.get("key_facts"),
                        "important_sections": attachment.get("important_sections"),
                        "table_summaries": attachment.get("table_summaries"),
                        "warnings": attachment.get("warnings"),
                    }
                    for attachment in material.get("attachments") or []
                    if isinstance(attachment, dict)
                ],
            }
        )
    auxiliary_payload = []
    for material in auxiliary_source:
        auxiliary_payload.append(
            {
                "metadata": {key: material.get(key) for key in ["menu_code", "articleid", "title", "audittime", "areaname", "publicorg", "projecttype", "category"]},
                "content_summary": material.get("content_summary") or material.get("summary"),
                "relation_to_primary": material.get("relation_to_primary"),
                "relevance_score": material.get("relevance_score"),
                "relevant_snippets": material.get("relevant_snippets"),
                "usable_points": material.get("usable_points"),
            }
        )
    payload = {
        "primary_materials": primary_payload,
        "auxiliary_materials": auxiliary_payload,
        "combined_key_facts": pack.get("combined_key_facts"),
        "report_focus": pack.get("report_focus"),
        "warnings": pack.get("warnings"),
        "generation_guidance": pack.get("generation_guidance"),
    }
    return len(json.dumps(payload, ensure_ascii=False, sort_keys=True))


def _attachment_count(materials: list[dict[str, Any]]) -> int:
    return sum(len([item for item in material.get("attachments") or [] if isinstance(item, dict)]) for material in materials)


def _dify_input_strategy(
    *,
    original_pack_chars: int,
    dify_relevant_input_chars: int | None = None,
    primary_source: list[dict[str, Any]],
    auxiliary_source: list[dict[str, Any]],
) -> str:
    thresholds = _dify_input_thresholds()
    relevant_chars = dify_relevant_input_chars if dify_relevant_input_chars is not None else original_pack_chars
    if len(primary_source) >= 2:
        return "staged_generation"
    if any(_is_attachment_led_primary_material(material) for material in primary_source):
        return "attachment_led"
    if _pack_has_table_heavy_primary(primary_source):
        return "table_heavy"
    if relevant_chars <= thresholds["full_input"]:
        return "full_input"
    if relevant_chars <= thresholds["safe_compact"]:
        return "light_compact"
    if relevant_chars <= thresholds["hard_limit"]:
        return "safe_compact"
    return "staged_generation"


def _dify_input_strategy_type(input_strategy: str) -> str:
    return {
        "full_input": "size_based",
        "light_compact": "size_based",
        "safe_compact": "size_based",
        "attachment_led": "structure_based",
        "table_heavy": "structure_based",
        "staged_generation": "complexity_based",
    }.get(input_strategy, "size_based")


def _dify_input_strategy_description(input_strategy: str) -> str:
    return {
        "full_input": (
            "完整证据输入策略：关键证据量未超过安全阈值，系统尽量完整保留主材料正文、"
            "主材料附件摘要、表格摘要和辅助材料摘要。"
        ),
        "light_compact": (
            "轻压缩输入策略：关键证据量处于中等区间，系统只清理重复模板、免责声明和低价值噪声，"
            "保留核心规则、操作步骤、表格列名、产品范围和价格字段。"
        ),
        "safe_compact": (
            "安全压缩策略：输入接近 Dify 变量上限，系统优先保留主材料和核心附件，压缩辅助材料与低价值重复内容。"
        ),
        "attachment_led": (
            "附件主导型生成策略：主材料正文较短，但核心附件已解析出可用摘要或表格结构。"
            "该策略不是压缩模式，未超限时仍会尽量完整保留主附件证据。"
        ),
        "table_heavy": (
            "大表格结构化输入策略：附件行数较多或包含企业、产品、价格、采购量等关键字段。"
            "系统不输入全量行，而是输入字段结构、统计摘要和样例值。"
        ),
        "staged_generation": (
            "复杂材料分段生成策略：多主材料或证据量超过单次 Dify 输入安全范围，"
            "应先按材料梳理规则，再综合生成报告。"
        ),
    }.get(input_strategy, "Dify 自适应输入策略。")


def _dify_input_strategy_basis(input_strategy: str, *, relevant_chars: int, thresholds: dict[str, int], primary_source: list[dict[str, Any]]) -> str:
    if input_strategy == "staged_generation" and len(primary_source) >= 2:
        return "multi_primary_materials"
    if input_strategy == "attachment_led":
        return "short_primary_body_with_usable_core_attachment"
    if input_strategy == "table_heavy":
        return "primary_table_rows_exceed_table_heavy_threshold"
    if input_strategy == "full_input":
        return "dify_relevant_input_chars_within_full_input_threshold"
    if input_strategy == "light_compact":
        return "dify_relevant_input_chars_within_light_compact_threshold"
    if input_strategy == "safe_compact":
        return "near_hard_limit"
    if input_strategy == "staged_generation":
        return "dify_relevant_input_chars_exceeds_hard_limit"
    return "adaptive_input_strategy_default"


def _omitted_content_entry(
    omitted_type: str,
    *,
    reason: str,
    risk: str = "",
    manual_review: bool = False,
    affects_primary_detail: bool = False,
    affects_core_attachment_detail: bool = False,
    affects_auxiliary_detail: bool = False,
    **extra: Any,
) -> dict[str, Any]:
    entry: dict[str, Any] = {
        "type": omitted_type,
        "reason": reason,
        "risk": risk,
        "manual_review": bool(manual_review),
        "affects_primary_detail": bool(affects_primary_detail),
        "affects_core_attachment_detail": bool(affects_core_attachment_detail),
        "affects_auxiliary_detail": bool(affects_auxiliary_detail),
    }
    entry.update({key: value for key, value in extra.items() if value not in (None, "")})
    return entry


def _refresh_dify_char_fields(compact: dict[str, Any]) -> int:
    previous = -1
    current = len(json.dumps(compact, ensure_ascii=False, sort_keys=True))
    for _ in range(6):
        compact["compact_pack_chars"] = current
        compact["final_dify_input_chars"] = current
        next_size = len(json.dumps(compact, ensure_ascii=False, sort_keys=True))
        if next_size == current or next_size == previous:
            current = next_size
            compact["compact_pack_chars"] = current
            compact["final_dify_input_chars"] = current
            return current
        previous = current
        current = next_size
    compact["compact_pack_chars"] = current
    compact["final_dify_input_chars"] = current
    return current


def _target_report_length_guidance(primary_source: list[dict[str, Any]], auxiliary_source: list[dict[str, Any]]) -> str:
    primary_count = len(primary_source)
    auxiliary_count = len(auxiliary_source)
    if primary_count == 1 and auxiliary_count == 0:
        material = primary_source[0]
        if _material_content_text_length(material) < 800 and not material.get("attachments"):
            return "800-1200字；正文短且无附件时完整覆盖事实，分析克制，不硬扩写。"
        if _is_attachment_led_primary_material(material):
            return "1500-2500字；正文短但附件为主时，按附件表格、产品/价格/操作步骤展开。"
    if primary_count >= 2 or auxiliary_count >= 2:
        return "3000-5000字；多主材料或2+n场景按主材料分段、辅助材料作背景，并允许更多表格。"
    return "1800-3000字；按公告要点、规则、附件细节、企业影响和建议补齐。"


def _evidence_budget_for_strategy(input_strategy: str, *, primary_count: int, auxiliary_count: int) -> dict[str, Any]:
    if input_strategy == "attachment_led":
        budget = {"primary_body_pct": 18, "primary_attachment_pct": 57, "auxiliary_pct": 10, "guidance_warnings_pct": 15}
    elif input_strategy == "table_heavy":
        budget = {"primary_body_pct": 18, "table_summary_pct": 52, "attachment_sections_pct": 20, "auxiliary_pct": 10}
    elif input_strategy == "safe_compact":
        budget = {"primary_body_and_rules_pct": 30, "primary_core_attachment_pct": 45, "auxiliary_pct": 10, "guidance_warnings_omitted_pct": 15}
    elif input_strategy == "light_compact":
        budget = {"primary_body_pct": 35, "primary_attachment_pct": 35, "auxiliary_pct": 15, "guidance_warnings_pct": 15}
    elif input_strategy == "staged_generation":
        budget = {
            "per_primary_material": True,
            "primary_material_count": primary_count,
            "auxiliary_background_only": True,
            "auxiliary_material_count": auxiliary_count,
        }
    else:
        budget = {"primary_body": "full", "primary_attachment": "full", "auxiliary": "summary_and_relevant_snippets", "compression": "none"}
    return budget


def _compact_attachment_for_dify(attachment: dict[str, Any], *, role: str = "primary", preserve_detail: bool = False) -> dict[str, Any]:
    table_summaries: list[dict[str, Any]] = []
    if role == "auxiliary":
        table_limit = 2 if attachment.get("core_attachment") else 1
        table_header_limit = 18
        table_key_column_limit = 12
        table_summary_limit = 300
        table_business_limit = 180
    else:
        table_limit = 5 if attachment.get("core_attachment") else 3
        table_header_limit = 30
        table_key_column_limit = 20
        table_summary_limit = 500
        table_business_limit = 300
    if preserve_detail:
        table_limit = max(table_limit, len(list(attachment.get("table_summaries") or [])))
        table_header_limit = 80
        table_key_column_limit = 60
    for table in list(attachment.get("table_summaries") or [])[:table_limit]:
        if not isinstance(table, dict):
            continue
        summary = str(table.get("summary") or "")
        business_value = str(table.get("business_value") or "")
        table_summaries.append(
            {
                "sheet_name": table.get("sheet_name") or "",
                "rows": table.get("rows") or 0,
                "columns_count": table.get("columns_count") or 0,
                "headers": list(table.get("headers") or [])[:table_header_limit],
                "key_columns": list(table.get("key_columns") or [])[:table_key_column_limit],
                "enterprise_columns": list(table.get("enterprise_columns") or [])[:table_key_column_limit],
                "product_columns": list(table.get("product_columns") or [])[:table_key_column_limit],
                "registration_cert_columns": list(table.get("registration_cert_columns") or [])[:table_key_column_limit],
                "medical_insurance_code_columns": list(table.get("medical_insurance_code_columns") or [])[:table_key_column_limit],
                "price_columns": list(table.get("price_columns") or [])[:table_key_column_limit],
                "purchase_volume_columns": list(table.get("purchase_volume_columns") or [])[:table_key_column_limit],
                "selected_status_columns": list(table.get("selected_status_columns") or [])[:table_key_column_limit],
                "region_columns": list(table.get("region_columns") or [])[:table_key_column_limit],
                "group_columns": list(table.get("group_columns") or [])[:table_key_column_limit],
                "contains_enterprise": bool(table.get("contains_enterprise")),
                "contains_product": bool(table.get("contains_product")),
                "contains_registration_cert": bool(table.get("contains_registration_cert")),
                "contains_medical_insurance_code": bool(table.get("contains_medical_insurance_code")),
                "contains_specification": bool(table.get("contains_specification")),
                "contains_price": bool(table.get("contains_price")),
                "contains_selected_status": bool(table.get("contains_selected_status")),
                "contains_purchase_volume": bool(table.get("contains_purchase_volume")),
                "table_heavy": bool(table.get("table_heavy") or _table_summary_is_heavy(table)),
                "field_stats": copy.deepcopy(table.get("field_stats") or {}),
                "evidence_value_score": int(table.get("evidence_value_score") or 0),
                "data_completeness_hint": table.get("data_completeness_hint") or "",
                "recommended_report_usage": table.get("recommended_report_usage") or (
                    "Use this structured table summary to analyze product scope, enterprise scope, price or purchase-volume fields; do not invent row-level values that are not provided."
                ),
                "summary": summary if preserve_detail else _limit_text(summary, table_summary_limit),
                "business_value": business_value if preserve_detail else _limit_text(business_value, table_business_limit),
            }
        )
    if role == "auxiliary":
        summary_limit = 500
        key_fact_limit = 8
        section_limit = 4
        section_chars = 220
    else:
        summary_limit = 1400 if attachment.get("core_attachment") else 700
        key_fact_limit = 18
        section_limit = 8
        section_chars = 360
    if preserve_detail:
        summary_limit = max(summary_limit, len(str(attachment.get("summary") or "")))
        key_fact_limit = max(key_fact_limit, len(list(attachment.get("key_facts") or [])))
        section_limit = max(section_limit, len(list(attachment.get("important_sections") or [])))
        section_chars = 1200
    compact = {
        "articleattid": attachment.get("articleattid") or "",
        "filename": attachment.get("filename") or "",
        "fileext": attachment.get("fileext") or "",
        "filesize": attachment.get("filesize") or 0,
        "core_attachment": bool(attachment.get("core_attachment")),
        "business_type": attachment.get("business_type") or "普通附件",
        "core_attachment_unavailable": bool(attachment.get("core_attachment_unavailable")),
        "parse_status": attachment.get("parse_status") or "metadata_only",
        "download_status": attachment.get("download_status") or "",
        "cache_status": attachment.get("cache_status") or "",
        "text_length": attachment.get("text_length") or 0,
        "summary": str(attachment.get("summary") or "") if preserve_detail else _limit_text(attachment.get("summary"), summary_limit),
        "key_facts": list(attachment.get("key_facts") or [])[:key_fact_limit],
        "important_sections": [
            (item if preserve_detail else _limit_text(item, section_chars))
            for item in list(attachment.get("important_sections") or [])[:section_limit]
        ],
        "table_summaries": table_summaries,
    }
    compact["evidence_value_score"] = _evidence_value_score_attachment(attachment, role=role)
    return compact


def _compact_material_for_dify(
    material: dict[str, Any],
    role: str,
    *,
    content_limit: int | None = None,
    attachment_limit: int | None = None,
    preserve_detail: bool = False,
) -> dict[str, Any]:
    content_limit = content_limit if content_limit is not None else (5200 if role == "primary" else 1200)
    attachment_limit = attachment_limit if attachment_limit is not None else (8 if role == "primary" else 5)
    if preserve_detail:
        content_limit = max(content_limit, len(str(material.get("content_text") or "")))
        attachment_limit = max(attachment_limit, len(list(material.get("attachments") or [])))
    attachments = sorted(
        [item for item in list(material.get("attachments") or [])[: (100 if preserve_detail else 12)] if isinstance(item, dict)],
        key=lambda item: (not bool(item.get("core_attachment")), int(item.get("sortnum") or 0), str(item.get("articleattid") or "")),
    )
    summary_limit = max(500, len(str(material.get("summary") or ""))) if preserve_detail else 500
    content_summary_limit = max(700, len(str(material.get("content_summary") or ""))) if preserve_detail else 700
    key_fact_limit = max(16, len(list(material.get("key_facts") or []))) if preserve_detail else 16
    compact = {
        "material_role": material.get("material_role") or role,
        "evidence_source_type": "attachment_led_primary"
        if role == "primary" and _is_attachment_led_primary_material(material)
        else ("primary_body" if role == "primary" else "auxiliary_reference"),
        "menu_code": material.get("menu_code") or "",
        "articleid": material.get("articleid") or "",
        "title": material.get("title") or "",
        "audittime": material.get("audittime") or "",
        "updatetime": material.get("updatetime") or "",
        "menu_name": material.get("menu_name") or "",
        "source": material.get("source") or "",
        "sourceurl": material.get("sourceurl") or "",
        "areaname": material.get("areaname") or "",
        "publicorg": material.get("publicorg") or "",
        "projectphase": material.get("projectphase") or "",
        "projecttype": material.get("projecttype") or "",
        "dl_project_type": material.get("dl_project_type") or "",
        "category": material.get("category") or "",
        "referencenumber": material.get("referencenumber") or "",
        "policytype": material.get("policytype") or "",
        "belongproject": material.get("belongproject") or "",
        "projectabbreviation": material.get("projectabbreviation") or "",
        "summary": _limit_text(material.get("summary"), summary_limit),
        "content_text": str(material.get("content_text") or "")
        if len(str(material.get("content_text") or "")) <= content_limit
        else _limit_text(material.get("content_text"), content_limit),
        "content_text_length": material.get("content_text_length") or len(str(material.get("content_text") or "")),
        "content_summary": _limit_text(material.get("content_summary"), content_summary_limit),
        "key_facts": list(material.get("key_facts") or [])[:key_fact_limit],
        "attachments": [_compact_attachment_for_dify(item, role=role, preserve_detail=preserve_detail) for item in attachments[:attachment_limit]],
        "warnings": [_limit_text(item, 220) for item in list(material.get("warnings") or [])[:6]],
        "evidence_value_score": _evidence_value_score_material(material, role=role),
    }
    if role == "primary":
        compact.update(
            {
                "important_passages": list(material.get("important_passages") or [])[:8],
                "policy_rules": [_limit_text(item, 300) for item in list(material.get("policy_rules") or [])[:6]],
                "price_rules": [_limit_text(item, 300) for item in list(material.get("price_rules") or [])[:6]],
                "time_requirements": [_limit_text(item, 300) for item in list(material.get("time_requirements") or [])[:6]],
                "product_scope": [_limit_text(item, 300) for item in list(material.get("product_scope") or [])[:6]],
                "enterprise_requirements": [_limit_text(item, 300) for item in list(material.get("enterprise_requirements") or [])[:6]],
                "execution_requirements": [_limit_text(item, 300) for item in list(material.get("execution_requirements") or [])[:6]],
                "attachment_summaries": list(material.get("attachment_summaries") or [])[:8],
            }
        )
    else:
        compact.update(
            {
                "relation_to_primary": material.get("relation_to_primary") or "",
                "relevance_score": material.get("relevance_score") or 0,
                "usable_points": list(material.get("usable_points") or [])[:8],
                "relevant_snippets": list(material.get("relevant_snippets") or [])[:6],
                "attachment_summaries": [
                    _compact_attachment_for_dify(item, role="auxiliary")
                    for item in list(material.get("attachment_summaries") or [])[:3]
                    if isinstance(item, dict)
                ],
            }
        )
    return compact


def _compact_primary_content_limit(primary_source: list[dict[str, Any]], auxiliary_source: list[dict[str, Any]], original_pack_chars: int, max_chars: int) -> int:
    primary_count = max(1, len(primary_source))
    if primary_count == 1 and original_pack_chars <= max_chars * 1.1:
        return 20000
    if primary_count == 1:
        return 16000
    if primary_count == 2:
        return 10000
    return 7600


def _compact_evidence_pack_for_dify(pack: dict[str, Any], max_chars: int | None = None) -> dict[str, Any]:
    original_pack_chars = len(json.dumps(pack, ensure_ascii=False, sort_keys=True))
    primary_source = [item for item in list(pack.get("primary_materials") or [])[:3] if isinstance(item, dict)]
    auxiliary_source = [item for item in list(pack.get("auxiliary_materials") or [])[:10] if isinstance(item, dict)]
    thresholds = _dify_input_thresholds()
    hard_limit = thresholds["hard_limit"]
    target_max_chars = max_chars or hard_limit
    dify_relevant_input_chars = _dify_relevant_input_chars(pack, primary_source, auxiliary_source)
    input_strategy = _dify_input_strategy(
        original_pack_chars=original_pack_chars,
        dify_relevant_input_chars=dify_relevant_input_chars,
        primary_source=primary_source,
        auxiliary_source=auxiliary_source,
    )
    preserve_detail = input_strategy in {"full_input", "attachment_led"}
    primary_content_limit = _compact_primary_content_limit(primary_source, auxiliary_source, dify_relevant_input_chars, target_max_chars)
    if input_strategy == "safe_compact":
        primary_content_limit = min(primary_content_limit, 12000)
    elif input_strategy == "table_heavy":
        primary_content_limit = min(primary_content_limit, 9000)
    primary = [
        _compact_material_for_dify(
            item,
            "primary",
            content_limit=primary_content_limit,
            attachment_limit=10 if len(primary_source) == 1 else 8,
            preserve_detail=preserve_detail,
        )
        for item in primary_source
    ]
    auxiliary_content_limit = 900
    auxiliary_attachment_limit = 4
    if input_strategy == "safe_compact":
        auxiliary_content_limit = 600
        auxiliary_attachment_limit = 2
    elif input_strategy in {"attachment_led", "table_heavy"}:
        auxiliary_content_limit = 700
        auxiliary_attachment_limit = 2
    auxiliary = [
        _compact_material_for_dify(item, "auxiliary", content_limit=auxiliary_content_limit, attachment_limit=auxiliary_attachment_limit, preserve_detail=False)
        for item in auxiliary_source
    ]
    omitted_content: list[dict[str, Any]] = []
    compression_reason: list[str] = []
    primary_detail_preserved = True
    core_attachment_detail_preserved = True
    auxiliary_detail_preserved = True
    full_row_level_detail_preserved = True
    if input_strategy != "full_input":
        omitted_content.append(
            _omitted_content_entry(
                "attachment_full_text",
                reason="附件全文和 Excel 全量行数据不进入 Dify，仅保留摘要、关键事实、字段统计和结构化表格信息。",
                risk="模型不能逐行核验未输入的附件明细；涉及产品级核验时需人工打开附件。",
                manual_review=input_strategy in {"safe_compact", "table_heavy", "staged_generation"},
                affects_core_attachment_detail=False,
            )
        )
    if input_strategy == "light_compact":
        compression_reason.append("light_compact_duplicate_noise_removed")
    if input_strategy == "safe_compact":
        compression_reason.append("safe_compact_primary_core_first")
        compression_reason.append("auxiliary_content_trimmed")
        auxiliary_detail_preserved = False
        omitted_content.append(
            _omitted_content_entry(
                "auxiliary_content_tail",
                reason="输入接近 Dify 变量限制，辅助材料只保留摘要、相关片段和相关性分数。",
                risk="辅助材料细节可能未完整进入 Dify，但主材料和核心附件优先保留。",
                manual_review=False,
                affects_auxiliary_detail=True,
            )
        )
    if input_strategy == "table_heavy":
        compression_reason.append("table_full_rows_not_included")
        full_row_level_detail_preserved = False
        omitted_content.append(
            _omitted_content_entry(
                "excel_full_rows",
                reason="表格行数较多，未将全量行输入 Dify，仅保留字段结构、统计摘要和样例值。",
                risk="模型无法逐行核验全部产品明细，如需精确产品级核验应人工打开附件。",
                manual_review=True,
                affects_core_attachment_detail=False,
            )
        )
    if any(len([item for item in list(material.get("attachments") or []) if isinstance(item, dict)]) > 5 for material in list(pack.get("auxiliary_materials") or []) if isinstance(material, dict)):
        auxiliary_detail_preserved = False
        omitted_content.append(
            _omitted_content_entry(
                "auxiliary_attachment_overflow",
                reason="辅助材料附件较多，Dify 精简版仅保留核心或排序靠前的结构化摘要，避免工作流超时。",
                risk="辅助材料附件细节不会作为主材料结论依据。",
                manual_review=False,
                affects_auxiliary_detail=True,
            )
        )
    if input_strategy == "staged_generation":
        compression_reason.append("staged_generation_complexity_compact")
    unavailable_core_attachments = [
        attachment
        for material in [*primary_source, *auxiliary_source]
        for attachment in material.get("attachments") or []
        if isinstance(attachment, dict) and attachment.get("core_attachment") and not _attachment_is_usable_for_dify(attachment)
    ]
    for attachment in unavailable_core_attachments:
        core_attachment_detail_preserved = False
        omitted_content.append(
            _omitted_content_entry(
                "core_attachment_unavailable",
                filename=str(attachment.get("filename") or ""),
                reason="核心附件下载或解析失败，仅保留元数据和失败状态。",
                risk="报告不能展开该附件中的产品、企业、价格或中选结果明细。",
                manual_review=True,
                affects_core_attachment_detail=True,
            )
        )
    compact: dict[str, Any] = {
        "pack_variant": "dify_compact",
        "input_strategy": input_strategy,
        "input_strategy_type": _dify_input_strategy_type(input_strategy),
        "input_strategy_description": _dify_input_strategy_description(input_strategy),
        "strategy_basis": _dify_input_strategy_basis(
            input_strategy,
            relevant_chars=dify_relevant_input_chars,
            thresholds=thresholds,
            primary_source=primary_source,
        ),
        "attachment_led": input_strategy == "attachment_led",
        "table_heavy": input_strategy == "table_heavy",
        "pack_id": pack.get("pack_id") or "",
        "created_at": pack.get("created_at") or "",
        "pack_version": pack.get("pack_version") or "2.0",
        "source": pack.get("source") or "database_selection",
        "primary_materials": primary,
        "auxiliary_materials": auxiliary,
        "combined_key_facts": list(pack.get("combined_key_facts") or [])[:30],
        "report_focus": list(pack.get("report_focus") or [])[:12],
        "warnings": [_limit_text(item, 260) for item in list(pack.get("warnings") or [])[:8]],
        "generation_guidance": copy.deepcopy(pack.get("generation_guidance") or {}),
    }
    compact["evidence_budget"] = _evidence_budget_for_strategy(input_strategy, primary_count=len(primary_source), auxiliary_count=len(auxiliary_source))
    compact["generation_guidance"]["target_report_length"] = _target_report_length_guidance(primary_source, auxiliary_source)
    compact["generation_guidance"]["mandatory_sections"] = [
        "导语",
        "公告要点",
        "核心规则拆解表",
        "时间节点/操作步骤",
        "附件表格与产品价格细节",
        "企业影响",
        "风险提示",
        "企业操作建议",
    ]
    if input_strategy == "full_input":
        compact["generation_guidance"]["detail_policy"] = "small_input_no_compression"
        compact["generation_guidance"]["generation_mode"] = "single_pass_full_evidence"
        compact["generation_guidance"]["detail_instruction"] = (
            "输入证据量未超过安全阈值，主材料正文和主附件解析结果已尽量完整保留；"
            "不要因为 evidence_pack 有 compact 标记就写得过短；"
            "报告篇幅应来自事实覆盖和证据支撑型分析，不得无依据扩写。"
        )
    elif input_strategy == "attachment_led":
        compact["generation_guidance"]["detail_policy"] = "attachment_led_full_core_attachment"
        compact["generation_guidance"]["generation_mode"] = "single_pass_attachment_led"
        compact["generation_guidance"]["detail_instruction"] = (
            "主材料正文较短，核心信息在主材料附件；应以核心附件 summary、key_facts、important_sections、table_summaries 为主体依据，"
            "展开产品范围、企业、价格、时间节点、执行规则或中选结果；不得因为正文短而只生成简短摘要。"
        )
        compact["generation_guidance"]["target_report_length"] = "1500-2500字；正文短但附件为主时，按附件表格、产品/价格/操作步骤展开。"
    elif input_strategy == "table_heavy":
        compact["generation_guidance"]["detail_policy"] = "table_heavy_structured_summary"
        compact["generation_guidance"]["generation_mode"] = "single_pass_large_table_structured_summary"
        compact["generation_guidance"]["detail_instruction"] = (
            "附件为大表格或关键字段表格，系统未输入全量行；已输入字段结构、关键列、统计摘要和样例值。"
            "报告应基于这些结构化信息分析附件价值，不得编造未提供的具体产品、企业、价格。"
        )
    elif input_strategy == "staged_generation":
        compact["generation_guidance"]["detail_policy"] = "large_input_staged_compact"
        compact["generation_guidance"]["generation_mode"] = "per_material_then_synthesis_staged_generation"
        compact["generation_guidance"]["detail_instruction"] = (
            "证据包较大或多主材料多附件，先分别梳理每个主材料和核心附件规则，再综合成报告；"
            "辅助材料只作背景，不能覆盖主材料结论。"
        )
    elif input_strategy == "safe_compact":
        compact["generation_guidance"]["detail_policy"] = "safe_compact_primary_core_first"
        compact["generation_guidance"]["generation_mode"] = "single_pass_safe_compact"
        compact["generation_guidance"]["detail_instruction"] = (
            "输入接近变量限制，系统优先保留主材料正文、关键规则、核心附件摘要和 table_summaries；"
            "辅助材料只作背景，不得把辅助材料当作主材料事实。"
        )
    else:
        compact["generation_guidance"]["detail_policy"] = "medium_input_light_compact"
        compact["generation_guidance"]["generation_mode"] = "single_pass_light_compact"
        compact["generation_guidance"]["detail_instruction"] = (
            "只清理重复模板和低价值噪声，保留核心规则、操作步骤、表格列名、产品范围和价格字段。"
        )
    compact["generation_guidance"]["attachment_reporting_rule"] = (
        "附件未解析、仅元数据、下载失败、解析失败等材料完整性提示不得混入正式分析段落；"
        "材料完整性限制只写入 generation_warnings，不写入 report_markdown；"
        "如核心附件不可用，应避免编造企业、产品、价格、中选结果。"
    )
    compact["generation_guidance"]["attachment_led_primary_rule"] = (
        "当主材料正文较短但 evidence_source_type=attachment_led_primary，且核心附件已解析为摘要或表格结构时，"
        "应以该主材料的核心附件摘要、key_facts、important_sections、table_summaries 作为主体依据展开报告。"
    )
    compact["primary_evidence"] = {
        "materials": [
            {
                "menu_code": item.get("menu_code"),
                "articleid": item.get("articleid"),
                "title": item.get("title"),
                "content_summary": item.get("content_summary"),
            }
            for item in primary
        ],
        "key_facts": list((pack.get("primary_evidence") or {}).get("key_facts") or [])[:30],
        "important_passages": list((pack.get("primary_evidence") or {}).get("important_passages") or [])[:12],
        "attachment_summaries": list((pack.get("primary_evidence") or {}).get("attachment_summaries") or [])[:8],
    }
    compact["auxiliary_evidence"] = {
        "summaries": list((pack.get("auxiliary_evidence") or {}).get("summaries") or [])[:8],
        "relevant_snippets": list((pack.get("auxiliary_evidence") or {}).get("relevant_snippets") or [])[:10],
        "comparison_points": list((pack.get("auxiliary_evidence") or {}).get("comparison_points") or [])[:10],
    }
    compact["attachment_evidence"] = {
        "parsed_summaries": [
            _compact_attachment_for_dify(item, role="auxiliary")
            for item in sorted(
                [entry for entry in list((pack.get("attachment_evidence") or {}).get("parsed_summaries") or []) if isinstance(entry, dict)],
                key=lambda entry: (not bool(entry.get("core_attachment")), str(entry.get("articleattid") or ""), str(entry.get("filename") or "")),
            )[:6]
        ],
        "table_summaries": [
            {
                "sheet_name": table.get("sheet_name") or "",
                "rows": table.get("rows") or 0,
                "columns_count": table.get("columns_count") or 0,
                "headers": list(table.get("headers") or [])[:18],
                "key_columns": list(table.get("key_columns") or [])[:12],
                "summary": _limit_text(table.get("summary"), 300),
                "business_value": _limit_text(table.get("business_value"), 180),
            }
            for table in list((pack.get("attachment_evidence") or {}).get("table_summaries") or [])[:5]
            if isinstance(table, dict)
        ],
        "warnings": [],
    }
    compact["dify_compacted"] = True
    compact["compression_applied"] = input_strategy in {"light_compact", "safe_compact", "staged_generation"}
    compact["compression_reason"] = list(dict.fromkeys(compression_reason))
    compact["primary_detail_preserved"] = primary_detail_preserved
    compact["core_attachment_detail_preserved"] = core_attachment_detail_preserved
    compact["auxiliary_detail_preserved"] = auxiliary_detail_preserved
    compact["full_row_level_detail_preserved"] = full_row_level_detail_preserved
    compact["detail_preserved"] = primary_detail_preserved and core_attachment_detail_preserved
    compact["original_pack_chars"] = original_pack_chars
    compact["strategy_basis_chars"] = dify_relevant_input_chars
    compact["dify_relevant_input_chars"] = dify_relevant_input_chars
    compact["hard_limit_chars"] = hard_limit
    compact["full_input_max_chars"] = thresholds["full_input"]
    compact["safe_compact_max_chars"] = thresholds["safe_compact"]
    compact["strategy_thresholds"] = {
        "hard_limit_chars": hard_limit,
        "full_input_max_chars": thresholds["full_input"],
        "safe_compact_max_chars": thresholds["safe_compact"],
    }
    _refresh_dify_char_fields(compact)
    compact["primary_evidence_score"] = sum(int(item.get("evidence_value_score") or 0) for item in primary)
    compact["auxiliary_evidence_score"] = sum(int(item.get("evidence_value_score") or 0) for item in auxiliary)
    compact["attachment_evidence_score"] = sum(
        int(attachment.get("evidence_value_score") or 0)
        for material in [*primary, *auxiliary]
        for attachment in material.get("attachments") or []
        if isinstance(attachment, dict)
    )
    compact["table_evidence_score"] = sum(
        int(table.get("evidence_value_score") or 0)
        for material in [*primary, *auxiliary]
        for attachment in material.get("attachments") or []
        if isinstance(attachment, dict)
        for table in attachment.get("table_summaries") or []
        if isinstance(table, dict)
    )
    compact["compression_strategy"] = [
        "小输入不压缩正文、附件摘要和附件表格摘要",
        "中等输入只清理重复公告模板、免责声明和低价值技术 warning",
        "接近 Dify 限制时优先保留主材料和核心附件，压缩辅助材料",
        "大表格不输入全量行，保留字段结构、字段统计、样例值和报告使用建议",
        "大输入或复杂多材料先按材料保留核心规则，再综合生成",
        "保留主材料正文、关键事实、核心附件摘要、表格列名、产品/价格字段和操作步骤",
        "辅助材料优先保留相关片段和可用要点，不能反客为主",
    ]
    compact["omitted_content"] = omitted_content
    compact["dify_compaction_note"] = "Structured compact evidence pack for Dify variable-size limits. Use ?full=true for full evidence pack."

    secondary_blocks_removed = False
    while len(json.dumps(compact, ensure_ascii=False, sort_keys=True)) > target_max_chars:
        changed = False
        if not secondary_blocks_removed and any(key in compact for key in ["primary_evidence", "auxiliary_evidence", "attachment_evidence"]):
            compact.pop("primary_evidence", None)
            compact.pop("auxiliary_evidence", None)
            compact.pop("attachment_evidence", None)
            compression_reason.append("secondary_evidence_blocks_removed")
            compact["omitted_content"].append(
                _omitted_content_entry(
                    "secondary_evidence_blocks",
                    reason="超过 Dify 变量长度限制，已移除重复的二级证据汇总块。",
                    risk="不影响主材料正文和主附件主体信息，因为 primary_materials 中仍保留核心证据。",
                    manual_review=False,
                )
            )
            secondary_blocks_removed = True
            compact["compression_applied"] = True
            compact["compression_reason"] = list(dict.fromkeys(compression_reason))
            _refresh_dify_char_fields(compact)
            continue
        for material in compact.get("primary_materials", []):
            text = str(material.get("content_text") or "")
            if len(text) > 2500:
                material["content_text"] = _limit_text(text, max(2500, len(text) - 1000))
                primary_detail_preserved = False
                compression_reason.append("primary_content_trimmed")
                compact["omitted_content"].append(
                    _omitted_content_entry(
                        "primary_content_tail",
                        reason="主材料正文过长，已保留前部核心正文和结构化要点。",
                        risk="主材料尾部细节可能未完整进入 Dify，报告应优先依据保留正文、关键事实和结构化规则。",
                        manual_review=True,
                        affects_primary_detail=True,
                    )
                )
                changed = True
            for attachment in material.get("attachments") or []:
                summary = str(attachment.get("summary") or "")
                floor = 900 if attachment.get("core_attachment") else 400
                if len(summary) > floor:
                    attachment["summary"] = _limit_text(summary, floor)
                    compression_reason.append("attachment_summary_trimmed")
                    affects_core = bool(attachment.get("core_attachment"))
                    if affects_core:
                        core_attachment_detail_preserved = False
                    compact["omitted_content"].append(
                        _omitted_content_entry(
                            "attachment_summary_tail",
                            reason="附件摘要过长，已截断摘要尾部，但保留 table_summaries、key_facts 和字段统计。",
                            risk="核心附件细节可能不完整，报告应优先依据结构化摘要和字段统计，避免编造行级明细。",
                            manual_review=affects_core,
                            affects_core_attachment_detail=affects_core,
                        )
                    )
                    changed = True
        for material in compact.get("auxiliary_materials", []):
            text = str(material.get("content_text") or "")
            if len(text) > 600:
                material["content_text"] = _limit_text(text, 600)
                auxiliary_detail_preserved = False
                compression_reason.append("auxiliary_content_trimmed")
                compact["omitted_content"].append(
                    _omitted_content_entry(
                        "auxiliary_content_tail",
                        reason="辅助材料正文过长，优先保留相关片段和摘要。",
                        risk="辅助材料细节可能未完整进入 Dify，不能作为主材料结论依据。",
                        manual_review=False,
                        affects_auxiliary_detail=True,
                    )
                )
                changed = True
        if changed:
            compact["compression_applied"] = True
            compact["compression_reason"] = list(dict.fromkeys(compression_reason))
            compact["primary_detail_preserved"] = primary_detail_preserved
            compact["core_attachment_detail_preserved"] = core_attachment_detail_preserved
            compact["auxiliary_detail_preserved"] = auxiliary_detail_preserved
            compact["detail_preserved"] = primary_detail_preserved and core_attachment_detail_preserved
            _refresh_dify_char_fields(compact)
            continue
        compact["omitted_content"].append(
            _omitted_content_entry(
                "hard_limit_residual_overflow",
                reason="证据包已执行所有保守压缩动作后仍接近或超过 Dify 变量限制，应走分段生成或人工复核。",
                risk="最终输入可能仍接近 Dify 上限，报告质量需人工复核。",
                manual_review=True,
                affects_primary_detail=not primary_detail_preserved,
                affects_core_attachment_detail=not core_attachment_detail_preserved,
                affects_auxiliary_detail=not auxiliary_detail_preserved,
            )
        )
        compression_reason.append("final_compact_pack_exceeded_target_max_chars")
        compact["compression_applied"] = True
        compact["compression_reason"] = list(dict.fromkeys(compression_reason))
        break
    compact["primary_detail_preserved"] = primary_detail_preserved
    compact["core_attachment_detail_preserved"] = core_attachment_detail_preserved
    compact["auxiliary_detail_preserved"] = auxiliary_detail_preserved
    compact["full_row_level_detail_preserved"] = full_row_level_detail_preserved
    compact["detail_preserved"] = primary_detail_preserved and core_attachment_detail_preserved
    compact["compression_reason"] = list(dict.fromkeys(compression_reason))
    _refresh_dify_char_fields(compact)
    compact["compression_ratio"] = round(original_pack_chars / compact["final_dify_input_chars"], 4) if compact["final_dify_input_chars"] else None
    _refresh_dify_char_fields(compact)
    compact["compression_ratio"] = round(original_pack_chars / compact["final_dify_input_chars"], 4) if compact["final_dify_input_chars"] else None
    return compact


def _analysis_run_dir() -> Path:
    configured = (os.getenv("ANALYSIS_RUN_DIR") or "").strip()
    if configured:
        return Path(configured)
    return _database_evidence_pack_dir().parent / "analysis_runs"


def _safe_run_id(run_id: str) -> str:
    if not re.fullmatch(r"run_[A-Za-z0-9_-]{8,80}", run_id or ""):
        raise HTTPException(status_code=404, detail="analysis run not found")
    return run_id


def _analysis_run_path(run_id: str) -> Path:
    return _analysis_run_dir() / f"{_safe_run_id(run_id)}.json"


def _make_analysis_run_id(pack_id: str) -> str:
    digest = hashlib.sha256(f"{pack_id}:{datetime.now().isoformat()}:{uuid.uuid4().hex}".encode("utf-8")).hexdigest()[:10]
    return f"run_{datetime.now().strftime('%Y%m%d')}_{digest}"


def _write_analysis_run(record: dict[str, Any]) -> None:
    run_id = str(record.get("run_id") or "")
    path = _analysis_run_path(run_id)
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(record, ensure_ascii=False, indent=2), encoding="utf-8")
    logger.info(
        "analysis_run_saved run_id=%s pack_id=%s status=%s path=%s",
        run_id,
        record.get("pack_id") or "",
        record.get("status") or "",
        path,
    )


def _read_analysis_run(run_id: str) -> dict[str, Any]:
    path = _analysis_run_path(run_id)
    if not path.exists():
        raise HTTPException(status_code=404, detail="analysis run not found")
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except Exception as exc:  # noqa: BLE001
        logger.warning("analysis_run_read_failed run_id=%s error_type=%s", run_id, exc.__class__.__name__)
        raise HTTPException(status_code=500, detail="analysis run read failed") from exc


class DifyWorkflowError(Exception):
    def __init__(self, code: str, message: str, detail: str = "", status_code: int = 502):
        super().__init__(message)
        self.code = code
        self.message = message
        self.detail = detail
        self.status_code = status_code


def _dify_config() -> dict[str, Any]:
    base_url = (os.getenv("DIFY_BASE_URL") or "").strip().rstrip("/")
    api_key = (os.getenv("DIFY_WORKFLOW_API_KEY") or "").strip()
    endpoint = (os.getenv("DIFY_REPORT_WORKFLOW_ENDPOINT") or "/workflows/run").strip()
    response_mode = (os.getenv("DIFY_RESPONSE_MODE") or "blocking").strip() or "blocking"
    user = (os.getenv("DIFY_USER") or "analysis_frontend").strip() or "analysis_frontend"
    timeout_seconds = _env_float("DIFY_TIMEOUT_SECONDS", 600.0)
    max_attempts = max(1, _env_int("DIFY_WORKFLOW_MAX_ATTEMPTS", 3))
    retry_backoff_seconds = max(0.0, _env_float("DIFY_WORKFLOW_RETRY_BACKOFF_SECONDS", 2.0))
    staged_timeout_seconds = max(30.0, _env_float("DIFY_STAGED_TIMEOUT_SECONDS", 300.0))
    staged_max_attempts = max(1, _env_int("DIFY_STAGED_MAX_ATTEMPTS", 1))
    if not base_url or not api_key:
        logger.warning("dify_configuration_incomplete base_url_configured=%s api_key_configured=%s", bool(base_url), bool(api_key))
        raise DifyWorkflowError("DIFY_NOT_CONFIGURED", "Dify API 配置不完整", status_code=500)
    return {
        "base_url": base_url,
        "api_key": api_key,
        "endpoint": endpoint,
        "response_mode": response_mode,
        "user": user,
        "timeout_seconds": timeout_seconds,
        "max_attempts": max_attempts,
        "retry_backoff_seconds": retry_backoff_seconds,
        "staged_timeout_seconds": staged_timeout_seconds,
        "staged_max_attempts": staged_max_attempts,
    }


def _dify_workflow_url(config: dict[str, Any]) -> str:
    return f"{str(config['base_url']).rstrip('/')}/{str(config['endpoint']).lstrip('/')}"


def _dify_input_strategy_from_pack(pack: dict[str, Any] | None) -> str:
    if not isinstance(pack, dict):
        return ""
    guidance = pack.get("generation_guidance") if isinstance(pack.get("generation_guidance"), dict) else {}
    strategy = str(pack.get("input_strategy") or guidance.get("input_strategy") or "").strip()
    if strategy:
        return strategy
    primary_source = [item for item in pack.get("primary_materials") or [] if isinstance(item, dict)]
    auxiliary_source = [item for item in pack.get("auxiliary_materials") or [] if isinstance(item, dict)]
    evidence_chars = _dify_relevant_input_chars(pack, primary_source, auxiliary_source)
    return _dify_input_strategy(
        original_pack_chars=evidence_chars,
        dify_relevant_input_chars=evidence_chars,
        primary_source=primary_source,
        auxiliary_source=auxiliary_source,
    )


def _dify_request_policy(config: dict[str, Any], pack: dict[str, Any] | None = None) -> dict[str, Any]:
    strategy = _dify_input_strategy_from_pack(pack)
    timeout_seconds = float(config["timeout_seconds"])
    max_attempts = int(config.get("max_attempts") or 1)
    if strategy == "staged_generation":
        timeout_seconds = min(timeout_seconds, float(config.get("staged_timeout_seconds") or timeout_seconds))
        max_attempts = min(max_attempts, int(config.get("staged_max_attempts") or max_attempts))
    return {
        "input_strategy": strategy,
        "timeout_seconds": max(1.0, timeout_seconds),
        "max_attempts": max(1, max_attempts),
        "retry_backoff_seconds": float(config.get("retry_backoff_seconds") or 0),
    }


def _post_dify_workflow_once(url: str, config: dict[str, Any], payload: dict[str, Any], timeout_seconds: float) -> dict[str, Any]:
    with httpx.Client(timeout=timeout_seconds) as client:
        response = client.post(
            url,
            headers={
                "Authorization": f"Bearer {config['api_key']}",
                "Content-Type": "application/json",
            },
            json=payload,
        )
        response.raise_for_status()
        parsed = response.json()
        return parsed if isinstance(parsed, dict) else {}


def _post_dify_workflow_with_wall_timeout(
    url: str,
    config: dict[str, Any],
    payload: dict[str, Any],
    timeout_seconds: float,
) -> dict[str, Any]:
    holder: dict[str, Any] = {}

    def target() -> None:
        try:
            holder["value"] = _post_dify_workflow_once(url, config, payload, timeout_seconds)
        except BaseException as exc:  # noqa: BLE001
            holder["exc"] = exc

    worker = threading.Thread(target=target, name="dify-workflow-call", daemon=True)
    worker.start()
    worker.join(timeout_seconds)
    if worker.is_alive():
        raise httpx.TimeoutException(f"Dify workflow wall timeout after {timeout_seconds:.1f}s")
    if "exc" in holder:
        raise holder["exc"]
    value = holder.get("value")
    return value if isinstance(value, dict) else {}


def _analysis_watchdog_timeout_seconds(pack: dict[str, Any] | None) -> float:
    try:
        policy = _dify_request_policy(_dify_config(), pack)
    except DifyWorkflowError:
        return 0.0
    timeout_seconds = float(policy["timeout_seconds"])
    watchdog_max_seconds = _env_float("DIFY_WATCHDOG_MAX_SECONDS", 300.0)
    if watchdog_max_seconds > 0:
        timeout_seconds = min(timeout_seconds, watchdog_max_seconds)
    return max(1.0, timeout_seconds) + max(0.0, _env_float("DIFY_WATCHDOG_GRACE_SECONDS", 15.0))


def _parse_json_object(value: Any) -> dict[str, Any] | None:
    if isinstance(value, dict):
        return value
    if isinstance(value, str):
        text = value.strip()
        if not text:
            return None
        try:
            parsed = json.loads(text)
        except json.JSONDecodeError:
            return None
        return parsed if isinstance(parsed, dict) else None
    return None


def _parse_json_list(value: Any) -> list[Any]:
    if value is None:
        return []
    if isinstance(value, list):
        return value
    if isinstance(value, str):
        text = value.strip()
        if not text:
            return []
        try:
            parsed = json.loads(text)
        except json.JSONDecodeError:
            return [text]
        if isinstance(parsed, list):
            return parsed
        if parsed is None:
            return []
        return [parsed]
    return [value]


def _extract_dify_outputs(response_json: dict[str, Any]) -> tuple[str, dict[str, Any], str]:
    workflow_run_id = str(response_json.get("workflow_run_id") or "")
    data = response_json.get("data") if isinstance(response_json.get("data"), dict) else {}
    if not workflow_run_id:
        workflow_run_id = str(data.get("id") or data.get("workflow_run_id") or "")
    if str(data.get("status") or "").lower() in {"failed", "stopped"}:
        raise DifyWorkflowError("DIFY_WORKFLOW_FAILED", "Dify 工作流执行失败", str(data.get("error") or data.get("status") or ""))
    outputs = data.get("outputs") if isinstance(data.get("outputs"), dict) else response_json.get("outputs")
    output_dict = _parse_json_object(outputs) or {}
    candidates: list[dict[str, Any] | None] = [
        output_dict,
        _parse_json_object(output_dict.get("result")),
        _parse_json_object(output_dict.get("output")),
        _parse_json_object(output_dict.get("answer")),
        _parse_json_object(output_dict.get("text")),
        _parse_json_object(output_dict.get("report")),
    ]
    for value in output_dict.values():
        parsed = _parse_json_object(value)
        if parsed:
            candidates.append(parsed)
    for candidate in candidates:
        if candidate and candidate.get("report_markdown"):
            return workflow_run_id, candidate, str(data.get("status") or "")
    if str(data.get("status") or "").lower() in {"succeeded", "finished", ""}:
        return workflow_run_id, output_dict, str(data.get("status") or "")
    raise DifyWorkflowError("DIFY_INVALID_RESPONSE", "Dify 返回内容缺少 report_markdown", "outputs did not contain report_markdown")


def _normalize_dify_result(response_json: dict[str, Any], pack_id: str) -> dict[str, Any]:
    workflow_run_id, output, workflow_status = _extract_dify_outputs(response_json)
    quality_check = _parse_json_object(output.get("quality_check")) or {"passed": None, "issues": []}
    status = str(output.get("status") or "").strip() or "finished"
    if status not in {"finished", "needs_manual_review", "failed"}:
        status = "finished" if workflow_status.lower() in {"succeeded", "finished", ""} else workflow_status.lower()
    warnings = _parse_json_list(output.get("generation_warnings") or output.get("warnings"))
    return {
        "workflow_run_id": workflow_run_id,
        "status": status,
        "pack_id": str(output.get("pack_id") or pack_id),
        "report_title": str(output.get("report_title") or ""),
        "report_markdown": _clean_model_output(str(output.get("report_markdown") or "")),
        "version": int(output.get("version") or 1),
        "quality_check": quality_check,
        "generation_warnings": warnings,
        "warnings": warnings,
        "remaining_issues": _parse_json_list(output.get("remaining_issues")),
    }


def _is_unusable_report_markdown(markdown: Any) -> bool:
    text = str(markdown or "").strip()
    if not text:
        return True
    compact = re.sub(r"\s+", "", text)
    if compact in {"...", "…", "......", "。", "."}:
        return True
    return len(text) < 8


def _is_fragmentary_dify_report_markdown(markdown: Any, pack: dict[str, Any]) -> bool:
    text = str(markdown or "").strip()
    if _is_unusable_report_markdown(text):
        return True
    has_intro_or_first_section = bool(re.search(r"(^|\n)\s*(?:#{1,6}\s+)?(?:导语|一[、.．]|1[、.．])", text))
    has_report_structure = bool(re.search(r"(^|\n)\s*(?:#{1,6}\s+)?(?:导语|一[、.．]|1[、.．]|二[、.．])", text))
    starts_midstream = bool(re.match(r"^\s*(?:[-*]\s+|\|)", text))
    starts_after_first_section = bool(re.match(r"^\s*(?:#{1,6}\s+)?(?:二|三|四|五|六|七|八|九|十|2|3|4|5|6|7|8|9|10)[、.．]", text))
    if starts_after_first_section and not has_intro_or_first_section:
        return True
    if len(text) < 500 and (starts_midstream or not has_report_structure):
        return True
    diagnostics = build_pack_diagnostics(pack)
    weighted_chars = int(diagnostics.get("weighted_evidence_chars") or diagnostics.get("total_content_chars") or 0)
    if weighted_chars < 1000:
        return False
    if starts_midstream and not has_intro_or_first_section:
        return True
    if len(text) < 700 and (starts_midstream or not has_report_structure):
        return True
    return False


def _fallback_text_snippet(text: str, limit: int = 380) -> str:
    clean = _normalize_text(text)
    if not clean:
        return ""
    if len(clean) <= limit:
        return clean
    return f"{clean[:limit].rstrip()}..."


def _fallback_report_from_pack(pack: dict[str, Any]) -> tuple[str, str, list[str]]:
    primary = [item for item in pack.get("primary_materials") or [] if isinstance(item, dict)]
    auxiliary = [item for item in pack.get("auxiliary_materials") or [] if isinstance(item, dict)]
    main = primary[0] if primary else {}
    title = str(main.get("title") or "材料分析报告").strip()
    report_title = title if title.endswith("分析报告") else f"{title}分析报告"
    warnings = ["报告生成结果过短或疑似只返回修订片段，已根据证据包生成保守兜底报告，请人工复核。"]

    lines: list[str] = [
        "## 导语",
        f"本报告基于已选择的主材料《{title or '未命名材料'}》整理。"
        "以下内容按当前可读取材料进行保守梳理，重点呈现原文已披露的规则、执行要求和企业关注点。",
    ]
    meta_parts = []
    for label, field in [
        ("发布时间", "audittime"),
        ("地区", "areaname"),
        ("发布机构", "publicorg"),
        ("项目阶段", "projectphase"),
        ("项目类型", "projecttype"),
        ("分类", "category"),
    ]:
        value = _clean_inline_text(main.get(field) or "")
        if value:
            meta_parts.append(f"{label}：{value}")
    if meta_parts:
        lines.extend(["", "主材料基础信息：", *[f"- {item}" for item in meta_parts]])

    summary = _clean_inline_text(main.get("summary") or main.get("content_summary") or "")
    content_text = _normalize_text(main.get("content_text") or "")
    lines.extend(["", "## 一、核心内容梳理"])
    if summary:
        lines.append(f"材料摘要显示：{summary}")
    snippet = _fallback_text_snippet(content_text, 700)
    if snippet:
        lines.append(f"从正文可见，材料主要内容包括：{snippet}")
    else:
        lines.append("当前公告正文较简略，具体规则主要见随文附件；以下分析以可读取的附件摘要和结构化要点为主要依据。")

    facts = [item for item in main.get("key_facts") or [] if isinstance(item, dict)]
    if facts:
        lines.extend(["", "可直接引用的结构化事实包括："])
        for fact in facts[:12]:
            name = _clean_inline_text(fact.get("name") or "")
            value = _clean_inline_text(fact.get("value") or "")
            if name and value:
                lines.append(f"- {name}：{value}")

    passages = main.get("important_passages") if isinstance(main.get("important_passages"), list) else []
    if passages:
        lines.extend(["", "## 二、可关注规则与执行要点"])
        for item in passages[:6]:
            if isinstance(item, dict) and item.get("text"):
                lines.append(f"- {item.get('text')}")

    lines.extend(["", "## 三、附件要点补充"])
    attachments = [item for item in main.get("attachments") or [] if isinstance(item, dict)]
    core_unavailable = False
    attachment_detail_added = False
    if attachments:
        for attachment in attachments[:10]:
            filename = _clean_inline_text(attachment.get("filename") or "未命名附件")
            summary_text = _clean_inline_text(attachment.get("summary") or "")
            if summary_text:
                lines.append(f"- {filename}：{_fallback_text_snippet(summary_text, 900)}")
                attachment_detail_added = True
            key_facts = [str(item).strip() for item in attachment.get("key_facts") or [] if str(item).strip()]
            if key_facts:
                lines.append("  其中可关注事实包括：")
                lines.extend([f"  - {fact}" for fact in key_facts[:10]])
                attachment_detail_added = True
            sections = [item for item in attachment.get("important_sections") or [] if isinstance(item, dict)]
            if sections:
                lines.append("  附件分节信息显示：")
                for section in sections[:6]:
                    section_title = _clean_inline_text(section.get("title") or "相关章节")
                    section_summary = _fallback_text_snippet(str(section.get("summary") or ""), 220)
                    if section_summary:
                        lines.append(f"  - {section_title}：{section_summary}")
                        attachment_detail_added = True
            elif attachment.get("core_attachment"):
                core_unavailable = True
        if not attachment_detail_added:
            lines.append("随文附件未形成可用于展开分析的摘要，报告不展开产品、价格或项目明细。")
    else:
        lines.append("主材料未关联附件。")
    if auxiliary:
        lines.extend(["", "## 四、辅助材料参考"])
        lines.append("")
        lines.append("辅助材料仅用于背景补充和对照，不覆盖主材料结论：")
        for item in auxiliary[:5]:
            aux_title = _clean_inline_text(item.get("title") or "")
            aux_summary = _fallback_text_snippet(str(item.get("content_summary") or item.get("summary") or ""), 180)
            lines.append(f"- {aux_title}：{aux_summary or '暂无可用摘要'}")

    lines.extend(
        [
            "",
            "## 五、初步分析提示",
            "根据现有材料，企业侧应先核对公告标题、正文规则、附件摘要和表格字段是否覆盖自身产品或业务。"
            "若材料披露了执行时间、申报路径、产品范围、价格字段或机构责任，应据此形成内部核对清单；"
            "若材料未披露相应事实，报告不补充推断性结论。"
        ]
    )
    if core_unavailable:
        warnings.append("核心附件未形成可用摘要，报告未展开产品、企业、价格或中选结果明细。")
    return report_title, "\n\n".join(lines), warnings


def _repair_unusable_dify_result(result: dict[str, Any], pack: dict[str, Any]) -> dict[str, Any]:
    if not _is_fragmentary_dify_report_markdown(result.get("report_markdown"), pack):
        return result
    title, markdown, fallback_warnings = _fallback_report_from_pack(pack)
    warnings = [*list(result.get("warnings") or []), *list(result.get("generation_warnings") or []), *fallback_warnings]
    issue = {
        "issue_id": "Q_DIFY_FRAGMENTARY_REPORT",
        "severity": "high",
        "problem_type": "empty_or_fragmentary_report",
        "report_text": str(result.get("report_markdown") or ""),
        "source_basis": "evidence_pack",
        "fix_instruction": "Dify 返回正文过短、占位或疑似只包含修订片段，已启用后端兜底报告，仍需人工复核或重新运行工作流。",
    }
    repaired = dict(result)
    repaired.update(
        {
            "status": "needs_manual_review",
            "report_title": title if _is_unusable_report_markdown(result.get("report_title")) else result.get("report_title") or title,
            "report_markdown": markdown,
            "quality_check": {"passed": False, "issues": [issue]},
            "generation_warnings": warnings,
            "warnings": warnings,
            "remaining_issues": [issue],
        }
    )
    return repaired


def _apply_local_quality_gate_to_dify_result(result: dict[str, Any], pack: dict[str, Any]) -> dict[str, Any]:
    gated = dict(result)
    diagnostics = build_run_diagnostics(gated, pack, _compact_evidence_pack_for_dify(pack))
    quality_gate = diagnostics.get("quality_gate") if isinstance(diagnostics.get("quality_gate"), dict) else {}
    gated["quality_gate"] = quality_gate
    if quality_gate.get("deliverable_status") != "needs_manual_review":
        return gated

    issue = {
        "issue_id": "Q_LOCAL_QUALITY_GATE",
        "severity": "high",
        "problem_type": "local_quality_gate",
        "report_text": "",
        "source_basis": "evidence_pack",
        "fix_instruction": "本地质量门禁发现原文遵循、核心覆盖或分析深度问题，报告需人工复核后再交付。",
        "quality_gate": quality_gate,
    }
    quality_check = gated.get("quality_check") if isinstance(gated.get("quality_check"), dict) else {"passed": None, "issues": []}
    issues = list(quality_check.get("issues") or []) if isinstance(quality_check.get("issues"), list) else []
    if not any(isinstance(item, dict) and item.get("issue_id") == issue["issue_id"] for item in issues):
        issues.append(issue)
    quality_check = dict(quality_check)
    quality_check["passed"] = False
    quality_check["issues"] = issues
    remaining_issues = list(gated.get("remaining_issues") or [])
    if not any(isinstance(item, dict) and item.get("issue_id") == issue["issue_id"] for item in remaining_issues):
        remaining_issues.append(issue)
    warnings = list(gated.get("warnings") or gated.get("generation_warnings") or [])
    warning = "报告未通过原文遵循或分析深度门禁，已标记为需人工复核。"
    if warning not in warnings:
        warnings.append(warning)
    gated.update(
        {
            "status": "needs_manual_review",
            "quality_check": quality_check,
            "remaining_issues": remaining_issues,
            "warnings": warnings,
            "generation_warnings": warnings,
        }
    )
    return gated


def _fallback_result_from_dify_error(
    exc: DifyWorkflowError,
    pack: dict[str, Any],
    pack_id: str,
    *,
    apply_quality_gate: bool = True,
) -> dict[str, Any]:
    title, markdown, fallback_warnings = _fallback_report_from_pack(pack)
    issue = {
        "issue_id": "Q_DIFY_CALL_FAILED_FALLBACK",
        "severity": "high",
        "problem_type": exc.code,
        "report_text": "",
        "source_basis": "evidence_pack",
        "fix_instruction": "Dify 调用失败后已基于 evidence_pack 生成保守报告；该报告需人工复核，不得直接交付。",
        "error_detail": exc.detail,
    }
    warnings = [
        *fallback_warnings,
        f"{exc.message}，已启用保守兜底报告。",
    ]
    result = {
        "workflow_run_id": "",
        "status": "needs_manual_review",
        "pack_id": pack_id,
        "report_title": title,
        "report_markdown": markdown,
        "version": 1,
        "quality_check": {"passed": False, "issues": [issue]},
        "generation_warnings": warnings,
        "warnings": warnings,
        "remaining_issues": [issue],
        "dify_error_code": exc.code,
        "dify_error_message": exc.message,
        "dify_error_detail": exc.detail,
    }
    if apply_quality_gate:
        return _apply_local_quality_gate_to_dify_result(result, pack)
    result["quality_gate"] = {
        "deliverable_status": "needs_manual_review",
        "blocking_issue_codes": ["DIFY_TIMEOUT"],
        "summary_only_risk": False,
        "source_fidelity_score": 0,
        "analysis_depth_score": 0,
        "evidence_backed_analysis_count": 0,
    }
    return result


def _call_dify_workflow(pack_id: str, run_id: str, pack: dict[str, Any] | None = None) -> dict[str, Any]:
    config = _dify_config()
    policy = _dify_request_policy(config, pack)
    url = _dify_workflow_url(config)
    payload = {
        "inputs": {"pack_id": pack_id},
        "response_mode": config["response_mode"],
        "user": config["user"],
    }
    started = time.perf_counter()
    logger.info(
        "dify_workflow_call_started run_id=%s pack_id=%s url=%s input_strategy=%s timeout_seconds=%s max_attempts=%s",
        run_id,
        pack_id,
        url,
        policy.get("input_strategy") or "",
        policy["timeout_seconds"],
        policy["max_attempts"],
    )
    attempts = int(policy["max_attempts"])
    backoff = float(policy["retry_backoff_seconds"])
    response_json: dict[str, Any] | None = None
    for attempt in range(1, attempts + 1):
        try:
            response_json = _post_dify_workflow_with_wall_timeout(url, config, payload, float(policy["timeout_seconds"]))
            break
        except httpx.TimeoutException as exc:
            if attempt < attempts:
                logger.warning("dify_workflow_retry_timeout run_id=%s pack_id=%s attempt=%s/%s", run_id, pack_id, attempt, attempts)
                if backoff:
                    time.sleep(backoff * attempt)
                continue
            raise DifyWorkflowError("DIFY_TIMEOUT", "Dify 工作流调用超时", str(exc), status_code=504) from exc
        except httpx.HTTPStatusError as exc:
            status = exc.response.status_code
            detail = f"HTTP {status}: {_truncate(exc.response.text, 300)}"
            if status in {429, 502, 503, 504} and attempt < attempts:
                logger.warning("dify_workflow_retry_status run_id=%s pack_id=%s status=%s attempt=%s/%s", run_id, pack_id, status, attempt, attempts)
                if backoff:
                    time.sleep(backoff * attempt)
                continue
            if status in {429, 503}:
                raise DifyWorkflowError("DIFY_MODEL_BUSY", "Dify 上游模型服务繁忙", detail, status_code=502) from exc
            raise DifyWorkflowError("DIFY_CALL_FAILED", "调用 Dify 工作流失败", detail) from exc
        except httpx.HTTPError as exc:
            if attempt < attempts:
                logger.warning("dify_workflow_retry_http_error run_id=%s pack_id=%s attempt=%s/%s", run_id, pack_id, attempt, attempts)
                if backoff:
                    time.sleep(backoff * attempt)
                continue
            raise DifyWorkflowError("DIFY_CALL_FAILED", "调用 Dify 工作流失败", str(exc)) from exc
        except json.JSONDecodeError as exc:
            raise DifyWorkflowError("DIFY_INVALID_RESPONSE", "Dify 返回内容不是合法 JSON", str(exc)) from exc
    if response_json is None:
        raise DifyWorkflowError("DIFY_CALL_FAILED", "调用 Dify 工作流失败", "empty response")

    elapsed_ms = int((time.perf_counter() - started) * 1000)
    result = _normalize_dify_result(response_json, pack_id)
    logger.info(
        "dify_workflow_call_finished run_id=%s pack_id=%s workflow_run_id=%s status=%s elapsed_ms=%s",
        run_id,
        pack_id,
        result.get("workflow_run_id") or "",
        result.get("status") or "",
        elapsed_ms,
    )
    return result


def _make_revision_id(run_id: str) -> str:
    digest = hashlib.sha256(f"{run_id}:{datetime.now().isoformat()}:{uuid.uuid4().hex}".encode("utf-8")).hexdigest()[:10]
    return f"rev_{datetime.now().strftime('%Y%m%d')}_{digest}"


def _call_dify_revision_workflow(
    *,
    pack_id: str,
    run_id: str,
    feedback: str,
    current_report: str,
    evidence_pack: dict[str, Any],
    analysis_highlight: bool,
) -> dict[str, Any]:
    config = _dify_config()
    url = _dify_workflow_url(config)
    payload = {
        "inputs": {
            "mode": "user_feedback_revision",
            "pack_id": pack_id,
            "run_id": run_id,
            "feedback": feedback,
            "current_report": current_report,
            "evidence_pack": evidence_pack,
            "analysis_highlight": analysis_highlight,
        },
        "response_mode": config["response_mode"],
        "user": config["user"],
    }
    started = time.perf_counter()
    logger.info("dify_revision_call_started run_id=%s pack_id=%s feedback_chars=%s", run_id, pack_id, len(feedback))
    try:
        with httpx.Client(timeout=float(config["timeout_seconds"])) as client:
            response = client.post(
                url,
                headers={
                    "Authorization": f"Bearer {config['api_key']}",
                    "Content-Type": "application/json",
                },
                json=payload,
            )
            response.raise_for_status()
            response_json = response.json()
    except httpx.TimeoutException as exc:
        raise DifyWorkflowError("DIFY_TIMEOUT", "Dify 修订工作流调用超时", str(exc), status_code=504) from exc
    except httpx.HTTPStatusError as exc:
        detail = f"HTTP {exc.response.status_code}: {_truncate(exc.response.text, 300)}"
        raise DifyWorkflowError("DIFY_CALL_FAILED", "调用 Dify 修订工作流失败", detail) from exc
    except httpx.HTTPError as exc:
        raise DifyWorkflowError("DIFY_CALL_FAILED", "调用 Dify 修订工作流失败", str(exc)) from exc
    except json.JSONDecodeError as exc:
        raise DifyWorkflowError("DIFY_INVALID_RESPONSE", "Dify 修订返回内容不是合法 JSON", str(exc)) from exc
    result = _normalize_dify_result(response_json, pack_id)
    logger.info(
        "dify_revision_call_finished run_id=%s pack_id=%s status=%s elapsed_ms=%s",
        run_id,
        pack_id,
        result.get("status") or "",
        int((time.perf_counter() - started) * 1000),
    )
    return result


def _make_database_pack_id(primary_keys: list[tuple[str, str]], auxiliary_keys: list[tuple[str, str]], created_at: str) -> str:
    payload = {
        "primary": primary_keys,
        "auxiliary": auxiliary_keys,
        "created_at": created_at,
        "nonce": uuid.uuid4().hex,
    }
    digest = hashlib.sha256(json.dumps(payload, ensure_ascii=False, sort_keys=True).encode("utf-8")).hexdigest()[:10]
    return f"pack_{datetime.now().strftime('%Y%m%d')}_{digest}"


def _fetch_database_material_rows(keys: list[tuple[str, str]]) -> dict[tuple[str, str], dict[str, Any]]:
    if not keys:
        return {}
    placeholders = ", ".join(["(%s, %s)"] * len(keys))
    params: list[object] = []
    for menu_code, articleid in keys:
        params.extend([menu_code, articleid])
    fields = ", ".join(f"a.{field}" for field in ARTICLE_DETAIL_FIELDS)
    rows = _db_fetch_all(
        f"""
        SELECT {fields}
        FROM sample_article_wide a
        WHERE a.status = %s AND (a.menu_code, a.articleid) IN ({placeholders})
        """,
        [0, *params],
    )
    return {(str(row.get("menu_code") or ""), str(row.get("articleid") or "")): _normalize_db_row(row) for row in rows}


def _fetch_database_attachments(keys: list[tuple[str, str]]) -> dict[tuple[str, str], list[dict[str, Any]]]:
    if not keys:
        return {}
    placeholders = ", ".join(["(%s, %s)"] * len(keys))
    params: list[object] = []
    for menu_code, articleid in keys:
        params.extend([menu_code, articleid])
    rows = _db_fetch_all(
        f"""
        SELECT menu_code, articleid, {", ".join(ATTACHMENT_FIELDS)}
        FROM sample_article_attach
        WHERE (menu_code, articleid) IN ({placeholders})
        ORDER BY menu_code ASC, articleid ASC, sortnum ASC, uploadtime DESC
        """,
        params,
    )
    grouped: dict[tuple[str, str], list[dict[str, Any]]] = {}
    for row in rows:
        normalized = _normalize_db_row(row)
        key = (str(normalized.get("menu_code") or ""), str(normalized.get("articleid") or ""))
        grouped.setdefault(key, []).append(normalized)
    return grouped


def _content_summary(summary: str, content_text: str) -> str:
    clean_summary = _normalize_text(summary)
    if clean_summary:
        return _truncate(clean_summary, 500)
    return _truncate(_normalize_text(content_text), 500)


def _attachment_request_options(req: SelectionPreviewRequest) -> dict[str, Any]:
    user_cookie = (req.attachment_cookie or "").strip()
    user_headers = {
        str(key): str(value)
        for key, value in (req.attachment_headers or {}).items()
        if str(key).strip() and value is not None
    }
    enable_download = req.enable_attachment_download
    if enable_download is None:
        enable_download = _env_bool("ENABLE_ATTACHMENT_DOWNLOAD", True) or bool(user_cookie or user_headers)
    return {
        "enable_download": bool(enable_download),
        "force_refresh": bool(req.force_refresh_attachments),
        "user_cookie": user_cookie,
        "user_headers": user_headers,
    }


CORE_ATTACHMENT_KEYWORDS = [
    ("医疗服务价格项目", "医疗服务价格项目"),
    ("服务价格项目", "医疗服务价格项目"),
    ("价格项目", "医疗服务价格项目"),
    ("收费项目", "医疗服务价格项目"),
    ("医保支付政策", "医疗服务价格项目"),
    ("支付政策", "医疗服务价格项目"),
    ("项目目录", "医疗服务价格项目"),
    ("价格标准", "医疗服务价格项目"),
    ("规范整合", "医疗服务价格项目"),
    ("中选结果", "中选结果"),
    ("结果表", "中选结果"),
    ("明细表", "产品清单"),
    ("产品清单", "产品清单"),
    ("挂网清单", "产品清单"),
    ("申报目录", "申报目录"),
    ("目录", "申报目录"),
    ("价格表", "价格表"),
    ("采购文件", "采购文件"),
    ("分组", "产品清单"),
    ("映射关系", "产品清单"),
    ("企业名单", "企业名单"),
    ("配送关系", "配送关系"),
    ("采购量", "产品清单"),
    ("报量", "产品清单"),
]


def _attachment_core_metadata(filename: str) -> tuple[bool, str]:
    for keyword, business_type in CORE_ATTACHMENT_KEYWORDS:
        if keyword in filename:
            return True, business_type
    return False, "普通附件"


def _attachment_unavailable_for_report(status: str) -> bool:
    return status in {"metadata_only", "network_unreachable", "download_failed", "unsupported", "parse_failed"}


def _database_attachment_metadata(row: dict[str, Any], attachment_options: dict[str, Any] | None = None) -> dict[str, Any]:
    articleattid = str(row.get("articleattid") or "")
    filename = str(row.get("filename") or "")
    core_attachment, business_type = _attachment_core_metadata(filename)
    warnings: list[str] = []
    if row.get("fileerrortype"):
        warnings.append(f"附件存在异常标记: {row.get('fileerrortype')}")
    if not articleattid:
        warnings.append("附件缺少 articleattid，后续无法通过内网下载接口定位")
    metadata = {
        "articleattid": articleattid,
        "filename": filename,
        "filepath": row.get("filepath") or "",
        "fileext": row.get("fileext") or "",
        "filesize": row.get("filesize") or 0,
        "uploadtime": row.get("uploadtime") or "",
        "sortnum": row.get("sortnum") or 0,
        "download_method": "internal_attid",
        "download_url_template": "https://qx.eliancloud.cn/Common/EmailFileDownLoad?AttID={articleattid}",
        "core_attachment": core_attachment,
        "business_type": business_type,
        "core_attachment_unavailable": False,
        "cache_status": "cache_disabled",
        "retry_after_seconds": 0,
        "parse_status": "metadata_only",
        "download_status": "metadata_only",
        "download_auth_mode": "none",
        "stored_original_file": False,
        "temporary_file_used": False,
        "temporary_file_deleted": True,
        "text_length": 0,
        "summary": "",
        "key_facts": [],
        "important_sections": [],
        "table_summaries": [],
        "warnings": warnings,
    }
    options = attachment_options or {}
    enable_download = bool(options.get("enable_download", _env_bool("ENABLE_ATTACHMENT_DOWNLOAD", True)))
    force_refresh = bool(options.get("force_refresh"))
    if not enable_download:
        metadata["warnings"].append("附件下载未启用，仅保留元数据")
        metadata["core_attachment_unavailable"] = bool(core_attachment)
        return metadata

    cached, cache_status = load_cached_result(metadata, force_refresh=force_refresh)
    metadata["cache_status"] = cache_status
    if cached:
        metadata.update(cached)
        metadata["core_attachment"] = core_attachment
        metadata["business_type"] = business_type
        metadata["stored_original_file"] = False
        metadata["temporary_file_used"] = bool(metadata.get("temporary_file_used", False))
        metadata["core_attachment_unavailable"] = bool(core_attachment and _attachment_unavailable_for_report(str(metadata.get("parse_status") or "")))
        return metadata

    download = fetch_attachment_bytes(
        metadata,
        enable_download=enable_download,
        user_cookie=str(options.get("user_cookie") or ""),
        user_headers=options.get("user_headers") if isinstance(options.get("user_headers"), dict) else None,
    )
    metadata["download_status"] = download.download_status
    metadata["download_auth_mode"] = download.auth_mode
    metadata["warnings"].extend(download.warnings or [])
    if download.download_status != "downloaded" or not download.content:
        metadata["parse_status"] = download.download_status
        metadata["core_attachment_unavailable"] = bool(core_attachment and _attachment_unavailable_for_report(str(metadata["parse_status"])))
        store_cached_result(metadata, metadata)
        return metadata
    if not _env_bool("ENABLE_ATTACHMENT_PARSE", True):
        metadata["parse_status"] = "stream_parsed"
        metadata["warnings"].append("附件解析未启用，未提取摘要")
        metadata["core_attachment_unavailable"] = bool(core_attachment)
        store_cached_result(metadata, metadata)
        return metadata

    parsed = parse_attachment_bytes(download.content, str(metadata["filename"]), str(metadata["fileext"]), metadata["filesize"])
    statuses = list(parsed.get("parse_statuses") or [])
    metadata["parse_statuses"] = statuses
    metadata["parse_status"] = statuses[-1] if statuses else "parse_failed"
    metadata["text_length"] = int(parsed.get("text_length") or 0)
    metadata["summary"] = parsed.get("summary") or ""
    metadata["key_facts"] = list(parsed.get("key_facts") or [])
    metadata["important_sections"] = list(parsed.get("important_sections") or [])
    metadata["table_summaries"] = list(parsed.get("table_summaries") or [])
    metadata["warnings"].extend(parsed.get("warnings") or [])
    if metadata["parse_status"] not in {"unsupported", "parse_failed"}:
        metadata["download_status"] = "stream_parsed"
    metadata["core_attachment_unavailable"] = bool(core_attachment and _attachment_unavailable_for_report(str(metadata["parse_status"])))
    store_cached_result(metadata, metadata)
    return metadata


def _material_key_facts(row: dict[str, Any]) -> list[dict[str, str]]:
    facts: list[dict[str, str]] = []
    for field, label in [
        ("title", "标题"),
        ("audittime", "发布时间"),
        ("areaname", "地区"),
        ("publicorg", "发布机构"),
        ("projectphase", "项目阶段"),
        ("projecttype", "项目类型"),
        ("category", "分类"),
        ("referencenumber", "文号"),
        ("policytype", "政策类别"),
    ]:
        value = _clean_inline_text(row.get(field) or "")
        if value:
            facts.append({"name": label, "value": value})
    return facts


def _important_passages(content_text: str, keywords: list[str] | None = None, limit: int = 8) -> list[dict[str, str]]:
    words = keywords or ["申报", "价格", "采购", "企业", "产品", "执行", "周期", "规则", "医保编码", "注册证"]
    passages = []
    for sentence in re.split(r"(?<=[。！？；;])\s*|\n+", _normalize_text(content_text)):
        text = sentence.strip()
        if len(text) < 12:
            continue
        matched = [word for word in words if word and word in text]
        if matched:
            passages.append({"reason": f"包含关键词：{', '.join(matched[:3])}", "text": _truncate(text, 280)})
        if len(passages) >= limit:
            break
    return passages


def _passage_texts_by_keywords(content_text: str, keywords: list[str]) -> list[str]:
    return [item["text"] for item in _important_passages(content_text, keywords=keywords, limit=6)]


def _primary_topic_keywords(primary_rows: list[dict[str, Any]]) -> list[str]:
    keywords: list[str] = []
    for row in primary_rows:
        for field in ["title", "areaname", "projectphase", "projecttype", "category", "projectabbreviation"]:
            for token in re.split(r"[\s,，、（）()《》]+", str(row.get(field) or "")):
                token = token.strip()
                if len(token) >= 2:
                    keywords.append(token)
    return _unique(keywords)[:30]


def _auxiliary_relation(auxiliary: dict[str, Any], primary_keywords: list[str]) -> tuple[str, float, list[dict[str, str]]]:
    title = str(auxiliary.get("title") or "")
    content = str(auxiliary.get("content_text") or auxiliary.get("content_summary") or "")
    haystack = f"{title} {content}"
    matched = [word for word in primary_keywords if word and word in haystack]
    score = min(1.0, len(matched) / max(len(primary_keywords[:10]), 1))
    if score < 0.15:
        relation = "低相关"
    elif any(word in title for word in ["结果", "中选", "公布"]):
        relation = "结果公告"
    elif any(word in title for word in ["解读", "政策"]):
        relation = "政策解读"
    elif any(word in title for word in ["历史", "往年", "上一轮"]):
        relation = "历史项目"
    else:
        relation = "同类项目"
    snippets = [
        {"reason": f"与主材料关键词“{word}”相关", "text": _truncate(sentence, 260)}
        for word in matched[:5]
        for sentence in re.split(r"(?<=[。！？；;])\s*|\n+", content)
        if word in sentence and len(sentence.strip()) >= 12
    ][:5]
    return relation, round(score, 3), snippets


def _enhance_material_for_stage4(material: dict[str, Any], primary_keywords: list[str]) -> None:
    content_text = str(material.get("content_text") or "")
    attachments = list(material.get("attachments") or [])
    attachment_summaries = [
        {
            "articleattid": item.get("articleattid") or "",
            "filename": item.get("filename") or "",
            "parse_status": item.get("parse_status") or "metadata_only",
            "summary": item.get("summary") or "",
            "key_facts": item.get("key_facts") or [],
            "table_summaries": item.get("table_summaries") or [],
            "warnings": item.get("warnings") or [],
        }
        for item in attachments
    ]
    material["attachment_summaries"] = attachment_summaries
    if material.get("material_role") == "primary":
        material["important_passages"] = _important_passages(content_text)
        material["policy_rules"] = _passage_texts_by_keywords(content_text, ["政策", "规则", "要求", "执行"])
        material["price_rules"] = _passage_texts_by_keywords(content_text, ["价格", "报价", "申报价", "中选价"])
        material["time_requirements"] = _passage_texts_by_keywords(content_text, ["时间", "截止", "周期", "日期"])
        material["product_scope"] = _passage_texts_by_keywords(content_text, ["产品", "范围", "医保编码", "注册证"])
        material["enterprise_requirements"] = _passage_texts_by_keywords(content_text, ["企业", "申报", "资质"])
        material["execution_requirements"] = _passage_texts_by_keywords(content_text, ["执行", "配送", "采购", "医疗机构"])
    else:
        relation, score, snippets = _auxiliary_relation(material, primary_keywords)
        material["relation_to_primary"] = relation
        material["relevance_score"] = score
        material["relevant_snippets"] = snippets
        if relation == "低相关":
            material.setdefault("warnings", []).append("辅助材料与主材料相关性较低，不建议作为重点依据")
        material["usable_points"] = list(material.get("usable_points") or [])[:8]


def _build_stage4_pack_fields(
    primary_materials: list[dict[str, Any]],
    auxiliary_materials: list[dict[str, Any]],
    combined_key_facts: list[dict[str, str]],
    report_focus: list[str],
) -> dict[str, Any]:
    all_materials = [*primary_materials, *auxiliary_materials]
    attachment_items = [item for material in all_materials for item in material.get("attachments") or []]
    parsed_attachments = [
        item
        for item in attachment_items
        if item.get("parse_status") in {"stream_parsed", "temp_file_parsed", "parsed_text", "parsed_summary", "parsed_table_summary"}
        or "parsed_summary" in (item.get("parse_statuses") or [])
        or "parsed_table_summary" in (item.get("parse_statuses") or [])
    ]
    unusable = [
        item
        for item in attachment_items
        if item.get("parse_status") in {"metadata_only", "auth_required", "auth_failed", "download_failed", "unsupported", "parse_failed"}
    ]
    guidance_warnings = []
    if unusable:
        guidance_warnings.append("部分附件仅有元数据或不可解析，不能引用附件正文或表格内容。")
    return {
        "pack_version": "2.0",
        "primary_evidence": {
            "materials": primary_materials,
            "key_facts": combined_key_facts,
            "important_passages": [passage for material in primary_materials for passage in material.get("important_passages") or []],
            "attachment_summaries": [item for material in primary_materials for item in material.get("attachment_summaries") or []],
        },
        "auxiliary_evidence": {
            "summaries": [
                {
                    "menu_code": material.get("menu_code") or "",
                    "articleid": material.get("articleid") or "",
                    "title": material.get("title") or "",
                    "content_summary": material.get("content_summary") or "",
                    "relation_to_primary": material.get("relation_to_primary") or "",
                    "relevance_score": material.get("relevance_score") or 0,
                }
                for material in auxiliary_materials
            ],
            "relevant_snippets": [snippet for material in auxiliary_materials for snippet in material.get("relevant_snippets") or []],
            "comparison_points": [point for material in auxiliary_materials for point in material.get("usable_points") or []],
        },
        "attachment_evidence": {
            "parsed_summaries": parsed_attachments,
            "table_summaries": [summary for item in parsed_attachments for summary in item.get("table_summaries") or []],
            "warnings": [warning for item in attachment_items for warning in item.get("warnings") or []],
        },
        "generation_guidance": {
            "main_topic": report_focus[0] if report_focus else "",
            "suggested_report_focus": report_focus[:12],
            "suggested_report_style": "专业、克制、分析型，不大段照搬原文",
            "analysis_highlight_enabled": _env_bool("ENABLE_ANALYSIS_HIGHLIGHT", True),
            "do_not_use_as_basis": guidance_warnings,
        },
    }


def _build_database_material(
    *,
    role: str,
    row: dict[str, Any],
    attachments: list[dict[str, Any]],
    warnings: list[str],
    attachment_options: dict[str, Any] | None = None,
) -> dict[str, Any]:
    content_html = str(row.get("content") or "")
    content_text = _article_text_from_html(content_html)
    key = f"{row.get('menu_code')}/{row.get('articleid')}"
    if not content_html:
        warnings.append(f"{key} content 为空")
        logger.warning("database_material_empty_content key=%s", key)
    elif len(content_text) < 80:
        warnings.append(f"{key} 清洗后正文较短")
        logger.warning("database_material_short_content key=%s content_text_length=%s", key, len(content_text))

    base = {
        "material_role": role,
        "menu_code": row.get("menu_code") or "",
        "articleid": row.get("articleid") or "",
        "title": row.get("title") or "",
        "audittime": row.get("audittime") or "",
        "menu_name": row.get("menu_name") or "",
        "source": row.get("source") or "",
        "sourceurl": row.get("sourceurl") or "",
        "areaname": row.get("areaname") or "",
        "publicorg": row.get("publicorg") or "",
        "projectphase": row.get("projectphase") or "",
        "projecttype": row.get("projecttype") or "",
        "category": row.get("category") or "",
        "summary": row.get("summary") or "",
        "content_text": content_text,
        "content_summary": _content_summary(str(row.get("summary") or ""), content_text),
        "attachments": [_database_attachment_metadata(item, attachment_options) for item in attachments],
    }
    if role == "primary":
        base.update(
            {
                "updatetime": row.get("updatetime") or "",
                "dl_project_type": row.get("dl_project_type") or "",
                "referencenumber": row.get("referencenumber") or "",
                "policytype": row.get("policytype") or "",
                "belongproject": row.get("belongproject") or "",
                "projectabbreviation": row.get("projectabbreviation") or "",
                "content_html_length": len(content_html),
                "content_text_length": len(content_text),
                "key_facts": _material_key_facts(row),
            }
        )
    else:
        base["usable_points"] = _material_key_facts(row)
    return base


def _build_database_evidence_pack(req: SelectionPreviewRequest) -> dict[str, Any] | JSONResponse:
    primary_keys, auxiliary_keys, error = _validate_material_selection(req)
    if error is not None:
        return error

    all_keys = list(dict.fromkeys([*primary_keys, *auxiliary_keys]))
    logger.info(
        "analysis_prepare_requested primary_count=%s auxiliary_count=%s keys=%s",
        len(primary_keys),
        len(auxiliary_keys),
        ",".join(f"{menu_code}/{articleid}" for menu_code, articleid in all_keys),
    )

    rows_by_key = _fetch_database_material_rows(all_keys)
    missing = [key for key in all_keys if key not in rows_by_key]
    if missing:
        missing_text = ", ".join(f"{menu_code}/{articleid}" for menu_code, articleid in missing)
        logger.warning("analysis_prepare_material_missing keys=%s", missing_text)
        return _selection_error(422, "MATERIAL_NOT_FOUND", f"所选文章不存在或状态不可用: {missing_text}")

    attachments_by_key = _fetch_database_attachments(all_keys)
    attachment_options = _attachment_request_options(req)
    warnings: list[str] = []
    primary_materials = [
        _build_database_material(
            role="primary",
            row=rows_by_key[key],
            attachments=attachments_by_key.get(key, []),
            warnings=warnings,
            attachment_options=attachment_options,
        )
        for key in primary_keys
    ]
    auxiliary_materials = [
        _build_database_material(
            role="auxiliary",
            row=rows_by_key[key],
            attachments=attachments_by_key.get(key, []),
            warnings=warnings,
            attachment_options=attachment_options,
        )
        for key in auxiliary_keys
    ]
    primary_keywords = _primary_topic_keywords([rows_by_key[key] for key in primary_keys])
    for material in primary_materials:
        _enhance_material_for_stage4(material, primary_keywords)
    for material in auxiliary_materials:
        _enhance_material_for_stage4(material, primary_keywords)
    attachment_count = sum(len(item.get("attachments") or []) for item in [*primary_materials, *auxiliary_materials])
    combined_key_facts: list[dict[str, str]] = []
    for material in primary_materials:
        combined_key_facts.extend(material.get("key_facts") or [])
    report_focus = [
        _clean_inline_text(value)
        for material in primary_materials
        for value in [material.get("title"), material.get("areaname"), material.get("projectphase"), material.get("projecttype"), material.get("category")]
        if _clean_inline_text(value)
    ]
    created_at = datetime.now().isoformat(sep=" ", timespec="seconds")
    pack_id = _make_database_pack_id(primary_keys, auxiliary_keys, created_at)
    pack = {
        "pack_id": pack_id,
        "created_at": created_at,
        "source": "database_selection",
        "primary_materials": primary_materials,
        "auxiliary_materials": auxiliary_materials,
        "combined_key_facts": combined_key_facts,
        "report_focus": _unique(report_focus)[:20],
        "warnings": warnings,
    }
    pack.update(_build_stage4_pack_fields(primary_materials, auxiliary_materials, combined_key_facts, pack["report_focus"]))
    logger.info("analysis_prepare_materials_loaded pack_id=%s attachment_count=%s warnings=%s", pack_id, attachment_count, len(warnings))
    _write_database_evidence_pack(pack)
    return pack


@app.get("/records-ui")
def records_ui():
    path = Path(__file__).resolve().parent / "static" / "records.html"
    if not path.exists():
        raise HTTPException(status_code=404, detail="records UI not found")
    return FileResponse(path, media_type="text/html; charset=utf-8")


@app.get("/analysis-runs/{run_id}")
def analysis_run_ui(run_id: str):
    _safe_run_id(run_id)
    path = Path(__file__).resolve().parent / "static" / "analysis_run.html"
    if not path.exists():
        raise HTTPException(status_code=404, detail="analysis run UI not found")
    return FileResponse(path, media_type="text/html; charset=utf-8")


@app.get("/records", response_model=RecordListResponse)
def list_records(
    keyword: str = "",
    menu_code: str = "",
    areaname: str = "",
    projectphase: str = "",
    projecttype: str = "",
    start_date: str = "",
    end_date: str = "",
    page: int = Query(default=1, ge=1),
    page_size: int = Query(default=20, ge=1, le=100),
) -> RecordListResponse:
    where_sql, params = _build_records_where(
        keyword=keyword.strip(),
        menu_code=menu_code.strip(),
        areaname=areaname.strip(),
        projectphase=projectphase.strip(),
        projecttype=projecttype.strip(),
        start_date=start_date.strip(),
        end_date=end_date.strip(),
    )
    total_row = _db_fetch_one(f"SELECT COUNT(*) AS total FROM sample_article_wide a WHERE {where_sql}", params)
    total = int((total_row or {}).get("total") or 0)
    offset = (page - 1) * page_size
    selected_fields = ", ".join(f"a.{field}" for field in ARTICLE_LIST_FIELDS)
    group_fields = ", ".join(f"a.{field}" for field in ARTICLE_LIST_FIELDS)
    list_sql = f"""
        SELECT {selected_fields}, COUNT(att.articleattid) AS attachment_count
        FROM sample_article_wide a
        LEFT JOIN sample_article_attach att
          ON a.menu_code = att.menu_code AND a.articleid = att.articleid
        WHERE {where_sql}
        GROUP BY {group_fields}
        ORDER BY a.audittime DESC
        LIMIT %s OFFSET %s
    """
    rows = _db_fetch_all(list_sql, [*params, page_size, offset])
    return RecordListResponse(
        items=[_record_list_item(row) for row in rows],
        page=page,
        page_size=page_size,
        total=total,
        total_pages=(total + page_size - 1) // page_size if total else 0,
    )


@app.get("/records/{menu_code}/{articleid}", response_model=RecordDetailResponse)
def get_record_detail(menu_code: str, articleid: str) -> RecordDetailResponse:
    fields = ", ".join(ARTICLE_DETAIL_FIELDS)
    row = _db_fetch_one(
        f"""
        SELECT {fields}
        FROM sample_article_wide a
        WHERE a.status = %s AND a.menu_code = %s AND a.articleid = %s
        LIMIT 1
        """,
        [0, menu_code, articleid],
    )
    if not row:
        raise HTTPException(status_code=404, detail="record not found")
    attachments = _db_fetch_all(
        f"""
        SELECT {", ".join(ATTACHMENT_FIELDS)}
        FROM sample_article_attach
        WHERE menu_code = %s AND articleid = %s
        ORDER BY sortnum ASC, uploadtime DESC
        """,
        [menu_code, articleid],
    )
    normalized = _normalize_db_row(row)
    normalized["content_text"] = _article_text_from_html(str(normalized.get("content") or ""))
    normalized["attachments"] = [_normalize_db_row(item) for item in attachments]
    return RecordDetailResponse(**normalized)


@app.post("/analysis/selection/preview", response_model=SelectionPreviewResponse)
def preview_selection(req: SelectionPreviewRequest) -> SelectionPreviewResponse:
    if not (1 <= len(req.primary_materials) <= 3):
        raise HTTPException(status_code=422, detail="primary_materials 必须选择 1-3 条")
    if len(req.auxiliary_materials) > 10:
        raise HTTPException(status_code=422, detail="auxiliary_materials 最多 10 条")

    primary_keys = [_record_key(item) for item in req.primary_materials]
    auxiliary_keys = [_record_key(item) for item in req.auxiliary_materials]
    overlap = set(primary_keys) & set(auxiliary_keys)
    if overlap:
        raise HTTPException(status_code=422, detail="同一条文章不能同时作为主分析材料和辅助分析材料")

    all_keys = list(dict.fromkeys([*primary_keys, *auxiliary_keys]))
    records = _records_by_keys(all_keys)
    missing = [key for key in all_keys if key not in records]
    if missing:
        missing_text = ", ".join(f"{menu_code}/{articleid}" for menu_code, articleid in missing)
        raise HTTPException(status_code=422, detail=f"所选文章不存在或状态不可用: {missing_text}")

    def material_for(key: tuple[str, str]) -> RecordListItem:
        item = records[key]
        return item if isinstance(item, RecordListItem) else _record_list_item(item)

    return SelectionPreviewResponse(
        success=True,
        primary_materials=[material_for(key) for key in primary_keys],
        auxiliary_materials=[material_for(key) for key in auxiliary_keys],
    )


@app.post("/analysis/prepare", response_model=AnalysisPrepareResponse)
def prepare_analysis(req: SelectionPreviewRequest):
    result = _build_database_evidence_pack(req)
    if isinstance(result, JSONResponse):
        return result
    attachment_count = sum(
        len(material.get("attachments") or [])
        for material in [*(result.get("primary_materials") or []), *(result.get("auxiliary_materials") or [])]
    )
    return AnalysisPrepareResponse(
        success=True,
        pack_id=str(result.get("pack_id") or ""),
        primary_count=len(result.get("primary_materials") or []),
        auxiliary_count=len(result.get("auxiliary_materials") or []),
        attachment_count=attachment_count,
        warnings=list(result.get("warnings") or []),
        evidence_pack=result,
    )


@app.get("/analysis/packs/{pack_id}")
def get_analysis_pack(pack_id: str, full: bool = False) -> dict[str, Any]:
    pack = _read_database_evidence_pack(pack_id)
    if full:
        return pack
    compact = _compact_evidence_pack_for_dify(pack)
    logger.info(
        "analysis_pack_compacted_for_dify pack_id=%s full_chars=%s compact_chars=%s",
        pack_id,
        len(json.dumps(pack, ensure_ascii=False)),
        len(json.dumps(compact, ensure_ascii=False)),
    )
    return compact


@app.get("/analysis/packs/{pack_id}/summary")
def get_analysis_pack_summary(pack_id: str) -> dict[str, Any]:
    pack = _read_database_evidence_pack(pack_id)
    primary = list(pack.get("primary_materials") or [])
    auxiliary = list(pack.get("auxiliary_materials") or [])
    attachment_count = sum(len(material.get("attachments") or []) for material in [*primary, *auxiliary])
    return {
        "pack_id": pack.get("pack_id") or pack_id,
        "created_at": pack.get("created_at") or "",
        "source": pack.get("source") or "",
        "primary_count": len(primary),
        "auxiliary_count": len(auxiliary),
        "attachment_count": attachment_count,
        "primary_titles": [item.get("title") or "" for item in primary],
        "auxiliary_titles": [item.get("title") or "" for item in auxiliary],
        "warnings": list(pack.get("warnings") or []),
    }


@app.get("/analysis/packs/{pack_id}/diagnostics")
def get_analysis_pack_diagnostics(pack_id: str) -> dict[str, Any]:
    pack = _read_database_evidence_pack(pack_id)
    dify_pack = _compact_evidence_pack_for_dify(pack)
    return {
        "success": True,
        "pack_id": pack.get("pack_id") or pack_id,
        "diagnostics": build_pack_diagnostics(pack, dify_pack),
    }


@app.post("/analysis/cache/cleanup")
def cleanup_analysis_cache() -> dict[str, Any]:
    try:
        result = cleanup_cache()
    except Exception as exc:  # noqa: BLE001
        logger.warning("attachment_parse_cache_manual_cleanup_failed error_type=%s", exc.__class__.__name__)
        raise HTTPException(status_code=500, detail="attachment parse cache cleanup failed") from exc
    logger.info(
        "attachment_parse_cache_manual_cleanup deleted_files=%s freed_bytes=%s",
        result.get("deleted_files"),
        result.get("freed_bytes"),
    )
    return {"success": True, **result}


def _save_analysis_timeout_fallback_if_running(pack_id: str, run_id: str, pack: dict[str, Any] | None, timeout_seconds: float) -> None:
    try:
        current = _read_analysis_run(run_id)
    except HTTPException as exc:
        logger.warning("analysis_run_watchdog_missing run_id=%s pack_id=%s detail=%s", run_id, pack_id, exc.detail)
        return
    if current.get("status") != "running":
        return
    try:
        evidence_pack = pack or _read_database_evidence_pack(pack_id)
        exc = DifyWorkflowError(
            "DIFY_TIMEOUT",
            "Dify 工作流调用超时",
            f"watchdog timeout after {timeout_seconds:.1f}s",
            status_code=504,
        )
        fallback = _fallback_result_from_dify_error(exc, evidence_pack, pack_id, apply_quality_gate=False)
        warnings = [
            *list(current.get("warnings") or []),
            "Dify 工作流超过本地等待上限，系统已生成保守降级报告，需人工复核。",
            *list(fallback.get("warnings") or fallback.get("generation_warnings") or []),
        ]
        current.update(fallback)
        current.update(
            {
                "success": True,
                "status": "needs_manual_review",
                "updated_at": datetime.now().isoformat(sep=" ", timespec="seconds"),
                "error_message": "Dify 工作流调用超时：证据包较大或生成/质检耗时过长",
                "error_detail": exc.detail,
                "warnings": warnings,
                "generation_warnings": warnings,
            }
        )
        _write_analysis_run(current)
        logger.warning(
            "analysis_run_watchdog_timeout_fallback_saved run_id=%s pack_id=%s timeout_seconds=%s",
            run_id,
            pack_id,
            timeout_seconds,
        )
    except Exception as fallback_exc:  # noqa: BLE001
        logger.warning(
            "analysis_run_watchdog_fallback_failed run_id=%s pack_id=%s error_type=%s",
            run_id,
            pack_id,
            fallback_exc.__class__.__name__,
        )


def _parse_analysis_timestamp(value: Any) -> datetime | None:
    if not value:
        return None
    try:
        return datetime.fromisoformat(str(value).replace("Z", "+00:00")).replace(tzinfo=None)
    except ValueError:
        return None


def _maybe_finalize_timed_out_analysis_run(record: dict[str, Any]) -> dict[str, Any]:
    if record.get("status") != "running":
        return record
    run_id = str(record.get("run_id") or "")
    pack_id = str(record.get("pack_id") or "")
    if not run_id or not pack_id:
        return record
    created_at = _parse_analysis_timestamp(record.get("created_at"))
    if not created_at:
        return record
    try:
        pack = _read_database_evidence_pack(pack_id)
    except HTTPException as exc:
        logger.warning("analysis_run_timeout_check_pack_missing run_id=%s pack_id=%s detail=%s", run_id, pack_id, exc.detail)
        return record
    timeout_seconds = _analysis_watchdog_timeout_seconds(pack)
    if timeout_seconds <= 0:
        return record
    if (datetime.now() - created_at).total_seconds() < timeout_seconds:
        return record
    _save_analysis_timeout_fallback_if_running(pack_id, run_id, pack, timeout_seconds)
    try:
        return _read_analysis_run(run_id)
    except HTTPException:
        return record


def _execute_analysis_run_background(pack_id: str, run_id: str) -> None:
    try:
        record = _read_analysis_run(run_id)
    except HTTPException as exc:
        logger.warning("analysis_run_background_missing run_id=%s pack_id=%s detail=%s", run_id, pack_id, exc.detail)
        return

    logger.info("analysis_run_background_started run_id=%s pack_id=%s", run_id, pack_id)
    pack_for_policy: dict[str, Any] | None = None
    try:
        pack_for_policy = _read_database_evidence_pack(pack_id)
    except Exception as exc:  # noqa: BLE001
        logger.warning("analysis_run_pack_preread_failed run_id=%s pack_id=%s error_type=%s", run_id, pack_id, exc.__class__.__name__)
    watchdog: threading.Timer | None = None
    watchdog_timeout = _analysis_watchdog_timeout_seconds(pack_for_policy)
    if watchdog_timeout > 0:
        watchdog = threading.Timer(watchdog_timeout, _save_analysis_timeout_fallback_if_running, args=(pack_id, run_id, pack_for_policy, watchdog_timeout))
        watchdog.daemon = True
        watchdog.start()
    try:
        result = _call_dify_workflow(pack_id, run_id, pack_for_policy)
        try:
            pack = pack_for_policy or _read_database_evidence_pack(pack_id)
            result = _repair_unusable_dify_result(result, pack)
            result = _apply_local_quality_gate_to_dify_result(result, pack)
        except Exception as repair_exc:  # noqa: BLE001
            logger.warning(
                "analysis_run_repair_skipped run_id=%s pack_id=%s error_type=%s",
                run_id,
                pack_id,
                repair_exc.__class__.__name__,
            )
    except DifyWorkflowError as exc:
        warnings = list(record.get("warnings") or [])
        error_message = exc.message
        if exc.code == "DIFY_TIMEOUT":
            error_message = "Dify 工作流调用超时：证据包较大或生成/质检耗时过长"
            warnings.append("证据包较大时可能导致 Dify 超时，可减少辅助材料、缩短辅助附件摘要或稍后重试。")
        try:
            pack = pack_for_policy or _read_database_evidence_pack(pack_id)
            fallback = _fallback_result_from_dify_error(exc, pack, pack_id)
            merged_warnings = [*warnings, *list(fallback.get("warnings") or fallback.get("generation_warnings") or [])]
            record.update(fallback)
            record.update(
                {
                    "success": True,
                    "status": "needs_manual_review",
                    "updated_at": datetime.now().isoformat(sep=" ", timespec="seconds"),
                    "error_message": error_message,
                    "error_detail": exc.detail,
                    "warnings": merged_warnings,
                    "generation_warnings": merged_warnings,
                }
            )
            _write_analysis_run(record)
            logger.warning(
                "analysis_run_dify_failed_fallback_saved run_id=%s pack_id=%s code=%s detail=%s",
                run_id,
                pack_id,
                exc.code,
                _truncate(exc.detail, 200),
            )
            return
        except Exception as fallback_exc:  # noqa: BLE001
            logger.warning(
                "analysis_run_dify_fallback_failed run_id=%s pack_id=%s code=%s fallback_error_type=%s",
                run_id,
                pack_id,
                exc.code,
                fallback_exc.__class__.__name__,
            )
        record.update(
            {
                "success": False,
                "status": "failed",
                "updated_at": datetime.now().isoformat(sep=" ", timespec="seconds"),
                "error_message": error_message,
                "error_detail": exc.detail,
                "warnings": warnings,
            }
        )
        _write_analysis_run(record)
        logger.warning("analysis_run_failed run_id=%s pack_id=%s code=%s detail=%s", run_id, pack_id, exc.code, _truncate(exc.detail, 200))
        return
    except Exception as exc:  # noqa: BLE001
        record.update(
            {
                "success": False,
                "status": "failed",
                "updated_at": datetime.now().isoformat(sep=" ", timespec="seconds"),
                "error_message": "报告生成失败",
                "error_detail": exc.__class__.__name__,
            }
        )
        _write_analysis_run(record)
        logger.exception("analysis_run_unexpected_failed run_id=%s pack_id=%s", run_id, pack_id)
        return
    finally:
        if watchdog:
            watchdog.cancel()

    try:
        current_record = _read_analysis_run(run_id)
        if current_record.get("status") != "running":
            logger.warning(
                "analysis_run_late_dify_result_ignored run_id=%s pack_id=%s current_status=%s",
                run_id,
                pack_id,
                current_record.get("status") or "",
            )
            return
    except HTTPException:
        pass
    record.update(result)
    record["success"] = True
    record["updated_at"] = datetime.now().isoformat(sep=" ", timespec="seconds")
    _write_analysis_run(record)
    logger.info(
        "analysis_run_finished run_id=%s pack_id=%s workflow_run_id=%s status=%s",
        run_id,
        pack_id,
        record.get("workflow_run_id") or "",
        record.get("status") or "",
    )


@app.post("/analysis/run")
def run_analysis(req: AnalysisRunRequest):
    pack_id = req.pack_id.strip()
    try:
        _read_database_evidence_pack(pack_id)
    except HTTPException as exc:
        if exc.status_code == 404:
            return _analysis_error(404, "PACK_NOT_FOUND", "evidence_pack 不存在", str(exc.detail))
        return _analysis_error(exc.status_code, "PACK_READ_FAILED", "读取 evidence_pack 失败", str(exc.detail))

    run_id = _make_analysis_run_id(pack_id)
    now = datetime.now().isoformat(sep=" ", timespec="seconds")
    record: dict[str, Any] = {
        "success": True,
        "run_id": run_id,
        "pack_id": pack_id,
        "status": "running",
        "workflow_run_id": "",
        "report_title": "",
        "report_markdown": "",
        "quality_check": {"passed": None, "issues": []},
        "generation_warnings": [],
        "warnings": [],
        "remaining_issues": [],
        "version": 1,
        "created_at": now,
        "updated_at": now,
        "error_message": "",
    }
    _write_analysis_run(record)
    logger.info("analysis_run_started run_id=%s pack_id=%s", run_id, pack_id)

    threading.Thread(target=_execute_analysis_run_background, args=(pack_id, run_id), daemon=True).start()
    return AnalysisRunResponse(
        success=True,
        run_id=run_id,
        pack_id=pack_id,
        status="running",
        workflow_run_id=str(record.get("workflow_run_id") or ""),
        report_title=str(record.get("report_title") or ""),
        quality_passed=(record.get("quality_check") or {}).get("passed") if isinstance(record.get("quality_check"), dict) else None,
        version=int(record.get("version") or 1),
        warnings=list(record.get("warnings") or record.get("generation_warnings") or []),
    )


@app.get("/analysis/runs/{run_id}")
def get_analysis_run(run_id: str) -> dict[str, Any]:
    try:
        record = _read_analysis_run(run_id)
    except HTTPException as exc:
        if exc.status_code == 404:
            return _analysis_error(404, "RUN_NOT_FOUND", "analysis run 不存在", str(exc.detail))
        return _analysis_error(exc.status_code, "RUN_READ_FAILED", "读取 analysis run 失败", str(exc.detail))
    record = _maybe_finalize_timed_out_analysis_run(record)
    summary = dict(record)
    summary.pop("report_markdown", None)
    summary["quality_passed"] = (record.get("quality_check") or {}).get("passed") if isinstance(record.get("quality_check"), dict) else None
    summary["progress"] = build_run_progress(record)
    return summary


@app.get("/analysis/runs/{run_id}/diagnostics")
def get_analysis_run_diagnostics(run_id: str) -> dict[str, Any]:
    try:
        record = _read_analysis_run(run_id)
    except HTTPException as exc:
        if exc.status_code == 404:
            return _analysis_error(404, "RUN_NOT_FOUND", "analysis run 不存在", str(exc.detail))
        return _analysis_error(exc.status_code, "RUN_READ_FAILED", "读取 analysis run 失败", str(exc.detail))
    record = _maybe_finalize_timed_out_analysis_run(record)

    pack: dict[str, Any] | None = None
    pack_id = str(record.get("pack_id") or "")
    if pack_id:
        try:
            pack = _read_database_evidence_pack(pack_id)
        except HTTPException as exc:
            logger.warning("analysis_run_diagnostics_pack_missing run_id=%s pack_id=%s detail=%s", run_id, pack_id, exc.detail)

    return {
        "success": True,
        "run_id": str(record.get("run_id") or run_id),
        "pack_id": pack_id,
        "status": str(record.get("status") or ""),
        "diagnostics": build_run_diagnostics(record, pack, _compact_evidence_pack_for_dify(pack) if pack else None),
    }


@app.get("/analysis/runs/{run_id}/report", response_model=AnalysisRunReportResponse)
def get_analysis_run_report(run_id: str) -> AnalysisRunReportResponse | JSONResponse:
    try:
        record = _read_analysis_run(run_id)
    except HTTPException as exc:
        if exc.status_code == 404:
            return _analysis_error(404, "RUN_NOT_FOUND", "analysis run 不存在", str(exc.detail))
        return _analysis_error(exc.status_code, "RUN_READ_FAILED", "读取 analysis run 失败", str(exc.detail))
    record = _maybe_finalize_timed_out_analysis_run(record)
    if not record.get("report_markdown"):
        return _analysis_error(409, "REPORT_NOT_READY", "报告尚未生成完成")
    return AnalysisRunReportResponse(
        success=bool(record.get("success", True)),
        run_id=str(record.get("run_id") or run_id),
        pack_id=str(record.get("pack_id") or ""),
        report_title=str(record.get("report_title") or ""),
        report_markdown=_clean_model_output(str(record.get("report_markdown") or "")),
        quality_check=record.get("quality_check") if isinstance(record.get("quality_check"), dict) else {"passed": None, "issues": []},
        quality_gate=record.get("quality_gate") if isinstance(record.get("quality_gate"), dict) else {},
        version=int(record.get("version") or 1),
        warnings=list(record.get("warnings") or record.get("generation_warnings") or []),
        remaining_issues=list(record.get("remaining_issues") or []),
    )


@app.post("/analysis/runs/{run_id}/revise", response_model=AnalysisRunReviseResponse)
def revise_analysis_run(run_id: str, req: AnalysisRunReviseRequest) -> AnalysisRunReviseResponse | JSONResponse:
    if not _env_bool("ENABLE_USER_FEEDBACK_REVISION", True):
        return _analysis_error(403, "REVISION_DISABLED", "用户反馈修改功能未启用")
    try:
        record = _read_analysis_run(run_id)
    except HTTPException as exc:
        if exc.status_code == 404:
            return _analysis_error(404, "RUN_NOT_FOUND", "analysis run 不存在", str(exc.detail))
        return _analysis_error(exc.status_code, "RUN_READ_FAILED", "读取 analysis run 失败", str(exc.detail))
    current_report = str(record.get("report_markdown") or "").strip()
    if not current_report:
        return _analysis_error(409, "REPORT_NOT_READY", "报告尚未生成完成")
    pack_id = str(record.get("pack_id") or "")
    try:
        pack = _read_database_evidence_pack(pack_id)
    except HTTPException as exc:
        return _analysis_error(exc.status_code, "PACK_READ_FAILED", "读取 evidence_pack 失败", str(exc.detail))

    analysis_highlight = req.analysis_highlight if req.analysis_highlight is not None else _env_bool("ENABLE_ANALYSIS_HIGHLIGHT", True)
    revision_id = _make_revision_id(run_id)
    try:
        result = _call_dify_revision_workflow(
            pack_id=pack_id,
            run_id=run_id,
            feedback=req.feedback,
            current_report=current_report,
            evidence_pack=pack,
            analysis_highlight=bool(analysis_highlight),
        )
    except DifyWorkflowError as exc:
        logger.warning("analysis_revision_failed run_id=%s pack_id=%s code=%s detail=%s", run_id, pack_id, exc.code, _truncate(exc.detail, 200))
        return _analysis_error(exc.status_code, exc.code, exc.message, exc.detail)

    old_version = int(record.get("version") or 1)
    new_version = old_version + 1
    previous_versions = list(record.get("report_versions") or [])
    previous_versions.append(
        {
            "version": old_version,
            "report_title": record.get("report_title") or "",
            "report_markdown": current_report,
            "created_at": record.get("updated_at") or record.get("created_at") or "",
        }
    )
    report_title = str(result.get("report_title") or record.get("report_title") or "")
    report_markdown = _clean_model_output(str(result.get("report_markdown") or current_report))
    quality_check = result.get("quality_check") if isinstance(result.get("quality_check"), dict) else {"passed": None, "issues": []}
    warnings = list(result.get("warnings") or result.get("generation_warnings") or [])
    revision_record = {
        "revision_id": revision_id,
        "run_id": run_id,
        "version": new_version,
        "feedback": req.feedback,
        "mode": req.mode,
        "analysis_highlight": bool(analysis_highlight),
        "report_title": report_title,
        "report_markdown": report_markdown,
        "created_at": datetime.now().isoformat(sep=" ", timespec="seconds"),
        "quality_check": quality_check,
        "warnings": warnings,
    }
    record.update(
        {
            "report_versions": previous_versions[-5:],
            "revisions": [*list(record.get("revisions") or []), revision_record][-10:],
            "version": new_version,
            "report_title": report_title,
            "report_markdown": report_markdown,
            "quality_check": quality_check,
            "warnings": warnings,
            "generation_warnings": warnings,
            "user_feedback_modified": True,
            "last_revision": revision_record,
            "updated_at": revision_record["created_at"],
        }
    )
    _write_analysis_run(record)
    logger.info("analysis_revision_saved run_id=%s revision_id=%s version=%s feedback_chars=%s", run_id, revision_id, new_version, len(req.feedback))
    return AnalysisRunReviseResponse(
        success=True,
        revision_id=revision_id,
        run_id=run_id,
        pack_id=pack_id,
        version=new_version,
        report_title=report_title,
        report_markdown=report_markdown,
        quality_check=quality_check,
        warnings=warnings,
    )


@app.get("/analysis/runs/{run_id}/download")
def download_analysis_run_report(run_id: str):
    try:
        record = _read_analysis_run(run_id)
    except HTTPException as exc:
        if exc.status_code == 404:
            return _analysis_error(404, "RUN_NOT_FOUND", "analysis run 不存在", str(exc.detail))
        return _analysis_error(exc.status_code, "RUN_READ_FAILED", "读取 analysis run 失败", str(exc.detail))

    report_markdown = str(record.get("report_markdown") or "").strip()
    if not report_markdown:
        return _analysis_error(409, "REPORT_NOT_READY", "报告尚未生成完成")

    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    _cleanup_old_reports(REPORT_DIR)
    title = str(record.get("report_title") or "分析报告").strip() or "分析报告"
    version = str(record.get("version") or 1)
    report_hash = hashlib.sha256(report_markdown.encode("utf-8", errors="ignore")).hexdigest()[:12]
    filename = f"{_safe_filename(title)}_{_safe_filename(run_id)}_v{_safe_filename(version)}_{report_hash}.docx"
    path = REPORT_DIR / filename
    if not path.exists():
        _markdown_to_docx(report_markdown, path, title)
        logger.info("analysis_run_report_download_created run_id=%s filename=%s", run_id, filename)
    else:
        logger.info("analysis_run_report_download_cache_hit run_id=%s filename=%s", run_id, filename)
    return FileResponse(
        path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{quote(filename)}"},
    )


@app.post("/analyze", response_model=AnalyzeResponse)
async def analyze(req: AnalyzeRequest) -> AnalyzeResponse:
    original_url = str(req.url)
    url, url_warnings = _normalize_notice_url(original_url)
    _validate_url(url)
    started = time.perf_counter()
    logger.info(
        "analyze_started url=%s max_attachments=%s max_combined_chars=%s",
        _safe_url_for_log(url),
        req.max_attachments,
        req.max_combined_chars,
    )

    warnings: list[str] = list(url_warnings)
    site_detail_task = asyncio.create_task(_try_known_site_detail(url, warnings))
    try:
        page_task = asyncio.create_task(_fetch_page(url))
        page = await page_task
    except httpx.HTTPError as exc:
        logger.warning(
            "analyze_page_fetch_failed url=%s error_type=%s error=%s",
            _safe_url_for_log(url),
            exc.__class__.__name__,
            str(exc)[:300],
        )
        page = _load_cached_page(url, warnings)
        if page is None:
            early_detail = await site_detail_task
            if early_detail:
                page = {"html": "", "final_url": url}
                warnings.append(f"公告网页 HTML 抓取失败，已使用站点详情接口兜底: {url} ({exc})")
        if page is None:
            page = {
                "html": (
                    "<html><head><title>公告网页抓取失败</title></head><body>"
                    f"<p>公告URL: {html_lib.escape(url)}</p>"
                    f"<p>抓取失败: {html_lib.escape(str(exc))}</p>"
                    "</body></html>"
                ),
                "final_url": url,
            }
            warnings.append(f"抓取状态: crawl_insufficient；公告网页抓取失败: {url} ({exc})")

    final_url = str(page.get("final_url") or url)
    html = str(page.get("html") or "")

    if "hugedomains.com/domain_profile" in final_url and "szgzy.com" in original_url:
        corrected_url = original_url.replace("://www.szgzy.com", "://www.szggzy.com").replace(
            "://szgzy.com", "://szggzy.com"
        )
        warnings.append(f"原 URL 跳转到域名售卖页，已尝试修正为: {corrected_url}")
        try:
            page = await _fetch_page(corrected_url)
            url = corrected_url
            final_url = str(page.get("final_url") or corrected_url)
            html = str(page.get("html") or "")
        except httpx.HTTPError as exc:
            logger.warning(
                "analyze_corrected_page_fetch_failed url=%s error_type=%s error=%s",
                _safe_url_for_log(corrected_url),
                exc.__class__.__name__,
                str(exc)[:300],
            )
            page = {
                "html": (
                    "<html><head><title>修正后的公告网页抓取失败</title></head><body>"
                    f"<p>公告URL: {html_lib.escape(corrected_url)}</p>"
                    f"<p>抓取失败: {html_lib.escape(str(exc))}</p>"
                    "</body></html>"
                ),
                "final_url": corrected_url,
            }
            warnings.append(f"抓取状态: crawl_insufficient；修正后的公告网页抓取失败: {corrected_url} ({exc})")
            url = corrected_url
            final_url = corrected_url
            html = str(page.get("html") or "")

    site_detail = await site_detail_task
    if not site_detail and final_url != url:
        site_detail = await _try_known_site_detail(final_url, warnings)
    if site_detail:
        html = html + "\n" + site_detail["html"]

    firecrawl_text = await _try_firecrawl(url, req.firecrawl_api_key or os.getenv("FIRECRAWL_API_KEY"), warnings)
    parsed_page = _parse_html(html, final_url)
    if site_detail and site_detail.get("title"):
        parsed_page["title"] = str(site_detail["title"])
    page_text = firecrawl_text or parsed_page["text"]
    if site_detail:
        page_text = _normalize_text(
            "\n".join(
                [
                    f"标题: {site_detail.get('title', '')}",
                    f"发布时间: {site_detail.get('release_time', '')}",
                    f"栏目: {site_detail.get('node_list', '')}",
                    f"信息来源: {site_detail.get('source', '')}",
                    page_text,
                ]
            )
        )
    title = parsed_page["title"]

    links = _discover_attachment_links(html, final_url, req.max_attachments)
    attachments = await _collect_attachments(links, warnings)

    crawl_status = _crawl_status_from_warnings(warnings)
    if not crawl_status and _extract_elian_qx_article_id(url) and not site_detail:
        crawl_status = "login_required"
        warnings.append("抓取状态: login_required；公司器械集采准入系统页面仅返回前端壳，未能取得公告正文。")
    if not crawl_status and _looks_like_guizhou_dynamic_shell(url, title, page_text, attachments):
        crawl_status = "crawl_insufficient"
        warnings.append("抓取状态: crawl_insufficient；贵州公共服务平台动态详情页未返回可解析公告正文。")
    if crawl_status:
        page_text = _normalize_text(
            "\n".join(
                [
                    f"抓取状态: {crawl_status}",
                    f"原始链接: {original_url}",
                    "系统未取得可用于正式分析的公告正文或附件。请提供可访问的原文链接、公司站点登录态或人工上传原文附件。",
                    page_text,
                ]
            )
        )

    publish_dates = _unique(_extract_publish_metadata_dates(html) + _extract_dates(" ".join([title, page_text])))
    regions = _unique(_extract_regions(" ".join([title, page_text])))
    combined_text = _build_combined_text(title, final_url, page_text, attachments)
    evidence_for_llm = _truncate(combined_text, req.max_combined_chars)

    if len(combined_text) > len(evidence_for_llm):
        warnings.append(f"证据包过长，已截断给 Dify: {len(combined_text)} -> {len(evidence_for_llm)} 字符")

    response_attachments = [_trim_attachment_for_response(attachment) for attachment in attachments]

    logger.info(
        "analyze_completed url=%s final_url=%s attachments_found=%s attachments_parsed=%s evidence_chars=%s warnings=%s elapsed_ms=%s",
        _safe_url_for_log(original_url),
        _safe_url_for_log(final_url),
        len(links),
        len(response_attachments),
        len(evidence_for_llm),
        len(warnings),
        int((time.perf_counter() - started) * 1000),
    )
    return AnalyzeResponse(
        source_url=original_url,
        final_url=final_url,
        title=title,
        publish_date_candidates=publish_dates,
        region_candidates=regions,
        page_text=_truncate(page_text, 8_000),
        attachments=response_attachments,
        combined_text="",
        evidence_for_llm=evidence_for_llm,
        warnings=warnings,
    )


@app.post("/analyze_v2", response_model=AnalyzeV2Response)
async def analyze_v2(req: AnalyzeV2Request) -> AnalyzeV2Response:
    original_url = str(req.url)
    url, url_warnings = _normalize_notice_url(original_url)
    _validate_url(url)

    run_id = str(uuid.uuid4())
    started = time.perf_counter()
    logger.info(
        "analyze_v2_started run_id=%s url=%s max_attachments=%s max_combined_chars=%s",
        run_id,
        _safe_url_for_log(url),
        req.max_attachments,
        req.max_combined_chars,
    )
    warnings: list[str] = list(url_warnings)
    direct_suffix = Path(urlparse(url).path.lower()).suffix
    if direct_suffix in ATTACHMENT_EXTENSIONS:
        attachments = await _collect_attachments([{"url": url, "text": "direct file url"}], warnings)
        if not attachments:
            return _analyze_v2_failure(
                source_url=original_url,
                run_id=run_id,
                error_type="unsupported_file",
                message="文件链接无法下载或解析，请确认文件可公开访问，或改为上传原始附件。",
                warnings=warnings,
                errors=["direct file parse failed"],
            )
        title = attachments[0].get("filename") or Path(urlparse(url).path).name or "direct-file"
        notice_summary = {
            "title": title,
            "publish_date": "",
            "source_agency": "",
            "document_name": title,
            "notice_type_hint": "direct_file",
        }
        evidence_pack = _build_evidence_pack(
            source_url=original_url,
            final_url=url,
            title=title,
            page_text="",
            notice_summary=notice_summary,
            attachments=attachments,
            warnings=warnings,
        )
        content_hash = _stable_content_hash(original_url, "", attachments)
        llm_input_text = _build_llm_input_text(evidence_pack, req.max_combined_chars)
        _write_evidence_cache(content_hash, {"run_id": run_id, "content_hash": content_hash, "evidence_pack": evidence_pack})
        logger.info(
            "analyze_v2_completed run_id=%s url=%s mode=direct_file attachments_parsed=%s evidence_chars=%s warnings=%s elapsed_ms=%s",
            run_id,
            _safe_url_for_log(original_url),
            len(attachments),
            len(llm_input_text),
            len(warnings),
            int((time.perf_counter() - started) * 1000),
        )
        return AnalyzeV2Response(
            success=True,
            run_id=run_id,
            evidence_id=content_hash,
            content_hash=content_hash,
            source_url=original_url,
            fetch_status={
                "main_page": "success",
                "attachments_found": 1,
                "attachments_downloaded": len(attachments),
                "attachments_parsed": len(attachments),
                "warnings": warnings,
                "errors": [],
            },
            notice_summary=notice_summary,
            evidence_pack=evidence_pack,
            llm_input_text=llm_input_text,
            raw_storage={"snapshot_id": content_hash, "raw_files": [title]},
        )
    site_detail_task = asyncio.create_task(_try_known_site_detail(url, warnings))

    page: dict[str, str] | None = None
    fetch_error: Exception | None = None
    try:
        page = await _fetch_page(url)
    except Exception as exc:  # noqa: BLE001
        fetch_error = exc
        page = _load_cached_page(url, warnings)
        if page is None:
            early_detail = await site_detail_task
            if early_detail:
                page = {"html": early_detail["html"], "final_url": url}
                warnings.append("主页面抓取失败，已使用站点详情接口兜底。")

    if page is None:
        error_type = _classify_fetch_error(fetch_error)
        return _analyze_v2_failure(
            source_url=original_url,
            run_id=run_id,
            error_type=error_type,
            message=_fetch_error_user_message(error_type, fetch_error),
            warnings=warnings,
            errors=[str(fetch_error or "fetch failed")],
        )

    final_url = str(page.get("final_url") or url)
    html = str(page.get("html") or "")
    site_detail = await site_detail_task
    if not site_detail and final_url != url:
        site_detail = await _try_known_site_detail(final_url, warnings)
    if site_detail:
        html = html + "\n" + site_detail["html"]

    firecrawl_text = await _try_firecrawl(url, req.firecrawl_api_key or os.getenv("FIRECRAWL_API_KEY"), warnings)
    parsed_page = _parse_html(html, final_url)
    if site_detail and site_detail.get("title"):
        parsed_page["title"] = str(site_detail["title"])
    page_text = firecrawl_text or parsed_page["text"]
    title = parsed_page["title"]

    links = _discover_attachment_links(html, final_url, req.max_attachments)
    attachments = await _collect_attachments(links, warnings)

    crawl_status = _crawl_status_from_warnings(warnings)
    meaningful_text = _normalize_text("\n".join([title, page_text]))
    if crawl_status or (len(meaningful_text) < 40 and not attachments):
        error_type = "empty_content" if not crawl_status else "parse_error"
        message = "未能取得可用于分析的公告正文或附件，请确认链接可公开访问，或上传原始 Word/PDF/Excel 附件。"
        return _analyze_v2_failure(
            source_url=original_url,
            run_id=run_id,
            error_type=error_type,
            message=message,
            warnings=warnings,
            errors=[crawl_status] if crawl_status else [],
        )

    publish_dates = _unique(_extract_publish_metadata_dates(html) + _extract_dates(" ".join([title, page_text])))
    notice_summary = {
        "title": title,
        "publish_date": publish_dates[0] if publish_dates else "",
        "source_agency": _guess_source_agency(html, page_text),
        "document_name": title,
        "notice_type_hint": "",
    }
    evidence_pack = _build_evidence_pack(
        source_url=original_url,
        final_url=final_url,
        title=title,
        page_text=page_text,
        notice_summary=notice_summary,
        attachments=attachments,
        warnings=warnings,
    )
    content_hash = _stable_content_hash(original_url, page_text, attachments)
    evidence_id = content_hash
    llm_input_text = _build_llm_input_text(evidence_pack, req.max_combined_chars)
    _write_evidence_cache(evidence_id, {"run_id": run_id, "content_hash": content_hash, "evidence_pack": evidence_pack})

    logger.info(
        "analyze_v2_completed run_id=%s url=%s final_url=%s attachments_found=%s attachments_parsed=%s evidence_chars=%s warnings=%s elapsed_ms=%s",
        run_id,
        _safe_url_for_log(original_url),
        _safe_url_for_log(final_url),
        len(links),
        len(attachments),
        len(llm_input_text),
        len(warnings),
        int((time.perf_counter() - started) * 1000),
    )
    return AnalyzeV2Response(
        success=True,
        run_id=run_id,
        evidence_id=evidence_id,
        content_hash=content_hash,
        source_url=original_url,
        fetch_status={
            "main_page": "success" if not fetch_error else "partial",
            "attachments_found": len(links),
            "attachments_downloaded": len(attachments),
            "attachments_parsed": len(attachments),
            "warnings": warnings,
            "errors": [],
        },
        notice_summary=notice_summary,
        evidence_pack=evidence_pack,
        llm_input_text=llm_input_text,
        raw_storage={"snapshot_id": evidence_id, "raw_files": [item.get("filename", "") for item in attachments]},
        message="",
    )


@app.post("/report/export", response_model=ExportReportResponse)
async def export_report(req: ExportReportRequest) -> ExportReportResponse:
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    _cleanup_old_reports(REPORT_DIR)
    try:
        report = _prepare_report_for_export(req)
        quality_issues = _quality_check_report(report)
        if req.strict_quality and quality_issues:
            raise ValueError("报告质量检查未通过：" + "；".join(quality_issues))
        filename = _unique_report_filename(_build_report_filename(report), REPORT_DIR)
        path = REPORT_DIR / filename
        _report_ir_to_docx(report, path, req.title)
        logger.info("report_export_completed filename=%s", filename)
    except ValueError as exc:
        logger.warning("report_export_failed error=%s", str(exc)[:500])
        raise HTTPException(status_code=422, detail=str(exc)) from exc
    return ExportReportResponse(filename=filename, download_url=_download_url(filename))


@app.post("/report/render", response_model=RenderReportResponse)
def render_report(req: RenderReportRequest) -> RenderReportResponse:
    try:
        report = _prepare_report_for_export(req)
        quality_issues = _quality_check_report(report)
        if req.strict_quality and quality_issues:
            raise ValueError("报告质量检查未通过：" + "；".join(quality_issues))
        logger.info("report_render_completed strict_quality=%s warnings=%s", req.strict_quality, len(quality_issues))
        return RenderReportResponse(success=True, report_ir=report, report_markdown=_report_ir_to_markdown(report))
    except ValueError as exc:
        logger.warning("report_render_failed error=%s", str(exc)[:500])
        error = _render_parse_error(req.markdown, str(exc))
        return RenderReportResponse(success=False, error=error, user_message=error)


@app.post("/report/render_v2", response_model=RenderReportResponse)
def render_report_v2(req: RenderReportRequest) -> RenderReportResponse:
    try:
        report = _prepare_report_for_export(req)
        quality_issues = _quality_check_report(report)
        if req.strict_quality and quality_issues:
            message = "报告结构校验未通过：" + "；".join(quality_issues)
            logger.warning("report_render_v2_quality_failed warnings=%s", len(quality_issues))
            return RenderReportResponse(success=False, error=message, user_message=message, quality_warnings=quality_issues)
        logger.info("report_render_v2_completed strict_quality=%s warnings=%s", req.strict_quality, len(quality_issues))
        return RenderReportResponse(
            success=True,
            report_ir=report,
            report_markdown=_report_ir_to_markdown(report),
            quality_warnings=quality_issues,
        )
    except ValueError as exc:
        logger.warning("report_render_v2_failed error=%s", str(exc)[:500])
        error = _render_parse_error(req.markdown, str(exc))
        return RenderReportResponse(success=False, error=error, user_message=error)


@app.post("/report/export_checked", response_model=CheckedExportReportResponse)
def export_report_checked(req: CheckedExportReportRequest) -> CheckedExportReportResponse:
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    _cleanup_old_reports(REPORT_DIR)
    try:
        qa = _qa_from_checked_export_request(req)
        report_for_local_qa = _prepare_report_for_export(req)
        local_qa = _quality_check_report_against_evidence(report_for_local_qa, req.evidence_text, req.history_text)
        qa = _merge_local_qa(qa, local_qa)
        qa = _suppress_allowed_calculated_percentage_qa(qa, report_for_local_qa)
        qa = _normalize_qa_status_for_workflow(qa)

        report = report_for_local_qa
        report_markdown = _report_ir_to_markdown(report)
        quality_issues = _quality_check_report(report)
        if quality_issues:
            _add_nonblocking_export_quality_issues(qa, quality_issues)
            qa = _normalize_qa_status_for_workflow(qa)
        qa_summary = _format_qa_summary(qa)
        if req.strict_quality and _workflow_qa_status(qa) != "pass":
            logger.warning(
                "report_export_checked_blocked qa_status=%s issues=%s unsupported_claims=%s history_leakage=%s missing_rules=%s",
                _workflow_qa_status(qa),
                len(qa.issues),
                len(qa.unsupported_claims),
                len(qa.history_leakage),
                len(qa.missing_rules),
            )
            return CheckedExportReportResponse(
                success=False,
                blocked=True,
                qa_summary=qa_summary,
                report_markdown=report_markdown,
            )
        filename = _unique_report_filename(_build_report_filename(report), REPORT_DIR)
        path = REPORT_DIR / filename
        _report_ir_to_docx(report, path, req.title)
        logger.info("report_export_checked_completed filename=%s qa_status=%s", filename, _workflow_qa_status(qa))
        return CheckedExportReportResponse(
            success=True,
            filename=filename,
            download_url=_download_url(filename),
            blocked=False,
            qa_summary=qa_summary,
            report_markdown=report_markdown,
        )
        if req.qa_output:
            try:
                qa = _parse_qa_output(req.qa_output)
            except ValueError as exc:
                qa = ReportQA(
                    status="needs_fix",
                    issues=[
                        ReportQAIssue(
                            severity="major",
                            category="qa_parse_error",
                            report_text=str(exc),
                            fix_instruction="检查质检模型输出格式，确保为合法 QA JSON。",
                        )
                    ],
                    fix_instructions=["检查质检模型输出格式，确保为合法 QA JSON。"],
                    summary="质检 JSON 解析失败，已作为缺陷提示保留，不阻断 Word 导出。",
                )
        else:
            qa = ReportQA()
        report_for_local_qa = _prepare_report_for_export(req)
        local_qa = _quality_check_report_against_evidence(report_for_local_qa, req.evidence_text, req.history_text)
        qa = _merge_local_qa(qa, local_qa)
        qa = _suppress_allowed_calculated_percentage_qa(qa, report_for_local_qa)

        report = report_for_local_qa
        report_markdown = _report_ir_to_markdown(report)
        quality_issues = _quality_check_report(report)
        _add_nonblocking_export_quality_issues(qa, quality_issues)
        qa_summary = _format_qa_summary(qa)
        filename = _unique_report_filename(_build_report_filename(report), REPORT_DIR)
        path = REPORT_DIR / filename
        _report_ir_to_docx(report, path, req.title)
    except ValueError as exc:
        logger.warning("report_export_checked_failed error=%s", str(exc)[:500])
        raise HTTPException(status_code=422, detail=str(exc)) from exc
    return CheckedExportReportResponse(
        success=True,
        filename=filename,
        download_url=_download_url(filename),
        blocked=False,
        qa_summary=qa_summary,
        report_markdown=report_markdown,
    )


@app.post("/report/qa", response_model=ReportQAParseResponse)
def parse_report_qa(req: ReportQAParseRequest) -> ReportQAParseResponse:
    try:
        qa = _parse_qa_output(req.qa_output)
    except ValueError as exc:
        logger.warning("report_qa_parse_failed error=%s qa_output_chars=%s", str(exc)[:500], len(req.qa_output or ""))
        qa = ReportQA(
            status="block",
            issues=[
                ReportQAIssue(
                    severity="blocker",
                    category="qa_parse_error",
                    report_text=str(exc),
                    fix_instruction="检查质检模型输出格式，确保为合法 QA JSON。",
                )
            ],
            fix_instructions=["检查质检模型输出格式，确保为合法 QA JSON。"],
            summary="质检 JSON 解析失败，已作为缺陷提示保留。",
        )

    if req.report_ir is not None:
        local_report = _normalize_report_ir(req.report_ir, fallback_title=DEFAULT_REPORT_TITLE)
    else:
        local_report = _markdown_to_report_ir(req.report_text) if req.report_text else ReportIR()
    local_qa = _quality_check_report_against_evidence(local_report, req.evidence_text, req.history_text)
    qa = _merge_local_qa(qa, local_qa)
    qa = _suppress_allowed_calculated_percentage_qa(qa, local_report)
    qa = _normalize_qa_status_for_workflow(qa)
    status = _workflow_qa_status(qa)
    logger.info(
        "report_qa_completed status=%s issues=%s unsupported_claims=%s history_leakage=%s missing_rules=%s",
        status,
        len(qa.issues),
        len(qa.unsupported_claims),
        len(qa.history_leakage),
        len(qa.missing_rules),
    )

    return ReportQAParseResponse(
        blocked=status == "block",
        needs_fix=status == "needs_fix",
        status=status,
        issues=qa.issues,
        unsupported_claims=qa.unsupported_claims,
        history_leakage=qa.history_leakage,
        missing_rules=qa.missing_rules,
        language_issues=qa.language_issues,
        fix_instructions=qa.fix_instructions,
        summary=qa.summary,
        qa=qa,
        qa_summary=_format_qa_summary(qa),
    )


def _cleanup_old_reports(directory: Path) -> None:
    retention_hours = int(os.getenv("REPORT_RETENTION_HOURS", "168"))
    cutoff = time.time() - retention_hours * 3600
    for path in directory.glob("*.docx"):
        try:
            if path.is_file() and path.stat().st_mtime < cutoff:
                path.unlink()
        except OSError:
            continue


def _configured_public_base_url() -> str:
    base_url = (os.getenv("PUBLIC_BASE_URL") or DEFAULT_PUBLIC_BASE_URL).strip().rstrip("/")
    if not base_url:
        base_url = DEFAULT_PUBLIC_BASE_URL
    return base_url


def _download_url(filename: str) -> str:
    base_url = _configured_public_base_url()
    return f"{base_url}/download/{quote(filename)}"


@app.get("/download/{filename}")
def download_report(filename: str):
    from fastapi.responses import FileResponse

    safe = Path(filename).name
    path = REPORT_DIR / safe
    if not path.exists():
        raise HTTPException(status_code=404, detail="report file not found")
    return FileResponse(
        path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{quote(safe)}"},
    )


def _validate_url(url: str) -> None:
    parsed = urlparse(url)
    if parsed.scheme not in {"http", "https"} or not parsed.netloc:
        raise HTTPException(status_code=400, detail="Only http/https URLs are supported")


def _safe_url_for_log(url: str) -> str:
    parsed = urlparse(url)
    if not parsed.scheme or not parsed.netloc:
        return url[:200]
    safe = f"{parsed.scheme}://{parsed.netloc}{parsed.path}"
    query_keys = sorted(parse_qs(parsed.query).keys())
    if query_keys:
        safe += "?keys=" + ",".join(query_keys[:12])
    return safe[:500]


def _normalize_notice_url(url: str) -> tuple[str, list[str]]:
    warnings: list[str] = []
    parsed = urlparse(url)
    host = parsed.netloc.lower()
    if host in {"www.szgzy.com", "szgzy.com"}:
        corrected = url.replace("://www.szgzy.com", "://www.szggzy.com").replace("://szgzy.com", "://szggzy.com")
        warnings.append(f"检测到疑似深圳公共资源交易网域名拼写错误，已自动修正为: {corrected}")
        return corrected, warnings
    return url, warnings


async def _fetch_page(url: str) -> dict[str, str]:
    async with httpx.AsyncClient(follow_redirects=True, timeout=30, headers={"User-Agent": USER_AGENT}) as client:
        resp = await client.get(url)
        resp.raise_for_status()
        return {"html": resp.text, "final_url": str(resp.url)}


async def _try_firecrawl(url: str, api_key: str | None, warnings: list[str]) -> str:
    if not api_key:
        return ""
    payload = {"url": url, "formats": ["markdown", "html"], "onlyMainContent": False}
    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
    try:
        async with httpx.AsyncClient(timeout=60) as client:
            resp = await client.post("https://api.firecrawl.dev/v1/scrape", json=payload, headers=headers)
            resp.raise_for_status()
            data = resp.json()
            item = data.get("data") or data
            return _normalize_text(str(item.get("markdown") or item.get("content") or ""))
    except Exception as exc:  # noqa: BLE001
        warnings.append(f"Firecrawl 抓取失败，已使用本地 HTML 解析兜底: {exc}")
        return ""


def _parse_html(html: str, base_url: str) -> dict[str, str]:
    soup = BeautifulSoup(html, "html.parser")
    for tag in soup(["script", "style", "noscript", "svg"]):
        tag.decompose()

    title = ""
    title_node = soup.select_one(".els-contentTitle, .contentTitle, .article-title, .detail-title, h1")
    if title_node:
        title = _normalize_text(title_node.get_text(" "))
    if soup.title and soup.title.string:
        title = title or _normalize_text(soup.title.string)
    if not title:
        h1 = soup.find("h1")
        title = _normalize_text(h1.get_text(" ")) if h1 else base_url

    content_node = _select_main_content_node(soup)
    text = _normalize_text(content_node.get_text("\n") if content_node else (soup.body or soup).get_text("\n"))
    return {"title": title, "text": text}


def _analyze_v2_failure(
    *,
    source_url: str,
    run_id: str,
    error_type: str,
    message: str,
    warnings: list[str],
    errors: list[str],
) -> AnalyzeV2Response:
    logger.warning(
        "analyze_v2_failed run_id=%s url=%s error_type=%s warnings=%s errors=%s message=%s",
        run_id,
        _safe_url_for_log(source_url),
        error_type,
        len(warnings),
        len(errors),
        message[:300],
    )
    return AnalyzeV2Response(
        success=False,
        run_id=run_id,
        source_url=source_url,
        fetch_status={
            "main_page": "failed",
            "attachments_found": 0,
            "attachments_downloaded": 0,
            "attachments_parsed": 0,
            "warnings": warnings,
            "errors": errors,
        },
        evidence_pack={"warnings": warnings},
        llm_input_text="",
        error_type=error_type,
        message=message,
    )


def _classify_fetch_error(exc: Exception | None) -> str:
    if isinstance(exc, httpx.HTTPStatusError):
        status_code = exc.response.status_code
        if status_code in {401, 403}:
            return "fetch_forbidden"
        if status_code == 404:
            return "empty_content"
    if isinstance(exc, (httpx.TimeoutException, TimeoutError)):
        return "timeout"
    return "unknown"


def _fetch_error_user_message(error_type: str, exc: Exception | None) -> str:
    if error_type == "fetch_forbidden":
        return "无法抓取该公告页面，目标网站可能限制访问或需要登录。请确认链接可公开访问，或上传原始公告附件。"
    if error_type == "timeout":
        return "公告页面抓取超时。请稍后重试，或上传原始 Word/PDF/Excel 附件。"
    if error_type == "empty_content":
        return "公告页面不存在或未返回有效内容。请检查 URL。"
    return f"无法抓取公告页面。请检查 URL 或上传原始附件。错误信息：{exc or 'unknown'}"


def _guess_source_agency(html: str, page_text: str) -> str:
    text = _normalize_text("\n".join([html, page_text]))
    patterns = [
        r"(?:信息来源|发布机构|发布单位|来源)\s*[:：]\s*([^\n<。；;]{2,80})",
        r"([\u4e00-\u9fff]{2,30}(?:医疗保障局|公共资源交易中心|医保局|采购中心|服务平台))",
    ]
    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            return _clean_inline_text(match.group(1))[:80]
    return ""


def _build_evidence_pack(
    *,
    source_url: str,
    final_url: str,
    title: str,
    page_text: str,
    notice_summary: dict[str, Any],
    attachments: list[dict[str, Any]],
    warnings: list[str],
) -> dict[str, Any]:
    mandatory_facts = _extract_mandatory_facts(title, page_text)
    key_rules = _extract_key_rules(page_text)
    attachment_items: list[dict[str, Any]] = []
    tables_summary: list[dict[str, Any]] = []
    source_index = [
        {"ref": "main_page", "type": "html", "title": title, "url": final_url},
    ]

    for index, attachment in enumerate(attachments, start=1):
        ref = f"attachment_{index}"
        table_summaries = [_summarize_table(table, ref) for table in attachment.get("tables") or []]
        tables_summary.extend(table_summaries)
        attachment_warnings = list(attachment.get("warnings") or [])
        if any((summary.get("row_count") or 0) > 30_000 for summary in table_summaries):
            attachment_warnings.append("大表已结构化摘要，完整明细以原始附件为准。")
        summary_text = _summarize_attachment_text(str(attachment.get("text") or ""))
        attachment_items.append(
            {
                "file_name": attachment.get("filename") or "",
                "url": attachment.get("url") or "",
                "file_type": Path(str(attachment.get("filename") or "")).suffix.lower().lstrip("."),
                "size": attachment.get("size_bytes") or 0,
                "parse_status": "success",
                "error": "",
                "summary": summary_text,
                "tables_summary": table_summaries,
                "evidence_refs": [ref],
                "warnings": attachment_warnings,
            }
        )
        source_index.append(
            {
                "ref": ref,
                "type": "attachment",
                "title": attachment.get("filename") or "",
                "url": attachment.get("url") or "",
            }
        )

    return {
        "source_url": source_url,
        "final_url": final_url,
        "notice_summary": notice_summary,
        "mandatory_facts": mandatory_facts,
        "key_rules": key_rules,
        "tables_summary": tables_summary,
        "attachments": attachment_items,
        "source_index": source_index,
        "warnings": warnings,
        "main_page_excerpt": _truncate(page_text, 8_000),
    }


def _extract_mandatory_facts(title: str, page_text: str) -> list[dict[str, str]]:
    text = _normalize_text("\n".join([title, page_text]))
    facts: list[dict[str, str]] = []
    for label, pattern in [
        ("title", r"^(.{4,120})"),
        ("publish_date", r"(20\d{2}[-/.年]\d{1,2}[-/.月]\d{1,2}日?)"),
        ("deadline", r"(20\d{2}[-/.年]\d{1,2}[-/.月]\d{1,2}日?[^。\n]{0,40}(?:前|截止|止))"),
    ]:
        match = re.search(pattern, text)
        if match:
            facts.append({"name": label, "value": _clean_inline_text(match.group(1)), "source_ref": "main_page"})
    return facts


def _extract_key_rules(page_text: str) -> list[dict[str, str]]:
    rules: list[dict[str, str]] = []
    paragraphs = [item.strip() for item in re.split(r"[\n。；;]", page_text) if len(item.strip()) >= 12]
    markers = [
        "采购",
        "申报",
        "报价",
        "挂网",
        "中选",
        "价格",
        "执行",
        "联动",
        "信用",
        "附件",
    ]
    for paragraph in paragraphs:
        if any(marker in paragraph for marker in markers):
            rules.append({"text": _truncate(_clean_inline_text(paragraph), 260), "source_ref": "main_page"})
        if len(rules) >= 40:
            break
    return rules


def _summarize_attachment_text(text: str) -> str:
    clean = _normalize_text(text)
    if not clean:
        return ""
    return _truncate(clean, 2_000)


def _summarize_table(table: dict[str, Any], source_ref: str) -> dict[str, Any]:
    rows = table.get("rows") or []
    if not isinstance(rows, list):
        rows = []
    headers = [str(value) for value in (table.get("columns") or (rows[0] if rows else []))]
    body = rows[1:] if rows and rows[0] == headers else rows[1:] if rows else []
    row_count = int(table.get("row_count") or len(rows))
    column_count = int(table.get("column_count") or max((len(row) for row in rows), default=len(headers)))
    warnings: list[str] = []
    if row_count > 5_000:
        warnings.append("超过5000行的大表已仅保留摘要、统计、关键行和业务口径。")
    if row_count > 30_000:
        warnings.append("大表已结构化摘要，完整明细以原始附件为准。")
    price_values = _detect_price_values(headers, body)
    important_rows = _select_important_rows(headers, body, limit=50 if row_count <= 30_000 else 20)
    return {
        "source_ref": source_ref,
        "sheet_name": table.get("sheet") or table.get("name") or "",
        "row_count": row_count,
        "column_count": column_count,
        "columns": headers[:80],
        "detected_business_columns": _detect_business_columns(headers),
        "product_count": _count_distinct_by_header(headers, body, ["产品", "品种", "耗材", "product"]),
        "company_count": _count_distinct_by_header(headers, body, ["企业", "公司", "company"]),
        "price_min": min(price_values) if price_values else None,
        "price_max": max(price_values) if price_values else None,
        "price_avg": round(sum(price_values) / len(price_values), 4) if price_values else None,
        "missing_value_summary": _missing_value_summary(headers, body),
        "category_summary": _category_summary(headers, body),
        "important_rows": important_rows,
        "warnings": warnings,
    }


def _detect_business_columns(headers: list[str]) -> list[str]:
    markers = ["产品", "品种", "企业", "注册证", "医保", "价格", "报价", "分类", "规格", "型号", "采购量", "省份"]
    return [header for header in headers if any(marker.lower() in header.lower() for marker in markers)]


def _header_index(headers: list[str], candidates: list[str]) -> int | None:
    for index, header in enumerate(headers):
        lower = header.lower()
        if any(candidate.lower() in lower for candidate in candidates):
            return index
    return None


def _count_distinct_by_header(headers: list[str], rows: list[list[Any]], candidates: list[str]) -> int:
    index = _header_index(headers, candidates)
    if index is None:
        return 0
    values = {str(row[index]).strip() for row in rows if index < len(row) and str(row[index]).strip()}
    return len(values)


def _detect_price_values(headers: list[str], rows: list[list[Any]]) -> list[float]:
    price_indexes = [index for index, header in enumerate(headers) if any(marker in header.lower() for marker in ["price", "价格", "报价", "限价"])]
    values: list[float] = []
    for row in rows:
        for index in price_indexes:
            if index >= len(row):
                continue
            match = re.search(r"-?\d+(?:\.\d+)?", str(row[index]).replace(",", ""))
            if match:
                try:
                    values.append(float(match.group(0)))
                except ValueError:
                    continue
    return values[:20_000]


def _missing_value_summary(headers: list[str], rows: list[list[Any]]) -> dict[str, int]:
    summary: dict[str, int] = {}
    for index, header in enumerate(headers[:40]):
        missing = 0
        for row in rows[:5_000]:
            if index >= len(row) or str(row[index]).strip() == "":
                missing += 1
        if missing:
            summary[header or f"column_{index + 1}"] = missing
    return summary


def _category_summary(headers: list[str], rows: list[list[Any]]) -> dict[str, int]:
    index = _header_index(headers, ["分类", "类别", "category"])
    if index is None:
        return {}
    counts: dict[str, int] = {}
    for row in rows[:20_000]:
        if index < len(row):
            value = str(row[index]).strip()
            if value:
                counts[value] = counts.get(value, 0) + 1
    return dict(sorted(counts.items(), key=lambda item: item[1], reverse=True)[:30])


def _select_important_rows(headers: list[str], rows: list[list[Any]], limit: int) -> list[dict[str, str]]:
    selected: list[list[Any]] = []
    for row in rows:
        if any(str(cell).strip() for cell in row):
            selected.append(row)
        if len(selected) >= limit:
            break
    result: list[dict[str, str]] = []
    for row in selected:
        item = {
            (headers[index] if index < len(headers) and headers[index] else f"column_{index + 1}"): _clean_inline_text(cell)
            for index, cell in enumerate(row[: len(headers) or len(row)])
            if _clean_inline_text(cell)
        }
        if item:
            result.append(item)
    return result


def _stable_content_hash(source_url: str, page_text: str, attachments: list[dict[str, Any]]) -> str:
    payload = {
        "source_url": source_url,
        "page_hash": hashlib.sha256(_normalize_text(page_text).encode("utf-8", errors="ignore")).hexdigest(),
        "attachments": [
            {
                "filename": item.get("filename") or "",
                "url": item.get("url") or "",
                "size": item.get("size_bytes") or 0,
                "text_hash": hashlib.sha256(str(item.get("text") or "").encode("utf-8", errors="ignore")).hexdigest(),
            }
            for item in attachments
        ],
    }
    return hashlib.sha256(json.dumps(payload, ensure_ascii=False, sort_keys=True).encode("utf-8")).hexdigest()[:32]


def _evidence_cache_dir() -> Path:
    return Path(os.getenv("EVIDENCE_CACHE_DIR", ".cache/evidence"))


def _write_evidence_cache(evidence_id: str, payload: dict[str, Any]) -> None:
    try:
        cache_dir = _evidence_cache_dir()
        cache_dir.mkdir(parents=True, exist_ok=True)
        (cache_dir / f"{evidence_id}.json").write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    except Exception:
        return


def _build_llm_input_text(evidence_pack: dict[str, Any], max_chars: int) -> str:
    sections = [
        "# evidence_pack",
        json.dumps(
            {
                "source_url": evidence_pack.get("source_url"),
                "notice_summary": evidence_pack.get("notice_summary"),
                "mandatory_facts": evidence_pack.get("mandatory_facts"),
                "key_rules": evidence_pack.get("key_rules"),
                "tables_summary": evidence_pack.get("tables_summary"),
                "attachments": evidence_pack.get("attachments"),
                "source_index": evidence_pack.get("source_index"),
                "warnings": evidence_pack.get("warnings"),
                "main_page_excerpt": evidence_pack.get("main_page_excerpt"),
            },
            ensure_ascii=False,
            indent=2,
        ),
    ]
    return _truncate("\n".join(sections), max_chars)


def _select_main_content_node(soup: BeautifulSoup):
    selectors = [
        ".els-messageinfo-box",
        ".detailPage",
        ".els-contentCon",
        ".contentCon",
        ".article-content",
        ".detail-content",
        ".TRS_Editor",
        "article",
    ]
    for selector in selectors:
        node = soup.select_one(selector)
        if node and len(_normalize_text(node.get_text(" "))) >= 80:
            return node
    return None


async def _try_szggzy_detail(url: str, warnings: list[str]) -> dict[str, str] | None:
    parsed = urlparse(url)
    if parsed.netloc.lower() not in {"www.szggzy.com", "szggzy.com"}:
        return None
    content_id = (parse_qs(parsed.query).get("contentId") or [""])[0]
    if not content_id:
        return None

    api_url = f"{parsed.scheme}://{parsed.netloc}/cms/api/v1/trade/content/detail?contentId={content_id}"
    try:
        async with httpx.AsyncClient(follow_redirects=True, timeout=30, headers={"User-Agent": USER_AGENT, "Referer": url}) as client:
            resp = await client.get(api_url)
            resp.raise_for_status()
            data = resp.json()
    except Exception as exc:  # noqa: BLE001
        warnings.append(f"深圳公共资源交易网详情接口读取失败: {api_url} ({exc})")
        return None

    if data.get("code") != 200 or not data.get("data"):
        warnings.append(f"深圳公共资源交易网详情接口未返回有效数据: {api_url}")
        return None

    item = data["data"]
    txt = str(item.get("txt") or "")
    release_time = str(item.get("releaseTime") or "")
    title = str(item.get("title") or "")
    source = str(item.get("source") or "")
    node_list = str(item.get("nodeList") or "")
    metadata_html = (
        f"<h1>{title}</h1>"
        f"<p>发布时间：{release_time}</p>"
        f"<p>信息来源：{source}</p>"
        f"<p>栏目：{node_list}</p>"
    )
    return {
        "title": title,
        "release_time": release_time,
        "source": source,
        "node_list": node_list,
        "html": metadata_html + txt,
    }


def _extract_xjbtylbz_news_id(url: str) -> str:
    parsed = urlparse(url)
    if parsed.netloc.lower() != "gwt.xjbtylbz.cn":
        return ""
    route = parsed.fragment or parsed.path
    match = re.search(r"(?:^|/)news/(\d+)(?:\D|$)", route)
    return match.group(1) if match else ""


def _xjbtylbz_latest_info_to_detail(item: dict[str, Any]) -> dict[str, str]:
    title = str(item.get("title") or "")
    release_time = str(item.get("crterTime") or "")
    content = str(item.get("content") or "")
    source = "\u65b0\u7586\u751f\u4ea7\u5efa\u8bbe\u5175\u56e2\u533b\u7597\u4fdd\u969c\u5c40\u7f51\u4e0a\u670d\u52a1\u5927\u5385"
    node_list = "\u62db\u91c7\u901a\u77e5"
    metadata_html = (
        f"<h1>{html_lib.escape(title)}</h1>"
        f"<p>\u53d1\u5e03\u65f6\u95f4: {html_lib.escape(release_time)}</p>"
        f"<p>\u4fe1\u606f\u6765\u6e90: {html_lib.escape(source)}</p>"
        f"<p>\u680f\u76ee: {html_lib.escape(node_list)}</p>"
    )
    return {
        "title": title,
        "release_time": release_time,
        "source": source,
        "node_list": node_list,
        "html": metadata_html + content,
    }


async def _try_xjbtylbz_detail(url: str, warnings: list[str]) -> dict[str, str] | None:
    news_id = _extract_xjbtylbz_news_id(url)
    if not news_id:
        return None

    parsed = urlparse(url)
    api_url = f"{parsed.scheme}://{parsed.netloc}/hsa-pss-admin/latestInfo/getAllLatestInfo"
    headers = {
        "User-Agent": USER_AGENT,
        "Origin": f"{parsed.scheme}://{parsed.netloc}",
        "Referer": f"{parsed.scheme}://{parsed.netloc}/hallEnter/",
    }
    try:
        async with httpx.AsyncClient(follow_redirects=True, timeout=30, headers=headers) as client:
            resp = await client.post(api_url, json={"infoId": news_id})
            resp.raise_for_status()
            data = resp.json()
    except Exception as exc:  # noqa: BLE001
        warnings.append(f"\u5175\u56e2\u533b\u4fdd\u62db\u91c7\u5e73\u53f0\u8be6\u60c5\u63a5\u53e3\u8bfb\u53d6\u5931\u8d25: {api_url} ({exc})")
        return None

    rows = ((data.get("data") or {}).get("rows") or []) if isinstance(data, dict) else []
    if str(data.get("code")) not in {"0", "200"} or not rows:
        warnings.append(f"\u5175\u56e2\u533b\u4fdd\u62db\u91c7\u5e73\u53f0\u8be6\u60c5\u63a5\u53e3\u672a\u8fd4\u56de\u6709\u6548\u6570\u636e: {api_url}")
        return None
    return _xjbtylbz_latest_info_to_detail(rows[0])


def _extract_elian_qx_article_id(url: str) -> str:
    parsed = urlparse(url)
    if parsed.netloc.lower() != "qx.eliancloud.cn":
        return ""
    query = parse_qs(parsed.query)
    return (query.get("articleid") or query.get("articleId") or [""])[0].strip()


def _elian_qx_headers(url: str) -> dict[str, str]:
    headers = {
        "User-Agent": USER_AGENT,
        "Content-Type": "application/json",
        "Referer": url,
    }
    token = (os.getenv("ELIAN_QX_TOKEN") or "").strip()
    cookie = (os.getenv("ELIAN_QX_COOKIE") or "").strip()
    if token:
        headers["Authorization"] = token if token.lower().startswith("bearer ") else f"Bearer {token}"
    if cookie:
        headers["Cookie"] = cookie
    return headers


def _is_elian_qx_auth_failed(data: Any) -> bool:
    if not isinstance(data, dict):
        return False
    code = str(data.get("code") or "")
    message = str(data.get("msg") or data.get("message") or "")
    return code == "401" or any(marker in message for marker in ["认证失败", "登录状态已过期", "无效的会话", "无法访问系统资源"])


def _elian_qx_article_to_detail(item: dict[str, Any]) -> dict[str, str]:
    title = str(item.get("title") or item.get("articleTitle") or item.get("noticeTitle") or "").strip()
    release_time = str(
        item.get("audittime")
        or item.get("auditTime")
        or item.get("releaseTime")
        or item.get("publishTime")
        or item.get("createTime")
        or ""
    ).strip()
    source = str(item.get("source") or item.get("sourceName") or item.get("author") or "器械集采准入系统").strip()
    node_list = str(item.get("noticeCategory") or item.get("category") or item.get("columnName") or "").strip()
    content = str(item.get("content") or item.get("articleContent") or item.get("txt") or item.get("body") or "")
    attachments = (
        item.get("attachments")
        or item.get("attachList")
        or item.get("articleAttachList")
        or item.get("bpArticleAttachList")
        or []
    )

    attachment_html: list[str] = []
    if isinstance(attachments, list):
        for attachment in attachments:
            if not isinstance(attachment, dict):
                continue
            name = str(
                attachment.get("filename")
                or attachment.get("fileName")
                or attachment.get("name")
                or attachment.get("attachName")
                or ""
            ).strip()
            href = str(attachment.get("url") or attachment.get("fileUrl") or attachment.get("downloadUrl") or "").strip()
            articleattid = str(attachment.get("articleattid") or attachment.get("articleAttachId") or "").strip()
            if not href and articleattid:
                href = f"https://qx.eliancloud.cn/prod-api/common/bpArticleattach/download?articleattid={quote(articleattid)}"
            if name or href:
                label = html_lib.escape(name or href)
                attachment_html.append(f'<p>附件: <a href="{html_lib.escape(href)}">{label}</a></p>' if href else f"<p>附件: {label}</p>")

    metadata_html = (
        f"<h1>{html_lib.escape(title)}</h1>"
        f"<p>发布时间: {html_lib.escape(release_time)}</p>"
        f"<p>信息来源: {html_lib.escape(source)}</p>"
        f"<p>栏目: {html_lib.escape(node_list)}</p>"
    )
    return {
        "title": title,
        "release_time": release_time,
        "source": source,
        "node_list": node_list,
        "html": metadata_html + content + "".join(attachment_html),
    }


async def _try_elian_qx_detail(url: str, warnings: list[str]) -> dict[str, str] | None:
    article_id = _extract_elian_qx_article_id(url)
    if not article_id:
        return None

    if not os.getenv("ELIAN_QX_TOKEN") and not os.getenv("ELIAN_QX_COOKIE"):
        if os.getenv("ELIAN_QX_USERNAME") or os.getenv("ELIAN_QX_PASSWORD"):
            warnings.append("公司器械集采准入系统检测到账号密码环境变量，但该站点登录含验证码/加密流程；请提供 ELIAN_QX_TOKEN 或 ELIAN_QX_COOKIE。")

    endpoints = [
        ("/prod-api/common/bpArticleContent/queryNewContent", {"articleid": article_id, "newHospitalInfoFlag": 1}),
        ("/prod-api/common/bpArticleContent/queryContent", {"articleid": article_id}),
        ("/prod-api/gwzr/zczxblarticlefinal/queryNewHospitalInformationDetail", {"articleid": article_id}),
    ]
    auth_failed = False
    async with httpx.AsyncClient(follow_redirects=True, timeout=30, headers=_elian_qx_headers(url)) as client:
        for endpoint, payload in endpoints:
            api_url = f"https://qx.eliancloud.cn{endpoint}"
            try:
                resp = await client.post(api_url, json=payload)
                resp.raise_for_status()
                data = resp.json()
            except Exception as exc:  # noqa: BLE001
                warnings.append(f"公司器械集采准入系统详情接口读取失败: {api_url} ({exc})")
                continue
            if _is_elian_qx_auth_failed(data):
                auth_failed = True
                continue
            item = data.get("data") if isinstance(data, dict) else None
            if isinstance(item, dict):
                detail = _elian_qx_article_to_detail(item)
                if detail.get("title") or _normalize_text(BeautifulSoup(detail.get("html") or "", "html.parser").get_text(" ")):
                    return detail
            warnings.append(f"公司器械集采准入系统详情接口未返回有效正文: {api_url}")
    if auth_failed:
        warnings.append("抓取状态: login_required；公司器械集采准入系统详情接口需要登录态，未能取得公告正文。")
    return None


async def _try_known_site_detail(url: str, warnings: list[str]) -> dict[str, str] | None:
    site_detail = await _try_szggzy_detail(url, warnings)
    if site_detail:
        return site_detail
    site_detail = await _try_xjbtylbz_detail(url, warnings)
    if site_detail:
        return site_detail
    return await _try_elian_qx_detail(url, warnings)


def _crawl_status_from_warnings(warnings: list[str]) -> str:
    text = "\n".join(warnings)
    if "login_required" in text:
        return "login_required"
    if "crawl_insufficient" in text:
        return "crawl_insufficient"
    return ""


def _looks_like_guizhou_dynamic_shell(url: str, title: str, page_text: str, attachments: list[dict[str, Any]]) -> bool:
    parsed = urlparse(url)
    compact = re.sub(r"\s+", "", f"{title}\n{page_text}")
    return (
        parsed.netloc.lower() == "fuwu.pubs.ylbzj.guizhou.gov.cn"
        and "hsa-pass-hallEnter" in url
        and not attachments
        and ("贵州医疗保障公共服务平台" in compact or len(compact) < 300)
    )


def _discover_attachment_links(html: str, base_url: str, limit: int) -> list[dict[str, str]]:
    if limit <= 0:
        return []

    soup = BeautifulSoup(html, "html.parser")
    candidates: list[dict[str, str]] = []
    seen: set[str] = set()

    for a in soup.find_all("a", href=True):
        href = str(a.get("href") or "").strip()
        if not href or href.startswith(("javascript:", "mailto:", "#")):
            continue
        abs_url = urljoin(base_url, href)
        text = _normalize_text(a.get_text(" "))
        parsed_link = urlparse(abs_url)
        lower_path = parsed_link.path.lower()
        query_values = " ".join(value for values in parse_qs(parsed_link.query).values() for value in values).lower()
        suffix = Path(lower_path).suffix
        query_has_file = any(ext in query_values for ext in ATTACHMENT_EXTENSIONS)
        looks_like_attachment = (
            suffix in ATTACHMENT_EXTENSIONS
            or query_has_file
            or any(word in text for word in ["附件", "下载", "表", "清单", "目录", "采购文件", "承诺书", "采购需求量"])
        )
        if text.strip() in {"附件下载", "下载", "下载中心"} and suffix not in ATTACHMENT_EXTENSIONS and not query_has_file:
            looks_like_attachment = False
        if looks_like_attachment and abs_url not in seen:
            seen.add(abs_url)
            candidates.append({"url": abs_url, "text": text})
        if len(candidates) >= limit:
            break

    return candidates


def _attachment_request_headers(url: str) -> dict[str, str]:
    headers = {"User-Agent": USER_AGENT}
    parsed = urlparse(url)
    if parsed.netloc.lower() == "gwt.xjbtylbz.cn:20002":
        headers["Referer"] = "https://gwt.xjbtylbz.cn/hallEnter/"
        headers["Origin"] = "https://gwt.xjbtylbz.cn"
    return headers


async def _collect_attachments(
    links: list[dict[str, str]],
    warnings: list[str],
    concurrency: int | None = None,
) -> list[dict[str, Any]]:
    if not links:
        return []
    worker_count = max(1, min(concurrency or int(os.getenv("ATTACHMENT_CONCURRENCY", "4")), len(links)))
    semaphore = asyncio.Semaphore(worker_count)

    async def collect(link: dict[str, str]) -> dict[str, Any] | None:
        async with semaphore:
            try:
                downloaded = await _download_attachment(link["url"])
                parsed = await asyncio.to_thread(_parse_attachment, downloaded)
                return {
                    "url": downloaded.url,
                    "filename": downloaded.filename,
                    "link_text": link.get("text", ""),
                    "content_type": downloaded.content_type,
                    "size_bytes": len(downloaded.content),
                    "text": parsed["text"],
                    "tables": parsed["tables"],
                    "warnings": parsed["warnings"],
                }
            except Exception as exc:  # noqa: BLE001 - keep extraction partial instead of failing the workflow
                logger.warning(
                    "attachment_failed url=%s error_type=%s error=%s",
                    _safe_url_for_log(link["url"]),
                    exc.__class__.__name__,
                    str(exc)[:300],
                )
                warnings.append(f"附件读取失败: {link['url']} ({exc})")
                return None

    results = await asyncio.gather(*(collect(link) for link in links))
    return [item for item in results if item is not None]


async def _download_attachment(url: str) -> DownloadedFile:
    _validate_url(url)
    cached = _load_cached_attachment(url)
    if cached is not None:
        return cached
    async with httpx.AsyncClient(follow_redirects=True, timeout=90, headers=_attachment_request_headers(url)) as client:
        resp = await client.get(url)
        resp.raise_for_status()
        content = resp.content
        if len(content) > MAX_ATTACHMENT_BYTES:
            logger.warning(
                "attachment_too_large url=%s size_bytes=%s limit_bytes=%s",
                _safe_url_for_log(str(resp.url)),
                len(content),
                MAX_ATTACHMENT_BYTES,
            )
            raise ValueError(f"附件超过大小限制 {MAX_ATTACHMENT_BYTES} bytes")
        filename = _filename_from_response(url, resp.headers)
        return DownloadedFile(url=str(resp.url), filename=filename, content_type=resp.headers.get("content-type", ""), content=content)


def _cache_key(url: str) -> str:
    return hashlib.sha1(url.encode("utf-8")).hexdigest()[:16]


def _load_cache_manifest(cache_dir: Path) -> dict[str, Any] | None:
    manifest_path = cache_dir / "manifest.json"
    if not manifest_path.exists():
        return None
    try:
        return json.loads(manifest_path.read_text(encoding="utf-8"))
    except Exception:
        return None


def _load_cached_page(url: str, warnings: list[str]) -> dict[str, str] | None:
    cache_dir = SITE_CACHE_DIR / _cache_key(url)
    manifest = _load_cache_manifest(cache_dir)
    if not manifest:
        return None
    page_file = cache_dir / str(manifest.get("page_file") or "page.html")
    if not page_file.exists():
        return None
    html = page_file.read_text(encoding="utf-8", errors="ignore")
    warnings.append(f"网页在线抓取失败，已使用本地缓存: {cache_dir.name}")
    return {"html": html, "final_url": str(manifest.get("final_url") or url)}


def _load_cached_attachment(url: str) -> DownloadedFile | None:
    for cache_dir in SITE_CACHE_DIR.iterdir() if SITE_CACHE_DIR.exists() else []:
        if not cache_dir.is_dir():
            continue
        manifest = _load_cache_manifest(cache_dir)
        if not manifest:
            continue
        attachment = (manifest.get("attachments") or {}).get(url)
        if not attachment:
            continue
        path = cache_dir / str(attachment.get("file") or "")
        if path.exists() and path.is_file():
            return DownloadedFile(
                url=url,
                filename=str(attachment.get("filename") or path.name),
                content_type=str(attachment.get("content_type") or ""),
                content=path.read_bytes(),
            )
    return None


def _filename_from_response(url: str, headers: httpx.Headers) -> str:
    disposition = headers.get("content-disposition", "")
    match = re.search(r"filename\*?=(?:UTF-8''|\"?)([^\";]+)", disposition, re.I)
    if match:
        return Path(match.group(1)).name
    path_name = Path(urlparse(url).path).name
    if path_name:
        return path_name
    digest = hashlib.sha1(url.encode("utf-8")).hexdigest()[:10]
    return f"attachment-{digest}"


def _parse_attachment(file: DownloadedFile) -> dict[str, Any]:
    suffix = Path(file.filename.lower()).suffix
    warnings: list[str] = []
    tables: list[dict[str, Any]] = []

    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(file.content)
        tmp_path = Path(tmp.name)

    try:
        if suffix == ".pdf":
            text, tables = _parse_pdf(tmp_path, warnings)
        elif suffix == ".docx":
            text, tables = _parse_docx(tmp_path)
        elif suffix == ".doc":
            text, tables = _parse_doc(tmp_path)
        elif suffix in {".xlsx", ".xlsm"}:
            text, tables = _parse_xlsx(tmp_path)
        elif suffix == ".xls":
            text, tables = _parse_xls(tmp_path)
        elif suffix == ".csv":
            text, tables = _parse_csv(file.content)
        else:
            text = _normalize_text(file.content.decode("utf-8", errors="ignore"))
        return {"text": text, "tables": tables, "warnings": warnings}
    finally:
        try:
            tmp_path.unlink(missing_ok=True)
        except Exception:
            pass


def _parse_pdf(path: Path, warnings: list[str]) -> tuple[str, list[dict[str, Any]]]:
    parts: list[str] = []
    tables: list[dict[str, Any]] = []
    if pdfplumber is not None:
        with pdfplumber.open(path) as pdf:
            for page_index, page in enumerate(pdf.pages, start=1):
                text = page.extract_text() or ""
                if text.strip():
                    parts.append(f"[PDF 第 {page_index} 页]\n{_normalize_text(text)}")
                try:
                    for table_index, table in enumerate(page.extract_tables() or [], start=1):
                        rows = _clean_rows(table)
                        if rows:
                            tables.append({"sheet": f"page-{page_index}", "name": f"table-{table_index}", "rows": rows})
                except Exception as exc:  # noqa: BLE001
                    warnings.append(f"PDF 第 {page_index} 页表格提取失败: {exc}")
        return "\n\n".join(parts), tables

    reader = PdfReader(str(path))
    for page_index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            parts.append(f"[PDF 第 {page_index} 页]\n{_normalize_text(text)}")
    warnings.append("当前运行镜像未安装 pdfplumber，PDF 表格仅作为页面文本抽取")
    return "\n\n".join(parts), tables


def _parse_docx(path: Path) -> tuple[str, list[dict[str, Any]]]:
    doc = Document(str(path))
    parts = [_normalize_text(p.text) for p in doc.paragraphs if _normalize_text(p.text)]
    tables: list[dict[str, Any]] = []
    for index, table in enumerate(doc.tables, start=1):
        rows = [[_normalize_text(cell.text) for cell in row.cells] for row in table.rows]
        rows = _clean_rows(rows)
        if rows:
            tables.append({"sheet": "document", "name": f"table-{index}", "rows": rows})
            parts.append(_rows_to_markdown(rows, f"Word 表格 {index}"))
    return "\n\n".join(parts), tables


def _parse_doc(path: Path) -> tuple[str, list[dict[str, Any]]]:
    result = subprocess.run(["antiword", str(path)], capture_output=True, text=True, timeout=60, check=False)
    if result.returncode != 0:
        raise ValueError(result.stderr.strip() or "antiword failed to parse .doc")
    return _normalize_text(result.stdout), []


def _parse_xlsx(path: Path) -> tuple[str, list[dict[str, Any]]]:
    wb = load_workbook(path, data_only=True, read_only=True)
    parts: list[str] = []
    tables: list[dict[str, Any]] = []
    for ws in wb.worksheets:
        rows = []
        for row in ws.iter_rows(values_only=True):
            rows.append([_cell_to_text(cell) for cell in row])
        rows = _clean_rows(rows)
        if rows:
            tables.append(
                {
                    "sheet": ws.title,
                    "name": ws.title,
                    "rows": rows[:500],
                    "row_count": len(rows),
                    "column_count": max((len(row) for row in rows), default=0),
                    "columns": rows[0] if rows else [],
                }
            )
            parts.append(_rows_to_markdown(rows[:80], f"Excel Sheet: {ws.title}"))
    return "\n\n".join(parts), tables


def _parse_xls(path: Path) -> tuple[str, list[dict[str, Any]]]:
    book = xlrd.open_workbook(str(path))
    parts: list[str] = []
    tables: list[dict[str, Any]] = []
    for sheet in book.sheets():
        rows = []
        for r in range(sheet.nrows):
            rows.append([_cell_to_text(sheet.cell_value(r, c)) for c in range(sheet.ncols)])
        rows = _clean_rows(rows)
        if rows:
            tables.append(
                {
                    "sheet": sheet.name,
                    "name": sheet.name,
                    "rows": rows[:500],
                    "row_count": len(rows),
                    "column_count": max((len(row) for row in rows), default=0),
                    "columns": rows[0] if rows else [],
                }
            )
            parts.append(_rows_to_markdown(rows[:80], f"Excel Sheet: {sheet.name}"))
    return "\n\n".join(parts), tables


def _parse_csv(content: bytes) -> tuple[str, list[dict[str, Any]]]:
    text = content.decode("utf-8-sig", errors="ignore")
    rows = _clean_rows(list(csv.reader(io.StringIO(text))))
    table = {
        "sheet": "csv",
        "name": "csv",
        "rows": rows[:500],
        "row_count": len(rows),
        "column_count": max((len(row) for row in rows), default=0),
        "columns": rows[0] if rows else [],
    }
    return _rows_to_markdown(rows[:120], "CSV"), [table] if rows else []


def _cell_to_text(value: Any) -> str:
    if value is None:
        return ""
    text = str(value)
    if text.endswith(".0") and text[:-2].isdigit():
        return text[:-2]
    return _normalize_text(text)


def _clean_rows(rows: list[list[Any]]) -> list[list[str]]:
    cleaned = []
    for row in rows:
        values = [_normalize_text(str(cell or "")) for cell in row]
        while values and values[-1] == "":
            values.pop()
        if any(values):
            cleaned.append(values)
    return cleaned


def _rows_to_markdown(rows: list[list[str]], title: str) -> str:
    if not rows:
        return ""
    width = max(len(row) for row in rows)
    normalized = [row + [""] * (width - len(row)) for row in rows]
    header = normalized[0]
    body = normalized[1:]
    lines = [f"### {title}", "| " + " | ".join(header) + " |", "| " + " | ".join(["---"] * width) + " |"]
    for row in body:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)


def _prepare_report_for_export(req: ExportReportRequest) -> ReportIR:
    if req.report_ir is not None:
        report = req.report_ir
    else:
        report = _parse_model_output_to_report_ir(req.markdown)
    report = _normalize_report_ir(report, fallback_title=req.title)
    return report


def _parse_model_output_to_report_ir(raw_text: str) -> ReportIR:
    tag_text = _remove_thinking_blocks(str(raw_text or ""))
    cleaned_text = _clean_model_output(tag_text)
    if not tag_text.strip() and not cleaned_text:
        raise ValueError("模型输出为空，无法导出报告")

    report_ir_errors: list[str] = []
    for report_ir_match in re.finditer(r"<report_ir\b[^>]*>([\s\S]*?)</report_ir>", tag_text, flags=re.I):
        try:
            return _report_ir_from_json(report_ir_match.group(1))
        except ValueError as exc:
            report_ir_errors.append(str(exc))

    stripped = _strip_json_code_fence(cleaned_text.strip())
    if stripped.startswith("{") and stripped.endswith("}"):
        try:
            return _report_ir_from_json(stripped)
        except ValueError:
            pass

    final_match = re.search(r"<final_report\b[^>]*>([\s\S]*?)</final_report>", tag_text, flags=re.I)
    if final_match:
        return _markdown_to_report_ir(final_match.group(1))

    if report_ir_errors:
        raise ValueError("ReportIR 解析失败：" + "；".join(report_ir_errors[:2]))
    raise ValueError("未找到 <report_ir> 结构化报告或 <final_report> 正文，已拒绝将模型原始回复直接写入 Word")


def _remove_thinking_blocks(text: str) -> str:
    patterns = [
        r"<think\b[^>]*>[\s\S]*?</think>",
        r"<思考\b[^>]*>[\s\S]*?</思考>",
        r"```(?:thinking|think|思考)[\s\S]*?```",
        r"`(?:thinking|think|思考)[\s\S]*?`",
    ]
    for pattern in patterns:
        text = re.sub(pattern, "", text, flags=re.I)
    return text


def _report_ir_from_json(text: str) -> ReportIR:
    payload = _strip_json_code_fence(text.strip())
    try:
        data = json.loads(payload)
    except json.JSONDecodeError as exc:
        repaired_payload = _repair_unescaped_quotes_in_json_strings(payload)
        if repaired_payload != payload:
            try:
                data = json.loads(repaired_payload)
            except json.JSONDecodeError:
                raise ValueError(f"ReportIR JSON 解析失败: {exc}") from exc
        else:
            raise ValueError(f"ReportIR JSON 解析失败: {exc}") from exc
    if isinstance(data, dict) and "report_ir" in data and isinstance(data["report_ir"], dict):
        data = data["report_ir"]
    if not isinstance(data, dict):
        raise ValueError("ReportIR JSON 顶层必须是对象")
    blocked_keys = {"reasoning", "analysis", "scratchpad", "thought", "debug", "raw_response", "chain_of_thought"}
    data = {key: value for key, value in data.items() if key not in blocked_keys}
    data = _coerce_report_ir_payload(data)
    try:
        return ReportIR.model_validate(data)
    except Exception as exc:  # noqa: BLE001
        raise ValueError(f"ReportIR 字段校验失败: {exc}") from exc


def _strip_json_code_fence(text: str) -> str:
    fence = re.match(r"^```(?:json)?\s*([\s\S]*?)\s*```$", text.strip(), flags=re.I)
    return fence.group(1).strip() if fence else text.strip()


def _repair_unescaped_quotes_in_json_strings(text: str) -> str:
    repaired: list[str] = []
    in_string = False
    escaped = False
    value = str(text or "")

    for index, char in enumerate(value):
        if not in_string:
            if char == '"':
                in_string = True
            repaired.append(char)
            continue

        if escaped:
            repaired.append(char)
            escaped = False
            continue

        if char == "\\":
            repaired.append(char)
            escaped = True
            continue

        if char == '"':
            if _looks_like_json_string_closer(value, index):
                in_string = False
                repaired.append(char)
            else:
                repaired.append('\\"')
            continue

        repaired.append(char)

    return "".join(repaired)


def _looks_like_json_string_closer(text: str, quote_index: int) -> bool:
    next_index, next_char = _next_non_whitespace_char(text, quote_index + 1)
    if not next_char:
        return True
    if next_char in {":", "}", "]"}:
        return True
    if next_char != ",":
        return False
    _, after_comma = _next_non_whitespace_char(text, next_index + 1)
    return not after_comma or after_comma in {'"', "{", "[", "}", "]"}


def _next_non_whitespace_char(text: str, start: int) -> tuple[int, str]:
    for index in range(start, len(text)):
        if not text[index].isspace():
            return index, text[index]
    return len(text), ""


def _coerce_report_ir_payload(data: dict[str, Any]) -> dict[str, Any]:
    coerced = dict(data)
    if "enterprise_tips" in coerced:
        coerced["enterprise_tips"] = _coerce_text_list(coerced["enterprise_tips"])
    return coerced


def _coerce_text_list(value: Any) -> list[str]:
    if value is None:
        return []
    items = value if isinstance(value, list) else [value]
    texts: list[str] = []
    for item in items:
        text = _coerce_text_item(item)
        if text:
            texts.append(text)
    return texts


def _coerce_text_item(item: Any) -> str:
    if isinstance(item, dict):
        for key in ("tip", "text", "content", "description", "value", "point", "summary"):
            value = item.get(key)
            if value not in (None, ""):
                return str(value)
        title = item.get("title") or item.get("heading")
        body = item.get("body") or item.get("detail") or item.get("details") or item.get("note")
        if title and body:
            return f"{title}：{body}"
        values = [str(value) for value in item.values() if value not in (None, "") and not isinstance(value, (dict, list))]
        return "；".join(values)
    return str(item or "")


def _markdown_to_report_ir(markdown: str) -> ReportIR:
    clean = _clean_model_output(markdown)
    lines = clean.splitlines()
    title = ""
    lead: list[str] = []
    sections: list[ReportSection] = []
    current: ReportSection | None = None
    index = 0

    def ensure_section() -> ReportSection:
        nonlocal current
        if current is None:
            current = ReportSection(heading="")
            sections.append(current)
        return current

    while index < len(lines):
        line = lines[index].rstrip()
        stripped = line.strip()
        if not stripped:
            index += 1
            continue
        heading_match = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading_match:
            heading = _strip_markdown(heading_match.group(2))
            if not title:
                title = heading
            else:
                current = ReportSection(heading=heading)
                sections.append(current)
            index += 1
            continue
        if stripped.startswith("|") and "|" in stripped[1:]:
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index].strip())
                index += 1
            table = _markdown_table_to_report_table(table_lines)
            if table:
                ensure_section().tables.append(table)
            continue
        if not title and _looks_like_report_title(stripped):
            title = _strip_markdown(stripped)
        elif current is None and len(lead) < 4:
            lead.append(_strip_markdown(stripped))
        else:
            ensure_section().paragraphs.append(_strip_markdown(stripped))
        index += 1

    return ReportIR(title=title, lead_paragraphs=lead, sections=sections)


def _markdown_table_to_report_table(table_lines: list[str]) -> ReportTable | None:
    rows: list[list[str]] = []
    for line in table_lines:
        cells = [_strip_markdown(cell.strip()) for cell in line.strip("|").split("|")]
        if cells and all(re.fullmatch(r":?-{3,}:?", cell or "") for cell in cells):
            continue
        rows.append(cells)
    if len(rows) < 2:
        return None
    width = max(len(row) for row in rows)
    normalized = [row + [""] * (width - len(row)) for row in rows]
    return ReportTable(headers=normalized[0], rows=normalized[1:])


def _normalize_report_ir(report: ReportIR, fallback_title: str) -> ReportIR:
    title = _clean_inline_text(report.title or fallback_title or DEFAULT_REPORT_TITLE)
    lead = [_clean_inline_text(item) for item in report.lead_paragraphs if _clean_inline_text(item)]
    sections: list[ReportSection] = []
    for section in report.sections:
        tables: list[ReportTable] = []
        for table in section.tables:
            headers = [_clean_inline_text(value) for value in table.headers]
            headers = [value for value in headers if value]
            rows = _normalize_table_rows(headers, table.rows)
            if headers and rows:
                tables.append(
                    ReportTable(
                        title=_clean_inline_text(table.title),
                        headers=headers,
                        rows=rows,
                        notes=[_clean_inline_text(note) for note in table.notes if _clean_inline_text(note)],
                    )
                )
        paragraphs = [_clean_inline_text(item) for item in section.paragraphs if _clean_inline_text(item)]
        highlights = [_clean_inline_text(item) for item in section.highlights if _clean_inline_text(item)]
        heading = _clean_inline_text(section.heading)
        if _is_glossary_heading(heading):
            paragraphs = _normalize_glossary_paragraphs(paragraphs)
        if heading or paragraphs or tables or highlights:
            sections.append(ReportSection(heading=heading, paragraphs=paragraphs, tables=tables, highlights=highlights))
    disclaimer = DEFAULT_DISCLAIMER
    suggested = _safe_filename(report.suggested_filename).removesuffix(".docx") if report.suggested_filename else ""
    return ReportIR(
        title=title,
        suggested_filename=suggested,
        notice_type=_clean_inline_text(report.notice_type),
        publish_date=_clean_inline_text(report.publish_date),
        source_agency=_clean_inline_text(report.source_agency),
        document_name=_clean_inline_text(report.document_name),
        lead_paragraphs=lead,
        sections=sections,
        enterprise_tips=[_clean_inline_text(item) for item in report.enterprise_tips if _clean_inline_text(item)],
        disclaimer=disclaimer,
    )


GLOSSARY_TERMS = [
    "同采购品种同类别同注册证",
    "非带量最低价",
    "带量最低价",
    "非中选产品",
    "地区最低价",
    "报价单元",
    "中选产品",
    "参考价",
]


def _is_glossary_heading(heading: str) -> bool:
    return "名词解释" in str(heading or "")


def _normalize_glossary_paragraphs(paragraphs: list[str]) -> list[str]:
    if not paragraphs:
        return []
    text = "\n".join(paragraphs)
    term_pattern = "|".join(re.escape(term) for term in sorted(GLOSSARY_TERMS, key=len, reverse=True))
    pattern = re.compile(rf"(?P<term>{term_pattern})(?=\s*(?:是指|指|定义为|[:：]))")
    matches = list(pattern.finditer(text))
    if not matches:
        return paragraphs

    normalized: list[str] = []
    intro = text[: matches[0].start()]
    normalized.extend(_clean_glossary_intro(intro))
    for index, match in enumerate(matches, start=1):
        end = matches[index].start() if index < len(matches) else len(text)
        raw_definition = text[match.end() : end]
        definition = _clean_glossary_definition(raw_definition)
        if definition:
            normalized.append(f"{index}. {match.group('term')}：{definition}")
    return normalized or paragraphs


def _clean_glossary_intro(text: str) -> list[str]:
    value = re.sub(r"^\s*\d+\s*[.、]\s*", "", str(text or "").strip())
    value = value.strip("。；;，, \n\t")
    if not value:
        return []
    return [value + "。"]


def _clean_glossary_definition(text: str) -> str:
    value = re.sub(r"^\s*\d+\s*[.、]\s*", "", str(text or "").strip())
    value = re.sub(r"^(是指|指|定义为)\s*", lambda m: m.group(1), value)
    value = value.lstrip(":： ")
    value = re.sub(r"\n+", " ", value)
    value = re.sub(r"\s+", " ", value).strip()
    return value


def _normalize_table_rows(headers: list[str], rows: list[list[Any]]) -> list[list[str]]:
    width = len(headers)
    if width <= 0:
        return []
    normalized: list[list[str]] = []
    for row in rows:
        raw_cells = list(row) if isinstance(row, list) else [row]
        cells = [_clean_inline_text(cell) for cell in raw_cells]
        if not any(cells):
            continue
        if len(cells) < width:
            cells.extend([""] * (width - len(cells)))
        elif len(cells) > width:
            cells = cells[: width - 1] + ["；".join(cell for cell in cells[width - 1 :] if cell)]
        normalized.append(cells)
    return normalized


def _clean_inline_text(text: Any) -> str:
    value = html_lib.unescape(str(text or ""))
    value = value.replace("\xa0", " ")
    value = re.sub(r"<br\s*/?>", "\n", value, flags=re.IGNORECASE)
    value = re.sub(r"\{\{[\s\S]*?\}\}", "", value)
    value = _clean_model_output(value)
    value = re.sub(r"(?im)^\s*(analysis|scratchpad|debug|chain-of-thought)\s*[:：].*$", "", value)
    value = re.sub(r"<[^>]+>", "", value)
    value = _strip_markdown(value)
    for source, replacement in MODELIZED_LANGUAGE_REPLACEMENTS.items():
        value = value.replace(source, replacement)
    value = re.sub(r"\n{3,}", "\n\n", value)
    value = re.sub(r"[ \t\r\f\v]+", " ", value)
    value = re.sub(r"《\s+", "《", value)
    value = re.sub(r"\s+》", "》", value)
    value = re.sub(r"“\s+", "“", value)
    value = re.sub(r"\s+”", "”", value)
    for phrase in ["严重失信", "第二轮接续", "医用耗材", "集中带量采购", "全国联审通办"]:
        value = re.sub(r"\s*".join(map(re.escape, phrase)), phrase, value)
    if re.fullmatch(r"\s*\{[\s\S]*\"(?:sections|report_ir|lead_paragraphs)\"[\s\S]*\}\s*", value):
        return ""
    return value.strip(" \n\r\t")


def _quality_check_report(report: ReportIR) -> list[str]:
    issues: list[str] = []
    text = _report_ir_text(report)
    tables = [table for section in report.sections for table in section.tables]

    if not report.title or len(report.title) < 6:
        issues.append("缺少清晰标题")
    if not report.lead_paragraphs and not any(section.paragraphs or section.tables or section.highlights for section in report.sections):
        issues.append("缺少正文内容")
    if not report.sections:
        issues.append("缺少报告章节")
    if re.search(r"<report_ir|</report_ir>|<final_report|</final_report>|```json|```|\{\{#?.*?\}\}", text, re.IGNORECASE):
        issues.append("报告中残留结构化标签、代码块或 Dify 变量")

    informal_terms = ["我认为", "大概", "应该是", "显然", "必然", "唯一选择", "全面利好", "重大颠覆"]
    found_terms = [term for term in informal_terms if term in text]
    if found_terms:
        issues.append("存在不适合正式报告的词语：" + "、".join(found_terms))

    blocked_terms = ["思考过程", "推理过程", "chain-of-thought", "analysis", "scratchpad", "debug", "<think", "</think"]
    found_blocked = [term for term in blocked_terms if term.lower() in text.lower()]
    if found_blocked:
        issues.append("存在模型思考或调试内容：" + "、".join(found_blocked))

    for table in tables:
        if not table.headers or any(not header.strip() for header in table.headers):
            issues.append(f"表格《{table.title or '未命名表格'}》存在空表头")
        width = len(table.headers)
        if any(len(row) != width for row in table.rows):
            issues.append(f"表格《{table.title or '未命名表格'}》存在行列错位")
    return issues


RICH_PROCUREMENT_KEYWORDS = [
    "采购文件",
    "集中带量采购",
    "接续采购",
    "最高有效申报价",
    "询价基准",
    "拟中选",
    "中选产品",
    "协议采购量",
    "价格联动",
    "非中选产品",
    "名词解释",
    "失信约束",
]
REPORT_DEPTH_MIN_CHARS = 1800
REPORT_DEPTH_MIN_SECTIONS = 4
PROCESS_NOTICE_MIN_SECTIONS = 3
PROCESS_NOTICE_REQUIRED_TOPICS = [
    ("调整内容", ["调整内容", "流程调整", "申报入口", "操作流程", "申报路径", "联审通办"]),
    ("影响分析", ["影响分析", "影响", "变化", "衔接", "申报节奏"]),
    ("企业关注点", ["企业关注", "操作建议", "关注点", "企业需", "提前核对", "申报材料"]),
]
EVIDENCE_MARKER_RE = re.compile(
    r"20\d{2}\s*年\s*\d{1,2}\s*月\s*\d{1,2}\s*日"
    r"|20\d{2}[-/.]\d{1,2}[-/.]\d{1,2}"
    r"|\d+(?:,\d{3})*(?:\.\d+)?\s*(?:元/个|元|万元|%|个工作日|年)"
    r"|[A-Z]{2,}-[A-Z0-9-]{3,}"
)
EVIDENCE_MARKER_SKIP_WORDS = ["未披露", "暂未公布", "以正式文件为准", "计算值", "按表内数据测算"]


def _extract_quality_evidence_text(evidence_text: str) -> str:
    value = str(evidence_text or "").strip()
    if not value:
        return ""
    try:
        data = json.loads(value)
    except Exception:  # noqa: BLE001 - evidence may be plain text.
        return value
    if not isinstance(data, dict):
        return value
    for key in ("evidence_for_llm", "combined_text", "page_text"):
        candidate = data.get(key)
        if isinstance(candidate, str) and candidate.strip():
            return candidate
    return value


def _quality_check_report_against_evidence(report: ReportIR, evidence_text: str, history_text: str = "") -> ReportQA:
    qa = ReportQA()
    report_text = _report_ir_text(report)
    clean_report_text = _clean_for_history_compare(report_text)
    prepared_evidence = _extract_quality_evidence_text(evidence_text)
    clean_evidence = _clean_for_history_compare(prepared_evidence)
    has_evidence = bool(clean_evidence)
    rich_evidence = has_evidence and _looks_like_rich_procurement_evidence(clean_evidence)
    process_notice = has_evidence and _looks_like_process_adjustment_notice(report, clean_evidence)

    crawl_status = _extract_crawl_status(clean_evidence)
    if crawl_status:
        qa.issues.append(
            ReportQAIssue(
                severity="major",
                category="crawl_insufficient",
                report_text=f"抓取状态: {crawl_status}",
                fix_instruction="当前链接未取得可用于正式分析的公告正文或附件；请提供可访问原文、公司站点登录态或人工上传原文附件后重新分析。",
            )
        )

    if rich_evidence:
        char_count = len(re.sub(r"\s+", "", clean_report_text))
        section_count = _report_quality_metrics(report)["meaningful_sections"]
        if char_count < REPORT_DEPTH_MIN_CHARS:
            qa.issues.append(
                ReportQAIssue(
                    severity="major",
                    category="report_depth",
                    report_text=f"报告正文约{char_count}字，证据较丰富但报告偏短。",
                    fix_instruction=(
                        f"补充导语、关键规则小节和分主题表格；采购文件类报告正文应接近人工分析稿，"
                        f"证据丰富时建议不少于{REPORT_DEPTH_MIN_CHARS}字。"
                    ),
                )
            )
        if section_count < REPORT_DEPTH_MIN_SECTIONS:
            qa.missing_rules.append(
                f"采购文件类报告结构偏少：当前有效小节{section_count}个，建议至少{REPORT_DEPTH_MIN_SECTIONS}个。"
            )
    elif process_notice:
        char_count = len(re.sub(r"\s+", "", clean_report_text))
        metrics = _report_quality_metrics(report)
        missing_topics = _missing_process_notice_topics(report_text)
        if metrics["meaningful_sections"] < PROCESS_NOTICE_MIN_SECTIONS or missing_topics:
            missing_text = "、".join(missing_topics) if missing_topics else "分主题小节"
            qa.issues.append(
                ReportQAIssue(
                    severity="major",
                    category="process_notice_depth",
                    report_text=(
                        f"流程调整类公告正文约{char_count}字，当前有效小节{metrics['meaningful_sections']}个，"
                        f"缺少{missing_text}。"
                    ),
                    fix_instruction=(
                        "流程调整/挂网操作调整类公告不能只写摘要；请补充调整内容、影响分析、"
                        f"企业关注点/操作建议，正文至少拆分为{PROCESS_NOTICE_MIN_SECTIONS}个有效小节。"
                    ),
                )
            )

    if has_evidence:
        qa.unsupported_claims.extend(_table_cells_missing_from_evidence(report, clean_evidence))
    qa.history_leakage.extend(_detect_history_leakage(clean_report_text, history_text, clean_evidence))

    if _qa_has_blocking_local_issues(qa):
        if crawl_status:
            qa.status = "block"
            qa.summary = "本地质检检测到抓取/数据读取问题，当前链接无法自动修复。"
        else:
            qa.status = "needs_fix"
            qa.summary = "本地质检发现报告深度、表格证据或历史边界问题。"
        for issue in [*qa.issues, *qa.unsupported_claims, *qa.history_leakage]:
            if issue.fix_instruction:
                qa.fix_instructions.append(issue.fix_instruction)
        for rule in qa.missing_rules:
            qa.fix_instructions.append(f"补充或拆分对应规则内容：{rule}")
        qa.fix_instructions = _unique([item for item in qa.fix_instructions if item])
    return qa


def _report_quality_metrics(report: ReportIR) -> dict[str, int]:
    sections = [
        section
        for section in report.sections
        if section.heading or section.paragraphs or section.tables or section.highlights
    ]
    return {
        "lead_paragraphs": len([item for item in report.lead_paragraphs if _clean_inline_text(item)]),
        "meaningful_sections": len(sections),
        "tables": sum(len(section.tables) for section in sections),
        "table_rows": sum(len(table.rows) for section in sections for table in section.tables),
    }


def _looks_like_rich_procurement_evidence(evidence_text: str) -> bool:
    text = str(evidence_text or "")
    hits = sum(1 for keyword in RICH_PROCUREMENT_KEYWORDS if keyword in text)
    return hits >= 4 and len(re.sub(r"\s+", "", text)) >= 300


def _looks_like_process_adjustment_notice(report: ReportIR, evidence_text: str) -> bool:
    text = "\n".join(
        [
            report.notice_type,
            report.title,
            report.document_name,
            report.source_agency,
            evidence_text,
        ]
    )
    compact = re.sub(r"\s+", "", text)
    process_hits = sum(
        1
        for keyword in [
            "流程调整",
            "操作流程",
            "申报挂网",
            "挂网操作",
            "申报入口",
            "申报路径",
            "联审通办",
            "提交申报",
        ]
        if keyword in compact
    )
    procurement_hits = sum(1 for keyword in ["医用耗材", "挂网", "申报"] if keyword in compact)
    return process_hits >= 1 and procurement_hits >= 2


def _missing_process_notice_topics(report_text: str) -> list[str]:
    compact = re.sub(r"\s+", "", str(report_text or ""))
    missing: list[str] = []
    for topic, keywords in PROCESS_NOTICE_REQUIRED_TOPICS:
        if not any(keyword in compact for keyword in keywords):
            missing.append(topic)
    return missing


def _table_cells_missing_from_evidence(report: ReportIR, evidence_text: str) -> list[ReportQAIssue]:
    normalized_evidence = _compact_evidence_marker(evidence_text)
    issues: list[ReportQAIssue] = []
    seen: set[str] = set()
    for section in report.sections:
        for table in section.tables:
            table_name = table.title or section.heading or "未命名表格"
            table_context = _table_unit_context(section, table)
            cells_with_rows = [
                *[(cell, []) for cell in table.headers],
                *[(cell, row) for row in table.rows for cell in row],
                *[(cell, []) for cell in table.notes],
            ]
            for cell, row in cells_with_rows:
                cell_text = _clean_inline_text(cell)
                if not cell_text or any(word in cell_text for word in EVIDENCE_MARKER_SKIP_WORDS):
                    continue
                for marker in _extract_evidence_markers(cell_text):
                    compact_marker = _compact_evidence_marker(marker)
                    if not compact_marker or compact_marker in seen:
                        continue
                    seen.add(compact_marker)
                    if _is_explicit_calculated_percentage(marker, section, table):
                        continue
                    if compact_marker in normalized_evidence:
                        continue
                    if _marker_matches_evidence_row_context(marker, row, table_context, evidence_text):
                        continue
                    issues.append(
                        ReportQAIssue(
                            severity="major",
                            category="table_evidence_mismatch",
                            report_text=f"{table_name}: {marker}",
                            source_quote="",
                            fix_instruction=f"核对表格《{table_name}》中的“{marker}”，删除无证据数值或改为原文披露内容。",
                        )
                    )
    return issues


def _table_unit_context(section: ReportSection, table: ReportTable) -> str:
    return "\n".join(
        _clean_inline_text(part)
        for part in [
            section.heading,
            table.title,
            *table.headers,
            *table.notes,
        ]
        if _clean_inline_text(part)
    )


def _marker_matches_evidence_row_context(
    marker: str,
    report_row: list[str],
    table_context: str,
    evidence_text: str,
) -> bool:
    equivalents = _numeric_marker_equivalents(marker, table_context)
    if not equivalents:
        return False
    row_keys = _row_context_keys(report_row)
    if not row_keys:
        return False
    evidence_lines = [line for line in str(evidence_text or "").splitlines() if line.strip()]
    for line in evidence_lines:
        compact_line = _compact_evidence_marker(line)
        if not any(key in compact_line for key in row_keys):
            continue
        if any(_line_has_numeric_token(line, equivalent) for equivalent in equivalents):
            return True
    return False


def _numeric_marker_equivalents(marker: str, table_context: str) -> list[str]:
    if "元" not in str(marker or ""):
        return []
    context = str(table_context or "")
    if not any(unit_hint in context for unit_hint in ["（元）", "(元)", "单位：元", "单位:元", "价格（元", "价格(元"]):
        return []
    equivalents: list[str] = []
    for match in re.finditer(r"(?<!\d)(\d+(?:,\d{3})*(?:\.\d+)?)(?:\s*)(?:万元|元(?:/个)?)", str(marker or "")):
        value = match.group(1).replace(",", "")
        if value:
            equivalents.append(value)
    return _unique(equivalents)


def _row_context_keys(row: list[str]) -> list[str]:
    keys: list[str] = []
    for cell in row:
        cell_text = _clean_inline_text(cell)
        if not cell_text:
            continue
        without_markers = cell_text
        for marker in _extract_evidence_markers(cell_text):
            without_markers = without_markers.replace(marker, "")
        compact = _compact_evidence_marker(without_markers)
        if len(compact) >= 3 and re.search(r"[\u4e00-\u9fffA-Za-z]", compact):
            keys.append(compact)
    return _unique(keys)


def _line_has_numeric_token(line: str, value: str) -> bool:
    normalized_line = str(line or "").replace(",", "")
    normalized_value = str(value or "").replace(",", "")
    if not normalized_value:
        return False
    pattern = rf"(?<![\d.]){re.escape(normalized_value)}(?:\.0+)?(?![\d.])"
    return bool(re.search(pattern, normalized_line))


def _extract_evidence_markers(text: str) -> list[str]:
    return _unique([match.group(0).strip() for match in EVIDENCE_MARKER_RE.finditer(str(text or ""))])


def _is_explicit_calculated_percentage(marker: str, section: ReportSection, table: ReportTable) -> bool:
    if "%" not in str(marker or ""):
        return False
    table_parts = [
        section.heading,
        *section.paragraphs,
        table.title,
        *table.headers,
        *(cell for row in table.rows for cell in row),
        *table.notes,
    ]
    context = "".join(_clean_inline_text(part) for part in table_parts)
    has_calculation_basis = any(keyword in context for keyword in ["按表内数据测算", "计算值", "测算"])
    has_percentage_column = any(keyword in context for keyword in ["占比", "比例", "百分比"])
    return has_calculation_basis and has_percentage_column


def _compact_evidence_marker(text: str) -> str:
    value = str(text or "")
    value = value.replace("\\r\\n", "").replace("\\n", "").replace("\\r", "").replace("\\t", "")
    return re.sub(r"[\s,，。；;：:（）()《》\"“”'、]", "", value)


def _extract_crawl_status(evidence_text: str) -> str:
    match = re.search(r"抓取状态\s*[:：]\s*([A-Za-z_]+)", str(evidence_text or ""))
    return match.group(1) if match else ""


def _qa_has_blocking_local_issues(qa: ReportQA) -> bool:
    hard_issues = [issue for issue in [*qa.issues, *qa.unsupported_claims, *qa.history_leakage] if issue.severity != "advisory"]
    return bool(hard_issues or qa.missing_rules)


def _merge_local_qa(base: ReportQA, local: ReportQA) -> ReportQA:
    if local.status == "block" or (local.status != "pass" and base.status == "pass"):
        base.status = local.status
    base.issues.extend(local.issues)
    base.unsupported_claims.extend(local.unsupported_claims)
    base.history_leakage.extend(local.history_leakage)
    base.missing_rules = _unique([*base.missing_rules, *local.missing_rules])
    base.language_issues.extend(local.language_issues)
    base.fix_instructions = _unique([*base.fix_instructions, *local.fix_instructions])
    if local.summary:
        model_summary = base.summary.strip()
        local_summary = local.summary.strip()
        if model_summary:
            base.summary = f"模型质检摘要：{model_summary}\n本地规则提示：{local_summary}"
        else:
            base.summary = f"本地规则提示：{local_summary}"
    return base


def _suppress_allowed_calculated_percentage_qa(qa: ReportQA, report: ReportIR) -> ReportQA:
    allowed_markers = _allowed_calculated_percentage_markers(report)
    if not allowed_markers:
        return qa

    qa.issues = [
        issue
        for issue in qa.issues
        if not _is_allowed_calculated_percentage_issue(issue, allowed_markers)
    ]
    qa.unsupported_claims = [
        issue
        for issue in qa.unsupported_claims
        if not _is_allowed_calculated_percentage_issue(issue, allowed_markers)
    ]
    qa.fix_instructions = [
        item
        for item in qa.fix_instructions
        if not _text_only_mentions_allowed_calculated_percentages(item, allowed_markers)
    ]
    if not _qa_has_blocking_local_issues(qa):
        qa.status = "pass"
        if qa.summary and any(marker in _compact_evidence_marker(qa.summary) for marker in allowed_markers):
            qa.summary = "通过；已忽略报告中明确标注测算口径的百分比派生值。"
    return qa


def _allowed_calculated_percentage_markers(report: ReportIR) -> set[str]:
    allowed: set[str] = set()
    for section in report.sections:
        for table in section.tables:
            cells = [*table.headers, *(cell for row in table.rows for cell in row), *table.notes]
            for cell in cells:
                for marker in _extract_evidence_markers(_clean_inline_text(cell)):
                    if _is_explicit_calculated_percentage(marker, section, table):
                        allowed.add(_compact_evidence_marker(marker))
    return allowed


def _is_allowed_calculated_percentage_issue(issue: ReportQAIssue, allowed_markers: set[str]) -> bool:
    text = f"{issue.category}\n{issue.report_text}\n{issue.fix_instruction}"
    if not (
        "unsupported" in issue.category
        or issue.category in {"table_evidence_mismatch", ""}
        or "无证据" in issue.fix_instruction
    ):
        return False
    return _text_only_mentions_allowed_calculated_percentages(text, allowed_markers)


def _text_only_mentions_allowed_calculated_percentages(text: str, allowed_markers: set[str]) -> bool:
    markers = [
        _compact_evidence_marker(marker)
        for marker in _extract_evidence_markers(str(text or ""))
        if "%" in marker
    ]
    return bool(markers) and all(marker in allowed_markers for marker in markers)


def _add_nonblocking_export_quality_issues(qa: ReportQA, quality_issues: list[str]) -> None:
    if not quality_issues:
        return
    if qa.status == "pass":
        qa.status = "needs_fix"
    qa.issues.extend(
        [
            ReportQAIssue(
                severity="major",
                category="export_quality",
                report_text=issue,
                fix_instruction=issue,
            )
            for issue in quality_issues
        ]
    )
    qa.fix_instructions = _unique([*qa.fix_instructions, *quality_issues])
    note = "导出前质量检查发现问题，已作为缺陷提示保留，不阻断 Word 导出。"
    if note not in qa.summary:
        qa.summary = f"{qa.summary}\n{note}".strip()


def _qa_from_checked_export_request(req: CheckedExportReportRequest) -> ReportQA:
    if req.qa_result:
        data = dict(req.qa_result)
        if req.qa_status:
            data["status"] = req.qa_status
        return _qa_from_dict(data)
    if req.qa_status:
        return _qa_from_dict({"status": req.qa_status, "summary": ""})
    if req.qa_output:
        try:
            return _parse_qa_output(req.qa_output)
        except ValueError as exc:
            return ReportQA(
                status="block",
                issues=[
                    ReportQAIssue(
                        severity="blocker",
                        category="qa_parse_error",
                        report_text=str(exc),
                        fix_instruction="质检 JSON 解析失败，需人工确认后再导出。",
                    )
                ],
                fix_instructions=["质检 JSON 解析失败，需人工确认后再导出。"],
                summary="质检 JSON 解析失败，已阻断 Word 导出。",
            )
    return ReportQA()


def _qa_from_dict(data: dict[str, Any]) -> ReportQA:
    normalized = {
        "status": _clean_inline_text(data.get("status") or "pass") or "pass",
        "issues": _normalize_qa_issues(data.get("issues"), default_severity="minor"),
        "unsupported_claims": _normalize_qa_issues(data.get("unsupported_claims"), default_severity="major"),
        "history_leakage": _normalize_qa_issues(data.get("history_leakage"), default_severity="major"),
        "missing_rules": [_clean_inline_text(item) for item in _ensure_list(data.get("missing_rules")) if _clean_inline_text(item)],
        "language_issues": _normalize_qa_issues(data.get("language_issues"), default_severity="minor"),
        "fix_instructions": [
            _clean_inline_text(item) for item in _ensure_list(data.get("fix_instructions")) if _clean_inline_text(item)
        ],
        "summary": _clean_inline_text(data.get("summary") or data.get("qa_summary") or ""),
    }
    return ReportQA.model_validate(normalized)


def _workflow_qa_status(qa: ReportQA) -> str:
    status = re.sub(r"[\s-]+", "_", _clean_inline_text(qa.status).lower())
    if status in {"block", "blocked", "fail", "failed", "needs_human", "manual_review"}:
        return "block"
    if status in {"needs_fix", "fix", "repair"}:
        return "needs_fix"
    severe_block = {"blocker", "critical", "fatal", "high", "阻断"}
    severe_major = {"major", "严重", "重大"}
    all_issues = [*qa.issues, *qa.unsupported_claims, *qa.history_leakage, *qa.language_issues]
    if any(_clean_inline_text(issue.severity).lower() in severe_block for issue in all_issues):
        return "block"
    if any(_clean_inline_text(issue.severity).lower() in severe_major for issue in all_issues):
        return "needs_fix"
    if qa.missing_rules:
        return "needs_fix"
    return "pass"


def _normalize_qa_status_for_workflow(qa: ReportQA) -> ReportQA:
    qa.status = _workflow_qa_status(qa)
    return qa


def _parse_qa_output(raw_text: str) -> ReportQA:
    text = _remove_thinking_blocks(str(raw_text or "")).strip()
    if not text:
        raise ValueError("质检模型输出为空，已阻止导出 Word")

    tag_match = re.search(r"<qa_report\b[^>]*>([\s\S]*?)</qa_report>", text, flags=re.I)
    payload_text = tag_match.group(1) if tag_match else text
    payload = _extract_first_json_object(_strip_json_code_fence(payload_text))
    if not payload:
        raise ValueError("质检模型未输出合法 JSON，已阻止导出 Word")

    try:
        data = json.loads(payload)
    except json.JSONDecodeError as exc:
        raise ValueError(f"质检 JSON 解析失败，已阻止导出 Word: {exc}") from exc
    if not isinstance(data, dict):
        raise ValueError("质检 JSON 顶层必须是对象，已阻止导出 Word")

    normalized = {
        "status": _clean_inline_text(data.get("status") or "pass") or "pass",
        "issues": _normalize_qa_issues(data.get("issues"), default_severity="minor"),
        "unsupported_claims": _normalize_qa_issues(data.get("unsupported_claims"), default_severity="major"),
        "history_leakage": _normalize_qa_issues(data.get("history_leakage"), default_severity="major"),
        "missing_rules": [_clean_inline_text(item) for item in _ensure_list(data.get("missing_rules")) if _clean_inline_text(item)],
        "language_issues": _normalize_qa_issues(data.get("language_issues"), default_severity="minor"),
        "fix_instructions": [
            _clean_inline_text(item) for item in _ensure_list(data.get("fix_instructions")) if _clean_inline_text(item)
        ],
        "summary": _clean_inline_text(data.get("summary") or ""),
    }
    return ReportQA.model_validate(normalized)


def _extract_first_json_object(text: str) -> str:
    value = str(text or "").strip()
    if value.startswith("{") and value.endswith("}"):
        return value
    fence_match = re.search(r"```(?:json)?\s*([\s\S]*?)\s*```", value, flags=re.I)
    if fence_match:
        fenced = fence_match.group(1).strip()
        if fenced.startswith("{") and fenced.endswith("}"):
            return fenced

    start = value.find("{")
    if start < 0:
        return ""
    depth = 0
    in_string = False
    escaped = False
    for index in range(start, len(value)):
        char = value[index]
        if in_string:
            if escaped:
                escaped = False
            elif char == "\\":
                escaped = True
            elif char == '"':
                in_string = False
            continue
        if char == '"':
            in_string = True
        elif char == "{":
            depth += 1
        elif char == "}":
            depth -= 1
            if depth == 0:
                return value[start : index + 1]
    return ""


def _normalize_qa_issues(value: Any, default_severity: str) -> list[dict[str, str]]:
    issues: list[dict[str, str]] = []
    for item in _ensure_list(value):
        if isinstance(item, dict):
            issue = {
                "severity": _clean_inline_text(item.get("severity") or default_severity) or default_severity,
                "category": _clean_inline_text(item.get("category") or ""),
                "report_text": _clean_inline_text(item.get("report_text") or item.get("text") or item.get("claim") or ""),
                "source_quote": _clean_inline_text(item.get("source_quote") or item.get("source") or ""),
                "fix_instruction": _clean_inline_text(item.get("fix_instruction") or item.get("fix") or ""),
            }
        else:
            issue = {
                "severity": default_severity,
                "category": "",
                "report_text": _clean_inline_text(item),
                "source_quote": "",
                "fix_instruction": "",
            }
        if any(issue.values()):
            issues.append(issue)
    return issues


def _ensure_list(value: Any) -> list[Any]:
    if value is None:
        return []
    if isinstance(value, list):
        return value
    return [value]


def _qa_blocks_export(qa: ReportQA) -> bool:
    status = re.sub(r"[\s-]+", "_", _clean_inline_text(qa.status).lower())
    if status in {"needs_fix", "block", "blocked", "fail", "failed", "needs_human", "manual_review", "需人工确认"}:
        return True
    severe = {"major", "blocker", "critical", "high", "fatal", "严重", "重大", "阻断"}
    for issue in [*qa.unsupported_claims, *qa.history_leakage]:
        if _clean_inline_text(issue.severity).lower() in severe:
            return True
    return False


def _format_qa_summary(qa: ReportQA) -> str:
    lines = [f"质检状态：{qa.status or 'pass'}"]
    if qa.summary:
        lines.append(f"质检摘要：{qa.summary}")
    counters = [
        ("事实依据问题", len(qa.unsupported_claims)),
        ("历史知识泄漏", len(qa.history_leakage)),
        ("缺失规则", len(qa.missing_rules)),
        ("语言问题", len(qa.language_issues)),
    ]
    lines.append("问题统计：" + "；".join(f"{label}{count}项" for label, count in counters))
    detail_groups = [
        ("一般问题", qa.issues),
        ("事实依据问题", qa.unsupported_claims),
        ("历史知识泄漏", qa.history_leakage),
        ("语言问题", qa.language_issues),
    ]
    for label, issues in detail_groups:
        if not issues:
            continue
        details = []
        for issue in issues[:5]:
            category = _clean_inline_text(issue.category) or "未分类"
            text = _clean_inline_text(issue.report_text)
            details.append(f"{category}: {text}" if text else category)
        lines.append(f"{label}明细：" + "；".join(details))
    if qa.missing_rules:
        lines.append("缺失规则明细：" + "；".join(qa.missing_rules[:5]))
    if qa.fix_instructions:
        lines.append("修复建议：" + "；".join(qa.fix_instructions[:5]))
    return "\n".join(lines)


def _detect_history_leakage(report_text: str, history_text: str, evidence_text: str) -> list[ReportQAIssue]:
    report = _clean_for_history_compare(report_text)
    history = _clean_for_history_compare(history_text)
    evidence = _clean_for_history_compare(evidence_text)
    if not report or not history:
        return []

    issues: list[ReportQAIssue] = []
    heading_pattern = re.compile(r"(?m)^\s*(?:#{1,6}\s*)?(历史对照|历史分析|历史知识|既往项目分析|历史承接|历史回顾)\s*$")
    heading_match = heading_pattern.search(report)
    if heading_match:
        issues.append(
            ReportQAIssue(
                severity="major",
                category="history_heading",
                report_text=heading_match.group(1),
                fix_instruction="删除历史分析新增标题，将历史承接内容自然并入相关段落后半段。",
            )
        )

    for fact in _history_fact_markers(history):
        if fact in report and fact not in evidence:
            issues.append(
                ReportQAIssue(
                    severity="major",
                    category="history_fact_leakage",
                    report_text=fact,
                    fix_instruction="该信息仅见于历史分析稿，不能写成本次公告事实；如需保留，只能使用带限定语的历史观察且不得包含具体规则或数值。",
                )
            )
    return _dedupe_qa_issues(issues)


def _clean_for_history_compare(text: str) -> str:
    value = _remove_thinking_blocks(str(text or ""))
    value = html_lib.unescape(value).replace("\xa0", " ")
    value = re.sub(r"\{\{[\s\S]*?\}\}", "", value)
    value = re.sub(r"<[^>]+>", "", value)
    value = re.sub(r"[ \t\r\f\v]+", " ", value)
    value = re.sub(r"\n{3,}", "\n\n", value)
    return value.strip()


def _history_fact_markers(text: str) -> list[str]:
    patterns = [
        r"采购周期[为：:\s]*[^，。；;\n]{0,20}?(?:\d+|一|二|两|三|四|五|六|七|八|九|十)年",
        r"最高有效申报价[为：:\s]*[^，。；;\n]{0,30}",
        r"询价基准[为：:\s]*[^，。；;\n]{0,30}",
        r"参考价[为：:\s]*[^，。；;\n]{0,30}",
        r"\d+(?:\.\d+)?\s*元",
        r"\d+(?:\.\d+)?\s*%",
        r"\d+\s*个工作日",
        r"20\d{2}\s*年\s*\d{1,2}\s*月\s*\d{1,2}\s*日",
        r"20\d{2}[-/.]\d{1,2}[-/.]\d{1,2}",
    ]
    markers: list[str] = []
    for pattern in patterns:
        markers.extend(match.group(0).strip() for match in re.finditer(pattern, text))
    return [marker for marker in _unique(markers) if len(marker) >= 2]


def _dedupe_qa_issues(issues: list[ReportQAIssue]) -> list[ReportQAIssue]:
    seen: set[tuple[str, str]] = set()
    result: list[ReportQAIssue] = []
    for issue in issues:
        key = (issue.category, issue.report_text)
        if key in seen:
            continue
        seen.add(key)
        result.append(issue)
    return result


def _report_ir_text(report: ReportIR) -> str:
    parts = [report.title, report.suggested_filename, report.notice_type, report.publish_date, report.source_agency, report.document_name]
    parts.extend(report.lead_paragraphs)
    for section in report.sections:
        parts.append(section.heading)
        parts.extend(section.paragraphs)
        parts.extend(section.highlights)
        for table in section.tables:
            parts.append(table.title)
            parts.extend(table.headers)
            parts.extend(cell for row in table.rows for cell in row)
            parts.extend(table.notes)
    parts.extend(report.enterprise_tips)
    parts.append(report.disclaimer)
    return "\n".join(str(part) for part in parts if part)


def _render_parse_error(raw_text: str, original_error: str) -> str:
    text = str(raw_text or "")
    if re.search(r"<report_ir\b[^>]*>", text, flags=re.I) and not re.search(r"</report_ir>", text, flags=re.I):
        return "ReportIR 未闭合，模型输出可能被截断；请减少输入证据或重新生成报告"
    if re.search(r"<final_report\b[^>]*>", text, flags=re.I) and not re.search(r"</final_report>", text, flags=re.I):
        return "final_report 未闭合，模型输出可能被截断；请减少输入证据或重新生成报告"
    return original_error


def _report_ir_to_markdown(report: ReportIR) -> str:
    lines: list[str] = []
    title = _clean_inline_text(report.title) or DEFAULT_REPORT_TITLE
    lines.extend([f"# {title}", ""])

    for paragraph in report.lead_paragraphs:
        text = _clean_inline_text(paragraph)
        if text:
            lines.extend([text, ""])

    for section in report.sections:
        heading = _clean_inline_text(section.heading)
        if heading:
            lines.extend([f"## {heading}", ""])
        for paragraph in section.paragraphs:
            text = _clean_inline_text(paragraph)
            if text:
                lines.extend([text, ""])
        for highlight in section.highlights:
            text = _clean_inline_text(highlight)
            if text:
                lines.append(f"- {text}")
        if section.highlights:
            lines.append("")
        for table in section.tables:
            rendered = _report_table_to_markdown(table)
            if rendered:
                lines.extend([rendered, ""])

    tips = [_clean_inline_text(tip) for tip in report.enterprise_tips if _clean_inline_text(tip)]
    if tips:
        lines.extend(["## 企业关注点", ""])
        for tip in tips:
            lines.append(f"- {tip}")
        lines.append("")

    markdown = "\n".join(lines)
    markdown = re.sub(r"\n{3,}", "\n\n", markdown)
    return markdown.strip()


def _report_table_to_markdown(table: ReportTable) -> str:
    headers = [_markdown_table_cell(header) for header in table.headers if _clean_inline_text(header)]
    rows = [[_markdown_table_cell(cell) for cell in row] for row in table.rows]
    if not headers or not rows:
        return ""
    width = len(headers)
    normalized_rows = [(row + [""] * width)[:width] for row in rows]
    lines: list[str] = []
    title = _clean_inline_text(table.title)
    if title:
        lines.extend([f"### {title}", ""])
    lines.append("| " + " | ".join(headers) + " |")
    lines.append("| " + " | ".join(["---"] * width) + " |")
    for row in normalized_rows:
        lines.append("| " + " | ".join(row) + " |")
    notes = [_clean_inline_text(note) for note in table.notes if _clean_inline_text(note)]
    for note in notes:
        lines.append(f"- {note}")
    return "\n".join(lines)


def _markdown_table_cell(value: str) -> str:
    text = _clean_inline_text(value)
    text = text.replace("\\", "\\\\").replace("|", "\\|")
    return text.replace("\n", "<br>")


def _report_ir_to_docx(report: ReportIR, path: Path, fallback_title: str) -> None:
    doc = Document()
    _configure_docx_document(doc)
    _add_report_header_footer(doc, report)
    _add_text_watermark(doc, REPORT_WATERMARK_TEXT)

    title = report.title or fallback_title or DEFAULT_REPORT_TITLE
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title)
    title_run.bold = True
    title_run.font.name = "SimSun"
    title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    title_run.font.size = Pt(18)

    for paragraph in report.lead_paragraphs:
        _add_body_paragraph(doc, paragraph)

    for section in report.sections:
        if section.heading:
            doc.add_heading(section.heading, level=1)
        for paragraph in section.paragraphs:
            _add_body_paragraph(doc, paragraph)
        for highlight in section.highlights:
            _add_highlight_paragraph(doc, highlight)
        for table in section.tables:
            _add_report_table(doc, table)

    if report.enterprise_tips:
        doc.add_heading("企业关注点", level=1)
        for tip in report.enterprise_tips:
            _add_body_paragraph(doc, tip)

    if report.disclaimer:
        doc.add_page_break()
        _add_disclaimer_heading(doc)
        _add_disclaimer_paragraph(doc, DEFAULT_DISCLAIMER)

    doc.save(path)


def _configure_docx_document(doc: Document) -> None:
    section = doc.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.35
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size in [("Heading 1", 14), ("Heading 2", 12)]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        style.font.bold = True
        style.font.size = Pt(size)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)


def _add_report_header_footer(doc: Document, report: ReportIR) -> None:
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = report.document_name or DEFAULT_REPORT_TITLE
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_font(header, size=9, color=RGBColor(120, 120, 120))

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("生成日期: " + (report.publish_date or "未披露") + "  |  第 ")
    _add_page_number(footer)
    footer.add_run(" 页")
    _set_paragraph_font(footer, size=9, color=RGBColor(120, 120, 120))


def _add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def _add_text_watermark(doc: Document, text: str) -> None:
    if not text:
        return
    escaped = html_lib.escape(text)
    watermark_xml = (
        '<w:pict xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'xmlns:v="urn:schemas-microsoft-com:vml" '
        'xmlns:o="urn:schemas-microsoft-com:office:office">'
        '<v:shape id="PowerPlusWaterMarkObject" o:spid="_x0000_s1025" type="#_x0000_t136" '
        'style="position:absolute;margin-left:0;margin-top:0;width:420pt;height:120pt;'
        'rotation:315;z-index:-251654144;mso-position-horizontal:center;mso-position-vertical:center" '
        'fillcolor="#d9d9d9" stroked="f">'
        '<v:fill opacity=".12"/>'
        f'<v:textpath style="font-family:宋体;font-size:1pt" string="{escaped}"/>'
        "</v:shape></w:pict>"
    )
    paragraph = doc.sections[0].header.paragraphs[0]
    run = paragraph.add_run()
    run._r.append(parse_xml(watermark_xml))


def _add_body_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Cm(0.74)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph.paragraph_format.line_spacing = 1.35
    _add_rich_text_runs(paragraph, text)


def _add_highlight_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    run.font.color.rgb = RGBColor(192, 0, 0)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(11)


def _add_disclaimer_heading(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run("声  明")
    run.bold = True
    run.font.name = "KaiTi"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "楷体")
    run.font.size = Pt(22)


def _add_disclaimer_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Cm(0.74)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph.paragraph_format.line_spacing = 1.35
    run = paragraph.add_run(text)
    run.font.name = "SimSun"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(16)


def _add_rich_text_runs(paragraph, text: str) -> None:
    parts = re.split(
        r"(\*\*.*?\*\*|<red>.*?</red>|<blue>.*?</blue>|<span(?:\s+[^>]*)?class=[\"']analysis-highlight[\"'][^>]*>.*?</span>|<span\s+style=[\"'][^\"']*(?:fff3cd|b8860b)[^\"']*[\"'][^>]*>.*?</span>)",
        text,
        flags=re.I,
    )
    for part in parts:
        if not part:
            continue
        is_bold = part.startswith("**") and part.endswith("**")
        is_red = part.startswith("<red>") and part.endswith("</red>")
        is_blue = part.startswith("<blue>") and part.endswith("</blue>")
        is_analysis_highlight = bool(re.match(r"<span\b", part, flags=re.I) and re.search(r"analysis-highlight|fff3cd|b8860b", part, flags=re.I))
        value = part[2:-2] if is_bold else part[5:-6] if is_red else part[6:-7] if is_blue else re.sub(r"</?span\b[^>]*>", "", part, flags=re.I)
        run = paragraph.add_run(value)
        run.bold = is_bold or is_red or is_blue
        if is_red:
            run.font.color.rgb = RGBColor(255, 0, 0)
        if is_blue:
            run.font.color.rgb = RGBColor(0, 112, 192)
        if is_analysis_highlight and _env_bool("ENABLE_ANALYSIS_HIGHLIGHT", True):
            run.font.color.rgb = RGBColor(184, 134, 11)
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        run.font.name = "Calibri"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(11)


def _set_paragraph_font(paragraph, size: int, color: RGBColor | None = None) -> None:
    for run in paragraph.runs:
        run.font.name = "Calibri"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(size)
        if color is not None:
            run.font.color.rgb = color


def _add_report_table(doc: Document, report_table: ReportTable) -> None:
    if not report_table.headers or not report_table.rows:
        return
    rows = _normalize_table_rows(report_table.headers, report_table.rows)
    if not rows:
        return
    if report_table.title:
        caption = doc.add_paragraph(report_table.title)
        caption.runs[0].bold = True
        _set_paragraph_font(caption, size=10)
    width = len(report_table.headers)
    table = doc.add_table(rows=1, cols=width)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    header_cells = table.rows[0].cells
    for index, header in enumerate(report_table.headers):
        cell = header_cells[index]
        cell.text = header
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        _shade_cell(cell, "0070C0")
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_paragraph_font(paragraph, size=10, color=RGBColor(255, 255, 255))
            for run in paragraph.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cell = cells[index]
            cell.text = str(value)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _shade_cell(cell, "FFFFFF")
            for paragraph in cell.paragraphs:
                if _looks_numeric(value):
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _set_paragraph_font(paragraph, size=10)
    _merge_repeated_first_column(table)
    _repeat_table_header(table.rows[0])
    for note in report_table.notes:
        note_para = doc.add_paragraph("注：" + str(note))
        _set_paragraph_font(note_para, size=9, color=RGBColor(100, 100, 100))
    doc.add_paragraph()


def _shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def _repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def _looks_numeric(value: str) -> bool:
    return bool(re.fullmatch(r"[<>≤≥=≈~约\s\d,，.％%元万元亿元个件支年月日/-]+|[-—]", str(value).strip()))


def _markdown_to_docx(markdown: str, path: Path, title: str) -> None:
    markdown = _clean_model_output(markdown)
    doc = Document()
    _configure_docx_document(doc)
    _add_report_header_footer(doc, ReportIR(title=title or DEFAULT_REPORT_TITLE, document_name=title or DEFAULT_REPORT_TITLE))
    _add_text_watermark(doc, REPORT_WATERMARK_TEXT)

    report_title = title or DEFAULT_REPORT_TITLE
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(report_title)
    title_run.bold = True
    title_run.font.name = "KaiTi"
    title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "楷体")
    title_run.font.size = Pt(18)

    lines = markdown.splitlines()
    index = 0
    first_heading_seen = False
    while index < len(lines):
        line = lines[index].rstrip()
        stripped = line.strip()
        if not stripped:
            index += 1
            continue
        if stripped.startswith("```"):
            fence = stripped.lower()
            block = []
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("```"):
                block.append(lines[index])
                index += 1
            block_text = "\n".join(block)
            if "mermaid" in fence:
                image_path = _render_mermaid_flowchart(block_text)
                if image_path:
                    doc.add_picture(str(image_path), width=Inches(6.3))
                    try:
                        image_path.unlink(missing_ok=True)
                    except OSError:
                        pass
                else:
                    doc.add_paragraph(block_text)
            else:
                doc.add_paragraph(block_text)
            index += 1
            continue
        if stripped.startswith("|") and "|" in stripped[1:]:
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index].strip())
                index += 1
            _add_markdown_table(doc, table_lines)
            continue
        heading_match = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading_match:
            level = min(len(heading_match.group(1)), 4)
            heading_text = _strip_markdown(heading_match.group(2))
            if level == 1 and not first_heading_seen and _normalize_text(heading_text) == _normalize_text(report_title):
                first_heading_seen = True
                index += 1
                continue
            first_heading_seen = True
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(10)
            paragraph.paragraph_format.space_after = Pt(6)
            run = paragraph.add_run(heading_text)
            run.bold = True
            run.font.name = "Calibri"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            run.font.size = Pt(14 if level == 1 else 12)
            run.font.color.rgb = RGBColor(0, 112, 192 if level <= 2 else 160)
        elif stripped.startswith(("- ", "* ", "• ")):
            paragraph = doc.add_paragraph(style="List Bullet")
            _add_rich_text_runs(paragraph, stripped[2:])
        elif re.match(r"^\d+[.、]\s+", stripped):
            paragraph = doc.add_paragraph(style="List Number")
            _add_rich_text_runs(paragraph, re.sub(r"^\d+[.、]\s+", "", stripped))
        else:
            _add_body_paragraph(doc, stripped)
        index += 1
    doc.add_page_break()
    _add_disclaimer_heading(doc)
    _add_disclaimer_paragraph(doc, DEFAULT_DISCLAIMER)
    doc.save(path)


def _strip_generated_material_note_sections(text: str) -> str:
    lines = str(text or "").splitlines()
    output: list[str] = []
    skip_level: int | None = None

    for line in lines:
        heading_match = re.match(r"^\s*(#{1,6})\s+(.+?)\s*$", line)
        if heading_match:
            level = len(heading_match.group(1))
            heading_text = _strip_markdown(heading_match.group(2))
            compact_heading = re.sub(r"\s+", "", _normalize_text(heading_text))
            is_material_note = "资料说明" in compact_heading
            is_declaration = compact_heading in {"声明", "聲明"}
            if is_material_note or is_declaration:
                skip_level = level
                continue
            if skip_level is not None and level <= skip_level:
                skip_level = None

        if skip_level is None:
            output.append(line)

    return "\n".join(output).strip()


def _strip_technical_attachment_note_lines(text: str) -> str:
    technical_patterns = [
        r"OCR\s*识别",
        r"数字来源于.*OCR",
        r"OCR.*误差",
        r"OCR.*错误",
        r"附件\s*PDF\s*因",
        r"附件.*文件过大",
        r"因大小限制",
        r"大小限制",
        r"仅解析",
        r"元数据",
        r"metadata[_ -]?only",
        r"未解析",
        r"未能提取",
        r"系统未获取",
        r"解析失败",
        r"识别错误",
        r"识别有误",
        r"识别问题",
        r"原文.*识别",
        r"根据.*修正",
        r"暂无法确认",
        r"以(?:正式|官方)文件为准",
        r"资料说明\s*[:：]",
    ]
    technical_union = "|".join(f"(?:{pattern})" for pattern in technical_patterns)
    output: list[str] = []
    for line in str(text or "").splitlines():
        stripped_line = line.strip()
        compact = _normalize_text(_strip_markdown(stripped_line))
        if not compact:
            output.append(line)
            continue
        if re.match(r"^\s*(?:>\s*)?(?:注|资料说明)\s*[:：]", stripped_line) and any(
            re.search(pattern, compact, flags=re.I) for pattern in technical_patterns
        ):
            continue
        cleaned = line
        cleaned = re.sub(rf"（[^（）]*(?:{technical_union})[^（）]*）", "", cleaned, flags=re.I)
        cleaned = re.sub(rf"\([^()]*({technical_union})[^()]*\)", "", cleaned, flags=re.I)
        cleaned = re.sub(rf"[^。！？!?；;\n]*({technical_union})[^。！？!?；;\n]*[。！？!?；;]?", "", cleaned, flags=re.I)
        cleaned = re.sub(r"\s{2,}", " ", cleaned).rstrip()
        if _normalize_text(_strip_markdown(cleaned)):
            output.append(cleaned)
    return "\n".join(output).strip()


def _clean_model_output(text: str) -> str:
    text = str(text or "")
    text = text.replace("\ufeff", "")

    cleanup_patterns = [
        r"<think\b[^>]*>[\s\S]*?</think>",
        r"<思考\b[^>]*>[\s\S]*?</思考>",
        r"```(?:thinking|think|思考)[\s\S]*?```",
        r"`(?:thinking|think|思考)[\s\S]*?`",
    ]
    for pattern in cleanup_patterns:
        text = re.sub(pattern, "", text, flags=re.I)

    marker_patterns = [
        r"^\s*(?:#{1,6}\s*)?(?:最终报告|报告正文|正式报告|正文)\s*[:：]?\s*$",
        r"^\s*(?:以下为|以下是)?(?:最终报告|报告正文|正式报告)\s*[:：]\s*$",
    ]
    lines = text.splitlines()
    marker_index: int | None = None
    for index, line in enumerate(lines):
        if any(re.match(pattern, line.strip(), flags=re.I) for pattern in marker_patterns):
            marker_index = index
    if marker_index is not None:
        text = "\n".join(lines[marker_index + 1 :])

    internal_line_patterns = [
        r"^\s*我们需要.*$",
        r"^\s*我准备.*$",
        r"^\s*先.*梳理.*$",
        r"^\s*注意.*不要输出.*$",
        r"^\s*任务[:：].*$",
        r"^\s*分析[:：].*$",
        r"^\s*思路[:：].*$",
        r"^\s*推理过程[:：].*$",
        r"^\s*模型推理.*$",
        r"^\s*提示词.*$",
        r"^\s*调试信息.*$",
        r"^\s*Based on the.*$",
        r"^\s*We need to.*$",
        r"^\s*I need to.*$",
        r"^\s*I will.*$",
    ]
    for pattern in internal_line_patterns:
        text = re.sub(pattern, "", text, flags=re.I | re.M)

    text = re.sub(r"</?think\b[^>]*>", "", text, flags=re.I)
    text = re.sub(r"</?思考\b[^>]*>", "", text, flags=re.I)

    lines = text.splitlines()
    first_report_line = None
    for index, line in enumerate(lines):
        stripped = line.strip()
        if not stripped:
            continue
        if re.match(r"^#{1,2}\s+\S", stripped):
            first_report_line = index
            break
        if _looks_like_report_title(stripped):
            first_report_line = index
            break
    if first_report_line and first_report_line > 0:
        text = "\n".join(lines[first_report_line:])

    text = _strip_generated_material_note_sections(text)
    text = _strip_technical_attachment_note_lines(text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _looks_like_report_title(line: str) -> bool:
    if len(line) > 80 or len(line) < 8:
        return False
    banned_prefixes = (
        "好的",
        "以下",
        "根据",
        "请根据",
        "输入材料",
        "公告 URL",
        "抓取服务",
        "用户要求",
        "写作要求",
    )
    if line.startswith(banned_prefixes):
        return False
    title_keywords = [
        "集采",
        "采购",
        "接续",
        "带量",
        "中选",
        "联动",
        "启动",
        "规则",
        "支架",
        "耗材",
    ]
    return any(keyword in line for keyword in title_keywords)


def _render_mermaid_flowchart(source: str) -> Path | None:
    nodes, edges = _parse_mermaid_flowchart(source)
    if not nodes:
        return None

    font_regular = _load_cjk_font(25)
    font_small = _load_cjk_font(20)
    font_title = _load_cjk_font(26)
    width = 1500
    top_margin = 70
    side_margin = 90
    row_gap = 125
    ranks = _flowchart_ranks(nodes, edges)
    rows: list[list[dict[str, str]]] = []
    for rank in sorted(set(ranks.values())):
        rows.append([node for node in nodes if ranks.get(node["id"], 0) == rank])

    rendered_by_id: dict[str, dict[str, Any]] = {}
    y = top_margin
    for row in rows:
        count = max(1, len(row))
        box_width = min(1080, max(360, int((width - side_margin * 2 - 55 * (count - 1)) / count)))
        box_gap = 55
        row_width = box_width * count + box_gap * (count - 1)
        x = (width - row_width) // 2
        row_items: list[dict[str, Any]] = []
        row_height = 0
        for node in row:
            lines = _wrap_for_image(node["label"], font_regular, box_width - 72)
            box_height = max(92, 34 * len(lines) + 48)
            item = {
                **node,
                "lines": lines,
                "x0": x,
                "y0": y,
                "x1": x + box_width,
                "y1": y + box_height,
            }
            row_items.append(item)
            rendered_by_id[node["id"]] = item
            row_height = max(row_height, box_height)
            x += box_width + box_gap
        for item in row_items:
            delta = row_height - (item["y1"] - item["y0"])
            item["y0"] += delta // 2
            item["y1"] += delta // 2
        y += row_height + row_gap

    height = max(360, y + 35)
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)

    for edge in edges:
        source = rendered_by_id.get(edge["from"])
        target = rendered_by_id.get(edge["to"])
        if not source or not target:
            continue
        start = ((source["x0"] + source["x1"]) // 2, source["y1"] + 8)
        end = ((target["x0"] + target["x1"]) // 2, target["y0"] - 8)
        mid_y = (start[1] + end[1]) // 2
        points = [start, (start[0], mid_y), (end[0], mid_y), end]
        draw.line(points, fill="#6B7280", width=4, joint="curve")
        draw.polygon([(end[0], end[1] + 10), (end[0] - 12, end[1] - 10), (end[0] + 12, end[1] - 10)], fill="#6B7280")
        label = edge.get("label", "")
        if label:
            label_lines = _wrap_for_image(label, font_small, 430)
            label_text = " / ".join(label_lines[:2])
            bbox = draw.textbbox((0, 0), label_text, font=font_small)
            pad_x, pad_y = 14, 7
            label_x = (start[0] + end[0]) // 2 + 20
            label_y = mid_y - (bbox[3] - bbox[1]) // 2
            draw.rounded_rectangle(
                [
                    label_x - pad_x,
                    label_y - pad_y,
                    label_x + (bbox[2] - bbox[0]) + pad_x,
                    label_y + (bbox[3] - bbox[1]) + pad_y,
                ],
                radius=10,
                fill="#FFFFFF",
                outline="#D1D5DB",
                width=2,
            )
            draw.text((label_x, label_y), label_text, fill="#374151", font=font_small)

    for item in rendered_by_id.values():
        shape = item.get("shape") or "rect"
        fill = "#F4F9FF" if shape != "diamond" else "#FFF8E8"
        outline = "#1D74D8" if shape != "diamond" else "#D98A00"
        if shape == "diamond":
            cx = (item["x0"] + item["x1"]) // 2
            cy = (item["y0"] + item["y1"]) // 2
            points = [(cx, item["y0"]), (item["x1"], cy), (cx, item["y1"]), (item["x0"], cy)]
            draw.polygon(points, fill=fill, outline=outline)
            draw.line(points + [points[0]], fill=outline, width=4)
        else:
            draw.rounded_rectangle(
                [item["x0"], item["y0"], item["x1"], item["y1"]],
                radius=22,
                fill=fill,
                outline=outline,
                width=4,
            )

        text_y = item["y0"] + 24
        for line in item["lines"]:
            bbox = draw.textbbox((0, 0), line, font=font_regular)
            text_x = item["x0"] + ((item["x1"] - item["x0"]) - (bbox[2] - bbox[0])) / 2
            draw.text((text_x, text_y), line, fill="#172033", font=font_regular)
            text_y += 34

    footer = "由报告导出服务根据 Mermaid 规则流程图生成"
    footer_bbox = draw.textbbox((0, 0), footer, font=font_title)
    draw.text((width - footer_bbox[2] - 40, height - footer_bbox[3] - 22), footer, fill="#9CA3AF", font=font_title)

    output = Path(tempfile.gettempdir()) / f"rule-flow-{uuid.uuid4().hex[:8]}.png"
    image.save(output)
    return output


def _flowchart_ranks(nodes: list[dict[str, str]], edges: list[dict[str, str]]) -> dict[str, int]:
    node_ids = [node["id"] for node in nodes]
    ranks = {node_id: 0 for node_id in node_ids}
    for _ in range(max(1, len(node_ids))):
        changed = False
        for edge in edges:
            source = edge.get("from", "")
            target = edge.get("to", "")
            if source in ranks and target in ranks and ranks[target] < ranks[source] + 1:
                ranks[target] = ranks[source] + 1
                changed = True
        if not changed:
            break
    for index, node_id in enumerate(node_ids):
        ranks.setdefault(node_id, index)
    return ranks


def _parse_mermaid_flowchart(source: str) -> tuple[list[dict[str, str]], list[dict[str, str]]]:
    nodes: list[dict[str, str]] = []
    edges: list[dict[str, str]] = []
    seen: set[str] = set()

    def add_node(node_id: str, label: str, shape: str = "rect") -> None:
        clean_id = node_id.strip()
        if not clean_id:
            return
        clean_label = _strip_mermaid_label(label or clean_id)
        if clean_id in seen:
            for node in nodes:
                if node["id"] == clean_id and node["label"] == node["id"] and clean_label != clean_id:
                    node["label"] = clean_label
                    node["shape"] = shape
            return
        seen.add(clean_id)
        nodes.append({"id": clean_id, "label": clean_label, "shape": shape})

    for raw_line in source.splitlines():
        line = raw_line.strip().rstrip(";")
        if not line or line.startswith("%%") or re.match(r"^(flowchart|graph)\b", line, flags=re.I):
            continue
        edge_match = re.match(r"(.+?)\s*(?:-->|---|==>)\s*(?:\|([^|]+)\|\s*)?(.+)$", line)
        if edge_match:
            left_id, left_label, left_shape = _parse_mermaid_node(edge_match.group(1))
            right_id, right_label, right_shape = _parse_mermaid_node(edge_match.group(3))
            add_node(left_id, left_label, left_shape)
            add_node(right_id, right_label, right_shape)
            edges.append({"from": left_id, "to": right_id, "label": _strip_mermaid_label(edge_match.group(2) or "")})
            continue
        node_id, label, shape = _parse_mermaid_node(line)
        add_node(node_id, label, shape)

    return nodes[:18], edges


def _parse_mermaid_node(token: str) -> tuple[str, str, str]:
    token = token.strip()
    diamond = re.match(r"^([A-Za-z][\w-]*)\s*\{\s*(.+?)\s*\}$", token)
    if diamond:
        return diamond.group(1), diamond.group(2), "diamond"
    bracket = re.match(r"^([A-Za-z][\w-]*)\s*\[\s*(.+?)\s*\]$", token)
    if bracket:
        return bracket.group(1), bracket.group(2), "rect"
    rounded = re.match(r"^([A-Za-z][\w-]*)\s*\(\(\s*(.+?)\s*\)\)$", token)
    if rounded:
        return rounded.group(1), rounded.group(2), "rect"
    plain = re.match(r"^([A-Za-z][\w-]*)", token)
    if plain:
        return plain.group(1), plain.group(1), "rect"
    fallback = re.sub(r"\W+", "_", token)[:20] or "node"
    return fallback, token, "rect"


def _strip_mermaid_label(text: str) -> str:
    text = text.strip().strip("\"'`")
    text = text.replace("<br/>", "\n").replace("<br>", "\n")
    return _strip_markdown(text)


def _edge_label(edges: list[dict[str, str]], source_id: str, target_id: str) -> str:
    for edge in edges:
        if edge["from"] == source_id and edge["to"] == target_id:
            return edge.get("label", "")
    return ""


def _load_cjk_font(size: int) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/opentype/noto/NotoSerifCJK-Regular.ttc",
        "C:/Windows/Fonts/msyh.ttc",
        "C:/Windows/Fonts/simhei.ttf",
    ]
    for candidate in candidates:
        try:
            return ImageFont.truetype(candidate, size=size)
        except OSError:
            continue
    return ImageFont.load_default()


def _wrap_for_image(text: str, font: ImageFont.ImageFont, max_width: int) -> list[str]:
    lines: list[str] = []
    for paragraph in str(text).splitlines() or [""]:
        current = ""
        for char in paragraph:
            trial = current + char
            bbox = font.getbbox(trial)
            if current and bbox[2] - bbox[0] > max_width:
                lines.append(current)
                current = char
            else:
                current = trial
        if current:
            lines.append(current)
    return lines or [""]


def _add_markdown_table(doc: Document, table_lines: list[str]) -> None:
    rows = []
    for line in table_lines:
        cells = [cell.strip() for cell in line.strip("|").split("|")]
        if cells and all(re.fullmatch(r":?-{3,}:?", cell or "") for cell in cells):
            continue
        rows.append(cells)
    if not rows:
        return
    width = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=width)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for r, row in enumerate(rows):
        for c in range(width):
            cell = table.cell(r, c)
            value = row[c] if c < len(row) else ""
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _clear_cell_paragraphs(cell)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if r == 0 or _looks_numeric(value) else WD_ALIGN_PARAGRAPH.LEFT
            _add_rich_text_runs(paragraph, value)
            if r == 0:
                _shade_cell(cell, "0070C0")
                _set_paragraph_font(paragraph, size=10, color=RGBColor(255, 255, 255))
                for run in paragraph.runs:
                    run.bold = True
            else:
                _shade_cell(cell, "F8FBFF" if r % 2 == 0 else "FFFFFF")
                _set_paragraph_font(paragraph, size=10)
    _merge_repeated_first_column(table)
    _repeat_table_header(table.rows[0])
    doc.add_paragraph()


def _clear_cell_paragraphs(cell) -> None:
    for paragraph in cell.paragraphs:
        for run in list(paragraph.runs):
            paragraph._p.remove(run._r)


def _merge_repeated_first_column(table) -> None:
    if len(table.columns) < 2 or len(table.rows) < 3:
        return
    start = 1
    previous = table.cell(1, 0).text.strip()
    for index in range(2, len(table.rows) + 1):
        current = table.cell(index, 0).text.strip() if index < len(table.rows) else None
        if current == previous and current:
            continue
        end = index - 1
        if previous and end > start:
            merged = table.cell(start, 0).merge(table.cell(end, 0))
            merged.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in merged.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _set_paragraph_font(paragraph, size=10)
        start = index
        previous = current or ""


def _strip_markdown(text: str) -> str:
    text = re.sub(r"```[a-zA-Z0-9_-]*", "", str(text or ""))
    text = text.replace("```", "")
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"</?(?:red|blue)>", "", text)
    text = re.sub(r"</?span\b[^>]*>", "", text, flags=re.I)
    text = re.sub(r"`([^`]*)`", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"^\s{0,3}#{1,6}\s*", "", text)
    text = re.sub(r"^\s*\|?\s*:?-{3,}:?(?:\s*\|\s*:?-{3,}:?)*\s*\|?\s*$", "", text)
    return text.strip()


def _safe_filename(text: str) -> str:
    return _sanitize_filename(text).removesuffix(".docx")


def _sanitize_filename(name: str) -> str:
    text = _clean_inline_text(name)
    text = re.sub(r"\{\{[\s\S]*?\}\}", "", text)
    text = text.replace("水印", "")
    text = re.sub(r'[\\/:*?"<>|]', "", text)
    text = re.sub(r"[\r\n\t]+", " ", text)
    text = re.sub(r"\s+", " ", text)
    text = _strip_markdown(text)
    text = text.strip(" .。．、，,；;：:—-_")
    if text.lower().endswith(".docx"):
        text = text[:-5]
    if not text:
        text = "医药器械采购项目分析报告_" + datetime.now().strftime("%Y%m%d%H%M%S")
    text = text[:75].strip(" .。．、，,；;：:—-_")
    return f"{text}.docx"


def _compact_date_for_filename(value: str) -> str:
    text = str(value or "")
    match = re.search(r"(20\d{2})[-/.年]\s*(\d{1,2})[-/.月]\s*(\d{1,2})", text)
    if not match:
        return ""
    year, month, day = match.groups()
    return f"{year}{int(month):02d}{int(day):02d}"


def _infer_filename_region(report: ReportIR, base: str) -> str:
    text = "\n".join([base, report.title, report.document_name, report.source_agency, report.notice_type])
    for region in ["新疆兵团", "河南", "山东", "广东", "国家", "全国"]:
        if region in text:
            return "国家" if region == "全国" else region
    if "兵团" in text and "新疆" in text:
        return "新疆兵团"
    return ""


def _condense_filename_core(text: str) -> str:
    core = _clean_inline_text(text)
    core = re.sub(r"^关于", "", core)
    core = re.sub(r"(的)?(通知|公告|采购文件|征求意见稿|项目分析|分析报告)$", "", core)
    core = re.sub(r"[《》“”\"']", "", core)
    return core[:34].strip(" .。．、，,；;：:—-_") or "医药器械采购项目"


def _build_report_filename(report: ReportIR) -> str:
    if report.suggested_filename:
        return _sanitize_filename(report.suggested_filename)

    source = report.title or report.document_name
    if source:
        region = _infer_filename_region(report, source)
        core = _condense_filename_core(source)
        date = _compact_date_for_filename(report.publish_date)
        if region or date:
            prefix = f"【{region}】" if region else ""
            suffix = f"【{date}】" if date else ""
            return _sanitize_filename(f"{prefix}{core}—项目分析{suffix}")
        return _sanitize_filename(source)

    return _sanitize_filename("")


def _unique_report_filename(filename: str, directory: Path) -> str:
    clean = _sanitize_filename(filename)
    path = directory / clean
    if not path.exists():
        return clean
    stem = path.stem
    suffix = path.suffix
    counter = 1
    while True:
        candidate = f"{stem}({counter}){suffix}"
        if not (directory / candidate).exists():
            return candidate
        counter += 1


def _build_combined_text(title: str, final_url: str, page_text: str, attachments: list[dict[str, Any]]) -> str:
    evidence_text_parts = [title, final_url, page_text]
    for attachment in attachments:
        evidence_text_parts.append(str(attachment.get("filename") or ""))
        evidence_text_parts.append(str(attachment.get("link_text") or ""))
        evidence_text_parts.append(str(attachment.get("text") or ""))
    rule_hints = _rule_completeness_hints("\n".join(evidence_text_parts))
    parts = [
        "# 网页正文",
        f"标题: {title}",
        f"URL: {final_url}",
        page_text,
        rule_hints,
    ]
    for attachment in attachments:
        parts.extend(
            [
                "\n# 附件",
                f"文件名: {attachment['filename']}",
                f"来源: {attachment['url']}",
                f"链接文字: {attachment.get('link_text') or ''}",
                str(attachment.get("text") or ""),
            ]
        )
        for table in attachment.get("tables") or []:
            rows = table.get("rows") or []
            parts.append(_rows_to_markdown(rows[:80], f"{attachment['filename']} / {table.get('name') or table.get('sheet')}"))
    return "\n\n".join(part for part in parts if part)


def _rule_completeness_hints(text: str) -> str:
    found = [keyword for keyword in RULE_COMPLETENESS_KEYWORDS if keyword in str(text or "")]
    if not found:
        return ""
    unique_found = _unique(found)
    return (
        "# 规则完整性保留提示\n"
        "输入材料中出现以下规则章节或关键词。生成报告时如原文有明确内容，应尽量保留为独立小节或分主题表格，"
        "不要压缩成概括段落，也不要编造原文未披露的信息：\n"
        + "\n".join(f"- {keyword}" for keyword in unique_found)
    )


def _trim_attachment_for_response(attachment: dict[str, Any]) -> dict[str, Any]:
    trimmed = dict(attachment)
    trimmed["text"] = _truncate(str(trimmed.get("text") or ""), 1_500)
    trimmed["tables"] = [
        {
            **table,
            "rows": (table.get("rows") or [])[:8],
        }
        for table in (trimmed.get("tables") or [])
    ]
    return trimmed


def _extract_dates(text: str) -> list[str]:
    patterns = [
        r"20\d{2}[-/.年]\d{1,2}[-/.月]\d{1,2}日?",
        r"20\d{2}\s*年\s*\d{1,2}\s*月\s*\d{1,2}\s*日",
    ]
    results: list[str] = []
    for pattern in patterns:
        results.extend(re.findall(pattern, text))
    return results[:20]


def _extract_publish_metadata_dates(html: str) -> list[str]:
    patterns = [
        r"发布时间\s*[：:]\s*(20\d{2}[-/.]\d{1,2}[-/.]\d{1,2})",
        r"发布时间\s*[：:]\s*(20\d{2}\s*年\s*\d{1,2}\s*月\s*\d{1,2}\s*日)",
    ]
    results: list[str] = []
    for pattern in patterns:
        results.extend(re.findall(pattern, html, flags=re.S))
    return results[:5]


def _extract_regions(text: str) -> list[str]:
    regions = [
        "北京", "天津", "上海", "重庆", "河北", "山西", "辽宁", "吉林", "黑龙江", "江苏", "浙江", "安徽", "福建",
        "江西", "山东", "河南", "湖北", "湖南", "广东", "海南", "四川", "贵州", "云南", "陕西", "甘肃", "青海",
        "内蒙古", "广西", "西藏", "宁夏", "新疆", "全国", "兵团",
    ]
    return [region for region in regions if region in text]


def _normalize_text(text: str) -> str:
    text = text.replace("\xa0", " ")
    text = re.sub(r"[ \t\r\f\v]+", " ", text)
    text = re.sub(r"\n\s*\n\s*\n+", "\n\n", text)
    return text.strip()


def _truncate(text: str, limit: int) -> str:
    if len(text) <= limit:
        return text
    head = int(limit * 0.65)
    tail = limit - head
    return text[:head] + "\n\n...【中间内容因长度限制被截断】...\n\n" + text[-tail:]


def _unique(values: list[str]) -> list[str]:
    seen = set()
    result = []
    for value in values:
        value = _normalize_text(value)
        if value and value not in seen:
            seen.add(value)
            result.append(value)
    return result

