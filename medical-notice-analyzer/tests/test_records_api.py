from __future__ import annotations

import os
import shutil
import tempfile
import time
import unittest
import json
from unittest.mock import patch

from docx import Document
from docx.oxml.ns import qn
from fastapi.testclient import TestClient
from openpyxl import Workbook

import app.main as main_module
from app.attachment_cache import cache_key
from app.attachment_fetcher import AttachmentDownloadResult, build_attachment_auth_headers, fetch_attachment_bytes
import app.attachment_parser as attachment_parser_module
from app.attachment_parser import parse_attachment_bytes
from app.diagnostics import build_pack_diagnostics, build_run_diagnostics


def list_row(**overrides):
    row = {
        "menu_code": "m1",
        "articleid": "a1",
        "title": "Stent procurement notice",
        "audittime": "2026-06-01 10:00:00",
        "menu_name": "Procurement",
        "areaname": "山东",
        "source": "Source",
        "publicorg": "Agency",
        "projectphase": "申报",
        "projecttype": "耗材",
        "category": "医用耗材",
        "sourceurl": "https://example.com/a1",
        "summary": "summary",
        "attachment_count": 2,
        "list_attach_filename": "first.docx",
    }
    row.update(overrides)
    return row


def detail_row(**overrides):
    row = {
        **list_row(),
        "updatetime": "2026-06-02 10:00:00",
        "dl_project_type": "带量",
        "referencenumber": "REF-1",
        "policytype": "policy",
        "belongproject": "project",
        "projectabbreviation": "abbr",
        "content": "<h1>公告标题</h1><script>bad()</script><p>申报时间为2026年6月1日。</p><table><tr><td>价格规则</td></tr></table>",
    }
    row.update(overrides)
    return row


def attachment_row(**overrides):
    row = {
        "articleattid": "att1",
        "filename": "file.docx",
        "filepath": "/files/file.docx",
        "fileext": ".docx",
        "filesize": 1234,
        "uploadtime": "2026-06-01 11:00:00",
        "sortnum": 1,
        "fileerrortype": "",
    }
    row.update(overrides)
    return row


def wait_until(predicate, timeout: float = 2.0):
    deadline = time.time() + timeout
    last_value = None
    while time.time() < deadline:
        last_value = predicate()
        if last_value:
            return last_value
        time.sleep(0.02)
    return last_value


class RecordsApiTests(unittest.TestCase):
    def setUp(self) -> None:
        self.client = TestClient(main_module.app)

    def test_records_list_filters_paginates_and_omits_content(self) -> None:
        calls = []

        def fake_fetch_one(sql: str, params: list[object]):
            calls.append(("one", sql, params))
            return {"total": 2}

        def fake_fetch_all(sql: str, params: list[object]):
            calls.append(("all", sql, params))
            return [
                list_row(articleid="a1", title="Stent A"),
                list_row(articleid="a2", title="Stent B", attachment_count=0),
            ]

        with patch.object(main_module, "_db_fetch_one", fake_fetch_one, create=True), patch.object(
            main_module, "_db_fetch_all", fake_fetch_all, create=True
        ):
            response = self.client.get(
                "/records",
                params={
                    "keyword": "stent",
                    "menu_code": "menu-01",
                    "areaname": "山东",
                    "projectphase": "申报",
                    "projecttype": "耗材",
                    "start_date": "2026-06-01",
                    "end_date": "2026-06-30",
                    "page": 2,
                    "page_size": 20,
                },
            )

        self.assertEqual(response.status_code, 200)
        body = response.json()
        self.assertEqual(body["page"], 2)
        self.assertEqual(body["page_size"], 20)
        self.assertEqual(body["total"], 2)
        self.assertEqual(body["total_pages"], 1)
        self.assertEqual(len(body["items"]), 2)
        self.assertNotIn("content", body["items"][0])

        list_sql = calls[1][1]
        list_params = calls[1][2]
        self.assertIn("a.status = %s", list_sql)
        self.assertIn("LEFT JOIN", list_sql)
        self.assertIn("GROUP BY", list_sql)
        self.assertIn("ORDER BY a.audittime DESC", list_sql)
        self.assertIn("LIMIT %s OFFSET %s", list_sql)
        self.assertIn("%stent%", list_params)
        self.assertEqual(list_params[-2:], [20, 20])

    def test_records_ui_serves_static_page(self) -> None:
        response = self.client.get("/records-ui")

        self.assertEqual(response.status_code, 200)
        self.assertIn("text/html", response.headers["content-type"])
        self.assertIn("数据库文章选材", response.text)
        self.assertIn("/analysis/prepare", response.text)
        self.assertIn("/analysis/run", response.text)
        self.assertIn("run-analysis-report", response.text)
        self.assertNotIn("enableAttachmentDownload", response.text)
        self.assertNotIn("attachmentCookie", response.text)
        self.assertNotIn("attachmentHeadersJson", response.text)
        self.assertIn("复制 pack_id", response.text)
        self.assertIn("查看证据包摘要", response.text)
        self.assertNotIn("查看完整 evidence_pack</a>", response.text)

    def test_records_list_tolerates_nullable_optional_fields(self) -> None:
        with patch.object(main_module, "_db_fetch_one", return_value={"total": 1}, create=True), patch.object(
            main_module,
            "_db_fetch_all",
            return_value=[list_row(projectphase=None, list_attach_filename=None, attachment_count=None)],
            create=True,
        ):
            response = self.client.get("/records?page=1&page_size=5")

        self.assertEqual(response.status_code, 200)
        item = response.json()["items"][0]
        self.assertEqual(item["projectphase"], "")
        self.assertEqual(item["list_attach_filename"], "")
        self.assertEqual(item["attachment_count"], 0)

    def test_record_detail_returns_clean_text_and_attachments(self) -> None:
        def fake_fetch_one(sql: str, params: list[object]):
            self.assertIn("sample_article_wide", sql)
            self.assertEqual(params, [0, "m1", "a1"])
            return {
                **list_row(),
                "updatetime": "2026-06-02 10:00:00",
                "dl_project_type": "带量",
                "referencenumber": "REF-1",
                "policytype": "policy",
                "belongproject": "project",
                "projectabbreviation": "abbr",
                "content": "<h1>Title</h1><p>Hello <b>world</b></p>",
            }

        def fake_fetch_all(sql: str, params: list[object]):
            self.assertIn("sample_article_attach", sql)
            self.assertEqual(params, ["m1", "a1"])
            return [
                {
                    "articleattid": "att1",
                    "filename": "file.docx",
                    "filepath": "/files/file.docx",
                    "fileext": ".docx",
                    "filesize": 1234,
                    "uploadtime": "2026-06-01 11:00:00",
                    "sortnum": 1,
                }
            ]

        with patch.object(main_module, "_db_fetch_one", fake_fetch_one, create=True), patch.object(
            main_module, "_db_fetch_all", fake_fetch_all, create=True
        ):
            response = self.client.get("/records/m1/a1")

        self.assertEqual(response.status_code, 200)
        body = response.json()
        self.assertEqual(body["menu_code"], "m1")
        self.assertIn("Hello world", body["content_text"])
        self.assertEqual(body["attachments"][0]["filename"], "file.docx")

    def test_selection_preview_validates_limits_overlap_and_missing_records(self) -> None:
        with patch.object(main_module, "_records_by_keys", return_value={}, create=True):
            response = self.client.post(
                "/analysis/selection/preview",
                json={"primary_materials": [], "auxiliary_materials": []},
            )
        self.assertEqual(response.status_code, 422)
        self.assertIn("primary_materials", response.text)

        with patch.object(main_module, "_records_by_keys", return_value={}, create=True):
            response = self.client.post(
                "/analysis/selection/preview",
                json={
                    "primary_materials": [{"menu_code": "m1", "articleid": "a1"}],
                    "auxiliary_materials": [{"menu_code": "m1", "articleid": "a1"}],
                },
            )
        self.assertEqual(response.status_code, 422)
        self.assertIn("同时作为主分析材料和辅助分析材料", response.text)

        with patch.object(main_module, "_records_by_keys", return_value={}, create=True):
            response = self.client.post(
                "/analysis/selection/preview",
                json={
                    "primary_materials": [{"menu_code": "m1", "articleid": "missing"}],
                    "auxiliary_materials": [],
                },
            )
        self.assertEqual(response.status_code, 422)
        self.assertIn("不存在", response.text)

    def test_selection_preview_returns_primary_and_auxiliary_records(self) -> None:
        records = {
            ("m1", "a1"): list_row(menu_code="m1", articleid="a1", title="Primary"),
            ("m2", "a2"): list_row(menu_code="m2", articleid="a2", title="Auxiliary", areaname="北京"),
        }

        with patch.object(main_module, "_records_by_keys", return_value=records, create=True):
            response = self.client.post(
                "/analysis/selection/preview",
                json={
                    "primary_materials": [{"menu_code": "m1", "articleid": "a1"}],
                    "auxiliary_materials": [{"menu_code": "m2", "articleid": "a2"}],
                    "enable_attachment_download": False,
                },
            )

        self.assertEqual(response.status_code, 200)
        body = response.json()
        self.assertTrue(body["success"])
        self.assertEqual(body["primary_materials"][0]["title"], "Primary")
        self.assertEqual(body["auxiliary_materials"][0]["areaname"], "北京")

    def test_analysis_prepare_builds_and_persists_database_evidence_pack(self) -> None:
        def fake_fetch_all(sql: str, params: list[object]):
            if "sample_article_wide" in sql:
                return [
                    detail_row(menu_code="m1", articleid="a1", title="Primary", summary="primary summary"),
                    detail_row(menu_code="m2", articleid="a2", title="Aux", summary="", content="<p>辅助正文内容足够用于摘要。</p>"),
                ]
            if "sample_article_attach" in sql:
                return [
                    attachment_row(menu_code="m1", articleid="a1", articleattid="att1", filename="primary.pdf", fileext=".pdf"),
                    attachment_row(menu_code="m2", articleid="a2", articleattid="att2", filename="aux.xlsx", fileext=".xlsx"),
                ]
            return []

        with tempfile.TemporaryDirectory() as tmpdir, patch.object(
            main_module, "_db_fetch_all", fake_fetch_all, create=True
        ), patch.object(main_module, "_database_evidence_pack_dir", return_value=main_module.Path(tmpdir), create=True):
            response = self.client.post(
                "/analysis/prepare",
                json={
                    "primary_materials": [{"menu_code": "m1", "articleid": "a1"}],
                    "auxiliary_materials": [{"menu_code": "m2", "articleid": "a2"}],
                    "enable_attachment_download": False,
                },
            )

            self.assertEqual(response.status_code, 200)
            body = response.json()
            self.assertTrue(body["success"])
            self.assertTrue(body["pack_id"].startswith("pack_"))
            self.assertEqual(body["primary_count"], 1)
            self.assertEqual(body["auxiliary_count"], 1)
            self.assertEqual(body["attachment_count"], 2)

            pack = body["evidence_pack"]
            self.assertEqual(pack["source"], "database_selection")
            self.assertEqual(pack["primary_materials"][0]["material_role"], "primary")
            self.assertEqual(pack["auxiliary_materials"][0]["material_role"], "auxiliary")
            self.assertIn("公告标题", pack["primary_materials"][0]["content_text"])
            self.assertNotIn("<script>", pack["primary_materials"][0]["content_text"])
            self.assertEqual(pack["primary_materials"][0]["content_summary"], "primary summary")
            attachment = pack["primary_materials"][0]["attachments"][0]
            self.assertEqual(attachment["download_method"], "internal_attid")
            self.assertEqual(attachment["parse_status"], "metadata_only")
            self.assertIn("{articleattid}", attachment["download_url_template"])
            self.assertEqual(pack["pack_version"], "2.0")
            self.assertIn("primary_evidence", pack)
            self.assertIn("auxiliary_evidence", pack)
            self.assertIn("attachment_evidence", pack)
            self.assertIn("generation_guidance", pack)
            self.assertIn("relation_to_primary", pack["auxiliary_materials"][0])

            pack_response = self.client.get(f"/analysis/packs/{body['pack_id']}")
            self.assertEqual(pack_response.status_code, 200)
            self.assertEqual(pack_response.json()["pack_id"], body["pack_id"])

            summary_response = self.client.get(f"/analysis/packs/{body['pack_id']}/summary")
            self.assertEqual(summary_response.status_code, 200)
            self.assertEqual(summary_response.json()["attachment_count"], 2)

    def test_attachment_download_without_cookie_attempts_request_and_handles_rejection(self) -> None:
        with patch.dict(
            main_module.os.environ,
            {
                "ENABLE_ATTACHMENT_DOWNLOAD": "true",
                "ATTACHMENT_COOKIE": "",
                "ATTACHMENT_HEADERS_JSON": "",
            },
        ), patch("app.attachment_fetcher.httpx.Client") as client_cls:
            client_cls.return_value.__enter__.return_value.get.return_value = main_module.httpx.Response(403, text="forbidden")
            result = fetch_attachment_bytes({"articleattid": "att-auth", "filename": "a.pdf"})

        self.assertEqual(result.download_status, "download_failed")
        self.assertEqual(result.auth_mode, "none")
        self.assertIsNone(result.content)

    def test_attachment_headers_include_cookie(self) -> None:
        with patch.dict(
            main_module.os.environ,
            {
                "ATTACHMENT_COOKIE": "SESSION=secret",
                "ATTACHMENT_HEADERS_JSON": '{"X-Test":"ok"}',
            },
        ):
            headers, mode = build_attachment_auth_headers()

        self.assertEqual(mode, "configured_cookie")
        self.assertEqual(headers["Cookie"], "SESSION=secret")
        self.assertEqual(headers["X-Test"], "ok")

    def test_docx_and_excel_attachment_parsers_return_summaries(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            doc_path = main_module.Path(tmpdir) / "notice.docx"
            doc = Document()
            doc.add_paragraph("申报时间为2026年6月，企业需要关注价格规则。")
            table = doc.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "企业名称"
            table.cell(0, 1).text = "产品名称"
            table.cell(1, 0).text = "企业A"
            table.cell(1, 1).text = "产品B"
            doc.save(doc_path)

            wb = Workbook()
            ws = wb.active
            ws.title = "产品清单"
            ws.append(["企业名称", "产品名称", "注册证号", "医保编码", "申报价"])
            ws.append(["企业A", "产品B", "证号1", "码1", 10])
            excel_path = main_module.Path(tmpdir) / "list.xlsx"
            wb.save(excel_path)

            doc_result = parse_attachment_bytes(doc_path.read_bytes(), "notice.docx", ".docx", 1024)
            excel_result = parse_attachment_bytes(excel_path.read_bytes(), "list.xlsx", ".xlsx", 2048)

        self.assertIn("parsed_summary", doc_result["parse_statuses"])
        self.assertIn("申报时间", doc_result["summary"])
        self.assertIn("parsed_table_summary", excel_result["parse_statuses"])
        self.assertEqual(excel_result["table_summaries"][0]["sheet_name"], "产品清单")
        self.assertIn("企业名称", excel_result["table_summaries"][0]["key_columns"])
        self.assertTrue(excel_result["table_summaries"][0]["contains_enterprise"])
        self.assertTrue(excel_result["table_summaries"][0]["contains_product"])
        self.assertTrue(excel_result["table_summaries"][0]["contains_price"])

    def test_doc_attachment_parser_converts_with_libreoffice_and_reuses_docx_parser(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            converted_docx = main_module.Path(tmpdir) / "converted.docx"
            doc = Document()
            doc.add_paragraph("申报时间为2026年6月，企业需要关注价格规则。")
            table = doc.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "企业名称"
            table.cell(0, 1).text = "产品名称"
            table.cell(1, 0).text = "企业A"
            table.cell(1, 1).text = "产品B"
            doc.save(converted_docx)

            def fake_run(command, capture_output, text, timeout, check):
                outdir = main_module.Path(command[command.index("--outdir") + 1])
                shutil.copyfile(converted_docx, outdir / "legacy.docx")
                return main_module.subprocess.CompletedProcess(command, 0, "", "")

            with patch("app.attachment_parser.shutil.which", return_value="/usr/bin/libreoffice"), patch(
                "app.attachment_parser.subprocess.run", side_effect=fake_run
            ):
                result = parse_attachment_bytes(b"legacy-binary-doc", "legacy.doc", ".doc", 1024)

        self.assertIn("temp_file_parsed", result["parse_statuses"])
        self.assertIn("parsed_text", result["parse_statuses"])
        self.assertIn("parsed_summary", result["parse_statuses"])
        self.assertIn("申报时间", result["summary"])
        self.assertTrue(result["table_summaries"][0]["contains_enterprise"])
        self.assertTrue(result["table_summaries"][0]["contains_product"])

    def test_doc_attachment_parser_returns_structured_failure_when_converter_missing(self) -> None:
        with patch("app.attachment_parser.shutil.which", return_value=None):
            result = parse_attachment_bytes(b"legacy-binary-doc", "legacy.doc", ".doc", 1024)

        self.assertEqual(result["parse_statuses"], ["parse_failed"])
        self.assertIn("LibreOffice", result["summary"])
        self.assertTrue(result["warnings"])

    def test_excel_parser_detects_header_row_below_title_row(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            workbook_path = main_module.Path(tmpdir) / "price-list.xlsx"
            wb = Workbook()
            ws = wb.active
            ws.title = "价格表"
            ws.append(["某某项目中选结果价格表", "", "", "", ""])
            ws.append(["企业名称", "产品名称", "注册证号", "医保编码", "中选价"])
            ws.append(["企业A", "产品B", "证号1", "C001", "12.50"])
            wb.save(workbook_path)

            result = parse_attachment_bytes(workbook_path.read_bytes(), "price-list.xlsx", ".xlsx", workbook_path.stat().st_size)

        table = result["table_summaries"][0]
        self.assertEqual(table["headers"][:5], ["企业名称", "产品名称", "注册证号", "医保编码", "中选价"])
        self.assertIn("企业名称", table["key_columns"])
        self.assertTrue(table["contains_enterprise"])
        self.assertTrue(table["contains_product"])
        self.assertTrue(table["contains_registration_cert"])
        self.assertTrue(table["contains_medical_insurance_code"])
        self.assertTrue(table["contains_price"])

    def test_csv_parser_streams_and_detects_header_row_below_title_row(self) -> None:
        rows = ["某某项目产品清单,,,,", "企业名称,产品名称,注册证号,医保编码,采购量"]
        rows.extend(f"企业{i},产品{i},证号{i},C{i},{i}" for i in range(200))
        content = "\n".join(rows).encode("utf-8-sig")

        result = parse_attachment_bytes(content, "list.csv", ".csv", len(content))

        table = result["table_summaries"][0]
        self.assertEqual(table["rows"], 202)
        self.assertEqual(table["headers"][:5], ["企业名称", "产品名称", "注册证号", "医保编码", "采购量"])
        self.assertTrue(table["contains_enterprise"])
        self.assertTrue(table["contains_purchase_volume"])
        self.assertEqual(table["sample_rows_count"], 20)

    def test_excel_parser_adds_field_stats_and_table_heavy_hint(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            workbook_path = main_module.Path(tmpdir) / "selected-result.xlsx"
            wb = Workbook()
            ws = wb.active
            ws.title = "selected_result"
            ws.append(["enterprise", "product", "registration_cert", "medical_insurance_code", "selected_price", "purchase_volume", "selected_status", "region"])
            for index in range(20):
                ws.append(
                    [
                        f"Enterprise {index % 4}",
                        f"Product {index}",
                        f"CERT-{index}",
                        f"MI-{index}",
                        str(10 + index * 0.5),
                        str(100 + index),
                        "selected" if index % 2 == 0 else "pending",
                        "Region A",
                    ]
                )
            wb.save(workbook_path)

            with patch.dict(os.environ, {"TABLE_HEAVY_ROW_THRESHOLD": "10", "TABLE_STATS_MAX_SCAN_ROWS": "12"}, clear=False):
                result = parse_attachment_bytes(workbook_path.read_bytes(), "selected-result.xlsx", ".xlsx", workbook_path.stat().st_size)

        table = result["table_summaries"][0]
        self.assertTrue(table["table_heavy"])
        self.assertIn("enterprise", table["enterprise_columns"])
        self.assertIn("selected_price", table["price_columns"])
        self.assertIn("purchase_volume", table["purchase_volume_columns"])
        self.assertEqual(table["field_stats"]["enterprise"]["non_empty_count"], 12)
        self.assertLessEqual(table["field_stats"]["enterprise"]["unique_count"], 4)
        self.assertEqual(table["field_stats"]["price"]["min"], 10.0)
        self.assertEqual(table["field_stats"]["price"]["max"], 15.5)
        self.assertIn("recommended_report_usage", table)
        self.assertGreater(table["evidence_value_score"], 0)

    def test_old_url_attachment_discovery_includes_text_html_and_zip_files(self) -> None:
        html = """
        <a href="/notice.txt">附件 TXT</a>
        <a href="/guide.html">附件 HTML</a>
        <a href="/archive.zip">附件 ZIP</a>
        """
        links = main_module._discover_attachment_links(html, "https://example.com/page", 10)

        urls = {item["url"] for item in links}
        self.assertIn("https://example.com/notice.txt", urls)
        self.assertIn("https://example.com/guide.html", urls)
        self.assertIn("https://example.com/archive.zip", urls)

    def test_attachment_parse_cache_reuses_success_and_marks_core_attachment(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            workbook_path = main_module.Path(tmpdir) / "core.xlsx"
            wb = Workbook()
            ws = wb.active
            ws.title = "产品清单"
            ws.append(["企业名称", "产品名称", "注册证号", "医保编码", "中选价", "采购量"])
            ws.append(["企业A", "产品A", "国械注准1", "C001", "12.50", "100"])
            wb.save(workbook_path)
            content = workbook_path.read_bytes()

            row = attachment_row(
                articleattid="att-core-cache",
                filename="中选结果产品清单.xlsx",
                fileext=".xlsx",
                filesize=len(content),
                uploadtime="2026-06-01 11:00:00",
            )
            env = {
                "ENABLE_ATTACHMENT_PARSE_CACHE": "true",
                "ATTACHMENT_PARSE_CACHE_DIR": str(main_module.Path(tmpdir) / "cache"),
            }
            with patch.dict(os.environ, env, clear=False), patch(
                "app.main.fetch_attachment_bytes",
                return_value=AttachmentDownloadResult("downloaded", "none", content=content, content_type="", filename=row["filename"], warnings=[]),
            ) as fetch:
                first = main_module._database_attachment_metadata(row, {"enable_download": True})
                second = main_module._database_attachment_metadata(row, {"enable_download": True})

        self.assertEqual(fetch.call_count, 1)
        self.assertTrue(first["core_attachment"])
        self.assertEqual(first["business_type"], "中选结果")
        self.assertEqual(first["cache_status"], "cache_miss")
        self.assertEqual(second["cache_status"], "cache_hit_success")
        self.assertTrue(second["table_summaries"][0]["contains_enterprise"])
        self.assertTrue(second["table_summaries"][0]["contains_product"])
        self.assertTrue(second["table_summaries"][0]["contains_price"])
        self.assertFalse(second["stored_original_file"])

    def test_attachment_parse_failure_cache_is_short_and_force_refreshable(self) -> None:
        row = attachment_row(articleattid="att-fail-cache", filename="价格表.xlsx", fileext=".xlsx", filesize=123, uploadtime="2026")
        with tempfile.TemporaryDirectory() as tmpdir:
            env = {
                "ENABLE_ATTACHMENT_PARSE_CACHE": "true",
                "ATTACHMENT_PARSE_CACHE_DIR": str(main_module.Path(tmpdir) / "cache"),
                "ATTACHMENT_PARSE_CACHE_FAILURE_TTL_MINUTES": "10",
            }
            with patch.dict(os.environ, env, clear=False), patch(
                "app.main.fetch_attachment_bytes",
                return_value=AttachmentDownloadResult("network_unreachable", "none", warnings=["timeout"]),
            ) as fetch:
                first = main_module._database_attachment_metadata(row, {"enable_download": True})
                second = main_module._database_attachment_metadata(row, {"enable_download": True})
                refreshed = main_module._database_attachment_metadata(row, {"enable_download": True, "force_refresh": True})

        self.assertEqual(fetch.call_count, 2)
        self.assertEqual(first["parse_status"], "network_unreachable")
        self.assertEqual(first["cache_status"], "cache_miss")
        self.assertEqual(second["cache_status"], "cache_hit_failure_short")
        self.assertGreaterEqual(second["retry_after_seconds"], 1)
        self.assertEqual(refreshed["cache_status"], "force_refreshed")
        self.assertTrue(refreshed["core_attachment_unavailable"])

    def test_medical_service_price_pdf_is_core_attachment(self) -> None:
        row = attachment_row(
            articleattid="att-service-price",
            filename="安庆市医疗保障局关于规范整合呼吸系统、神经系统等十二类医疗服务价格项目的通知.pdf",
            fileext=".pdf",
            filesize=123,
        )

        result = main_module._database_attachment_metadata(row, {"enable_download": False})

        self.assertTrue(result["core_attachment"])
        self.assertEqual(result["business_type"], "医疗服务价格项目")
        self.assertTrue(result["core_attachment_unavailable"])

    def test_large_pdf_extracts_limited_text_summary_instead_of_metadata_placeholder(self) -> None:
        extracted = (
            "安庆市医疗保障局关于规范整合呼吸系统、神经系统等十二类医疗服务价格项目的通知。"
            "本次规范整合涉及呼吸系统、神经系统、泌尿系统等医疗服务价格项目，"
            "明确项目内涵、计价单位、医保支付类别和执行要求。"
        ) * 12

        with patch.dict(os.environ, {"ATTACHMENT_MAX_PARSE_MB": "0.001"}, clear=False), patch.object(
            attachment_parser_module,
            "_parse_pdf",
            return_value=extracted,
        ) as parse_pdf:
            result = parse_attachment_bytes(b"%PDF large placeholder", "安庆医疗服务价格项目.pdf", ".pdf", 39_319_671)

        parse_pdf.assert_called_once()
        self.assertIn("too_large_summary_only", result["parse_statuses"])
        self.assertIn("parsed_summary", result["parse_statuses"])
        self.assertGreater(result["text_length"], 800)
        self.assertIn("医疗服务价格项目", result["summary"])
        self.assertGreater(len(result["important_sections"]), 0)

    def test_large_scanned_pdf_uses_ocr_fallback_when_text_extraction_is_empty(self) -> None:
        ocr_text = (
            "安庆市医疗保障局规范整合十二类医疗服务价格项目。"
            "附件列明项目名称、项目内涵、计价单位、医保支付类别和执行要求。"
        ) * 18

        with patch.dict(os.environ, {"ATTACHMENT_MAX_PARSE_MB": "0.001"}, clear=False), patch.object(
            attachment_parser_module,
            "_parse_pdf",
            return_value="",
        ), patch.object(attachment_parser_module, "_ocr_pdf_text", return_value=(ocr_text, ["OCR completed"])):
            result = parse_attachment_bytes(b"%PDF scanned placeholder", "安庆医疗服务价格项目.pdf", ".pdf", 39_319_671)

        self.assertIn("too_large_summary_only", result["parse_statuses"])
        self.assertIn("parsed_summary", result["parse_statuses"])
        self.assertGreater(result["text_length"], 800)
        self.assertIn("医保支付类别", result["summary"])
        self.assertTrue(any("OCR completed" in warning for warning in result["warnings"]))

    def test_large_pdf_without_extractable_text_is_parse_failed_not_fake_summary(self) -> None:
        with patch.dict(os.environ, {"ATTACHMENT_MAX_PARSE_MB": "0.001"}, clear=False), patch.object(
            attachment_parser_module,
            "_parse_pdf",
            return_value="",
        ), patch.object(attachment_parser_module, "_ocr_pdf_text", return_value=("", ["OCR unavailable"])):
            result = parse_attachment_bytes(b"%PDF scanned placeholder", "安庆医疗服务价格项目.pdf", ".pdf", 39_319_671)

        self.assertIn("parse_failed", result["parse_statuses"])
        self.assertEqual(result["text_length"], 0)
        self.assertNotIn("parsed_summary", result["parse_statuses"])

    def test_attachment_cache_key_includes_parser_version(self) -> None:
        row = attachment_row(articleattid="att-doc-cache", filename="legacy.doc", fileext=".doc", filesize=123, uploadtime="2026")

        with patch.dict(os.environ, {"ATTACHMENT_PARSER_VERSION": "old-parser"}, clear=False):
            old_key = cache_key(row)
        with patch.dict(os.environ, {"ATTACHMENT_PARSER_VERSION": "doc-libreoffice-parser"}, clear=False):
            new_key = cache_key(row)

        self.assertNotEqual(old_key, new_key)

    def test_run_diagnostics_uses_wps_like_visible_word_count(self) -> None:
        markdown = "# Report\n\n**\u4ef7\u683c\u89c4\u5219**: \u4f01\u4e1a impact 2026-06-16, 12.5%. <span class=\"analysis-highlight\">\u9700\u5173\u6ce8</span>"

        diagnostics = build_run_diagnostics(
            {
                "report_markdown": markdown,
                "quality_check": {"passed": True, "issues": []},
                "generation_warnings": [],
                "remaining_issues": [],
            },
            {"primary_materials": [], "auxiliary_materials": [], "warnings": []},
        )

        count = diagnostics["report"]["report_markdown_chars"]
        self.assertLess(count, len(markdown))
        self.assertEqual(count, 13)

    def test_analysis_prepare_downloads_attachment_by_default_without_cookie(self) -> None:
        content = "default intranet attachment content with price rule and product scope".encode("utf-8")
        with patch.dict(
            main_module.os.environ,
            {
                "ENABLE_ATTACHMENT_PARSE": "true",
                "ENABLE_ATTACHMENT_PARSE_CACHE": "false",
                "ATTACHMENT_COOKIE": "",
                "ATTACHMENT_HEADERS_JSON": "",
                "ELIAN_QX_COOKIE": "",
            },
        ), patch.object(main_module, "_fetch_database_material_rows", return_value={("m1", "a1"): detail_row()}, create=True), patch.object(
            main_module,
            "_fetch_database_attachments",
            return_value={("m1", "a1"): [attachment_row(articleattid="att-default", filename="notice.txt", fileext=".txt")]},
            create=True,
        ), patch.object(
            main_module,
            "fetch_attachment_bytes",
            return_value=AttachmentDownloadResult(
                "downloaded",
                "none",
                content=content,
                content_type="text/plain",
                filename="notice.txt",
                warnings=[],
            ),
        ) as fetch_mock, tempfile.TemporaryDirectory() as tmpdir, patch.object(
            main_module, "_database_evidence_pack_dir", return_value=main_module.Path(tmpdir), create=True
        ):
            response = self.client.post(
                "/analysis/prepare",
                json={"primary_materials": [{"menu_code": "m1", "articleid": "a1"}], "auxiliary_materials": []},
            )

        self.assertEqual(response.status_code, 200)
        fetch_mock.assert_called_once()
        _, kwargs = fetch_mock.call_args
        self.assertTrue(kwargs["enable_download"])
        self.assertEqual(kwargs["user_cookie"], "")
        self.assertEqual(kwargs["user_headers"], {})
        attachment = response.json()["evidence_pack"]["primary_materials"][0]["attachments"][0]
        self.assertEqual(attachment["download_auth_mode"], "none")
        self.assertIn("parsed_summary", attachment["parse_statuses"])
        self.assertGreater(attachment["text_length"], 0)
        self.assertFalse(attachment["stored_original_file"])

    def test_fetch_attachment_marks_network_unreachable(self) -> None:
        with patch.dict(main_module.os.environ, {"ENABLE_ATTACHMENT_DOWNLOAD": "true"}), patch(
            "app.attachment_fetcher.httpx.Client"
        ) as client_cls:
            client_cls.return_value.__enter__.return_value.get.side_effect = main_module.httpx.ConnectError("no route")
            result = fetch_attachment_bytes({"articleattid": "att1", "filename": "a.pdf"})

        self.assertEqual(result.download_status, "network_unreachable")
        self.assertEqual(result.auth_mode, "none")

    def test_analysis_prepare_can_disable_attachment_download_per_request(self) -> None:
        with patch.dict(
            main_module.os.environ,
            {
                "ENABLE_ATTACHMENT_DOWNLOAD": "true",
                "ENABLE_ATTACHMENT_PARSE": "true",
                "ATTACHMENT_COOKIE": "",
                "ATTACHMENT_HEADERS_JSON": "",
                "ELIAN_QX_COOKIE": "",
            },
        ), patch.object(main_module, "_fetch_database_material_rows", return_value={("m1", "a1"): detail_row()}, create=True), patch.object(
            main_module, "_fetch_database_attachments", return_value={("m1", "a1"): [attachment_row(articleattid="att-auth")]}, create=True
        ), tempfile.TemporaryDirectory() as tmpdir, patch.object(
            main_module, "_database_evidence_pack_dir", return_value=main_module.Path(tmpdir), create=True
        ):
            response = self.client.post(
                "/analysis/prepare",
                json={
                    "primary_materials": [{"menu_code": "m1", "articleid": "a1"}],
                    "auxiliary_materials": [],
                    "enable_attachment_download": False,
                },
            )

        self.assertEqual(response.status_code, 200)
        attachment = response.json()["evidence_pack"]["primary_materials"][0]["attachments"][0]
        self.assertEqual(attachment["parse_status"], "metadata_only")
        self.assertEqual(attachment["download_auth_mode"], "none")
        self.assertFalse(attachment["stored_original_file"])

    def test_analysis_prepare_uses_request_cookie_to_download_and_parse_attachment(self) -> None:
        content = "申报时间为2026年6月，企业需要关注价格规则和产品范围。".encode("utf-8")
        with patch.dict(
            main_module.os.environ,
            {
                "ENABLE_ATTACHMENT_DOWNLOAD": "false",
                "ENABLE_ATTACHMENT_PARSE": "true",
                "ENABLE_ATTACHMENT_PARSE_CACHE": "false",
                "ATTACHMENT_COOKIE": "",
                "ATTACHMENT_HEADERS_JSON": "",
                "ELIAN_QX_COOKIE": "",
            },
        ), patch.object(main_module, "_fetch_database_material_rows", return_value={("m1", "a1"): detail_row()}, create=True), patch.object(
            main_module,
            "_fetch_database_attachments",
            return_value={("m1", "a1"): [attachment_row(articleattid="att-txt", filename="notice.txt", fileext=".txt")]},
            create=True,
        ), patch.object(
            main_module,
            "fetch_attachment_bytes",
            return_value=AttachmentDownloadResult(
                "downloaded",
                "user_cookie",
                content=content,
                content_type="text/plain",
                filename="notice.txt",
                warnings=[],
            ),
        ) as fetch_mock, tempfile.TemporaryDirectory() as tmpdir, patch.object(
            main_module, "_database_evidence_pack_dir", return_value=main_module.Path(tmpdir), create=True
        ):
            response = self.client.post(
                "/analysis/prepare",
                json={
                    "primary_materials": [{"menu_code": "m1", "articleid": "a1"}],
                    "auxiliary_materials": [],
                    "enable_attachment_download": True,
                    "attachment_cookie": "SESSION=temp-cookie",
                    "attachment_headers": {"X-Test": "ok"},
                },
            )

        self.assertEqual(response.status_code, 200)
        fetch_mock.assert_called_once()
        _, kwargs = fetch_mock.call_args
        self.assertTrue(kwargs["enable_download"])
        self.assertEqual(kwargs["user_cookie"], "SESSION=temp-cookie")
        self.assertEqual(kwargs["user_headers"], {"X-Test": "ok"})
        attachment = response.json()["evidence_pack"]["primary_materials"][0]["attachments"][0]
        self.assertEqual(attachment["download_auth_mode"], "user_cookie")
        self.assertIn("parsed_summary", attachment["parse_statuses"])
        self.assertIn("申报时间", attachment["summary"])
        self.assertGreater(attachment["text_length"], 0)
        self.assertFalse(attachment["stored_original_file"])

    def test_analysis_prepare_returns_structured_validation_error(self) -> None:
        response = self.client.post(
            "/analysis/prepare",
            json={"primary_materials": [], "auxiliary_materials": []},
        )

        self.assertEqual(response.status_code, 422)
        body = response.json()
        self.assertFalse(body["success"])
        self.assertEqual(body["error"]["code"], "INVALID_SELECTION")

    def test_analysis_prepare_rejects_missing_database_records(self) -> None:
        with patch.object(main_module, "_db_fetch_all", return_value=[], create=True):
            response = self.client.post(
                "/analysis/prepare",
                json={
                    "primary_materials": [{"menu_code": "m1", "articleid": "missing"}],
                    "auxiliary_materials": [],
                },
            )

        self.assertEqual(response.status_code, 422)
        body = response.json()
        self.assertFalse(body["success"])
        self.assertEqual(body["error"]["code"], "MATERIAL_NOT_FOUND")

    def test_get_analysis_pack_returns_404_for_missing_pack(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir, patch.object(
            main_module, "_database_evidence_pack_dir", return_value=main_module.Path(tmpdir), create=True
        ):
            response = self.client.get("/analysis/packs/pack_missing")

        self.assertEqual(response.status_code, 404)
        self.assertIn("not found", response.text)

    def test_get_analysis_pack_defaults_to_dify_safe_compact_response(self) -> None:
        long_attachment_summary = "附件摘要" * 7000
        pack = {
            "pack_id": "pack_20260615_bigpack1",
            "created_at": "2026-06-15 10:00:00",
            "pack_version": "2.0",
            "source": "database_selection",
            "primary_materials": [
                {
                    "material_role": "primary",
                    "menu_code": "m1",
                    "articleid": "a1",
                    "title": "Big notice",
                    "content_text": "正文" * 12000,
                    "content_text_length": 24000,
                    "content_summary": "summary",
                    "attachments": [
                        {
                            "articleattid": "att1",
                            "filename": "中选结果价格表.pdf",
                            "fileext": ".pdf",
                            "core_attachment": True,
                            "business_type": "中选结果",
                            "parse_status": "parsed_summary",
                            "text_length": 50000,
                            "summary": long_attachment_summary,
                            "important_sections": ["重点内容" * 500],
                            "table_summaries": [],
                            "warnings": ["附件仅解析到元数据层面，不应混入正式正文"],
                        }
                    ],
                }
            ],
            "auxiliary_materials": [],
            "warnings": [],
            "primary_evidence": {"attachment_summaries": [{"summary": long_attachment_summary}]},
            "attachment_evidence": {"parsed_summaries": [{"summary": long_attachment_summary}], "warnings": []},
            "generation_guidance": {"do_not_use_as_basis": []},
        }
        with tempfile.TemporaryDirectory() as tmpdir, patch.object(
            main_module, "_database_evidence_pack_dir", return_value=main_module.Path(tmpdir), create=True
        ):
            main_module._write_database_evidence_pack(pack)
            compact_response = self.client.get("/analysis/packs/pack_20260615_bigpack1")
            full_response = self.client.get("/analysis/packs/pack_20260615_bigpack1?full=true")

        self.assertEqual(compact_response.status_code, 200)
        compact = compact_response.json()
        self.assertTrue(compact["dify_compacted"])
        self.assertLess(len(json.dumps(compact, ensure_ascii=False)), 80000)
        self.assertEqual(compact["input_strategy"], "full_input")
        self.assertEqual(compact["primary_materials"][0]["attachments"][0]["summary"], long_attachment_summary)
        self.assertLessEqual(compact["dify_relevant_input_chars"], compact["full_input_max_chars"])
        self.assertEqual(compact["pack_variant"], "dify_compact")
        self.assertIn("compression_strategy", compact)
        self.assertIn("omitted_content", compact)
        self.assertEqual(compact["primary_materials"][0]["attachments"][0]["business_type"], "中选结果")
        self.assertNotIn("附件仅解析到元数据层面", json.dumps(compact["primary_materials"][0]["attachments"][0], ensure_ascii=False))
        guidance_text = json.dumps(compact["generation_guidance"], ensure_ascii=False)
        self.assertNotIn("资料说明", guidance_text)
        self.assertNotIn("声  明", guidance_text)
        self.assertIn("generation_warnings", guidance_text)
        self.assertEqual(full_response.status_code, 200)
        self.assertGreater(len(json.dumps(full_response.json(), ensure_ascii=False)), 80000)

    def test_dify_compact_preserves_single_primary_material_when_under_limit(self) -> None:
        content_text = "primary rules price enterprise execution " * 260
        pack = {
            "pack_id": "pack_single_primary",
            "created_at": "2026-06-16 10:00:00",
            "source": "database_selection",
            "primary_materials": [
                {
                    "material_role": "primary",
                    "menu_code": "m1",
                    "articleid": "a1",
                    "title": "Single primary material",
                    "content_text": content_text,
                    "content_text_length": len(content_text),
                    "content_summary": "primary summary",
                    "key_facts": [{"name": "execution date", "value": "2026-06-01"}],
                    "important_passages": [{"reason": "price rule", "text": "Products use selected prices."}],
                    "attachments": [
                        {
                            "articleattid": "att1",
                            "filename": "selected-result.xlsx",
                            "core_attachment": True,
                            "business_type": "中选结果",
                            "parse_status": "parsed_table_summary",
                            "summary": "Attachment contains enterprise, product and selected price fields.",
                            "key_facts": ["enterprise/product/selected price columns exist"],
                            "table_summaries": [
                                {
                                    "sheet_name": "result",
                                    "rows": 120,
                                    "columns_count": 8,
                                    "headers": ["enterprise", "product", "selected_price"],
                                    "key_columns": ["enterprise", "product", "selected_price"],
                                    "summary": "Selected result table contains enterprise, product and price fields.",
                                    "business_value": "Supports selected result and price analysis.",
                                }
                            ],
                        }
                    ],
                    "warnings": [],
                }
            ],
            "auxiliary_materials": [],
            "combined_key_facts": [],
            "report_focus": ["Single primary material"],
            "warnings": [],
            "generation_guidance": {},
        }

        compact = main_module._compact_evidence_pack_for_dify(pack)
        primary = compact["primary_materials"][0]

        self.assertEqual(primary["content_text"], content_text)
        self.assertGreater(len(primary["content_text"]), 8000)
        self.assertNotIn("primary_content_tail", {item["type"] for item in compact["omitted_content"]})
        self.assertLess(compact["compact_pack_chars"], 65000)

    def test_dify_input_uses_full_detail_for_small_pack(self) -> None:
        long_summary = "Attachment table describes enterprise product selected price operation path. " * 70
        long_table_summary = "Rows include enterprise name product name registration certificate selected price. " * 45
        pack = {
            "pack_id": "pack_small_full_input",
            "created_at": "2026-06-17 12:00:00",
            "primary_materials": [
                {
                    "material_role": "primary",
                    "menu_code": "m1",
                    "articleid": "a1",
                    "title": "Small notice with rich attachment",
                    "content_text": "Short body says deadline is 2026-07-01 and details are in attachment.",
                    "content_text_length": 68,
                    "attachments": [
                        {
                            "articleattid": "att1",
                            "filename": "result-table.xlsx",
                            "core_attachment": True,
                            "business_type": "selected result",
                            "parse_status": "parsed_table_summary",
                            "summary": long_summary,
                            "key_facts": ["deadline 2026-07-01", "attachment contains operation path"],
                            "important_sections": ["operation path: login, maintain product, submit review"],
                            "table_summaries": [
                                {
                                    "sheet_name": "result",
                                    "rows": 500,
                                    "columns_count": 9,
                                    "headers": ["enterprise", "product", "registration_cert", "selected_price", "operation"],
                                    "key_columns": ["enterprise", "product", "selected_price", "operation"],
                                    "summary": long_table_summary,
                                    "business_value": "Supports product, price and operation-step reporting.",
                                }
                            ],
                        }
                    ],
                }
            ],
            "auxiliary_materials": [],
            "generation_guidance": {},
        }

        compact = main_module._compact_evidence_pack_for_dify(pack)
        attachment = compact["primary_materials"][0]["attachments"][0]
        table = attachment["table_summaries"][0]

        self.assertEqual(compact["input_strategy"], "attachment_led")
        self.assertTrue(compact["attachment_led"])
        self.assertFalse(compact["compression_applied"])
        self.assertEqual(attachment["summary"], long_summary)
        self.assertEqual(table["summary"], long_table_summary)
        self.assertEqual(compact["generation_guidance"]["detail_policy"], "attachment_led_full_core_attachment")
        self.assertIn("1500-2500", compact["generation_guidance"]["target_report_length"])
        diagnostics = build_pack_diagnostics(pack, compact)
        self.assertEqual(diagnostics["input_strategy"], "attachment_led")

    def test_dify_input_uses_essential_size_not_duplicate_bulk_for_full_input(self) -> None:
        content_text = "primary deadline price operation step " * 900
        attachment_summary = "core attachment selected result enterprise product price " * 420
        duplicated_evidence = {"attachment_summaries": [{"summary": "duplicate secondary evidence " * 5000}]}
        pack = {
            "pack_id": "pack_duplicate_bulk_small_essential",
            "primary_materials": [
                {
                    "material_role": "primary",
                    "menu_code": "m1",
                    "articleid": "p1",
                    "title": "Primary notice",
                    "content_text": content_text,
                    "content_text_length": len(content_text),
                    "attachments": [
                        {
                            "articleattid": "att1",
                            "filename": "selected-result.xlsx",
                            "core_attachment": True,
                            "business_type": "selected result",
                            "parse_status": "parsed_summary",
                            "summary": attachment_summary,
                            "key_facts": ["selected price and enterprise fields are present"],
                            "important_sections": ["operation path and deadline are listed in the attachment"],
                            "table_summaries": [],
                        }
                    ],
                }
            ],
            "auxiliary_materials": [
                {
                    "menu_code": "m2",
                    "articleid": "a1",
                    "title": "Auxiliary reference",
                    "content_text": "auxiliary full body should not drive strategy " * 2000,
                    "content_summary": "auxiliary summary only",
                    "relation_to_primary": "background reference",
                    "relevance_score": 0.6,
                    "relevant_snippets": ["similar price operation"],
                    "usable_points": ["background only"],
                    "attachments": [],
                }
            ],
            "primary_evidence": duplicated_evidence,
            "attachment_evidence": duplicated_evidence,
            "generation_guidance": {},
        }

        compact = main_module._compact_evidence_pack_for_dify(pack)

        self.assertGreater(compact["original_pack_chars"], 65000)
        self.assertEqual(compact["input_strategy"], "full_input")
        self.assertFalse(compact["compression_applied"])
        self.assertEqual(compact["primary_materials"][0]["content_text"], content_text)
        self.assertEqual(compact["primary_materials"][0]["attachments"][0]["summary"], attachment_summary)
        self.assertLessEqual(compact["dify_relevant_input_chars"], compact["full_input_max_chars"])

    def test_dify_input_uses_safe_compact_near_hard_limit_and_keeps_core_attachment(self) -> None:
        primary_text = "primary procurement rule deadline execution " * 900
        core_summary = "core attachment enterprise product selected price purchase volume " * 550
        aux_body = "auxiliary historical background should be compacted " * 800
        pack = {
            "pack_id": "pack_safe_compact_near_limit",
            "primary_materials": [
                {
                    "menu_code": "m1",
                    "articleid": "p1",
                    "title": "Primary notice",
                    "content_text": primary_text,
                    "content_text_length": len(primary_text),
                    "attachments": [
                        {
                            "articleattid": "att1",
                            "filename": "core-price-table.xlsx",
                            "core_attachment": True,
                            "business_type": "price table",
                            "parse_status": "parsed_table_summary",
                            "summary": core_summary,
                            "table_summaries": [
                                {
                                    "sheet_name": "result",
                                    "rows": 4800,
                                    "columns_count": 8,
                                    "headers": ["enterprise", "product", "selected_price", "purchase_volume"],
                                    "key_columns": ["enterprise", "product", "selected_price", "purchase_volume"],
                                    "summary": "table summary " * 120,
                                    "business_value": "supports price and purchase volume analysis",
                                }
                            ],
                        }
                    ],
                }
            ],
            "auxiliary_materials": [
                {
                    "menu_code": "m2",
                    "articleid": "a1",
                    "title": "Auxiliary",
                    "content_text": aux_body,
                    "content_summary": "auxiliary summary",
                    "relation_to_primary": "background only",
                    "relevance_score": 0.4,
                    "relevant_snippets": ["similar historical rule"],
                    "usable_points": ["background only"],
                }
            ],
            "generation_guidance": {},
        }

        compact = main_module._compact_evidence_pack_for_dify(pack)

        self.assertEqual(compact["input_strategy"], "safe_compact")
        self.assertLessEqual(compact["compact_pack_chars"], compact["hard_limit_chars"])
        self.assertTrue(compact["primary_materials"][0]["attachments"][0]["table_summaries"])
        self.assertLessEqual(len(compact["auxiliary_materials"][0]["content_text"]), 700)
        self.assertIn("auxiliary_content_tail", {item["type"] for item in compact["omitted_content"]})
        self.assertIn("safe_compact", compact["generation_guidance"]["detail_policy"])

    def test_dify_input_marks_single_primary_large_table_as_table_heavy(self) -> None:
        pack = {
            "pack_id": "pack_table_heavy",
            "primary_materials": [
                {
                    "menu_code": "m1",
                    "articleid": "p1",
                    "title": "Large table notice",
                    "content_text": "primary body is long enough to avoid attachment-led mode " * 30,
                    "content_text_length": 1600,
                    "attachments": [
                        {
                            "articleattid": "att1",
                            "filename": "large-selected-result.csv",
                            "core_attachment": True,
                            "business_type": "selected result",
                            "parse_status": "parsed_table_summary",
                            "summary": "large table summary",
                            "table_summaries": [
                                {
                                    "sheet_name": "CSV",
                                    "rows": 30000,
                                    "columns_count": 9,
                                    "headers": ["enterprise", "product", "registration_cert", "selected_price", "purchase_volume"],
                                    "key_columns": ["enterprise", "product", "selected_price", "purchase_volume"],
                                    "contains_enterprise": True,
                                    "contains_product": True,
                                    "contains_price": True,
                                    "contains_purchase_volume": True,
                                    "field_stats": {
                                        "enterprise": {"columns": ["enterprise"], "non_empty_count": 3000, "unique_count": 200}
                                    },
                                    "summary": "large selected result table",
                                    "business_value": "supports enterprise product price and purchase volume analysis",
                                }
                            ],
                        }
                    ],
                }
            ],
            "auxiliary_materials": [],
            "generation_guidance": {},
        }

        compact = main_module._compact_evidence_pack_for_dify(pack)

        self.assertEqual(compact["input_strategy"], "table_heavy")
        self.assertTrue(compact["table_heavy"])
        self.assertIn("large_table_structured_summary", compact["generation_guidance"]["generation_mode"])
        table = compact["primary_materials"][0]["attachments"][0]["table_summaries"][0]
        self.assertEqual(table["field_stats"]["enterprise"]["unique_count"], 200)
        self.assertIn("recommended_report_usage", table)

    def test_dify_input_marks_large_multi_material_pack_for_staged_generation(self) -> None:
        attachments = [
            {
                "articleattid": f"att{i}",
                "filename": f"table-{i}.xlsx",
                "core_attachment": True,
                "parse_status": "parsed_table_summary",
                "summary": "enterprise product price rule " * 450,
                "table_summaries": [
                    {
                        "sheet_name": "result",
                        "rows": 1000,
                        "columns_count": 12,
                        "headers": ["enterprise", "product", "price", "province", "deadline"],
                        "key_columns": ["enterprise", "product", "price", "deadline"],
                        "summary": "selected product price and deadline table " * 180,
                        "business_value": "large table",
                    }
                ],
            }
            for i in range(6)
        ]
        pack = {
            "pack_id": "pack_large_staged",
            "primary_materials": [
                {"menu_code": "m", "articleid": "p1", "title": "Primary 1", "content_text": "primary one " * 9000, "attachments": attachments[:3]},
                {"menu_code": "m", "articleid": "p2", "title": "Primary 2", "content_text": "primary two " * 9000, "attachments": attachments[3:]},
            ],
            "auxiliary_materials": [],
            "generation_guidance": {},
        }

        compact = main_module._compact_evidence_pack_for_dify(pack)

        self.assertEqual(compact["input_strategy"], "staged_generation")
        self.assertTrue(compact["compression_applied"])
        self.assertLess(compact["compact_pack_chars"], 65000)
        self.assertIn("per_material_then_synthesis", compact["generation_guidance"]["generation_mode"])
        self.assertIn("3000-5000", compact["generation_guidance"]["target_report_length"])

    def test_dify_input_marks_two_primary_pack_for_staged_generation(self) -> None:
        pack = {
            "pack_id": "pack_two_primary_staged",
            "primary_materials": [
                {"menu_code": "m", "articleid": "p1", "title": "Primary 1", "content_text": "primary one " * 4000, "attachments": []},
                {"menu_code": "m", "articleid": "p2", "title": "Primary 2", "content_text": "primary two " * 4000, "attachments": []},
            ],
            "auxiliary_materials": [],
        }

        compact = main_module._compact_evidence_pack_for_dify(pack)

        self.assertEqual(compact["input_strategy"], "staged_generation")
        self.assertEqual(main_module._dify_input_strategy_from_pack(pack), "staged_generation")

    def test_pack_diagnostics_marks_short_body_with_rich_core_attachment_as_attachment_led(self) -> None:
        pack = {
            "pack_id": "pack_attachment_led",
            "primary_materials": [
                {
                    "menu_code": "m1",
                    "articleid": "a1",
                    "title": "安徽医疗服务价格项目通知",
                    "content_text": "正文较短",
                    "content_text_length": 4,
                    "attachments": [
                        {
                            "articleattid": "att1",
                            "filename": "十二类医疗服务价格项目目录.xlsx",
                            "core_attachment": True,
                            "parse_status": "parsed_table_summary",
                            "summary": "附件列明呼吸系统、神经系统等十二类医疗服务价格项目、医保支付类别、项目内涵和计价单位。" * 20,
                            "key_facts": ["附件是本通知的核心依据，正文主要承担发布说明作用。"],
                            "table_summaries": [
                                {
                                    "sheet_name": "价格项目目录",
                                    "rows": 1500,
                                    "columns_count": 12,
                                    "headers": ["项目编码", "项目名称", "项目内涵", "计价单位", "医保支付类别"],
                                    "key_columns": ["项目名称", "项目内涵", "医保支付类别"],
                                    "summary": "该表为十二类医疗服务价格项目目录，覆盖项目编码、名称、内涵、计价单位和医保支付类别。" * 20,
                                    "business_value": "可作为报告梳理价格项目整合范围和医保支付政策的核心依据。" * 10,
                                }
                            ],
                        }
                    ],
                }
            ],
            "auxiliary_materials": [],
        }

        diagnostics = build_pack_diagnostics(pack, main_module._compact_evidence_pack_for_dify(pack))
        codes = {item["code"] for item in diagnostics["diagnosis"]}

        self.assertIn("ATTACHMENT_LED_PRIMARY_MATERIAL", codes)
        self.assertNotIn("EVIDENCE_PRIMARY_TOO_SHORT", codes)
        self.assertTrue(diagnostics["primary_materials"][0]["attachment_led"])
        self.assertGreater(diagnostics["primary_materials"][0]["primary_attachment_evidence_chars"], 800)

    def test_dify_compact_limits_auxiliary_attachment_bulk(self) -> None:
        attachments = []
        for index in range(14):
            attachments.append(
                {
                    "articleattid": f"att-{index}",
                    "filename": f"{index}.price-table.doc",
                    "fileext": ".doc",
                    "parse_status": "parsed_summary",
                    "download_status": "stream_parsed",
                    "summary": "auxiliary attachment summary " * 300,
                    "table_summaries": [
                        {
                            "sheet_name": "Sheet1",
                            "rows": 200,
                            "columns_count": 10,
                            "headers": ["code", "item", "price", "insurance"],
                            "key_columns": ["price", "insurance"],
                            "summary": "table summary " * 100,
                            "business_value": "background only",
                        }
                    ],
                }
            )
        pack = {
            "pack_id": "pack_aux_bulk",
            "primary_materials": [{"menu_code": "m", "articleid": "p", "title": "Primary", "content_text": "primary body " * 500}],
            "auxiliary_materials": [
                {
                    "menu_code": "m",
                    "articleid": "a",
                    "title": "Auxiliary",
                    "content_text": "auxiliary body " * 800,
                    "content_summary": "auxiliary summary",
                    "attachments": attachments,
                    "attachment_summaries": attachments,
                }
            ],
        }

        compact = main_module._compact_evidence_pack_for_dify(pack)
        auxiliary = compact["auxiliary_materials"][0]

        self.assertLessEqual(len(auxiliary["attachments"]), 5)
        self.assertTrue(all(len(item["summary"]) <= 520 for item in auxiliary["attachments"]))
        self.assertLess(compact["compact_pack_chars"], 65000)
        omitted_types = {item["type"] for item in compact["omitted_content"]}
        self.assertIn("auxiliary_attachment_overflow", omitted_types)

    def test_pack_diagnostics_reports_evidence_level_and_attachment_status(self) -> None:
        pack = {
            "pack_id": "pack_20260612_diag1234",
            "created_at": "2026-06-12 10:00:00",
            "primary_materials": [
                {
                    "menu_code": "m1",
                    "articleid": "a1",
                    "title": "Primary",
                    "summary": "summary",
                    "content_text": "A" * 1200,
                    "content_text_length": 1200,
                    "key_facts": [{"name": "fact"}],
                    "attachments": [{"articleattid": "att1", "parse_status": "metadata_only"}],
                }
            ],
            "auxiliary_materials": [
                {
                    "menu_code": "m2",
                    "articleid": "a2",
                    "title": "Aux",
                    "summary": "",
                    "content_text": "B" * 300,
                    "attachments": [],
                }
            ],
            "combined_key_facts": [],
            "warnings": ["content warning"],
        }
        with tempfile.TemporaryDirectory() as tmpdir, patch.object(
            main_module, "_database_evidence_pack_dir", return_value=main_module.Path(tmpdir), create=True
        ):
            main_module._write_database_evidence_pack(pack)
            response = self.client.get("/analysis/packs/pack_20260612_diag1234/diagnostics")

        self.assertEqual(response.status_code, 200)
        body = response.json()
        diagnostics = body["diagnostics"]
        self.assertEqual(diagnostics["estimated_evidence_level"], "medium")
        self.assertEqual(diagnostics["attachment_count"], 1)
        self.assertEqual(diagnostics["metadata_only_attachment_count"], 1)
        self.assertEqual(diagnostics["primary_content_chars"], 1200)
        self.assertEqual(diagnostics["auxiliary_content_chars"], 300)
        self.assertGreater(diagnostics["weighted_evidence_chars"], 0)
        self.assertEqual(diagnostics["primary_attachment_summary_chars"], 0)
        self.assertIn("raw_total_content_chars", diagnostics)
        self.assertIn("dify_compact_pack_chars", diagnostics)
        self.assertIn("ATTACHMENTS_NOT_PARSED", {item["code"] for item in diagnostics["diagnosis"]})
        self.assertIn("PACK_WARNINGS_EXIST", {item["code"] for item in diagnostics["diagnosis"]})

    def test_analysis_run_page_progress_and_diagnostics(self) -> None:
        pack = {
            "pack_id": "pack_20260612_rundiag1",
            "created_at": "2026-06-12 10:00:00",
            "primary_materials": [
                {
                    "menu_code": "m1",
                    "articleid": "a1",
                    "title": "Primary",
                    "summary": "short",
                    "content_text": "A" * 6000,
                    "content_text_length": 6000,
                    "key_facts": [],
                    "attachments": [{"articleattid": "att1", "parse_status": "metadata_only"}],
                }
            ],
            "auxiliary_materials": [],
            "warnings": [],
        }
        run = {
            "success": True,
            "run_id": "run_20260612_diag1234",
            "pack_id": "pack_20260612_rundiag1",
            "status": "finished",
            "workflow_run_id": "wf-1",
            "report_title": "Short report",
            "report_markdown": "Too short.",
            "version": 1,
            "quality_check": {"passed": True, "issues": []},
            "generation_warnings": [],
            "remaining_issues": [],
            "created_at": "2026-06-12 10:01:00",
            "updated_at": "2026-06-12 10:02:00",
        }
        with tempfile.TemporaryDirectory() as tmpdir:
            root = main_module.Path(tmpdir)
            pack_dir = root / "packs"
            run_dir = root / "runs"
            with patch.object(main_module, "_database_evidence_pack_dir", return_value=pack_dir, create=True), patch.object(
                main_module, "_analysis_run_dir", return_value=run_dir, create=True
            ):
                main_module._write_database_evidence_pack(pack)
                main_module._write_analysis_run(run)

                page_response = self.client.get("/analysis-runs/run_20260612_diag1234")
                status_response = self.client.get("/analysis/runs/run_20260612_diag1234")
                diagnostics_response = self.client.get("/analysis/runs/run_20260612_diag1234/diagnostics")

        self.assertEqual(page_response.status_code, 200)
        self.assertIn("报告生成详情", page_response.text)
        self.assertEqual(status_response.status_code, 200)
        self.assertEqual(status_response.json()["progress"]["percent"], 100)
        self.assertEqual(diagnostics_response.status_code, 200)
        codes = {item["code"] for item in diagnostics_response.json()["diagnostics"]["diagnosis"]}
        self.assertIn("REPORT_TOO_SHORT_FOR_WEIGHTED_EVIDENCE", codes)
        self.assertIn("ATTACHMENTS_NOT_PARSED", codes)
        evidence = diagnostics_response.json()["diagnostics"]["evidence"]
        self.assertIn("weighted_evidence_chars", evidence)
        self.assertIn("raw_total_content_chars", evidence)
        self.assertIn("dify_compact_pack_chars", evidence)
        self.assertIn("primary_attachment_summary_chars", evidence)

    def test_run_diagnostics_reports_missing_material_coverage_items(self) -> None:
        pack = {
            "pack_id": "pack_coverage",
            "primary_materials": [
                {
                    "menu_code": "m1",
                    "articleid": "a1",
                    "title": "Coverage notice",
                    "audittime": "2026-06-01",
                    "areaname": "云南省",
                    "publicorg": "医保局",
                    "content_text": (
                        "执行时间为2026年6月1日。产品范围包括药物球囊和泌尿介入类耗材。"
                        "价格规则为按中选价格挂网。企业需要维护产品信息。医疗机构按中选结果采购。"
                        "未按要求维护可能暂停挂网。"
                    ),
                    "content_text_length": 83,
                    "key_facts": [{"name": "执行时间", "value": "2026年6月1日"}],
                    "important_passages": [{"reason": "执行", "text": "医疗机构按中选结果采购。"}],
                    "price_rules": ["价格规则为按中选价格挂网。"],
                    "time_requirements": ["执行时间为2026年6月1日。"],
                    "product_scope": ["产品范围包括药物球囊和泌尿介入类耗材。"],
                    "enterprise_requirements": ["企业需要维护产品信息。"],
                    "execution_requirements": ["医疗机构按中选结果采购。"],
                    "attachments": [
                        {
                            "articleattid": "att1",
                            "filename": "中选结果表.xlsx",
                            "core_attachment": True,
                            "parse_status": "parsed_table_summary",
                            "summary": "附件包含企业、产品和中选价字段。",
                            "table_summaries": [
                                {
                                    "sheet_name": "中选结果",
                                    "headers": ["企业名称", "产品名称", "中选价"],
                                    "key_columns": ["企业名称", "产品名称", "中选价"],
                                    "summary": "中选结果表包含企业、产品和中选价。",
                                    "business_value": "可用于价格和产品范围分析。",
                                }
                            ],
                        }
                    ],
                }
            ],
            "auxiliary_materials": [],
            "warnings": [],
        }
        record = {
            "report_title": "Coverage report",
            "report_markdown": "本报告仅说明执行时间为2026年6月1日。",
            "quality_check": {"passed": True, "issues": []},
            "generation_warnings": [],
            "remaining_issues": [],
            "version": 1,
        }

        diagnostics = build_run_diagnostics(record, pack, main_module._compact_evidence_pack_for_dify(pack))

        coverage = diagnostics["coverage"]
        self.assertLess(coverage["coverage_score"], 80)
        self.assertTrue(coverage["is_report_too_short_by_coverage"])
        missing_labels = {item["label"] for item in coverage["missing_items"]}
        self.assertIn("价格规则", missing_labels)
        self.assertIn("产品范围", missing_labels)
        self.assertIn("附件表格摘要", missing_labels)
        self.assertIn("REPORT_MISSING_CORE_COVERAGE", {item["code"] for item in diagnostics["diagnosis"]})

    def test_run_diagnostics_reports_source_fidelity_and_blocks_unsupported_facts(self) -> None:
        pack = {
            "primary_materials": [
                {
                    "title": "天津市执行医用耗材集采结果的通知",
                    "content_text": "天津市医保局明确自2026年6月1日起执行医用耗材集采结果，企业需关注产品范围和执行时间。",
                    "summary": "天津市执行医用耗材集采结果。",
                    "attachments": [],
                }
            ],
            "auxiliary_materials": [],
        }
        record = {
            "status": "finished",
            "report_title": "天津市执行医用耗材集采结果分析",
            "report_markdown": (
                "## 导语\n天津市医保局明确自2026年7月1日起执行医用耗材集采结果。\n\n"
                "## 企业影响分析\n企业需要关注产品范围和执行时间。"
            ),
            "quality_check": {"passed": True, "issues": []},
            "generation_warnings": [],
            "remaining_issues": [],
            "version": 1,
        }

        diagnostics = build_run_diagnostics(record, pack, main_module._compact_evidence_pack_for_dify(pack))

        gate = diagnostics["quality_gate"]
        self.assertEqual(gate["deliverable_status"], "needs_manual_review")
        self.assertLess(gate["source_fidelity_score"], 100)
        self.assertGreaterEqual(gate["unsupported_fact_count"], 1)
        self.assertIn("UNSUPPORTED_FACT", {item["code"] for item in gate["blocking_issues"]})

    def test_run_diagnostics_detects_summary_only_report(self) -> None:
        pack = {
            "primary_materials": [
                {
                    "title": "河南调整医用耗材申报挂网操作流程的通知",
                    "content_text": "河南省调整医用耗材申报挂网操作流程，企业通过联审通办提交申报。",
                    "summary": "调整申报挂网操作流程。",
                    "attachments": [],
                }
            ],
            "auxiliary_materials": [],
        }
        record = {
            "status": "finished",
            "report_title": "河南调整医用耗材申报挂网操作流程分析",
            "report_markdown": "河南省调整医用耗材申报挂网操作流程，企业通过联审通办提交申报。",
            "quality_check": {"passed": True, "issues": []},
            "generation_warnings": [],
            "remaining_issues": [],
            "version": 1,
        }

        diagnostics = build_run_diagnostics(record, pack, main_module._compact_evidence_pack_for_dify(pack))

        gate = diagnostics["quality_gate"]
        self.assertTrue(gate["summary_only_risk"])
        self.assertLess(gate["analysis_depth_score"], 50)
        self.assertEqual(gate["deliverable_status"], "needs_manual_review")
        self.assertIn("SUMMARY_ONLY_REPORT", {item["code"] for item in gate["blocking_issues"]})

    def test_run_diagnostics_counts_evidence_backed_analysis(self) -> None:
        pack = {
            "primary_materials": [
                {
                    "title": "天津市执行医用耗材集采结果的通知",
                    "content_text": "天津市医保局明确自2026年6月1日起执行医用耗材集采结果，企业需关注产品范围和执行时间。",
                    "summary": "天津市执行医用耗材集采结果。",
                    "attachments": [],
                }
            ],
            "auxiliary_materials": [],
        }
        record = {
            "status": "finished",
            "report_title": "天津市执行医用耗材集采结果分析",
            "report_markdown": (
                "## 导语\n天津市医保局明确自2026年6月1日起执行医用耗材集采结果。\n\n"
                "## 影响分析\n该通知对企业的主要影响在于需要围绕产品范围和执行时间调整供货安排。\n\n"
                "## 风险提示\n若企业未及时核对医用耗材集采结果，可能影响执行衔接。\n\n"
                "## 企业建议\n建议企业按通知要求核对天津市执行范围和时间节点。"
            ),
            "quality_check": {"passed": True, "issues": []},
            "generation_warnings": [],
            "remaining_issues": [],
            "version": 1,
        }

        diagnostics = build_run_diagnostics(record, pack, main_module._compact_evidence_pack_for_dify(pack))

        gate = diagnostics["quality_gate"]
        self.assertFalse(gate["summary_only_risk"])
        self.assertGreaterEqual(gate["analysis_depth_score"], 60)
        self.assertGreaterEqual(gate["evidence_backed_analysis_count"], 2)

    def test_local_quality_gate_downgrades_finished_result_with_unsupported_fact(self) -> None:
        pack = {
            "primary_materials": [
                {
                    "title": "天津市执行医用耗材集采结果的通知",
                    "content_text": "天津市医保局明确自2026年6月1日起执行医用耗材集采结果。",
                    "attachments": [],
                }
            ],
            "auxiliary_materials": [],
        }
        result = {
            "status": "finished",
            "report_title": "天津市执行医用耗材集采结果分析",
            "report_markdown": "## 导语\n天津市医保局明确自2026年7月1日起执行医用耗材集采结果。",
            "quality_check": {"passed": True, "issues": []},
            "generation_warnings": [],
            "warnings": [],
            "remaining_issues": [],
        }

        gated = main_module._apply_local_quality_gate_to_dify_result(result, pack)

        self.assertEqual(gated["status"], "needs_manual_review")
        self.assertFalse(gated["quality_check"]["passed"])
        self.assertEqual(gated["quality_gate"]["deliverable_status"], "needs_manual_review")
        self.assertIn("Q_LOCAL_QUALITY_GATE", {item["issue_id"] for item in gated["remaining_issues"]})

    def test_analysis_run_progress_is_smooth_while_running(self) -> None:
        with patch("app.diagnostics.datetime") as datetime_mock:
            datetime_mock.fromisoformat.side_effect = main_module.datetime.fromisoformat
            datetime_mock.now.return_value = main_module.datetime.fromisoformat("2026-06-12 10:00:30")
            progress = main_module.build_run_progress(
                {
                    "status": "running",
                    "created_at": "2026-06-12 10:00:00",
                    "updated_at": "2026-06-12 10:00:00",
                    "version": 1,
                    "quality_check": {"passed": None, "issues": []},
                }
            )

        self.assertGreaterEqual(progress["percent"], 65)
        self.assertLessEqual(progress["percent"], 82)
        self.assertEqual(progress["current_step"], "Dify 工作流执行中")
        by_key = {step["key"]: step for step in progress["steps"]}
        self.assertEqual(by_key["create_run"]["status"], "finished")
        self.assertEqual(by_key["fetch_pack"]["status"], "finished")
        self.assertEqual(by_key["generate_report"]["status"], "running")
        self.assertEqual(by_key["quality_check"]["status"], "pending")
        self.assertEqual(by_key["revision"]["status"], "pending")

    def test_analysis_run_progress_finished_does_not_duplicate_middle_timestamps(self) -> None:
        progress = main_module.build_run_progress(
            {
                "status": "finished",
                "created_at": "2026-06-12 10:00:00",
                "updated_at": "2026-06-12 10:02:00",
                "version": 1,
                "quality_check": {"passed": True, "issues": []},
            }
        )

        self.assertEqual(progress["percent"], 100)
        by_key = {step["key"]: step for step in progress["steps"]}
        self.assertEqual(by_key["create_run"]["timestamp"], "2026-06-12 10:00:00")
        self.assertEqual(by_key["fetch_pack"]["timestamp"], "2026-06-12 10:00:00")
        self.assertIsNone(by_key["generate_report"]["timestamp"])
        self.assertIsNone(by_key["quality_check"]["timestamp"])
        self.assertEqual(by_key["revision"]["status"], "skipped")
        self.assertIsNone(by_key["revision"]["timestamp"])
        self.assertEqual(by_key["final_output"]["timestamp"], "2026-06-12 10:02:00")

    def test_dify_timeout_error_gets_non_empty_fallback_report(self) -> None:
        record = {
            "success": True,
            "run_id": "run_timeout123",
            "pack_id": "pack_timeout123",
            "status": "running",
            "workflow_run_id": "",
            "report_title": "",
            "report_markdown": "",
            "quality_check": {"passed": None, "issues": []},
            "created_at": "2026-06-16 10:00:00",
            "updated_at": "2026-06-16 10:00:00",
        }
        pack = {
            "pack_id": "pack_timeout123",
            "primary_materials": [
                {
                    "title": "Large primary notice",
                    "content_text": "The notice states execution starts on 2026-07-01 and enterprises maintain products.",
                    "key_facts": [{"name": "execution date", "value": "2026-07-01"}],
                    "attachments": [
                        {
                            "filename": "price-table.xlsx",
                            "core_attachment": True,
                            "parse_status": "parsed_table_summary",
                            "summary": "Attachment includes enterprise, product, selected price and operation path.",
                            "key_facts": ["enterprise, product and selected price columns exist"],
                            "table_summaries": [
                                {
                                    "sheet_name": "result",
                                    "headers": ["enterprise", "product", "selected_price"],
                                    "key_columns": ["enterprise", "product", "selected_price"],
                                    "summary": "Selected result table contains enterprise product selected price.",
                                    "business_value": "Supports price and product detail reporting.",
                                }
                            ],
                        }
                    ],
                }
            ],
            "auxiliary_materials": [],
        }
        with tempfile.TemporaryDirectory() as tmpdir:
            root = main_module.Path(tmpdir)
            run_dir = root / "runs"
            pack_dir = root / "packs"
            with patch.object(main_module, "_analysis_run_dir", return_value=run_dir, create=True), patch.object(
                main_module, "_database_evidence_pack_dir", return_value=pack_dir, create=True
            ), patch.object(
            main_module,
            "_call_dify_workflow",
                side_effect=main_module.DifyWorkflowError("DIFY_TIMEOUT", "Dify workflow timed out", "timed out", status_code=504),
            create=True,
            ):
                main_module._write_database_evidence_pack(pack)
                main_module._write_analysis_run(record)
                main_module._execute_analysis_run_background("pack_timeout123", "run_timeout123")
                saved = main_module._read_analysis_run("run_timeout123")

        self.assertEqual(saved["status"], "needs_manual_review")
        self.assertTrue(saved["success"])
        self.assertGreater(len(saved["report_markdown"]), 300)
        self.assertEqual(saved["dify_error_code"], "DIFY_TIMEOUT")
        self.assertIn("Q_DIFY_CALL_FAILED_FALLBACK", {item["issue_id"] for item in saved["remaining_issues"]})

    def test_staged_analysis_run_watchdog_writes_fallback_before_late_dify_result(self) -> None:
        pack = {
            "pack_id": "pack_watchdog123",
            "input_strategy": "staged_generation",
            "primary_materials": [
                {"title": "Primary 1", "content_text": "第一份主材料说明执行规则、企业要求和时间节点。" * 20, "attachments": []},
                {"title": "Primary 2", "content_text": "第二份主材料说明挂网规则、价格要求和风险后果。" * 20, "attachments": []},
            ],
            "auxiliary_materials": [],
            "warnings": [],
        }
        record = {
            "success": True,
            "run_id": "run_watchdog123",
            "pack_id": "pack_watchdog123",
            "status": "running",
            "workflow_run_id": "",
            "report_title": "",
            "report_markdown": "",
            "quality_check": {"passed": None, "issues": []},
            "generation_warnings": [],
            "warnings": [],
            "remaining_issues": [],
            "version": 1,
        }
        config = {
            "base_url": "http://dify.local/v1",
            "api_key": "test-key",
            "endpoint": "/workflows/run",
            "response_mode": "blocking",
            "user": "test",
            "timeout_seconds": 1,
            "max_attempts": 1,
            "retry_backoff_seconds": 0,
            "staged_timeout_seconds": 1,
            "staged_max_attempts": 1,
        }

        def slow_call(pack_id: str, run_id: str, pack_arg: dict | None = None) -> dict:
            time.sleep(1.4)
            return {
                "workflow_run_id": "late-workflow",
                "status": "finished",
                "report_title": "late",
                "report_markdown": "# late",
                "quality_check": {"passed": True, "issues": []},
                "generation_warnings": [],
                "remaining_issues": [],
                "version": 1,
            }

        with tempfile.TemporaryDirectory() as tmpdir:
            root = main_module.Path(tmpdir)
            run_dir = root / "runs"
            pack_dir = root / "packs"
            with patch.object(main_module, "_analysis_run_dir", return_value=run_dir, create=True), patch.object(
                main_module, "_database_evidence_pack_dir", return_value=pack_dir, create=True
            ), patch.object(main_module, "_dify_config", return_value=config), patch.object(main_module, "_call_dify_workflow", slow_call, create=True), patch.dict(
                os.environ, {"DIFY_WATCHDOG_GRACE_SECONDS": "0"}
            ):
                main_module._write_database_evidence_pack(pack)
                main_module._write_analysis_run(record)
                main_module._execute_analysis_run_background("pack_watchdog123", "run_watchdog123")
                saved = main_module._read_analysis_run("run_watchdog123")

        self.assertEqual(saved["status"], "needs_manual_review")
        self.assertEqual(saved["dify_error_code"], "DIFY_TIMEOUT")
        self.assertNotEqual(saved["workflow_run_id"], "late-workflow")
        self.assertGreater(len(saved["report_markdown"]), 300)

    def test_full_input_analysis_run_watchdog_writes_fallback_before_late_dify_result(self) -> None:
        pack = {
            "pack_id": "pack_full_watchdog123",
            "input_strategy": "full_input",
            "primary_materials": [
                {"title": "Short original notice", "content_text": "notice facts and attachment instructions", "attachments": []},
            ],
            "auxiliary_materials": [],
            "warnings": [],
        }
        record = {
            "success": True,
            "run_id": "run_full_watchdog123",
            "pack_id": "pack_full_watchdog123",
            "status": "running",
            "workflow_run_id": "",
            "report_title": "",
            "report_markdown": "",
            "quality_check": {"passed": None, "issues": []},
            "generation_warnings": [],
            "warnings": [],
            "remaining_issues": [],
            "version": 1,
        }
        config = {
            "base_url": "http://dify.local/v1",
            "api_key": "test-key",
            "endpoint": "/workflows/run",
            "response_mode": "blocking",
            "user": "test",
            "timeout_seconds": 1,
            "max_attempts": 1,
            "retry_backoff_seconds": 0,
            "staged_timeout_seconds": 1,
            "staged_max_attempts": 1,
        }

        def slow_call(pack_id: str, run_id: str, pack_arg: dict | None = None) -> dict:
            time.sleep(1.4)
            return {
                "workflow_run_id": "late-full-workflow",
                "status": "finished",
                "report_title": "late",
                "report_markdown": "# late",
                "quality_check": {"passed": True, "issues": []},
                "generation_warnings": [],
                "remaining_issues": [],
                "version": 1,
            }

        with tempfile.TemporaryDirectory() as tmpdir:
            root = main_module.Path(tmpdir)
            run_dir = root / "runs"
            pack_dir = root / "packs"
            with patch.object(main_module, "_analysis_run_dir", return_value=run_dir, create=True), patch.object(
                main_module, "_database_evidence_pack_dir", return_value=pack_dir, create=True
            ), patch.object(main_module, "_dify_config", return_value=config), patch.object(main_module, "_call_dify_workflow", slow_call, create=True), patch.dict(
                os.environ, {"DIFY_WATCHDOG_GRACE_SECONDS": "0"}
            ):
                main_module._write_database_evidence_pack(pack)
                main_module._write_analysis_run(record)
                main_module._execute_analysis_run_background("pack_full_watchdog123", "run_full_watchdog123")
                saved = main_module._read_analysis_run("run_full_watchdog123")

        self.assertEqual(saved["status"], "needs_manual_review")
        self.assertEqual(saved["dify_error_code"], "DIFY_TIMEOUT")
        self.assertNotEqual(saved["workflow_run_id"], "late-full-workflow")
        self.assertGreater(len(saved["report_markdown"]), 300)

    def test_analysis_run_status_self_heals_timed_out_running_run(self) -> None:
        pack = {
            "pack_id": "pack_poll_watchdog123",
            "input_strategy": "staged_generation",
            "primary_materials": [
                {"title": "Long primary notice", "content_text": "procurement facts and price rules " * 100, "attachments": []},
                {"title": "Second primary notice", "content_text": "execution steps and enterprise duties " * 100, "attachments": []},
            ],
            "auxiliary_materials": [],
            "warnings": [],
        }
        record = {
            "success": True,
            "run_id": "run_poll_watchdog123",
            "pack_id": "pack_poll_watchdog123",
            "status": "running",
            "workflow_run_id": "",
            "report_title": "",
            "report_markdown": "",
            "quality_check": {"passed": None, "issues": []},
            "generation_warnings": [],
            "warnings": [],
            "remaining_issues": [],
            "version": 1,
            "created_at": "2026-01-01 00:00:00",
            "updated_at": "2026-01-01 00:00:00",
        }
        config = {
            "base_url": "http://dify.local/v1",
            "api_key": "test-key",
            "endpoint": "/workflows/run",
            "response_mode": "blocking",
            "user": "test",
            "timeout_seconds": 1,
            "max_attempts": 1,
            "retry_backoff_seconds": 0,
            "staged_timeout_seconds": 1,
            "staged_max_attempts": 1,
        }

        with tempfile.TemporaryDirectory() as tmpdir:
            root = main_module.Path(tmpdir)
            run_dir = root / "runs"
            pack_dir = root / "packs"
            with patch.object(main_module, "_analysis_run_dir", return_value=run_dir, create=True), patch.object(
                main_module, "_database_evidence_pack_dir", return_value=pack_dir, create=True
            ), patch.object(main_module, "_dify_config", return_value=config), patch.object(
                main_module,
                "_apply_local_quality_gate_to_dify_result",
                side_effect=AssertionError("timeout self-heal must not run the heavy quality gate"),
            ), patch.dict(
                os.environ, {"DIFY_WATCHDOG_GRACE_SECONDS": "0"}
            ):
                main_module._write_database_evidence_pack(pack)
                main_module._write_analysis_run(record)
                response = self.client.get("/analysis/runs/run_poll_watchdog123")
                saved = main_module._read_analysis_run("run_poll_watchdog123")

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()["status"], "needs_manual_review")
        self.assertEqual(saved["dify_error_code"], "DIFY_TIMEOUT")
        self.assertGreater(len(saved["report_markdown"]), 300)

    def test_call_dify_workflow_retries_timeout_before_success(self) -> None:
        calls = {"count": 0}

        class FakeResponse:
            def raise_for_status(self) -> None:
                return None

            def json(self) -> dict:
                return {
                    "workflow_run_id": "wf-retry",
                    "data": {
                        "status": "succeeded",
                        "outputs": {
                            "report_markdown": "## \u5bfc\u8bed\nretry success report",
                            "report_title": "retry report",
                        },
                    },
                }

        class FakeClient:
            def __init__(self, timeout: float):
                self.timeout = timeout

            def __enter__(self):
                return self

            def __exit__(self, exc_type, exc, tb) -> None:
                return None

            def post(self, *args, **kwargs):
                calls["count"] += 1
                if calls["count"] == 1:
                    raise main_module.httpx.TimeoutException("temporary timeout")
                return FakeResponse()

        config = {
            "base_url": "http://dify.local/v1",
            "api_key": "test-key",
            "endpoint": "/workflows/run",
            "response_mode": "blocking",
            "user": "test",
            "timeout_seconds": 1,
            "max_attempts": 2,
            "retry_backoff_seconds": 0,
        }

        with patch.object(main_module, "_dify_config", return_value=config), patch.object(main_module.httpx, "Client", FakeClient):
            result = main_module._call_dify_workflow("pack_retry123", "run_retry123")

        self.assertEqual(calls["count"], 2)
        self.assertEqual(result["workflow_run_id"], "wf-retry")
        self.assertEqual(result["report_title"], "retry report")

    def test_call_dify_workflow_caps_staged_generation_timeout_and_attempts(self) -> None:
        calls = {"count": 0, "timeouts": []}

        class FakeClient:
            def __init__(self, timeout: float):
                calls["timeouts"].append(timeout)

            def __enter__(self):
                return self

            def __exit__(self, exc_type, exc, tb) -> None:
                return None

            def post(self, *args, **kwargs):
                calls["count"] += 1
                raise main_module.httpx.TimeoutException("staged timeout")

        config = {
            "base_url": "http://dify.local/v1",
            "api_key": "test-key",
            "endpoint": "/workflows/run",
            "response_mode": "blocking",
            "user": "test",
            "timeout_seconds": 600,
            "max_attempts": 3,
            "retry_backoff_seconds": 0,
            "staged_timeout_seconds": 300,
            "staged_max_attempts": 1,
        }

        with patch.object(main_module, "_dify_config", return_value=config), patch.object(main_module.httpx, "Client", FakeClient):
            with self.assertRaises(main_module.DifyWorkflowError) as caught:
                main_module._call_dify_workflow("pack_staged123", "run_staged123", {"input_strategy": "staged_generation"})

        self.assertEqual(caught.exception.code, "DIFY_TIMEOUT")
        self.assertEqual(calls["count"], 1)
        self.assertEqual(calls["timeouts"], [300])

    def test_call_dify_workflow_has_wall_clock_timeout_for_blocking_connection(self) -> None:
        calls = {"count": 0}

        class FakeClient:
            def __init__(self, timeout: float):
                self.timeout = timeout

            def __enter__(self):
                return self

            def __exit__(self, exc_type, exc, tb) -> None:
                return None

            def post(self, *args, **kwargs):
                calls["count"] += 1
                time.sleep(2)
                raise AssertionError("wall timeout should return before this point")

        config = {
            "base_url": "http://dify.local/v1",
            "api_key": "test-key",
            "endpoint": "/workflows/run",
            "response_mode": "blocking",
            "user": "test",
            "timeout_seconds": 1,
            "max_attempts": 1,
            "retry_backoff_seconds": 0,
            "staged_timeout_seconds": 1,
            "staged_max_attempts": 1,
        }

        started = time.perf_counter()
        with patch.object(main_module, "_dify_config", return_value=config), patch.object(main_module.httpx, "Client", FakeClient):
            with self.assertRaises(main_module.DifyWorkflowError) as caught:
                main_module._call_dify_workflow("pack_blocking123", "run_blocking123", {"input_strategy": "staged_generation"})

        self.assertEqual(caught.exception.code, "DIFY_TIMEOUT")
        self.assertEqual(calls["count"], 1)
        self.assertLess(time.perf_counter() - started, 1.8)

    def test_analysis_run_calls_dify_and_persists_report(self) -> None:
        calls: list[tuple[str, str]] = []
        report_markdown = "\n\n".join(
            [
                "# Report title",
                "## 导语",
                "This report is a complete mocked report used to verify that a normal Dify response is persisted without fallback repair.",
                "## 一、Core Findings",
                "The selected material contains enough structured content for a normal report body. "
                "This paragraph intentionally has enough length and report structure so the fragment guard does not treat it as a partial revision output.",
                "## 二、Analysis",
                "The generated report includes a stable title, multiple sections, and a readable body. "
                "It should remain finished because this test is checking persistence of a valid workflow response.",
            ]
        )

        def fake_call(pack_id: str, run_id: str, pack: dict | None = None):
            calls.append((pack_id, run_id))
            return {
                "workflow_run_id": "wf-run-1",
                "status": "finished",
                "report_title": "Report title",
                "report_markdown": report_markdown,
                "version": 1,
                "quality_check": {"passed": True, "issues": []},
                "generation_warnings": ["metadata only"],
                "remaining_issues": [],
            }

        with tempfile.TemporaryDirectory() as tmpdir, patch.object(
            main_module, "_analysis_run_dir", return_value=main_module.Path(tmpdir), create=True
        ), patch.object(
            main_module, "_read_database_evidence_pack", return_value={"pack_id": "pack_20260612_abcdef1234"}, create=True
        ), patch.object(main_module, "_call_dify_workflow", fake_call, create=True):
            response = self.client.post("/analysis/run", json={"pack_id": "pack_20260612_abcdef1234"})

            self.assertEqual(response.status_code, 200)
            body = response.json()
            self.assertTrue(body["success"])
            self.assertTrue(body["run_id"].startswith("run_"))
            self.assertEqual(body["pack_id"], "pack_20260612_abcdef1234")
            self.assertEqual(body["status"], "running")

            wait_until(lambda: self.client.get(f"/analysis/runs/{body['run_id']}").json().get("status") == "finished")
            self.assertEqual(calls, [("pack_20260612_abcdef1234", body["run_id"])])

            status_response = self.client.get(f"/analysis/runs/{body['run_id']}")
            self.assertEqual(status_response.status_code, 200)
            self.assertEqual(status_response.json()["workflow_run_id"], "wf-run-1")
            self.assertEqual(status_response.json()["status"], "finished")

            report_response = self.client.get(f"/analysis/runs/{body['run_id']}/report")
            self.assertEqual(report_response.status_code, 200)
            report_body = report_response.json()
            self.assertEqual(report_body["report_markdown"], report_markdown)
            self.assertTrue(report_body["quality_check"]["passed"])
            self.assertIn("quality_gate", report_body)

    def test_analysis_run_download_exports_markdown_docx(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            run_dir = main_module.Path(tmpdir) / "runs"
            report_dir = main_module.Path(tmpdir) / "reports"
            run_dir.mkdir()
            with patch.object(main_module, "_analysis_run_dir", return_value=run_dir, create=True):
                main_module._write_analysis_run(
                    {
                        "success": True,
                        "run_id": "run_test1234",
                        "pack_id": "pack_test",
                        "status": "finished",
                        "report_title": "测试报告",
                        "report_markdown": "# 测试报告\n\n正文内容",
                        "version": 1,
                        "quality_check": {"passed": True, "issues": []},
                    }
                )
            with patch.object(main_module, "_analysis_run_dir", return_value=run_dir, create=True), patch.object(
                main_module, "REPORT_DIR", report_dir
            ):
                response = self.client.get("/analysis/runs/run_test1234/download")

        self.assertEqual(response.status_code, 200)
        self.assertIn(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            response.headers["content-type"],
        )
        self.assertGreater(len(response.content), 1000)

    def test_analysis_run_download_reuses_existing_docx_for_same_run_version(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            run_dir = main_module.Path(tmpdir) / "runs"
            report_dir = main_module.Path(tmpdir) / "reports"
            run_dir.mkdir()
            calls: list[tuple[str, str]] = []

            def fake_markdown_to_docx(markdown: str, path: main_module.Path, title: str) -> None:
                calls.append((markdown, title))
                path.write_bytes(b"PK\x03\x04cached docx content")

            with patch.object(main_module, "_analysis_run_dir", return_value=run_dir, create=True):
                main_module._write_analysis_run(
                    {
                        "success": True,
                        "run_id": "run_cache1234",
                        "pack_id": "pack_test",
                        "status": "finished",
                        "report_title": "cached report",
                        "report_markdown": "# cached report\n\nbody",
                        "version": 2,
                        "quality_check": {"passed": True, "issues": []},
                    }
                )
            with patch.object(main_module, "_analysis_run_dir", return_value=run_dir, create=True), patch.object(
                main_module, "REPORT_DIR", report_dir
            ), patch.object(main_module, "_markdown_to_docx", side_effect=fake_markdown_to_docx):
                first = self.client.get("/analysis/runs/run_cache1234/download")
                second = self.client.get("/analysis/runs/run_cache1234/download")

        self.assertEqual(first.status_code, 200)
        self.assertEqual(second.status_code, 200)
        self.assertEqual(len(calls), 1)
        self.assertEqual(first.content, second.content)

    def test_markdown_docx_uses_manual_report_style_markers(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            path = main_module.Path(tmpdir) / "styled.docx"
            main_module._markdown_to_docx(
                "\n".join(
                    [
                        "## 重点内容",
                        "<blue>企业需关注申报截止时间。</blue>",
                        "",
                        "| 挂网状态 | 序号 | 对应情形 |",
                        "| --- | --- | --- |",
                        "| 暂停挂网 | 1 | 普通内容 |",
                        "| 暂停挂网 | 2 | <red>未进行价格联动的产品暂停挂网。</red> |",
                    ]
                ),
                path,
                "测试报告",
            )
            doc = Document(path)

        text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
        self.assertIn("声  明", text)
        self.assertIn("企业需关注申报截止时间。", text)
        self.assertEqual(len(doc.tables), 1)
        header_shading = doc.tables[0].cell(0, 0)._tc.tcPr.find(qn("w:shd"))
        self.assertEqual(header_shading.get(qn("w:fill")), "0070C0")
        red_runs = [
            run
            for row in doc.tables[0].rows
            for cell in row.cells
            for paragraph in cell.paragraphs
            for run in paragraph.runs
            if "未进行价格联动" in run.text
        ]
        self.assertTrue(red_runs)
        self.assertEqual(str(red_runs[0].font.color.rgb), "FF0000")

    def test_markdown_docx_cleans_analysis_highlight_span(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir, patch.dict(main_module.os.environ, {"ENABLE_ANALYSIS_HIGHLIGHT": "false"}):
            path = main_module.Path(tmpdir) / "highlight.docx"
            main_module._markdown_to_docx(
                '事实内容。<span class="analysis-highlight">这是分析内容。</span>',
                path,
                "测试报告",
            )
            doc = Document(path)

        text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
        self.assertIn("这是分析内容。", text)
        self.assertNotIn("analysis-highlight", text)
        self.assertNotIn("<span", text)

    def test_clean_model_output_removes_model_generated_material_note_and_declaration(self) -> None:
        markdown = "\n".join(
            [
                "# 测试报告",
                "",
                "正文第一段。",
                "",
                "## 资料说明",
                "",
                "本报告基于当前已解析的公告正文、附件摘要及结构化信息生成。",
                "",
                "## 声  明",
                "",
                "本报告仅供参考。",
            ]
        )

        cleaned = main_module._clean_model_output(markdown)

        self.assertIn("正文第一段。", cleaned)
        self.assertNotIn("资料说明", cleaned)
        self.assertNotIn("声  明", cleaned)

    def test_clean_model_output_removes_technical_attachment_parse_notes(self) -> None:
        markdown = "\n".join(
            [
                "# 安庆市医疗服务价格项目分析报告",
                "",
                "## 一、政策背景",
                "",
                "本次通知规范整合多类医疗服务价格项目，医疗机构需按新规则执行。",
                "",
                "> 注：数字来源于附件OCR识别，可能存在误差。",
                "",
                "| 支付类别 | 数量 | 说明 |",
                "| --- | --- | --- |",
                "| 乙类 | 数量待确认（OCR识别错误） | 以正式文件为准 |",
                "",
                "附件PDF因文件过大未解析，未能提取具体细则。",
                "该段是正常分析，应当保留。",
                "",
                "资料说明：附件仅解析到元数据层面。",
            ]
        )

        cleaned = main_module._clean_model_output(markdown)

        self.assertIn("本次通知规范整合多类医疗服务价格项目", cleaned)
        self.assertIn("该段是正常分析", cleaned)
        self.assertNotIn("OCR", cleaned)
        self.assertNotIn("附件PDF因", cleaned)
        self.assertNotIn("元数据", cleaned)
        self.assertNotIn("未解析", cleaned)
        self.assertNotIn("资料说明", cleaned)

    def test_clean_model_output_keeps_analysis_when_technical_note_is_same_line(self) -> None:
        markdown = "\n".join(
            [
                "# \u5b89\u5e86\u5e02\u533b\u7597\u670d\u52a1\u4ef7\u683c\u9879\u76ee\u5206\u6790\u62a5\u544a",
                "",
                "\u5b89\u5e86\u5e02\u5bf9\u5341\u4e8c\u7c7b\u533b\u7597\u670d\u52a1\u4ef7\u683c\u9879\u76ee\u8fdb\u884c\u89c4\u8303\u6574\u5408\uff0c\u533b\u7597\u673a\u6784\u9700\u6309\u65b0\u89c4\u5219\u66f4\u65b0\u6536\u8d39\u9879\u76ee\u3002\u6570\u5b57\u6765\u6e90\u4e8e\u9644\u4ef6OCR\u8bc6\u522b\uff0c\u53ef\u80fd\u5b58\u5728\u8bef\u5dee\u3002\u4f01\u4e1a\u5e94\u5173\u6ce8\u9879\u76ee\u5185\u6db5\u3001\u652f\u4ed8\u7c7b\u522b\u4e0e\u914d\u5957\u8017\u6750\u4f7f\u7528\u573a\u666f\u7684\u53d8\u5316\u3002",
            ]
        )

        cleaned = main_module._clean_model_output(markdown)

        self.assertIn("\u89c4\u8303\u6574\u5408", cleaned)
        self.assertIn("\u4f01\u4e1a\u5e94\u5173\u6ce8", cleaned)
        self.assertNotIn("OCR", cleaned)
        self.assertNotIn("\u8bef\u5dee", cleaned)

    def test_clean_model_output_removes_recognition_error_notes_in_tables(self) -> None:
        markdown = "\n".join(
            [
                "| 支付类别 | 项目数量 | 说明 |",
                "| --- | --- | --- |",
                "| 乙类 | 825项 | 个人先自付一定比例再报销。（注：原文此处存在识别错误，实际数量以正式文件为准。） |",
                "",
                "上述分类直接影响患者的自付水平和医院的收入结构。",
            ]
        )

        cleaned = main_module._clean_model_output(markdown)

        self.assertIn("乙类", cleaned)
        self.assertIn("上述分类直接影响", cleaned)
        self.assertNotIn("识别错误", cleaned)
        self.assertNotIn("正式文件为准", cleaned)

    def test_clean_model_output_removes_source_recognition_problem_notes(self) -> None:
        markdown = "\n".join(
            [
                "通知明确，将“颅内动脉瘤夹闭成形费”（注：原文为“来闭成形费”，根据医学常识修正）等项目纳入医保支付范围。",
                "部分项目按乙类管理（具体数量因原文识别问题暂无法确认），项目所标注价格为医保基金最高支付标准。",
                "医疗机构需完成价格库更新和收费流程调整。",
            ]
        )

        cleaned = main_module._clean_model_output(markdown)

        self.assertIn("通知明确", cleaned)
        self.assertIn("医疗机构需完成", cleaned)
        self.assertNotIn("根据医学常识", cleaned)
        self.assertNotIn("识别问题", cleaned)
        self.assertNotIn("暂无法确认", cleaned)

    def test_analysis_run_download_rejects_not_ready_report(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            run_dir = main_module.Path(tmpdir) / "runs"
            run_dir.mkdir()
            with patch.object(main_module, "_analysis_run_dir", return_value=run_dir, create=True):
                main_module._write_analysis_run(
                    {
                    "success": True,
                    "run_id": "run_pending1",
                    "pack_id": "pack_test",
                    "status": "running",
                    "report_markdown": "",
                    }
                )
                response = self.client.get("/analysis/runs/run_pending1/download")

        self.assertEqual(response.status_code, 409)
        self.assertEqual(response.json()["error"]["code"], "REPORT_NOT_READY")

    def test_normalize_dify_result_parses_json_string_fields(self) -> None:
        result = main_module._normalize_dify_result(
            {
                "workflow_run_id": "wf-1",
                "data": {
                    "status": "succeeded",
                    "outputs": {
                        "result": {
                            "status": "needs_manual_review",
                            "pack_id": "pack_test",
                            "report_title": "标题",
                            "report_markdown": "正文",
                            "version": 2,
                            "quality_check": '{"passed": false, "issues": [{"issue_id": "Q001"}]}',
                            "generation_warnings": '["附件仅元数据"]',
                            "remaining_issues": '[{"issue_id": "Q001"}]',
                        }
                    },
                },
            },
            "pack_test",
        )

        self.assertEqual(result["quality_check"]["passed"], False)
        self.assertEqual(result["generation_warnings"], ["附件仅元数据"])
        self.assertEqual(result["remaining_issues"][0]["issue_id"], "Q001")

    def test_unusable_dify_report_gets_fallback_from_evidence_pack(self) -> None:
        pack = {
            "pack_id": "pack_test",
            "primary_materials": [
                {
                    "title": "天津市医保局市卫生健康委市药监局关于执行集采结果有关工作的通知",
                    "audittime": "2026-06-15",
                    "areaname": "天津市",
                    "publicorg": "天津市医保局",
                    "projectphase": "执行",
                    "projecttype": "医用耗材",
                    "category": "集采",
                    "summary": "执行国家组织药物球囊、泌尿介入类和省际联盟相关医用耗材集采结果。",
                    "content_text": "本通知明确执行国家组织药物球囊、泌尿介入类和省际联盟组织硬脑膜补片类等医用耗材集采结果。医疗机构应按要求执行中选结果，企业需关注产品范围、执行时间和医保支付标准衔接。",
                    "key_facts": [{"name": "地区", "value": "天津市"}],
                    "attachments": [
                        {"filename": "附件1.pdf", "parse_status": "auth_required", "warnings": ["需要登录态"]},
                    ],
                }
            ],
            "auxiliary_materials": [],
            "warnings": [],
        }
        result = {
            "workflow_run_id": "wf-empty",
            "status": "finished",
            "pack_id": "pack_test",
            "report_title": "...",
            "report_markdown": "...",
            "version": 2,
            "quality_check": {"passed": True, "issues": []},
            "generation_warnings": [],
            "warnings": [],
            "remaining_issues": [],
        }

        repaired = main_module._repair_unusable_dify_result(result, pack)

        self.assertEqual(repaired["status"], "needs_manual_review")
        self.assertGreater(len(repaired["report_markdown"]), 300)
        self.assertIn("天津市医保局", repaired["report_title"])
        self.assertNotIn("由于本次自动生成结果未形成完整正文", repaired["report_markdown"])
        self.assertNotIn("供人工复核", repaired["report_markdown"])
        self.assertNotIn("Dify", repaired["report_markdown"])
        self.assertNotIn("后端兜底版本", repaired["report_markdown"])
        self.assertTrue(repaired["warnings"])
        self.assertFalse(repaired["quality_check"]["passed"])
        self.assertEqual(repaired["remaining_issues"][0]["issue_id"], "Q_DIFY_FRAGMENTARY_REPORT")

    def test_list_starting_dify_report_without_structure_gets_fallback(self) -> None:
        pack = {
            "pack_id": "pack_list_fragment",
            "primary_materials": [
                {
                    "title": "海南医药采购接待日活动通知",
                    "audittime": "2026-05-21",
                    "areaname": "海南省",
                    "content_text": "本通知说明医药采购接待日活动安排、报名方式、问题征集表提交和现场咨询流程。" * 60,
                    "summary": "活动面向企业收集业务需求和问题建议。",
                    "key_facts": [{"name": "活动流程", "value": "先解答邮件提交问题，再解答现场提问。"}],
                    "attachments": [
                        {
                            "filename": "业务需求和问题建议征集表.docx",
                            "summary": "附件为业务需求和问题建议征集表模板。",
                            "core_attachment": True,
                        }
                    ],
                }
            ],
            "auxiliary_materials": [],
            "warnings": [],
        }
        result = {
            "workflow_run_id": "wf-list-fragment",
            "status": "needs_manual_review",
            "pack_id": "pack_list_fragment",
            "report_title": "海南医药采购接待日活动通知分析报告",
            "report_markdown": "\n".join(
                [
                    "- **接待范围**：涵盖政策与经办领域",
                    "- **活动流程**：优先解答已通过邮件提交的问题",
                    "- **报名方式**：参会人员须扫描通知二维码报名",
                    "企业应关注材料提交和现场咨询安排。" * 80,
                ]
            ),
            "version": 1,
            "quality_check": {"passed": False, "issues": [{"issue_id": "Q001"}]},
            "generation_warnings": [],
            "warnings": [],
            "remaining_issues": [{"issue_id": "Q001"}],
        }

        repaired = main_module._repair_unusable_dify_result(result, pack)

        self.assertEqual(repaired["status"], "needs_manual_review")
        self.assertIn("## 导语", repaired["report_markdown"])
        self.assertEqual(repaired["remaining_issues"][0]["issue_id"], "Q_DIFY_FRAGMENTARY_REPORT")

    def test_dify_success_without_report_markdown_can_be_repaired(self) -> None:
        raw = {
            "workflow_run_id": "wf-missing-report",
            "data": {
                "status": "succeeded",
                "outputs": {
                    "quality_check": '{"passed": false, "issues": [{"issue_id": "Q_MISSING_REPORT"}]}',
                    "generation_warnings": '["missing report_markdown"]',
                },
            },
        }

        result = main_module._normalize_dify_result(raw, "pack_missing_report")

        self.assertEqual(result["workflow_run_id"], "wf-missing-report")
        self.assertEqual(result["pack_id"], "pack_missing_report")
        self.assertEqual(result["report_markdown"], "")
        self.assertEqual(result["quality_check"]["passed"], False)
        self.assertEqual(result["remaining_issues"], [])

    def test_fragmentary_dify_revision_gets_fallback_from_attachment_rich_pack(self) -> None:
        pack = {
            "primary_materials": [
                {
                    "title": "\u5b89\u5e86\u5e02\u89c4\u8303\u6574\u5408\u5341\u4e8c\u7c7b\u533b\u7597\u670d\u52a1\u4ef7\u683c\u9879\u76ee\u7684\u901a\u77e5",
                    "content_text": "\u8be6\u60c5\u8bf7\u89c1\u9644\u4ef6\u3002",
                    "attachments": [
                        {
                            "filename": "\u5b89\u5e86\u5e02\u5341\u4e8c\u7c7b\u533b\u7597\u670d\u52a1\u4ef7\u683c\u9879\u76ee.pdf",
                            "core_attachment": True,
                            "summary": "\u9644\u4ef6\u660e\u786e\u89c4\u8303\u6574\u5408\u547c\u5438\u7cfb\u7edf\u3001\u795e\u7ecf\u7cfb\u7edf\u7b49\u5341\u4e8c\u7c7b\u533b\u7597\u670d\u52a1\u4ef7\u683c\u9879\u76ee\uff0c\u81ea2026\u5e746\u67081\u65e5\u8d77\u6267\u884c\u3002"
                            "\u5176\u4e2d\u90e8\u5206\u9879\u76ee\u5b9e\u884c\u6700\u9ad8\u653f\u5e9c\u6307\u5bfc\u4ef7\uff0c\u90e8\u5206\u9879\u76ee\u5b9e\u884c\u5e02\u573a\u8c03\u8282\u4ef7\uff0c\u5e76\u533a\u5206\u7532\u7c7b\u3001\u4e59\u7c7b\u548c\u4e0d\u4e88\u652f\u4ed8\u9879\u76ee\u3002",
                            "key_facts": [
                                "\u6267\u884c\u65f6\u95f4\u4e3a2026\u5e746\u67081\u65e5\u3002",
                                "\u516c\u7acb\u533b\u7597\u673a\u6784\u9700\u5b8c\u6210\u4ef7\u683c\u516c\u793a\u548c\u6536\u8d39\u7cfb\u7edf\u66f4\u65b0\u3002",
                                "\u5e02\u573a\u8c03\u8282\u4ef7\u9879\u76ee\u9700\u5411\u533b\u4fdd\u90e8\u95e8\u5907\u6848\u3002",
                            ],
                            "important_sections": [
                                {"title": "\u4ef7\u683c\u7ba1\u7406", "summary": "\u9879\u76ee\u6309\u6700\u9ad8\u653f\u5e9c\u6307\u5bfc\u4ef7\u548c\u5e02\u573a\u8c03\u8282\u4ef7\u5206\u7c7b\u7ba1\u7406\u3002"},
                                {"title": "\u533b\u4fdd\u652f\u4ed8", "summary": "\u7532\u7c7b\u3001\u4e59\u7c7b\u548c\u4e0d\u4e88\u652f\u4ed8\u9879\u76ee\u9002\u7528\u4e0d\u540c\u533b\u4fdd\u652f\u4ed8\u89c4\u5219\u3002"},
                            ],
                        }
                    ],
                }
            ],
            "auxiliary_materials": [],
        }
        result = {
            "status": "needs_manual_review",
            "report_title": "\u5b89\u5e86\u5e02\u533b\u7597\u670d\u52a1\u4ef7\u683c\u9879\u76ee\u5206\u6790",
            "report_markdown": "| \u652f\u4ed8\u7c7b\u522b | \u9879\u76ee\u6570\u91cf |\n| --- | --- |\n| \u7532\u7c7b | 456 |",
            "quality_check": {"passed": False, "issues": []},
            "generation_warnings": [],
            "warnings": [],
        }

        repaired = main_module._repair_unusable_dify_result(result, pack)

        self.assertEqual(repaired["status"], "needs_manual_review")
        self.assertIn("\u9644\u4ef6\u8981\u70b9\u8865\u5145", repaired["report_markdown"])
        self.assertIn("\u6267\u884c\u65f6\u95f4\u4e3a2026\u5e746\u67081\u65e5", repaired["report_markdown"])
        self.assertGreater(len(repaired["report_markdown"]), 500)
        self.assertNotIn("evidence_pack", repaired["report_markdown"])
        self.assertNotIn("Dify", repaired["report_markdown"])
        self.assertNotIn("\u5143\u6570\u636e", repaired["report_markdown"])

    def test_report_starting_from_second_section_is_treated_as_fragment(self) -> None:
        pack = {
            "primary_materials": [
                {
                    "title": "安庆市规范整合十二类医疗服务价格项目的通知",
                    "content_text": "详情请见附件。",
                    "attachments": [
                        {
                            "filename": "价格项目.pdf",
                            "core_attachment": True,
                            "summary": "附件说明十二类医疗服务价格项目自2026年6月1日起执行，涉及政府指导价、市场调节价和医保支付分类。",
                            "key_facts": ["执行时间为2026年6月1日。"],
                        }
                    ],
                }
            ],
            "auxiliary_materials": [],
        }
        result = {
            "status": "needs_manual_review",
            "report_title": "安庆市医疗服务价格项目分析",
            "report_markdown": "### 二、价格管理规则\n\n通知明确了政府指导价和市场调节价两类机制。\n\n### 三、医保支付范围\n\n医保支付分类影响患者自付比例。",
            "quality_check": {"passed": False, "issues": []},
            "generation_warnings": [],
            "warnings": [],
        }

        repaired = main_module._repair_unusable_dify_result(result, pack)

        self.assertEqual(repaired["status"], "needs_manual_review")
        self.assertIn("## 导语", repaired["report_markdown"])
        self.assertIn("附件要点补充", repaired["report_markdown"])
        self.assertGreater(len(repaired["report_markdown"]), len(result["report_markdown"]))

    def test_user_feedback_revision_updates_latest_report_and_keeps_previous_version(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            root = main_module.Path(tmpdir)
            run_dir = root / "runs"
            pack_dir = root / "packs"
            run_dir.mkdir()
            pack_dir.mkdir()
            with patch.object(main_module, "_analysis_run_dir", return_value=run_dir, create=True), patch.object(
                main_module, "_database_evidence_pack_dir", return_value=pack_dir, create=True
            ):
                main_module._write_database_evidence_pack({"pack_id": "pack_20260615_revise1", "primary_materials": [], "auxiliary_materials": []})
                main_module._write_analysis_run(
                    {
                        "success": True,
                        "run_id": "run_20260615_revise1",
                        "pack_id": "pack_20260615_revise1",
                        "status": "finished",
                        "report_title": "旧标题",
                        "report_markdown": "旧正文",
                        "version": 1,
                        "quality_check": {"passed": True, "issues": []},
                        "created_at": "2026-06-15 09:00:00",
                        "updated_at": "2026-06-15 09:01:00",
                    }
                )

                def fake_revision(*args, **kwargs):
                    return {
                        "report_title": "新标题",
                        "report_markdown": '<span class="analysis-highlight">新增企业影响分析。</span>',
                        "warnings": [],
                    }

                with patch.object(main_module, "_call_dify_revision_workflow", fake_revision, create=True):
                    response = self.client.post(
                        "/analysis/runs/run_20260615_revise1/revise",
                        json={"feedback": "增加企业影响分析", "analysis_highlight": True},
                    )
                    report_response = self.client.get("/analysis/runs/run_20260615_revise1/report")
                    record = main_module._read_analysis_run("run_20260615_revise1")

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()["version"], 2)
        self.assertIn("analysis-highlight", response.json()["report_markdown"])
        self.assertEqual(report_response.json()["version"], 2)
        self.assertEqual(record["report_versions"][0]["version"], 1)
        self.assertEqual(record["revisions"][0]["feedback"], "增加企业影响分析")

    def test_records_ui_defaults_to_cached_prepare_and_has_reparse_button(self) -> None:
        html = (main_module.Path(main_module.__file__).resolve().parent / "static" / "records.html").read_text(encoding="utf-8")

        self.assertIn("force_refresh_attachments", html)
        self.assertIn("reparseAttachmentsBtn", html)
        self.assertIn("\u91cd\u65b0\u89e3\u6790\u9644\u4ef6\u5e76\u751f\u6210\u8bc1\u636e\u5305", html)
        self.assertIn(".doc", html)
        self.assertNotIn("force_refresh_attachments: true,\n      };", html)

    def test_analysis_run_rejects_missing_pack(self) -> None:
        with patch.object(
            main_module,
            "_read_database_evidence_pack",
            side_effect=main_module.HTTPException(status_code=404, detail="evidence pack not found"),
            create=True,
        ):
            response = self.client.post("/analysis/run", json={"pack_id": "pack_20260612_missing1"})

        self.assertEqual(response.status_code, 404)
        body = response.json()
        self.assertFalse(body["success"])
        self.assertEqual(body["error"]["code"], "PACK_NOT_FOUND")

    def test_analysis_run_returns_structured_error_when_dify_not_configured(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir, patch.dict(
            main_module.os.environ,
            {
                "DIFY_BASE_URL": "http://192.168.34.86/v1",
                "DIFY_WORKFLOW_API_KEY": "",
                "DIFY_REPORT_WORKFLOW_ENDPOINT": "/workflows/run",
            },
        ), patch.object(
            main_module, "_analysis_run_dir", return_value=main_module.Path(tmpdir), create=True
        ), patch.object(
            main_module, "_read_database_evidence_pack", return_value={"pack_id": "pack_20260612_abcdef1234"}, create=True
        ):
            response = self.client.post("/analysis/run", json={"pack_id": "pack_20260612_abcdef1234"})
            self.assertEqual(response.status_code, 200)
            body = response.json()
            self.assertTrue(body["success"])
            self.assertEqual(body["status"], "running")
            run_id = body["run_id"]

            wait_until(lambda: self.client.get(f"/analysis/runs/{run_id}").json().get("status") == "needs_manual_review")
            status_response = self.client.get(f"/analysis/runs/{run_id}")
            report_response = self.client.get(f"/analysis/runs/{run_id}/report")

        self.assertEqual(status_response.status_code, 200)
        status_body = status_response.json()
        self.assertTrue(status_body["success"])
        self.assertEqual(status_body["status"], "needs_manual_review")
        self.assertEqual(status_body["error_message"], "Dify API 配置不完整")
        self.assertEqual(status_body["dify_error_code"], "DIFY_NOT_CONFIGURED")
        self.assertEqual(report_response.status_code, 200)
        self.assertGreater(len(report_response.json()["report_markdown"]), 100)


if __name__ == "__main__":
    unittest.main()
